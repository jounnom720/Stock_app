# v5.13.7 안정화 리팩터링본 / 중복 함수 정리 / 버전 표기 통일 / 배포 안정성 점검
import io
import json
import math
import os
import re
import html
from pathlib import Path
from datetime import datetime, timedelta
from zoneinfo import ZoneInfo
from urllib.parse import urljoin
from concurrent.futures import ThreadPoolExecutor, as_completed

import pandas as pd
import numpy as np
import requests
import streamlit as st

try:
    import plotly.graph_objects as go
    from plotly.subplots import make_subplots
    PLOTLY_AVAILABLE = True
except Exception:
    go = None
    make_subplots = None
    PLOTLY_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except Exception:
    BeautifulSoup = None
    BS4_AVAILABLE = False

try:
    from streamlit_plotly_events import plotly_events
    PLOTLY_CLICK_AVAILABLE = True
except Exception:
    PLOTLY_CLICK_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except Exception:
    Document = None
    Pt = None
    Inches = None
    RGBColor = None
    WD_ALIGN_PARAGRAPH = None
    WD_TABLE_ALIGNMENT = None
    WD_CELL_VERTICAL_ALIGNMENT = None
    qn = None
    DOCX_AVAILABLE = False

APP_VERSION = "v5.14.4"  # 월간 리포트 본문 고도화 / html import 오류 수정 / 리포트 가독성 보완 / 실시간 시세 보정 유지


# -----------------------------------
# v5.13.7 안정화 메모
# - 중복 함수 정의 2건 정리: 야후실시간호가가져오기, 네이버시장지표목록가져오기
# - 최상단 버전 주석과 APP_VERSION 표기 통일
# - 기존 기능 로직은 유지하여 v5.13.5와의 실행 호환성 우선
# -----------------------------------

st.set_page_config(page_title=f"투자 분석 시스템 {APP_VERSION}", layout="wide")

# -----------------------------------
# v5.13.7 안정화 스타일 세트
# - 전체 폰트 볼드 완화
# - 모니터 카드 숫자/제목 계층 정리
# - 버튼/표/Metric 가독성 개선
# -----------------------------------
st.markdown(
    """
    <style>
    html, body, [class*="css"] {
        font-weight: 450;
        letter-spacing: -0.01em;
    }

    h1 {
        font-weight: 580 !important;
        letter-spacing: -0.025em;
    }

    h2, h3 {
        font-weight: 540 !important;
        letter-spacing: -0.02em;
    }

    h4, h5, h6 {
        font-weight: 500 !important;
    }

    p, span, div {
        font-weight: 450;
    }

    div[data-testid="stHorizontalBlock"] button[kind="primary"],
    div[data-testid="stHorizontalBlock"] button[kind="secondary"],
    section[data-testid="stSidebar"] .stButton > button,
    section[data-testid="stSidebar"] .stDownloadButton > button,
    .stButton > button,
    .stDownloadButton > button {
        font-weight: 500 !important;
        letter-spacing: -0.01em !important;
    }

    .stMetric label {
        font-size: 0.85rem !important;
        font-weight: 450 !important;
        color: #9ca3af !important;
    }

    .stMetric [data-testid="stMetricValue"] {
        font-weight: 540 !important;
        letter-spacing: -0.025em !important;
    }

    .stMetric [data-testid="stMetricDelta"] {
        font-weight: 500 !important;
    }

    .simple-market-label {
        font-weight: 500 !important;
    }

    .simple-market-title {
        font-weight: 520 !important;
        letter-spacing: -0.02em !important;
    }

    .simple-market-price {
        font-weight: 560 !important;
        letter-spacing: -0.035em !important;
    }

    .simple-market-delta {
        font-weight: 520 !important;
    }

    .top-monitor-title {
        font-weight: 580 !important;
    }

    .top-monitor-time {
        font-weight: 500 !important;
    }

    .flow-panel-title,
    .flow-value,
    .signal-main,
    .ratio-summary-main {
        font-weight: 560 !important;
    }

    .flow-name,
    .ratio-summary-title {
        font-weight: 500 !important;
    }

    thead tr th {
        font-weight: 500 !important;
    }

    tbody tr td {
        font-weight: 450 !important;
    }

    .block-container {
        padding-top: 2rem;
        padding-bottom: 1rem;
    }
    </style>
    """,
    unsafe_allow_html=True,
)


st.markdown(
    """
    <style>
    div[data-testid="stHorizontalBlock"] button[kind="primary"],
    div[data-testid="stHorizontalBlock"] button[kind="secondary"] {
        min-height: 48px;
        border-radius: 14px;
        font-weight: 520;
        font-size: 1.02rem;
        letter-spacing: -0.02em;
    }
    </style>
    """,
    unsafe_allow_html=True,
)


st.markdown(
    """
    <style>
    section[data-testid="stSidebar"] .stButton > button,
    section[data-testid="stSidebar"] .stDownloadButton > button {
        min-height: 42px;
        width: 100%;
        border-radius: 10px;
        font-weight: 500;
        white-space: normal;
        line-height: 1.25;
        padding: 0.55rem 0.7rem;
    }
    section[data-testid="stSidebar"] [data-testid="stFileUploader"] {
        width: 100%;
    }
    section[data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] {
        min-height: 112px;
        padding: 0.7rem;
    }
    section[data-testid="stSidebar"] h4,
    section[data-testid="stSidebar"] h5 {
        margin-top: 0.8rem;
        margin-bottom: 0.35rem;
    }
    section[data-testid="stSidebar"] .stCaption {
        line-height: 1.45;
    }
    </style>
    """,
    unsafe_allow_html=True,
)


if "price_refresh_token_v51" not in st.session_state:
    st.session_state["price_refresh_token_v51"] = 0


# -----------------------------------
# 시간대 고정: 한국 서울 시간
# -----------------------------------
try:
    from zoneinfo import ZoneInfo
    KST = ZoneInfo("Asia/Seoul")
except Exception:
    KST = None

def 서울현재시각():
    if KST is not None:
        return datetime.now(KST)
    return datetime.now()

def 서울현재시각ISO():
    return 서울현재시각().isoformat()


def 한국장중여부(기준시각=None):
    now = 기준시각 or 서울현재시각()
    try:
        ts = pd.Timestamp(now)
        if getattr(ts, "tzinfo", None) is None:
            if KST is not None:
                ts = ts.tz_localize("Asia/Seoul")
        else:
            if KST is not None:
                ts = ts.tz_convert("Asia/Seoul")
    except Exception:
        try:
            ts = pd.Timestamp.now(tz="Asia/Seoul")
        except Exception:
            ts = pd.Timestamp.now()

    try:
        if ts.weekday() >= 5:
            return False
        장시작 = ts.replace(hour=9, minute=0, second=0, microsecond=0)
        장종료 = ts.replace(hour=15, minute=30, second=0, microsecond=0)
        return 장시작 <= ts <= 장종료
    except Exception:
        return False

def 서울조회문자열(값=None, 포맷="조회 %Y-%m-%d %H:%M"):
    대상 = 값
    if 대상 is None:
        대상 = 서울현재시각()

    try:
        ts = pd.to_datetime(대상)
        if getattr(ts, "tzinfo", None) is None:
            if KST is not None:
                try:
                    ts = ts.tz_localize("Asia/Seoul")
                except Exception:
                    ts = ts
        else:
            if KST is not None:
                try:
                    ts = ts.tz_convert("Asia/Seoul")
                except Exception:
                    ts = ts
        return ts.strftime(포맷)
    except Exception:
        try:
            return pd.to_datetime(대상).strftime(포맷)
        except Exception:
            return str(대상)

st.markdown(
    """
    <style>
    div[data-testid="stHorizontalBlock"] button[kind="secondary"],
    div[data-testid="stHorizontalBlock"] button[kind="primary"] {
        min-height: 46px;
        border-radius: 12px;
        font-weight: 520;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

모바일모드 = st.query_params.get("mobile", "0") == "1"

def 모바일여부():
    return 모바일모드

if 모바일여부():
    st.title(f"📈 투자 분석 시스템 {APP_VERSION}")
    st.caption("모바일 조회용 간소화 화면")
else:
    st.title(f"📈 투자 분석 시스템 {APP_VERSION}")
    


if not PLOTLY_AVAILABLE:
    st.error("plotly가 설치되어 있지 않습니다. 터미널에서 'pip install plotly' 후 다시 실행해 주세요.")
    st.stop()


def 안전웹요청(url, params=None, timeout=10, attempts=2):
    마지막오류 = None
    for _ in range(attempts):
        try:
            응답 = requests.get(url, params=params, headers=USER_AGENT, timeout=timeout)
            응답.raise_for_status()
            return 응답
        except Exception as e:
            마지막오류 = e
    return None


USER_AGENT = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
        "(KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36"
    )
}

야후인덱스심볼 = {
    "1001": "^KS11",
    "2001": "^KQ11",
}

야후주요지표심볼 = {
    "USD/KRW": "KRW=X",
    "국제 금": "GC=F",
    "WTI": "CL=F",
    "브렌트유": "BZ=F",
    "미국 10년물 금리": "^TNX",
    "VIX": "^VIX",
}

지표대체우선순위 = {
    "USD/KRW": ["yahoo", "naver"],
    "국제 금": ["yahoo"],
    "WTI": ["yahoo"],
    "브렌트유": ["yahoo"],
    "미국 10년물 금리": ["yahoo"],
    "VIX": ["yahoo"],
}

# -----------------------------------
# 기본 설정
# -----------------------------------
주요자산 = {
    "코스피": {"구분": "index", "코드": "1001"},
    "코스닥": {"구분": "index", "코드": "2001"},
    "KODEX 200": {"구분": "etf", "코드": "069500"},
    "KODEX 코스닥150": {"구분": "etf", "코드": "229200"},
    "KODEX AI반도체핵심장비": {"구분": "etf", "코드": "471990"},
    "KODEX AI전력핵심설비": {"구분": "etf", "코드": "487240"},
    "TIGER 200": {"구분": "etf", "코드": "102110"},
    "삼성전자": {"구분": "stock", "코드": "005930"},
    "SK하이닉스": {"구분": "stock", "코드": "000660"},
    "에이피알": {"구분": "stock", "코드": "278470"},
}

관심종목 = {
    "069500": "KODEX 200",
    "229200": "KODEX 코스닥150",
    "471990": "KODEX AI반도체핵심장비",
    "487240": "KODEX AI전력핵심설비",
    "102110": "TIGER 200",
    "005930": "삼성전자",
    "000660": "SK하이닉스",
    "278470": "에이피알",
}


코드명매핑 = {값["코드"]: 이름 for 이름, 값 in 주요자산.items()}
이름코드매핑 = {이름: 코드 for 코드, 이름 in 코드명매핑.items()}

@st.cache_data(ttl=86400)
def 전체종목매핑가져오기():
    기본매핑 = {
        "069500": "KODEX 200",
        "229200": "KODEX 코스닥150",
        "471990": "KODEX AI반도체핵심장비",
        "487240": "KODEX AI전력핵심설비",
        "102110": "TIGER 200",
        "005930": "삼성전자",
        "000660": "SK하이닉스",
        "278470": "에이피알",
    }
    try:
        return {str(k).zfill(6): str(v).strip() for k, v in 기본매핑.items() if str(k).strip() and str(v).strip()}
    except Exception:
        return 기본매핑



def 공식종목명가져오기(종목코드):
    """앱 내부 기준으로 확정한 종목명입니다.
    한 번 잘못 등록된 동적 매핑보다 이 값을 우선합니다.
    """
    코드 = "" if pd.isna(종목코드) else re.sub(r"[^0-9]", "", str(종목코드)).zfill(6)
    if not 코드:
        return ""
    try:
        전체매핑 = 전체종목매핑가져오기()
        return str(전체매핑.get(코드, "")).strip() if isinstance(전체매핑, dict) else ""
    except Exception:
        return ""


def 종목매핑강제갱신(종목코드, 종목명, 구분=None):
    """공식명 또는 사용자가 확정한 이름으로 전역 매핑을 교체합니다."""
    global 주요자산, 관심종목, 코드명매핑, 이름코드매핑

    코드 = "" if pd.isna(종목코드) else re.sub(r"[^0-9]", "", str(종목코드)).zfill(6)
    이름 = 종목명이름정리(종목명)
    if not 코드 or len(코드) != 6 or not 코드.isdigit() or not 이름:
        return False

    기존이름 = 코드명매핑.get(코드, "")
    if 기존이름 and 기존이름 != 이름:
        이름코드매핑.pop(기존이름, None)
        주요자산.pop(기존이름, None)

    코드명매핑[코드] = 이름
    이름코드매핑[이름] = 코드
    관심종목[코드] = 이름
    주요자산[이름] = {"구분": 구분 or 종목구분추정(이름, 코드), "코드": 코드}
    return True


def 종목매핑수동등록(종목코드, 종목명, 구분=None):
    global 주요자산, 관심종목, 코드명매핑, 이름코드매핑

    코드 = "" if pd.isna(종목코드) else re.sub(r"[^0-9]", "", str(종목코드)).zfill(6)
    입력이름 = 종목명이름정리(종목명)
    공식이름 = 공식종목명가져오기(코드)
    이름 = 공식이름 or 입력이름

    if not 코드 or len(코드) != 6 or not 코드.isdigit() or not 이름:
        return False

    # 공식명이 있으면 사용자가 오타를 입력해도 공식명으로 강제 정리합니다.
    if 공식이름:
        return 종목매핑강제갱신(코드, 공식이름, 구분=구분)

    기존이름 = 코드명매핑.get(코드, "")
    if 기존이름 and 종목명이름정리(기존이름) != 이름:
        # 공식명이 없는 종목은 충돌을 막되, 기존 이름을 유지합니다.
        return False

    기존코드 = 이름코드매핑.get(이름, "")
    if 기존코드 and 기존코드 != 코드:
        return False

    return 종목매핑강제갱신(코드, 이름, 구분=구분)


def 종목명이름정리(종목명):
    이름 = "" if pd.isna(종목명) else str(종목명).strip()
    이름 = 이름.replace("\xa0", " ").replace("\u200b", "").strip()
    이름 = re.sub(r"\s+", " ", 이름)
    별칭매핑 = {
        "SK 하이닉스": "SK하이닉스",
        "sk 하이닉스": "SK하이닉스",
        "sk하이닉스": "SK하이닉스",
        "kodex 200": "KODEX 200",
        "kodex코스닥150": "KODEX 코스닥150",
        "kodex 코스닥150": "KODEX 코스닥150",
        "tiger 200": "TIGER 200",
        "tiger200": "TIGER 200",
        "TIGER200": "TIGER 200",
        "APR": "에이피알",
        "apr": "에이피알",
        "에이피알주식회사": "에이피알",
        "KODEX AI전략핵심설비": "KODEX AI전력핵심설비",
        "kodex ai전략핵심설비": "KODEX AI전력핵심설비",
    }
    return 별칭매핑.get(이름, 이름)


def 종목코드기준종목명(종목코드):
    코드 = "" if pd.isna(종목코드) else re.sub(r"[^0-9]", "", str(종목코드)).zfill(6)
    if not 코드:
        return ""

    공식이름 = 공식종목명가져오기(코드)
    if 공식이름:
        # 이전 세션/업로드에서 잘못 들어온 동적 매핑도 즉시 교정합니다.
        if 코드명매핑.get(코드) != 공식이름:
            종목매핑강제갱신(코드, 공식이름, 구분=종목구분추정(공식이름, 코드))
        return 공식이름

    if 코드 in 코드명매핑:
        return 코드명매핑[코드]
    return ""



def 종목명기준종목코드(종목명):
    이름 = 종목명이름정리(종목명)
    if 이름 in 이름코드매핑:
        return 이름코드매핑[이름]
    return ""


def 종목코드종목명불일치정보(종목코드, 종목명):
    코드 = "" if pd.isna(종목코드) else re.sub(r"[^0-9]", "", str(종목코드)).zfill(6)
    이름 = 종목명이름정리(종목명)

    if not 코드 or not 이름 or len(코드) != 6 or not 코드.isdigit():
        return None

    코드기준이름 = 종목코드기준종목명(코드)
    이름기준코드 = 종목명기준종목코드(이름)

    if 코드기준이름 and 종목명이름정리(코드기준이름) != 이름:
        return {
            "유형": "등록정보불일치",
            "입력코드": 코드,
            "입력이름": 이름,
            "코드기준이름": 코드기준이름,
            "이름기준코드": 이름기준코드,
        }

    if (not 코드기준이름) and 이름기준코드 and 이름기준코드 != 코드:
        return {
            "유형": "이름기준코드불일치",
            "입력코드": 코드,
            "입력이름": 이름,
            "코드기준이름": "",
            "이름기준코드": 이름기준코드,
        }

    return None


def 종목명자동보정(종목코드, 종목명=""):
    코드 = "" if pd.isna(종목코드) else re.sub(r"[^0-9]", "", str(종목코드)).zfill(6)
    이름 = 종목명이름정리(종목명)
    코드기준이름 = 종목코드기준종목명(코드)

    if 코드기준이름:
        if (not 이름) or 이름 == 코드 or 종목명이름정리(코드기준이름) != 이름:
            return 코드기준이름

    if 이름:
        return 이름
    return 코드

def 종목코드자동보정(종목명, 종목코드=""):
    이름 = 종목명이름정리(종목명)
    코드 = "" if pd.isna(종목코드) else str(종목코드).strip()
    숫자코드 = re.sub(r"[^0-9]", "", 코드)
    if len(숫자코드) == 6:
        return 숫자코드
    이름기준코드 = 종목명기준종목코드(이름)
    return 이름기준코드 if 이름기준코드 else (숫자코드.zfill(6) if 숫자코드 else "")


def 종목구분추정(종목명="", 종목코드=""):
    이름 = 종목명이름정리(종목명).upper()
    if any(키워드 in 이름 for 키워드 in ["KODEX", "TIGER", "KOSEF", "KBSTAR", "ARIRANG", "ACE", "SOL", "HANARO", "TIMEFOLIO", "PLUS"]):
        return "etf"
    return "stock"

def 종목구분판단(종목코드, 종목명=""):
    코드 = "" if pd.isna(종목코드) else re.sub(r"[^0-9]", "", str(종목코드)).zfill(6)
    이름 = 종목명이름정리(종목명)
    if 코드 in 야후인덱스심볼:
        return "index"
    if 코드 in 코드명매핑:
        등록이름 = 코드명매핑.get(코드, "")
        자산정보 = 주요자산.get(등록이름)
        if isinstance(자산정보, dict) and 자산정보.get("구분") in ["index", "etf", "stock"]:
            return 자산정보.get("구분")
    return 종목구분추정(이름, 코드)


def 동적종목매핑갱신(거래df):
    global 주요자산, 관심종목, 코드명매핑, 이름코드매핑

    if 거래df is None or 거래df.empty:
        return

    작업 = 거래df.copy()
    if "종목코드" not in 작업.columns:
        return

    if "종목명" not in 작업.columns:
        작업["종목명"] = ""

    작업["종목코드"] = 작업["종목코드"].apply(lambda 값: "" if pd.isna(값) else re.sub(r"[^0-9]", "", str(값)).zfill(6))
    작업["종목명"] = 작업["종목명"].apply(종목명이름정리)
    작업 = 작업[(작업["종목코드"] != "") & (작업["종목명"] != "")]
    if 작업.empty:
        return

    for _, 행 in 작업.drop_duplicates(subset=["종목코드", "종목명"]).iterrows():
        코드 = 행["종목코드"]
        이름 = 행["종목명"]
        if 코드 in ["1001", "2001"]:
            continue
        종목매핑수동등록(코드, 이름, 구분=종목구분추정(이름, 코드))


def 거래이력자동보정(df):
    보정 = df.copy()

    if 보정.empty:
        return pd.DataFrame(columns=["종목코드", "종목명", "거래일자", "거래구분", "거래수량", "거래단가", "운용사", "비고"])

    for 컬럼 in ["종목코드", "종목명", "운용사", "비고"]:
        if 컬럼 not in 보정.columns:
            보정[컬럼] = ""
    for 컬럼 in ["거래일자", "거래구분", "거래수량", "거래단가"]:
        if 컬럼 not in 보정.columns:
            보정[컬럼] = None

    보정["종목코드"] = 보정["종목코드"].apply(lambda 값: "" if pd.isna(값) else re.sub(r"[^0-9]", "", str(값)).zfill(6) if re.sub(r"[^0-9]", "", str(값)) else "")
    보정["종목명"] = 보정["종목명"].apply(종목명이름정리)
    보정["운용사"] = 보정["운용사"].apply(lambda 값: "" if pd.isna(값) else str(값).strip())
    보정["비고"] = 보정["비고"].apply(lambda 값: "" if pd.isna(값) else str(값).strip())

    def _행보정(행):
        코드 = 행.get("종목코드", "")
        이름 = 행.get("종목명", "")

        if 코드 and 이름:
            종목매핑수동등록(코드, 이름)

        if 코드 and not 이름:
            이름 = 종목코드기준종목명(코드) or 이름
        elif 이름 and not 코드:
            코드 = 종목명기준종목코드(이름) or 코드
        else:
            불일치 = 종목코드종목명불일치정보(코드, 이름)
            if 불일치 is None:
                코드 = 종목명기준종목코드(이름) or 코드
                이름 = 종목코드기준종목명(코드) or 이름
            elif 불일치.get("유형") == "등록정보불일치" and 불일치.get("코드기준이름"):
                이름 = 불일치.get("코드기준이름", 이름)

        행["종목코드"] = 코드 if (isinstance(코드, str) and len(코드) == 6 and 코드.isdigit()) else ""
        행["종목명"] = 이름
        return 행

    보정 = 보정.apply(_행보정, axis=1)

    보정["거래일자"] = pd.to_datetime(보정["거래일자"], errors="coerce").dt.date
    보정["거래구분"] = 보정["거래구분"].astype(str).str.strip().replace({"buy": "매수", "BUY": "매수", "Buy": "매수", "sell": "매도", "SELL": "매도", "Sell": "매도"})
    보정["거래구분"] = 보정["거래구분"].replace({"매입": "매수", "구매": "매수", "매각": "매도", "sell ": "매도", "buy ": "매수"})
    보정.loc[보정["거래구분"].isin(["", "None", "nan"]), "거래구분"] = ""

    보정["거래수량"] = pd.to_numeric(보정["거래수량"], errors="coerce")
    보정["거래단가"] = pd.to_numeric(보정["거래단가"], errors="coerce")

    return 보정



def 거래이력검증표생성(df):
    if df is None or df.empty:
        return pd.DataFrame(columns=["행", "점검항목", "현재값", "권장사항"])

    점검결과 = []
    작업 = 거래이력자동보정(df.reset_index(drop=True).copy())
    오늘 = datetime.today().date()

    for idx, 행 in 작업.iterrows():
        행번호 = idx + 1
        종목코드 = "" if pd.isna(행.get("종목코드")) else str(행.get("종목코드")).strip()
        종목명 = "" if pd.isna(행.get("종목명")) else str(행.get("종목명")).strip()
        거래일자 = 행.get("거래일자")
        거래구분 = "" if pd.isna(행.get("거래구분")) else str(행.get("거래구분")).strip()
        거래수량 = pd.to_numeric(pd.Series([행.get("거래수량")]), errors="coerce").fillna(0).iloc[0]
        거래단가 = pd.to_numeric(pd.Series([행.get("거래단가")]), errors="coerce").fillna(0).iloc[0]

        if 종목코드 == "" and 종목명 == "":
            점검결과.append({"행": 행번호, "점검항목": "종목 정보", "현재값": "공란", "권장사항": "종목코드 또는 종목명 입력"})

        if 종목코드 != "" and (len(종목코드) != 6 or not 종목코드.isdigit()):
            점검결과.append({"행": 행번호, "점검항목": "종목코드 형식", "현재값": 종목코드, "권장사항": "6자리 숫자로 입력"})

        불일치정보 = 종목코드종목명불일치정보(종목코드, 종목명)
        if 불일치정보 is not None:
            if 불일치정보.get("유형") == "등록정보불일치":
                권장 = f'코드 {불일치정보["입력코드"]}의 등록 종목명은 "{불일치정보["코드기준이름"]}" 입니다'
                if 불일치정보.get("이름기준코드"):
                    권장 += f' / "{불일치정보["입력이름"]}"의 등록 코드는 {불일치정보["이름기준코드"]}'
                점검결과.append({"행": 행번호, "점검항목": "종목코드-종목명 불일치", "현재값": f'{종목코드} / {종목명}', "권장사항": 권장})
            elif 불일치정보.get("유형") == "이름기준코드불일치":
                권장 = f'"{불일치정보["입력이름"]}"의 등록 코드는 {불일치정보["이름기준코드"]} 입니다'
                점검결과.append({"행": 행번호, "점검항목": "종목명 기준 코드 확인", "현재값": f'{종목코드} / {종목명}', "권장사항": 권장})

        변환일자 = pd.to_datetime(거래일자, errors="coerce")
        if pd.isna(변환일자):
            점검결과.append({"행": 행번호, "점검항목": "거래일자", "현재값": 거래일자, "권장사항": "YYYY-MM-DD 형식으로 입력"})
        elif 변환일자.date() > 오늘:
            점검결과.append({"행": 행번호, "점검항목": "미래 날짜", "현재값": str(거래일자), "권장사항": "오늘 또는 과거 날짜만 입력"})

        if 거래구분 not in ["매수", "매도"]:
            점검결과.append({"행": 행번호, "점검항목": "거래구분", "현재값": 거래구분, "권장사항": "매수 또는 매도만 입력"})

        if 거래수량 <= 0:
            점검결과.append({"행": 행번호, "점검항목": "거래수량", "현재값": 거래수량, "권장사항": "0보다 큰 수량 입력"})

        if 거래단가 <= 0:
            점검결과.append({"행": 행번호, "점검항목": "거래단가", "현재값": 거래단가, "권장사항": "0보다 큰 단가 입력"})

    정렬작업 = 작업.copy()
    정렬작업["_거래일자정렬"] = pd.to_datetime(정렬작업["거래일자"], errors="coerce")
    정렬작업["_원본행"] = 정렬작업.index + 1
    정렬작업 = 정렬작업.sort_values(["종목코드", "_거래일자정렬", "_원본행"])

    종목별보유수량 = {}

    for _, 행 in 정렬작업.iterrows():
        행번호 = int(행["_원본행"])
        종목코드 = "" if pd.isna(행.get("종목코드")) else str(행.get("종목코드")).strip()
        거래구분 = "" if pd.isna(행.get("거래구분")) else str(행.get("거래구분")).strip()
        거래수량 = pd.to_numeric(pd.Series([행.get("거래수량")]), errors="coerce").fillna(0).iloc[0]

        if not 종목코드 or 거래수량 <= 0 or 거래구분 not in ["매수", "매도"]:
            continue

        현재보유 = 종목별보유수량.get(종목코드, 0)

        if 거래구분 == "매수":
            종목별보유수량[종목코드] = 현재보유 + 거래수량
        else:
            if 거래수량 > 현재보유:
                점검결과.append({
                    "행": 행번호,
                    "점검항목": "초과매도",
                    "현재값": f"{거래수량}주 매도 / 보유 {현재보유}주",
                    "권장사항": "이전 거래이력 또는 수량 입력을 확인"
                })
            종목별보유수량[종목코드] = max(0, 현재보유 - 거래수량)

    return pd.DataFrame(점검결과)


종목별거래단가범위 = {
    "069500": {"최소": 50000, "최대": 120000, "이름": "KODEX 200"},
    "229200": {"최소": 10000, "최대": 40000, "이름": "KODEX 코스닥150"},
    "471990": {"최소": 10000, "최대": 50000, "이름": "KODEX AI반도체핵심장비"},
    "487240": {"최소": 10000, "최대": 80000, "이름": "KODEX AI전력핵심설비"},
    "005930": {"최소": 100000, "최대": 300000, "이름": "삼성전자"},
    "000660": {"최소": 500000, "최대": 1500000, "이름": "SK하이닉스"},
}



def 거래이력편집용자동보정(df):
    """
    편집 화면용 자동보정:
    - 입력 중인 행이 사라지지 않도록 원본 행을 최대한 유지
    - 종목코드 입력 시 종목명 자동 보정
    - 거래일자/수량/단가는 과도한 정규화 없이 편집 가능한 형태 유지
    """
    if df is None:
        return pd.DataFrame(columns=["종목코드", "종목명", "거래일자", "거래구분", "거래수량", "거래단가", "운용사", "비고"])

    작업 = df.copy()

    표준열 = ["종목코드", "종목명", "거래일자", "거래구분", "거래수량", "거래단가", "운용사", "비고"]
    for 열 in 표준열:
        if 열 not in 작업.columns:
            작업[열] = None if 열 in ["거래일자", "거래수량", "거래단가"] else ""

    if "_입력원본순서" not in 작업.columns:
        작업["_입력원본순서"] = range(len(작업))

    try:
        동적종목매핑갱신(작업)
    except Exception:
        pass

    작업["종목코드"] = 작업["종목코드"].apply(
        lambda 값: "" if pd.isna(값) else re.sub(r"[^0-9]", "", str(값)).zfill(6) if re.sub(r"[^0-9]", "", str(값)) else ""
    )
    작업["종목명"] = 작업.apply(
        lambda 행: 종목명자동보정(행.get("종목코드", ""), 행.get("종목명", "")),
        axis=1
    )
    작업["거래구분"] = 작업["거래구분"].apply(lambda 값: "" if pd.isna(값) else str(값).strip())
    작업["운용사"] = 작업["운용사"].apply(lambda 값: "" if pd.isna(값) else str(값).strip())
    작업["비고"] = 작업["비고"].apply(lambda 값: "" if pd.isna(값) else str(값).strip())

    try:
        dt_series = pd.to_datetime(작업["거래일자"], errors="coerce")
        작업["거래일자"] = dt_series.dt.date.where(dt_series.notna(), 작업["거래일자"])
    except Exception:
        pass

    작업 = 거래이력입력창정렬(작업)
    return 작업


def 거래이력계산대상추출(df):
    """
    계산 대상 추출:
    - 편집 중 빈 행은 제외
    - 종목코드/종목명/거래구분/수량/단가가 핵심적으로 유효한 행만 계산 대상으로 사용
    - 편집 화면 데이터는 유지하되, 계산용은 안정적으로 분리
    """
    if df is None:
        return pd.DataFrame(columns=["종목코드", "종목명", "거래일자", "거래구분", "거래수량", "거래단가", "운용사", "비고"])

    작업 = df.copy()

    표준열 = ["종목코드", "종목명", "거래일자", "거래구분", "거래수량", "거래단가", "운용사", "비고"]
    for 열 in 표준열:
        if 열 not in 작업.columns:
            작업[열] = None if 열 in ["거래일자", "거래수량", "거래단가"] else ""

    작업["종목코드"] = 작업["종목코드"].apply(lambda 값: "" if pd.isna(값) else str(값).strip())
    작업["종목명"] = 작업["종목명"].apply(lambda 값: "" if pd.isna(값) else str(값).strip())
    작업["거래구분"] = 작업["거래구분"].apply(lambda 값: "" if pd.isna(값) else str(값).strip())

    # 완전히 비어 있는 행 제외
    작업 = 작업.dropna(how="all")
    if 작업.empty:
        return 거래이력정규화(작업)

    # 계산 대상은 핵심 필드가 하나도 없는 행 제외
    핵심열 = ["종목코드", "종목명", "거래일자", "거래구분", "거래수량", "거래단가"]
    핵심값존재 = 작업[핵심열].apply(
        lambda row: any(str(v).strip() not in ["", "None", "nan", "NaT"] for v in row),
        axis=1
    )
    작업 = 작업.loc[핵심값존재].copy()
    if 작업.empty:
        return 거래이력정규화(작업)

    # 실제 계산 반영은 종목/거래구분/수량/단가가 유효한 행 중심
    작업["거래수량"] = pd.to_numeric(작업["거래수량"], errors="coerce")
    작업["거래단가"] = pd.to_numeric(작업["거래단가"], errors="coerce")

    계산대상마스크 = (
        (작업["종목코드"].astype(str).str.strip() != "") |
        (작업["종목명"].astype(str).str.strip() != "")
    ) & 작업["거래구분"].isin(["매수", "매도"]) & (작업["거래수량"].fillna(0) > 0) & (작업["거래단가"].fillna(0) > 0)

    작업 = 작업.loc[계산대상마스크].copy()
    return 거래이력정규화(작업)

def 거래이력이상치점검표생성(df):
    if df is None or df.empty:
        return pd.DataFrame(columns=["행", "점검항목", "현재값", "권장사항"])

    작업 = 거래이력자동보정(df.reset_index(drop=True).copy())
    점검결과 = []

    중복기준열 = ["종목코드", "종목명", "거래일자", "거래구분", "거래수량", "거래단가"]
    중복마스크 = 작업.duplicated(subset=중복기준열, keep=False)
    if 중복마스크.any():
        for idx, 행 in 작업.loc[중복마스크].iterrows():
            점검결과.append({
                "행": idx + 1,
                "점검항목": "중복거래 가능성",
                "현재값": f"{행.get('종목명', '')} / {행.get('거래일자', '')} / {행.get('거래구분', '')} / {행.get('거래수량', '')}주 / {행.get('거래단가', '')}",
                "권장사항": "같은 거래가 2번 입력되지 않았는지 확인"
            })

    for idx, 행 in 작업.iterrows():
        종목코드 = str(행.get("종목코드", "") or "").zfill(6)
        종목명 = 종목명자동보정(종목코드, 행.get("종목명", ""))
        거래단가 = pd.to_numeric(pd.Series([행.get("거래단가")]), errors="coerce").iloc[0]
        if pd.isna(거래단가) or 거래단가 <= 0:
            continue

        범위정보 = 종목별거래단가범위.get(종목코드)
        if 범위정보:
            최소값 = 범위정보["최소"]
            최대값 = 범위정보["최대"]
            if 거래단가 < 최소값 or 거래단가 > 최대값:
                점검결과.append({
                    "행": idx + 1,
                    "점검항목": "거래단가 범위 확인",
                    "현재값": f"{종목명} {거래단가:,.0f}",
                    "권장사항": f"{종목명}의 통상 입력 범위({최소값:,.0f}~{최대값:,.0f}원)와 크게 다르면 실제 체결단가를 다시 확인"
                })
        else:
            if 거래단가 < 100 or 거래단가 > 5000000:
                점검결과.append({
                    "행": idx + 1,
                    "점검항목": "거래단가 극단값 확인",
                    "현재값": f"{종목명} {거래단가:,.0f}",
                    "권장사항": "입력 자릿수 또는 실제 체결단가를 다시 확인"
                })

    if not 점검결과:
        return pd.DataFrame(columns=["행", "점검항목", "현재값", "권장사항"])

    결과df = pd.DataFrame(점검결과)
    return 결과df.drop_duplicates().sort_values(["행", "점검항목"]).reset_index(drop=True)


기본포트폴리오 = pd.DataFrame([
    {"종목코드": "069500", "종목명": "KODEX 200", "거래일자": "2026-02-20", "거래구분": "매수", "거래수량": 11, "거래단가": 86239, "운용사": "신한은행 IRP", "비고": ""},
    {"종목코드": "069500", "종목명": "KODEX 200", "거래일자": "2026-03-04", "거래구분": "매도", "거래수량": 11, "거래단가": 78475, "운용사": "신한은행 IRP", "비고": "미국의 이란 침공으로 주식 폭락, -10% 기준 손절"},
    {"종목코드": "069500", "종목명": "KODEX 200", "거래일자": "2026-03-05", "거래구분": "매수", "거래수량": 35, "거래단가": 84936, "비고": "폭락 이후 반등"},
    {"종목코드": "069500", "종목명": "KODEX 200", "거래일자": "2026-03-10", "거래구분": "매수", "거래수량": 59, "거래단가": 82900, "비고": "트럼프 이란전 조기 종료 발표, 주요 국제 유가 하락"},
    {"종목코드": "000660", "종목명": "SK하이닉스", "거래일자": "2026-03-10", "거래구분": "매수", "거래수량": 1, "거래단가": 941000, "비고": "트럼프 이란전 조기 종료 발표"},
    {"종목코드": "005930", "종목명": "삼성전자", "거래일자": "2026-03-10", "거래구분": "매수", "거래수량": 10, "거래단가": 189300, "비고": "트럼프 이란전 조기 종료 발표"},
    {"종목코드": "005930", "종목명": "삼성전자", "거래일자": "2026-03-11", "거래구분": "매수", "거래수량": 1, "거래단가": 189000, "비고": "주가 등락 안정 판단"},
    {"종목코드": "005930", "종목명": "삼성전자", "거래일자": "2026-03-11", "거래구분": "매수", "거래수량": 5, "거래단가": 191700, "비고": "주가 등락 안정 판단"},
    {"종목코드": "005930", "종목명": "삼성전자", "거래일자": "2026-03-13", "거래구분": "매수", "거래수량": 11, "거래단가": 180800, "비고": "주가 최근 최저, 주가등록 안정 판단"},
    {"종목코드": "069500", "종목명": "KODEX 200", "거래일자": "2026-03-13", "거래구분": "매수", "거래수량": 3, "거래단가": 83635, "비고": "정윤"},
    {"종목코드": "229200", "종목명": "KODEX 코스닥150", "거래일자": "2026-03-13", "거래구분": "매수", "거래수량": 2, "거래단가": 19970, "비고": "정윤"},
    {"종목코드": "229200", "종목명": "KODEX 코스닥150", "거래일자": "2026-03-16", "거래구분": "매수", "거래수량": 5, "거래단가": 19929, "비고": "코스닥 150 관심 증가 이광수대표 발언"},
    {"종목코드": "069500", "종목명": "KODEX 200", "거래일자": "2026-03-23", "거래구분": "매수", "거래수량": 23, "거래단가": 84031, "비고": "코덱스 -6.5% 하락에 분할매수 유리 판단"},
    {"종목코드": "000660", "종목명": "SK하이닉스", "거래일자": "2026-03-23", "거래구분": "매수", "거래수량": 1, "거래단가": 933000, "비고": "전일 종가 대비 주가 -7.15% 하락"},
    {"종목코드": "005930", "종목명": "삼성전자", "거래일자": "2026-03-23", "거래구분": "매수", "거래수량": 5, "거래단가": 186300, "비고": "전일 종가 대비 주가 -6.27% 하락"},
    {"종목코드": "069500", "종목명": "KODEX 200", "거래일자": "2026-03-24", "거래구분": "매수", "거래수량": 35, "거래단가": 83872, "비고": "전일 폭락 후 오늘 트럼프의 이란 공격 중지 발표로 유가, 환율 하락"},
])
기본포트폴리오["거래일자"] = pd.to_datetime(기본포트폴리오["거래일자"], errors="coerce").dt.date

시장지표네이버URL = {
    "USD/KRW": "https://finance.yahoo.com/quote/KRW%3DX",
    "국제 금": "https://finance.yahoo.com/quote/GC%3DF",
    "WTI": "https://finance.yahoo.com/quote/CL%3DF",
    "브렌트유": "https://finance.yahoo.com/quote/BZ%3DF",
    "미국 10년물 금리": "https://finance.yahoo.com/quote/%5ETNX",
    "VIX": "https://finance.yahoo.com/quote/%5EVIX",
}

목표비중저장파일 = "target_weights.json"
거래이력자동저장파일 = "trade_history_autosave.json"
최근업로드거래이력파일 = "trade_history_latest_uploaded.json"
최근업로드메타파일 = "trade_history_latest_uploaded_meta.json"
모니터관심종목저장파일 = "monitor_custom_assets_v54.json"

IRP비주식자산저장파일 = "integrated_non_stock_assets_v513.json"


def 기본IRP비주식자산표():
    오늘 = 서울현재시각().date().isoformat()
    return pd.DataFrame([
        {"계좌": "신한은행 IRP", "자산군": "TDF", "상품명": "TDF2035", "원금": 50000000, "평가금액": 51873538, "예상연수익률": 3.74, "만기일": "", "반영일자": "2026-04-30", "비고": "평가금액은 직접 입력"},
        {"계좌": "신한은행 IRP", "자산군": "TDF", "상품명": "TDF2045", "원금": 30000000, "평가금액": 31443846, "예상연수익률": 4.81, "만기일": "", "반영일자": "2026-04-30", "비고": "평가금액은 직접 입력"},
        {"계좌": "신한은행 IRP", "자산군": "정기예금", "상품명": "푸본현대생명 정기예금", "원금": 27499444, "평가금액": 27499444, "예상연수익률": 3.10, "만기일": "2027-02-11", "반영일자": "2026-04-30", "비고": "만기 유지 시 고정금리 기준"},
        {"계좌": "신한은행 IRP", "자산군": "현금성자산", "상품명": "현금성 대기자산", "원금": 5030813, "평가금액": 5030813, "예상연수익률": 2.30, "만기일": "", "반영일자": "2026-04-30", "비고": "예수금·MMDA 등 수동 입력"},
        {"계좌": "미래에셋증권", "자산군": "현금성자산", "상품명": "예수금", "원금": 17188280, "평가금액": 17188280, "예상연수익률": 0.0, "만기일": "", "반영일자": "2026-04-30", "비고": "CMA/예수금"},
    ])


def 날짜값_YYYYMMDD문자열(값):
    """날짜/일시 값을 화면 표시용 YYYY-MM-DD 문자열로 정리합니다.
    - 엑셀 업로드 시 들어오는 2026-05-06 00:00:00 형태를 2026-05-06으로 통일
    - 빈 값, NaT, nan, None은 공란으로 처리
    """
    if 값 is None:
        return ""
    try:
        if pd.isna(값):
            return ""
    except Exception:
        pass

    문자 = str(값).strip()
    if 문자 in ["", "NaT", "nat", "nan", "None"]:
        return ""

    try:
        변환 = pd.to_datetime(값, errors="coerce")
        if pd.isna(변환):
            return ""
        return 변환.strftime("%Y-%m-%d")
    except Exception:
        # 이미 YYYY-MM-DD로 시작하는 문자열이면 앞 10자리만 사용
        if re.match(r"^\d{4}-\d{2}-\d{2}", 문자):
            return 문자[:10]
        return 문자


def IRP비주식자산표준열맞추기(df):
    표준열 = ["계좌", "자산군", "상품명", "원금", "평가금액", "예상연수익률", "만기일", "반영일자", "비고"]
    작업 = pd.DataFrame() if df is None else pd.DataFrame(df).copy()

    # 통합 업로드 템플릿 호환: 수익률(%) / 기준일 컬럼명을 앱 내부명으로 변환
    컬럼변환 = {}
    if "수익률(%)" in 작업.columns and "예상연수익률" not in 작업.columns:
        컬럼변환["수익률(%)"] = "예상연수익률"
    if "기준일" in 작업.columns and "반영일자" not in 작업.columns:
        컬럼변환["기준일"] = "반영일자"
    if 컬럼변환:
        작업 = 작업.rename(columns=컬럼변환)

    for 열 in 표준열:
        if 열 not in 작업.columns:
            작업[열] = 0 if 열 in ["원금", "평가금액", "예상연수익률"] else ""

    작업 = 작업.dropna(how="all")
    작업 = 작업[표준열].copy()

    for 열 in ["계좌", "자산군", "상품명", "비고"]:
        작업[열] = 작업[열].apply(lambda 값: "" if pd.isna(값) else str(값).strip())
        작업[열] = 작업[열].replace({"NaT": "", "nan": "", "None": ""})

    # 날짜형 컬럼은 시간(00:00:00)이 표시되지 않도록 YYYY-MM-DD 문자열로 통일합니다.
    for 열 in ["만기일", "반영일자"]:
        작업[열] = 작업[열].apply(날짜값_YYYYMMDD문자열)

    for 열 in ["원금", "평가금액", "예상연수익률"]:
        작업[열] = pd.to_numeric(작업[열], errors="coerce").fillna(0.0)

    작업["계좌"] = 작업["계좌"].replace({"": "미지정 계좌", "신한 IRP": "신한은행 IRP", "미래에셋": "미래에셋증권"})
    작업["자산군"] = 작업["자산군"].replace("", "기타")
    작업["상품명"] = 작업["상품명"].replace("", "미입력 상품")
    작업 = 작업[(작업["원금"] > 0) | (작업["평가금액"] > 0) | (작업["상품명"].astype(str).str.strip() != "미입력 상품")].copy()
    return 작업.reset_index(drop=True)


def IRP비주식자산불러오기():
    if "irp_non_stock_assets_df_v512" in st.session_state:
        return IRP비주식자산표준열맞추기(st.session_state["irp_non_stock_assets_df_v512"])
    if os.path.exists(IRP비주식자산저장파일):
        try:
            with open(IRP비주식자산저장파일, "r", encoding="utf-8") as f:
                data = json.load(f)
            df = IRP비주식자산표준열맞추기(pd.DataFrame(data))
            st.session_state["irp_non_stock_assets_df_v512"] = df
            return df
        except Exception:
            pass
    df = 기본IRP비주식자산표()
    st.session_state["irp_non_stock_assets_df_v512"] = df
    return df


def IRP비주식자산저장(df):
    작업 = IRP비주식자산표준열맞추기(df)
    st.session_state["irp_non_stock_assets_df_v512"] = 작업
    저장용 = 작업.fillna("").to_dict(orient="records")
    return 안전JSON저장(저장용, IRP비주식자산저장파일)


def IRP비주식자산편집UI():
    st.markdown("### 계좌별 비주식·현금성 자산 관리")
    st.caption("TDF, 정기예금, 현금성 자산은 실시간 시세 조회 대신 원금과 평가금액을 직접 입력해 통합 자산에 반영합니다.")
    현재df = IRP비주식자산불러오기()
    with st.expander("계좌별 비주식·현금성 자산 입력/수정", expanded=False):
        편집df = st.data_editor(
            현재df,
            num_rows="dynamic",
            use_container_width=True,
            key="irp_non_stock_assets_editor_v513",
            column_config={
                "계좌": st.column_config.TextColumn("계좌"),
                "자산군": st.column_config.SelectboxColumn("자산군", options=["TDF", "정기예금", "현금성자산", "채권", "펀드", "기타"]),
                "상품명": st.column_config.TextColumn("상품명"),
                "원금": st.column_config.NumberColumn("원금", min_value=0, step=10000, format="%d"),
                "평가금액": st.column_config.NumberColumn("평가금액", min_value=0, step=10000, format="%d"),
                "예상연수익률": st.column_config.NumberColumn("예상연수익률(%)", step=0.1, format="%.2f"),
                "만기일": st.column_config.TextColumn("만기일"),
                "반영일자": st.column_config.TextColumn("반영일자"),
                "비고": st.column_config.TextColumn("비고"),
            },
        )
        버튼1, 버튼2, 버튼3 = st.columns([1.2, 1.2, 5])
        with 버튼1:
            if st.button("비주식 자산 저장", key="save_irp_non_stock_assets_v513", use_container_width=True):
                성공, 메시지 = IRP비주식자산저장(편집df)
                if 성공:
                    st.success("비주식·현금성 자산을 저장했습니다.")
                    st.rerun()
                else:
                    st.error(메시지)
        with 버튼2:
            if st.button("기본값 복원", key="reset_irp_non_stock_assets_v513", use_container_width=True):
                IRP비주식자산저장(기본IRP비주식자산표())
                st.success("기본 비주식·현금성 자산 표로 복원했습니다.")
                st.rerun()
        with 버튼3:
            st.caption("정기예금은 현재 평가를 보수적으로 원금 기준으로 두고, 만기 예상 이자는 비고로 관리하는 방식을 권장합니다.")
    return IRP비주식자산불러오기()


def 주식ETF자산요약행생성(보유포트폴리오):
    if 보유포트폴리오 is None or 보유포트폴리오.empty:
        return pd.DataFrame(columns=["계좌", "자산군", "상품명", "원금", "평가금액", "평가손익", "수익률", "비고"])
    작업 = 보유포트폴리오.copy()
    if "데이터상태" in 작업.columns:
        작업 = 작업[작업["데이터상태"].astype(str) == "정상"].copy()
    if 작업.empty:
        return pd.DataFrame(columns=["계좌", "자산군", "상품명", "원금", "평가금액", "평가손익", "수익률", "비고"])
    계좌값 = 작업["운용사"] if "운용사" in 작업.columns else pd.Series(["미래에셋/증권계좌"] * len(작업), index=작업.index)
    결과 = pd.DataFrame({
        "계좌": 계좌값.fillna("미래에셋/증권계좌").astype(str).replace("", "미래에셋/증권계좌"),
        "자산군": 작업.apply(lambda 행: "ETF" if 종목구분판단(행.get("종목코드", ""), 행.get("종목명", "")) == "etf" else "주식", axis=1),
        "상품명": 작업.get("종목명", ""),
        "원금": pd.to_numeric(작업.get("투자원금", 0), errors="coerce").fillna(0),
        "평가금액": pd.to_numeric(작업.get("평가금액", 0), errors="coerce").fillna(0),
        "평가손익": pd.to_numeric(작업.get("평가손익", 0), errors="coerce").fillna(0),
        "수익률": pd.to_numeric(작업.get("수익률", 0), errors="coerce").fillna(0),
        "비고": "실시간/준실시간 시세 반영",
    })
    return 결과


def IRP비주식자산요약행생성(irp_df):
    작업 = IRP비주식자산표준열맞추기(irp_df)
    작업 = 작업[(작업["원금"] > 0) | (작업["평가금액"] > 0)].copy()
    if 작업.empty:
        return pd.DataFrame(columns=["계좌", "자산군", "상품명", "원금", "평가금액", "평가손익", "수익률", "비고"])
    작업["평가손익"] = 작업["평가금액"] - 작업["원금"]
    작업["수익률"] = np.where(작업["원금"] != 0, 작업["평가손익"] / 작업["원금"] * 100, 0)
    return 작업[["계좌", "자산군", "상품명", "원금", "평가금액", "평가손익", "수익률", "비고"]].copy()


def 통합자산현황표생성(보유포트폴리오, irp_df):
    통합 = pd.concat([주식ETF자산요약행생성(보유포트폴리오), IRP비주식자산요약행생성(irp_df)], ignore_index=True)
    if 통합.empty:
        return 통합
    통합["원금"] = pd.to_numeric(통합["원금"], errors="coerce").fillna(0)
    통합["평가금액"] = pd.to_numeric(통합["평가금액"], errors="coerce").fillna(0)
    통합["평가손익"] = 통합["평가금액"] - 통합["원금"]
    통합["수익률"] = np.where(통합["원금"] != 0, 통합["평가손익"] / 통합["원금"] * 100, 0)
    총평가 = 통합["평가금액"].sum()
    통합["전체비중"] = np.where(총평가 != 0, 통합["평가금액"] / 총평가 * 100, 0)
    return 통합


def 통합자산현황UI(보유포트폴리오, irp_df):
    통합표 = 통합자산현황표생성(보유포트폴리오, irp_df)
    st.markdown("### 통합 자산 현황")
    if 통합표.empty:
        st.info("통합 자산 현황을 표시할 데이터가 없습니다.")
        return 통합표
    총원금 = 통합표["원금"].sum()
    총평가 = 통합표["평가금액"].sum()
    총손익 = 총평가 - 총원금
    총수익률 = (총손익 / 총원금 * 100) if 총원금 else 0
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("통합 원금", 금액표시(총원금))
    c2.metric("통합 평가액", 금액표시(총평가))
    c3.metric("통합 손익", 금액표시(총손익))
    c4.metric("통합 수익률", f"{총수익률:.2f}%")
    자산군요약 = 통합표.groupby("자산군", as_index=False).agg({"원금": "sum", "평가금액": "sum", "평가손익": "sum"})
    자산군요약["수익률"] = np.where(자산군요약["원금"] != 0, 자산군요약["평가손익"] / 자산군요약["원금"] * 100, 0)
    자산군요약["전체비중"] = np.where(총평가 != 0, 자산군요약["평가금액"] / 총평가 * 100, 0)
    자산군요약 = 자산군요약.sort_values("평가금액", ascending=False).reset_index(drop=True)
    숫자서식 = {"원금": 안전정수포맷, "평가금액": 안전정수포맷, "평가손익": 손익문자열, "수익률": 수익률문자열, "전체비중": lambda x: 안전소수포맷(x, 2)}
    st.caption("자산군별 요약")
    표데이터프레임(index_1부터(자산군요약.copy()).style.format(숫자서식).map(손익색상, subset=["평가손익"]).map(수익률색상, subset=["수익률"]), use_container_width=True)
    with st.expander("통합 자산 상세 보기", expanded=False):
        상세표시 = 통합표.sort_values(["계좌", "자산군", "평가금액"], ascending=[True, True, False]).reset_index(drop=True)
        표데이터프레임(index_1부터(상세표시).style.format(숫자서식).map(손익색상, subset=["평가손익"]).map(수익률색상, subset=["수익률"]), use_container_width=True)
    return 통합표



자동백업루트폴더 = "backup"
일일백업폴더 = os.path.join(자동백업루트폴더, "daily")
수정백업폴더 = os.path.join(자동백업루트폴더, "edit_history")
자동백업일일보관개수 = 7
자동백업수정보관개수 = 30


def 자동백업폴더준비():
    try:
        os.makedirs(일일백업폴더, exist_ok=True)
        os.makedirs(수정백업폴더, exist_ok=True)
    except Exception:
        pass


def 거래이력백업페이로드생성(df, backup_type="manual", reason="", source="app"):
    작업 = 거래이력편집용자동보정(df if df is not None else pd.DataFrame())
    return {
        "meta": {
            "backup_type": backup_type,
            "reason": reason,
            "source": source,
            "saved_at": 서울현재시각ISO(),
            "rows": int(len(작업)),
            "signature": 거래이력비교지문(작업),
            "app_version": APP_VERSION,
        },
        "data": 거래이력JSON변환(작업),
    }


def 자동백업파일저장(df, backup_type="daily", reason="", source="app"):
    자동백업폴더준비()
    작업 = 거래이력편집용자동보정(df if df is not None else pd.DataFrame())
    if 작업 is None:
        작업 = pd.DataFrame()

    저장시각 = 서울현재시각()
    시각문자 = 저장시각.strftime("%Y%m%d_%H%M%S")
    서명 = 거래이력비교지문(작업)
    짧은서명 = (str(abs(hash(서명))) if 서명 else "0")[-10:]
    폴더 = 일일백업폴더 if backup_type == "daily" else 수정백업폴더
    접두어 = "daily" if backup_type == "daily" else "edit"
    파일명 = f"{접두어}_{시각문자}_{짧은서명}.json"
    파일경로 = os.path.join(폴더, 파일명)

    try:
        with open(파일경로, "w", encoding="utf-8") as f:
            json.dump(거래이력백업페이로드생성(작업, backup_type=backup_type, reason=reason, source=source), f, ensure_ascii=False, indent=2)
        자동백업파일정리(backup_type)
        return True, 파일경로
    except Exception as e:
        return False, str(e)


def 자동백업파일목록가져오기(backup_type=None):
    자동백업폴더준비()
    대상 = []
    if backup_type in [None, "daily"]:
        대상.append(("daily", 일일백업폴더))
    if backup_type in [None, "edit"]:
        대상.append(("edit", 수정백업폴더))

    결과 = []
    for 종류, 폴더 in 대상:
        try:
            for 이름 in os.listdir(폴더):
                if not 이름.lower().endswith(".json"):
                    continue
                경로 = os.path.join(폴더, 이름)
                try:
                    with open(경로, "r", encoding="utf-8") as f:
                        payload = json.load(f)
                    meta = payload.get("meta", {}) if isinstance(payload, dict) else {}
                    결과.append({
                        "backup_type": 종류,
                        "path": 경로,
                        "file_name": 이름,
                        "saved_at": meta.get("saved_at", ""),
                        "rows": meta.get("rows", 0),
                        "reason": meta.get("reason", ""),
                        "source": meta.get("source", ""),
                        "signature": meta.get("signature", ""),
                        "size": os.path.getsize(경로) if os.path.exists(경로) else 0,
                    })
                except Exception:
                    continue
        except Exception:
            continue

    결과 = sorted(결과, key=lambda x: str(x.get("saved_at", "")), reverse=True)
    return 결과


def 자동백업파일정리(backup_type):
    자동백업폴더준비()
    if backup_type == "daily":
        폴더 = 일일백업폴더
        유지개수 = 자동백업일일보관개수
    else:
        폴더 = 수정백업폴더
        유지개수 = 자동백업수정보관개수

    try:
        파일들 = []
        for 이름 in os.listdir(폴더):
            if 이름.lower().endswith(".json"):
                경로 = os.path.join(폴더, 이름)
                파일들.append((os.path.getmtime(경로), 경로))
        파일들 = sorted(파일들, reverse=True)
        for _, 삭제경로 in 파일들[유지개수:]:
            try:
                os.remove(삭제경로)
            except Exception:
                pass
    except Exception:
        pass


def 자동백업불러오기(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            payload = json.load(f)
        데이터 = payload.get("data", payload)
        df = 거래이력표준열맞추기(pd.DataFrame(데이터))
        df = 거래이력자동보정(df)
        if "거래일자" in df.columns:
            df["거래일자"] = pd.to_datetime(df["거래일자"], errors="coerce").dt.date
        return 거래이력입력창정렬(df), payload.get("meta", {}) if isinstance(payload, dict) else {}
    except Exception as e:
        return None, {"error": str(e)}


def 자동백업복원적용(df, source_label="backup_restore"):
    편집df = 거래이력편집용자동보정(df if df is not None else pd.DataFrame())
    반영df, 변경됨, 저장성공, 저장메시지 = 거래이력세션반영(편집df, 저장강제=True, 자동저장허용=True)
    st.session_state["trade_history_source_v1"] = source_label
    st.session_state["trade_history_latest_upload_name_v1"] = source_label
    st.session_state["trade_history_latest_upload_time_v1"] = 서울현재시각ISO()
    st.session_state["price_refresh_token_v51"] = st.session_state.get("price_refresh_token_v51", 0) + 1
    st.session_state["manual_price_refresh_ts_v1"] = 서울현재시각ISO()
    시세관련캐시초기화()
    return 반영df, 변경됨, 저장성공, 저장메시지


def 자동백업일일실행(df):
    자동백업폴더준비()
    작업 = 거래이력편집용자동보정(df if df is not None else pd.DataFrame())
    오늘문자 = 서울현재시각().strftime("%Y-%m-%d")
    오늘서명 = 거래이력비교지문(작업)
    마지막일자 = st.session_state.get("backup_daily_last_date_v1", "")
    마지막서명 = st.session_state.get("backup_daily_last_signature_v1", "")

    if 오늘문자 == 마지막일자 and 오늘서명 == 마지막서명:
        return False, "이미 오늘 백업 완료"

    성공, 결과 = 자동백업파일저장(작업, backup_type="daily", reason="daily_startup", source="app_start")
    if 성공:
        st.session_state["backup_daily_last_date_v1"] = 오늘문자
        st.session_state["backup_daily_last_signature_v1"] = 오늘서명
        return True, 결과
    return False, 결과


def 자동백업수정전실행(이전df, 다음df=None, source="editor"):
    자동백업폴더준비()
    이전작업 = 거래이력편집용자동보정(이전df if 이전df is not None else pd.DataFrame())
    이전서명 = 거래이력비교지문(이전작업)
    다음서명 = 거래이력비교지문(거래이력편집용자동보정(다음df)) if 다음df is not None else ""

    if not 이전서명 or 이전서명 == 다음서명:
        return False, "변경 전 백업 불필요"

    마지막백업서명 = st.session_state.get("backup_last_edit_source_signature_v1", "")
    if 이전서명 == 마지막백업서명:
        return False, "동일 상태 이미 백업됨"

    성공, 결과 = 자동백업파일저장(이전작업, backup_type="edit", reason="before_edit", source=source)
    if 성공:
        st.session_state["backup_last_edit_source_signature_v1"] = 이전서명
        return True, 결과
    return False, 결과


def 자동백업관리UI(current_df, portfolio_df=None, holding_df=None):
    백업목록 = 자동백업파일목록가져오기()
    with st.expander("자동백업 관리", expanded=False):
        요약1, 요약2, 요약3 = st.columns(3)
        요약1.metric("일일 백업", f"{sum(1 for x in 백업목록 if x.get('backup_type') == 'daily')}개")
        요약2.metric("수정 전 백업", f"{sum(1 for x in 백업목록 if x.get('backup_type') == 'edit')}개")
        요약3.metric("현재 거래 건수", f"{len(current_df) if current_df is not None else 0}건")

        수동칸1, 수동칸2, 수동칸3 = st.columns([1.0, 1.25, 2.75])
        with 수동칸1:
            if st.button("지금 수동 백업", key="manual_backup_now_v1", use_container_width=True):
                성공, 결과 = 자동백업파일저장(current_df, backup_type="edit", reason="manual_backup", source="backup_ui")
                if 성공:
                    st.success("수동 백업을 저장했습니다.")
                    st.rerun()
                else:
                    st.error(f"수동 백업 저장 실패: {결과}")
        with 수동칸2:
            try:
                현재상태엑셀 = 통합백업엑셀저장바이트(current_df, portfolio_df=portfolio_df, holding_df=holding_df)
                st.download_button(
                    "현재 상태 xlsx",
                    data=현재상태엑셀,
                    file_name=백업엑셀파일명(),
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="backup_excel_download_current_v58",
                    use_container_width=True,
                )
            except Exception as e:
                st.warning(f"xlsx 백업 준비 실패: {e}")
        with 수동칸3:
            st.caption(f"백업 위치: {자동백업루트폴더}/daily, {자동백업루트폴더}/edit_history")

        if not 백업목록:
            st.info("아직 생성된 자동백업 파일이 없습니다.")
            return

        표시항목 = []
        경로매핑 = {}
        for 항목 in 백업목록:
            저장시각 = 항목.get("saved_at", "-")
            종류 = "일일" if 항목.get("backup_type") == "daily" else "수정전"
            사유 = 항목.get("reason", "") or "-"
            표시문자 = f"[{종류}] {저장시각} · {항목.get('rows', 0)}건 · {사유}"
            표시항목.append(표시문자)
            경로매핑[표시문자] = 항목.get("path")

        선택값 = st.selectbox("복원 또는 다운로드할 백업 선택", 표시항목, key="backup_select_v1")
        선택경로 = 경로매핑.get(선택값)
        복원df, 메타 = 자동백업불러오기(선택경로) if 선택경로 else (None, {})

        if 복원df is not None:
            st.caption(
                f"선택 백업 정보: 유형={메타.get('backup_type', '-')} / 저장시각={메타.get('saved_at', '-')} / 건수={메타.get('rows', 0)} / 사유={메타.get('reason', '-')}"
            )
            미리보기 = 거래이력표시용변환(복원df.head(10))
            표데이터프레임(미리보기, use_container_width=True)

            선택백업계산포트폴리오 = pd.DataFrame()
            선택백업보유포트폴리오 = pd.DataFrame()
            선택백업엑셀 = None
            try:
                선택백업계산포트폴리오 = 포트폴리오계산(
                    거래이력계산대상추출(복원df),
                    refresh_token=st.session_state.get("price_refresh_token_v51", 0)
                )
                선택백업보유포트폴리오 = 보유포트폴리오필터(선택백업계산포트폴리오)
                선택백업엑셀 = 통합백업엑셀저장바이트(
                    복원df,
                    portfolio_df=선택백업계산포트폴리오,
                    holding_df=선택백업보유포트폴리오,
                )
            except Exception:
                선택백업엑셀 = None

            복원칸1, 복원칸2, 복원칸3 = st.columns(3)
            with 복원칸1:
                if st.button("선택 백업 복원", key="restore_backup_btn_v1", use_container_width=True):
                    반영df, 변경됨, 저장성공, 저장메시지 = 자동백업복원적용(복원df, source_label="backup_restore")
                    if 저장성공:
                        st.success(f"백업을 복원했습니다. ({len(반영df)}건)")
                    else:
                        st.warning(f"복원은 되었지만 자동저장 실패: {저장메시지}")
                    st.rerun()
            with 복원칸2:
                try:
                    with open(선택경로, "rb") as f:
                        백업바이트 = f.read()
                    st.download_button(
                        "선택 백업 다운로드",
                        data=백업바이트,
                        file_name=os.path.basename(선택경로),
                        mime="application/json",
                        key="download_backup_btn_v1",
                        use_container_width=True,
                    )
                except Exception as e:
                    st.warning(f"선택 백업 다운로드 준비 실패: {e}")
            with 복원칸3:
                if 선택백업엑셀 is not None:
                    st.download_button(
                        "선택 백업 xlsx",
                        data=선택백업엑셀,
                        file_name=백업엑셀파일명(prefix="selected_backup"),
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key="download_backup_xlsx_btn_v58",
                        use_container_width=True,
                    )
                else:
                    st.caption("선택 백업 xlsx 변환 준비 중")
        else:
            st.warning(f"백업 파일을 읽지 못했습니다: {메타.get('error', '알 수 없는 오류')}")


def 안전JSON저장(data, file_path):
    temp_path = f"{file_path}.tmp"
    backup_path = f"{file_path}.bak"

    try:
        with open(temp_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

        if os.path.exists(file_path):
            try:
                import shutil
                shutil.copy(file_path, backup_path)
            except Exception:
                pass

        os.replace(temp_path, file_path)
        return True, "저장 완료"
    except Exception as e:
        try:
            if os.path.exists(temp_path):
                os.remove(temp_path)
        except Exception:
            pass
        return False, f"저장 실패: {e}"




def 거래이력JSON변환(df):
    저장df = df.copy()

    표준열 = ["종목코드", "종목명", "거래일자", "거래구분", "거래수량", "거래단가", "운용사", "비고", "_입력원본순서"]
    if "_입력원본순서" not in 저장df.columns:
        저장df["_입력원본순서"] = range(len(저장df))
    else:
        저장df["_입력원본순서"] = pd.to_numeric(저장df["_입력원본순서"], errors="coerce")
        저장df["_입력원본순서"] = 저장df["_입력원본순서"].fillna(pd.Series(range(len(저장df)), index=저장df.index)).astype(int)

    for 열 in ["종목코드", "종목명", "거래일자", "거래구분", "거래수량", "거래단가", "운용사", "비고"]:
        if 열 not in 저장df.columns:
            저장df[열] = ""

    저장df = 저장df[표준열].copy()
    저장df["거래일자"] = pd.to_datetime(저장df["거래일자"], errors="coerce").dt.strftime("%Y-%m-%d")
    저장df["종목코드"] = 저장df["종목코드"].apply(lambda 값: "" if pd.isna(값) else str(값).zfill(6))
    저장df = 저장df.fillna("")

    return 저장df.to_dict(orient="records")


def 자동저장불러오기(file_path):
    if not os.path.exists(file_path):
        return None

    try:
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        df = pd.DataFrame(data)
        df = 거래이력표준열맞추기(df)
        df = 거래이력자동보정(df)

        if "거래일자" in df.columns:
            df["거래일자"] = pd.to_datetime(df["거래일자"], errors="coerce").dt.date

        return 거래이력입력창정렬(df)
    except Exception:
        return None


def 거래이력자동저장실행(df):
    저장데이터 = 거래이력JSON변환(df)
    return 안전JSON저장(저장데이터, 거래이력자동저장파일)


def 최근업로드거래이력저장(df, 파일명=""):
    저장데이터 = 거래이력JSON변환(df)
    저장성공, 저장메시지 = 안전JSON저장(저장데이터, 최근업로드거래이력파일)
    메타성공, 메타메시지 = True, "저장 완료"
    메타정보 = {
        "파일명": 파일명 or "",
        "저장시각": 서울현재시각ISO(),
        "건수": len(df) if df is not None else 0,
    }
    if 저장성공:
        메타성공, 메타메시지 = 안전JSON저장(메타정보, 최근업로드메타파일)
    return (저장성공 and 메타성공), (저장메시지 if not 저장성공 else 메타메시지)


def 최근업로드거래이력불러오기():
    return 자동저장불러오기(최근업로드거래이력파일)


def 최근업로드메타불러오기():
    if not os.path.exists(최근업로드메타파일):
        return None
    try:
        with open(최근업로드메타파일, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, dict) else None
    except Exception:
        return None



def 모니터관심종목불러오기():
    if not os.path.exists(모니터관심종목저장파일):
        return []
    try:
        with open(모니터관심종목저장파일, "r", encoding="utf-8") as f:
            data = json.load(f)
        if not isinstance(data, list):
            return []
        결과 = []
        seen = set()
        for 값 in data:
            코드 = re.sub(r"[^0-9]", "", str(값)).zfill(6)
            if 코드 and len(코드) == 6 and 코드 not in seen:
                결과.append(코드)
                seen.add(코드)
        return 결과[:4]
    except Exception:
        return []

def 모니터관심종목저장(codes):
    정리 = []
    seen = set()
    for 값 in list(codes or []):
        코드 = re.sub(r"[^0-9]", "", str(값)).zfill(6)
        if 코드 and len(코드) == 6 and 코드 not in seen:
            정리.append(코드)
            seen.add(코드)
    정리 = 정리[:4]
    return 안전JSON저장(정리, 모니터관심종목저장파일)

def 세션모니터관심종목가져오기():
    return []
def 세션모니터관심종목저장(codes):
    정리 = []
    seen = set()
    for 값 in list(codes or []):
        코드 = re.sub(r"[^0-9]", "", str(값)).zfill(6)
        if 코드 and len(코드) == 6 and 코드 not in seen:
            정리.append(코드)
            seen.add(코드)
    정리 = 정리[:4]
    st.session_state["monitor_custom_codes_v53"] = 정리
    return 모니터관심종목저장(정리)


def 모니터추가옵션목록생성():
    결과 = []
    for 코드, 이름 in sorted(코드명매핑.items(), key=lambda x: str(x[1])):
        if 코드 not in ["1001", "2001"]:
            결과.append(f"{이름} ({코드})")
    return 결과

def 모니터추가선택동기화():
    선택표시 = st.session_state.get("monitor_add_select_v53", "")
    if 선택표시:
        m = re.search(r"\((\d{6})\)$", str(선택표시))
        if m:
            st.session_state["monitor_add_code_v53"] = m.group(1)
            return
    st.session_state["monitor_add_code_v53"] = ""

def 모니터추가코드동기화():
    코드입력값 = st.session_state.get("monitor_add_code_v53", "")
    코드 = re.sub(r"[^0-9]", "", str(코드입력값)).zfill(6) if str(코드입력값).strip() else ""
    st.session_state["monitor_add_code_v53"] = 코드
    if 코드 and len(코드) == 6:
        이름 = 종목코드기준종목명(코드) or 코드명매핑.get(코드) or ""
        if 이름:
            st.session_state["monitor_add_select_v53"] = f"{이름} ({코드})"
            return
    if not 코드:
        st.session_state["monitor_add_select_v53"] = ""

def 거래이력서명생성(df):
    try:
        return json.dumps(거래이력JSON변환(df), ensure_ascii=False, sort_keys=True)
    except Exception:
        try:
            return json.dumps(pd.DataFrame(df).fillna("").astype(str).to_dict(orient="records"), ensure_ascii=False, sort_keys=True)
        except Exception:
            return ""



def 거래이력비교지문(df):
    """
    거래이력 비교용 지문 문자열
    - 저장/비교 목적
    - 기존 코드의 거래이력서명생성과 동일 계열 역할
    """
    try:
        return 거래이력서명생성(df)
    except Exception:
        try:
            작업 = pd.DataFrame(df).copy()
            if 작업 is None:
                return ""
            return json.dumps(거래이력JSON변환(작업), ensure_ascii=False, sort_keys=True)
        except Exception:
            try:
                return json.dumps(pd.DataFrame(df).fillna("").astype(str).to_dict(orient="records"), ensure_ascii=False, sort_keys=True)
            except Exception:
                return ""

def 거래이력세션반영(df, 저장강제=False, 자동저장허용=True):
    편집df = 거래이력편집용자동보정(df)
    계산df = 거래이력계산대상추출(편집df)

    새서명 = 거래이력서명생성(편집df)
    이전서명 = st.session_state.get("trade_history_signature_v1", "")
    이전편집df = st.session_state.get("trade_history_editor_df_v1", pd.DataFrame()).copy() if "trade_history_editor_df_v1" in st.session_state else pd.DataFrame()
    변경됨 = (새서명 != 이전서명)

    if 변경됨 and 이전서명:
        자동백업수정전실행(이전편집df, 편집df, source="trade_session_apply")

    st.session_state["trade_history_editor_df_v1"] = 편집df.copy()
    st.session_state["trade_history_df_v22"] = 편집df.copy()
    st.session_state["trade_history_calc_df_v1"] = 계산df.copy()
    st.session_state["trade_history_signature_v1"] = 새서명
    st.session_state["trade_history_changed_v1"] = 변경됨

    저장성공 = True
    저장메시지 = "변경 없음"
    if 저장강제 or (변경됨 and 자동저장허용):
        저장성공, 저장메시지 = 거래이력자동저장실행(편집df)
        if 저장성공:
            st.session_state["trade_history_last_saved_signature_v1"] = 새서명

    return 편집df, 변경됨, 저장성공, 저장메시지


@st.cache_data(ttl=30, show_spinner=False)
def 거래이력통합점검표캐시(거래이력json문자열):
    try:
        원본 = json.loads(거래이력json문자열)
        작업df = 거래이력정규화(pd.DataFrame(원본))
    except Exception:
        return pd.DataFrame(columns=["행", "점검항목", "현재값", "권장사항"])

    입력검증표 = 거래이력검증표생성(작업df)
    이상치점검표 = 거래이력이상치점검표생성(작업df)
    통합점검표 = pd.concat([입력검증표, 이상치점검표], ignore_index=True) if not 이상치점검표.empty else 입력검증표.copy()
    if not 통합점검표.empty:
        통합점검표 = 통합점검표.drop_duplicates().reset_index(drop=True)
    return 통합점검표


def 거래이력열이름정리(열이름):
    if pd.isna(열이름):
        return ""
    이름 = str(열이름).strip()
    이름 = 이름.replace("\n", "").replace("\r", "").replace("\t", "")
    이름 = 이름.replace(" ", "").replace("_", "").replace("-", "")
    이름 = 이름.replace("(", "").replace(")", "")
    return 이름


def 거래이력셀문자정리(값):
    if pd.isna(값):
        return 값
    if isinstance(값, str):
        return 값.replace("\xa0", " ").replace("\u200b", "").strip()
    return 값


def 거래이력컬럼명매핑():
    return {
        "종목코드": "종목코드",
        "코드": "종목코드",
        "티커": "종목코드",
        "종목번호": "종목코드",
        "종목명": "종목명",
        "종목": "종목명",
        "종목이름": "종목명",
        "이름": "종목명",
        "거래일자": "거래일자",
        "일자": "거래일자",
        "날짜": "거래일자",
        "매매일": "거래일자",
        "거래날짜": "거래일자",
        "거래구분": "거래구분",
        "구분": "거래구분",
        "매매구분": "거래구분",
        "매수매도": "거래구분",
        "거래수량": "거래수량",
        "수량": "거래수량",
        "매매수량": "거래수량",
        "체결수량": "거래수량",
        "거래단가": "거래단가",
        "단가": "거래단가",
        "가격": "거래단가",
        "체결가": "거래단가",
        "체결단가": "거래단가",
        "매수가": "거래단가",
        "매도가": "거래단가",
        "비고": "비고",
        "메모": "비고",
        "참고": "비고",
        "노트": "비고",
        "운용사": "운용사",
        "계좌": "운용사",
        "증권사": "운용사",
        "계좌명": "운용사",
        "운용계좌": "운용사",
    }


def 거래이력원본정리(df):
    결과 = df.copy()
    결과.columns = [거래이력열이름정리(c) for c in 결과.columns]
    컬럼매핑 = 거래이력컬럼명매핑()
    rename_map = {}
    for col in 결과.columns:
        if col in 컬럼매핑:
            rename_map[col] = 컬럼매핑[col]
    결과 = 결과.rename(columns=rename_map)

    for col in 결과.columns:
        try:
            결과[col] = 결과[col].apply(거래이력셀문자정리)
        except Exception:
            pass

    return 결과


def 거래이력표준열맞추기(df):
    표준열 = ["종목코드", "종목명", "거래일자", "거래구분", "거래수량", "거래단가", "운용사", "비고"]
    결과 = 거래이력원본정리(df)

    if "_입력원본순서" not in 결과.columns:
        결과["_입력원본순서"] = range(len(결과))

    for 열 in 표준열:
        if 열 not in 결과.columns:
            결과[열] = None if 열 in ["거래일자", "거래수량", "거래단가"] else ""

    return 결과[표준열 + ["_입력원본순서"]]


def 업로드파일에서거래이력읽기(업로드파일):
    파일명 = (업로드파일.name or "").lower()

    if 파일명.endswith(".json"):
        내용 = json.load(업로드파일)
        if isinstance(내용, dict) and "data" in 내용:
            내용 = 내용["data"]
        return 거래이력표준열맞추기(pd.DataFrame(내용))

    if 파일명.endswith(".xlsx") or 파일명.endswith(".xls"):
        원본바이트 = 업로드파일.getvalue()
        try:
            xls = pd.ExcelFile(io.BytesIO(원본바이트))
            시트명 = "거래이력" if "거래이력" in xls.sheet_names else xls.sheet_names[0]
            읽기df = pd.read_excel(io.BytesIO(원본바이트), sheet_name=시트명, dtype=object)
        except Exception:
            읽기df = pd.read_excel(io.BytesIO(원본바이트), dtype=object)
        return 거래이력표준열맞추기(읽기df)

    원본바이트 = 업로드파일.getvalue()
    마지막오류 = None

    for 인코딩 in ["utf-8-sig", "utf-8", "cp949", "euc-kr"]:
        try:
            읽기df = pd.read_csv(io.BytesIO(원본바이트), encoding=인코딩, dtype=object)
            return 거래이력표준열맞추기(읽기df)
        except Exception as e:
            마지막오류 = e

    raise 마지막오류


def 업로드파일에서비주식자산읽기(업로드파일):
    """통합 엑셀 파일의 '비주식자산' 시트를 읽어 앱 내부 표준 컬럼으로 변환합니다."""
    파일명 = (업로드파일.name or "").lower()
    if not (파일명.endswith(".xlsx") or 파일명.endswith(".xls")):
        return None

    원본바이트 = 업로드파일.getvalue()
    try:
        xls = pd.ExcelFile(io.BytesIO(원본바이트))
        if "비주식자산" not in xls.sheet_names:
            return None
        읽기df = pd.read_excel(io.BytesIO(원본바이트), sheet_name="비주식자산", dtype=object)
        return IRP비주식자산표준열맞추기(읽기df)
    except Exception as e:
        raise e


def 통합엑셀업로드여부(업로드파일):
    try:
        파일명 = (업로드파일.name or "").lower()
        if not (파일명.endswith(".xlsx") or 파일명.endswith(".xls")):
            return False
        xls = pd.ExcelFile(io.BytesIO(업로드파일.getvalue()))
        return ("거래이력" in xls.sheet_names) and ("비주식자산" in xls.sheet_names)
    except Exception:
        return False


# -----------------------------------
# 표시용 함수
# -----------------------------------
def 금액표시(값):
    if pd.isna(값) or 값 is None:
        return "-"
    return f"{값:,.0f}원"


def 숫자표시(값, 소수점=0):
    if pd.isna(값) or 값 is None:
        return "-"
    if 소수점 == 0:
        return f"{값:,.0f}"
    return f"{값:,.{소수점}f}"


def 안전정수포맷(값):
    if pd.isna(값) or 값 is None:
        return "-"
    try:
        return f"{float(값):,.0f}"
    except Exception:
        return "-"


def 안전소수포맷(값, 소수점=2):
    if pd.isna(값) or 값 is None:
        return "-"
    try:
        return f"{float(값):,.{소수점}f}"
    except Exception:
        return "-"


def 비율표시(값):
    if pd.isna(값) or 값 is None:
        return "-"
    return f"{값:.2f}%"


def 증감문자열(값, suffix=""):
    if pd.isna(값) or 값 is None:
        return "-"
    if 값 > 0:
        return f"+{값:,.2f}{suffix}"
    return f"{값:,.2f}{suffix}"


def 기준일시표시문자열(기준일=None, 조회시각=None):
    기준문자 = "-"
    if 기준일 is not None and not pd.isna(기준일):
        try:
            기준문자 = pd.to_datetime(기준일).strftime("%Y-%m-%d")
        except Exception:
            기준문자 = str(기준일)
    조회문자 = ""
    if 조회시각 is not None and not pd.isna(조회시각):
        try:
            조회문자 = pd.to_datetime(조회시각).strftime("%Y-%m-%d %H:%M")
        except Exception:
            조회문자 = str(조회시각)
    return f"기준 {기준문자} · 조회 {조회문자}" if 조회문자 else f"기준 {기준문자}"


def 통화문자정리(값):
    if pd.isna(값) or 값 is None:
        return None
    if isinstance(값, (int, float)):
        return float(값)
    문자열 = str(값).strip()
    if 문자열 == "":
        return None
    문자열 = 문자열.replace("₩", "").replace("￦", "").replace("원", "").replace(",", "").strip()
    문자열 = re.sub(r"[^0-9.\-]", "", 문자열)
    if 문자열 in ["", ".", "-", "-."]:
        return None
    try:
        return float(문자열)
    except Exception:
        return None


def 거래단가표시문자열(값):
    숫자값 = 통화문자정리(값)
    if 숫자값 is None:
        return ""
    return f"₩ {int(round(숫자값)):,.0f}"


def 거래이력입력창정렬(df):
    if df is None:
        return pd.DataFrame(columns=["종목코드", "종목명", "거래일자", "거래구분", "거래수량", "거래단가", "비고"])

    결과 = df.copy()

    for 컬럼 in ["종목코드", "종목명", "거래일자", "거래구분", "거래수량", "거래단가", "운용사", "비고"]:
        if 컬럼 not in 결과.columns:
            결과[컬럼] = None if 컬럼 in ["거래일자", "거래수량", "거래단가"] else ""

    if "_입력원본순서" not in 결과.columns:
        결과["_입력원본순서"] = range(len(결과))

    결과 = 결과.dropna(how="all")
    결과["_입력원본순서"] = pd.to_numeric(
        결과["_입력원본순서"], errors="coerce"
    ).fillna(pd.Series(range(len(결과)), index=결과.index)).astype(int)

    # 표시 순서는 엑셀/입력 시트 원본 순서를 그대로 유지
    결과 = 결과.sort_values(by=["_입력원본순서"], ascending=[True], na_position="last", kind="stable")
    결과 = 결과.reset_index(drop=True)
    return 결과


def 거래이력정규화(df):
    if df is None:
        return pd.DataFrame(columns=["종목코드", "종목명", "거래일자", "거래구분", "거래수량", "거래단가", "운용사", "비고"])
    원본 = df.copy()
    if "_입력원본순서" not in 원본.columns:
        원본["_입력원본순서"] = range(len(원본))
    정규화 = 거래이력표준열맞추기(원본)
    정규화["_입력원본순서"] = 원본["_입력원본순서"].values
    정규화 = 거래이력자동보정(정규화)
    정규화 = 거래이력입력창정렬(정규화)
    return 정규화


def 보유포트폴리오필터(df):
    if df is None or df.empty:
        return pd.DataFrame(columns=df.columns if df is not None else [])
    결과 = df.copy()
    if "보유수량" not in 결과.columns:
        return 결과.iloc[0:0].copy()
    결과["보유수량"] = pd.to_numeric(결과["보유수량"], errors="coerce").fillna(0)
    결과 = 결과[결과["보유수량"] > 0].copy()
    if 결과.empty:
        return 결과
    if "평가금액" in 결과.columns:
        결과 = 결과.sort_values(["평가금액", "종목명"], ascending=[False, True])
    else:
        결과 = 결과.sort_values(["종목명", "종목코드"], ascending=[True, True])
    return 결과.reset_index(drop=True)


def 보유종목선택옵션생성(df):
    보유df = 보유포트폴리오필터(df)
    if 보유df.empty:
        return []
    옵션 = []
    for _, 행 in 보유df.iterrows():
        종목코드 = str(행.get("종목코드", "")).zfill(6)
        종목명 = 종목명자동보정(종목코드, 행.get("종목명", ""))
        옵션.append({
            "종목코드": 종목코드,
            "종목명": 종목명,
            "표시": f"{종목명} ({종목코드})",
        })
    return 옵션



def _거래이력캐시초기화():
    if "trade_editor_last_input_fp_v1" not in st.session_state:
        st.session_state["trade_editor_last_input_fp_v1"] = ""
    if "trade_editor_last_output_fp_v1" not in st.session_state:
        st.session_state["trade_editor_last_output_fp_v1"] = ""
    if "trade_calc_cache_key_v1" not in st.session_state:
        st.session_state["trade_calc_cache_key_v1"] = ""
    if "trade_calc_cache_df_v1" not in st.session_state:
        st.session_state["trade_calc_cache_df_v1"] = pd.DataFrame()
    if "trade_check_cache_df_v1" not in st.session_state:
        st.session_state["trade_check_cache_df_v1"] = pd.DataFrame()
    if "portfolio_cache_key_v1" not in st.session_state:
        st.session_state["portfolio_cache_key_v1"] = ""
    if "portfolio_cache_df_v1" not in st.session_state:
        st.session_state["portfolio_cache_df_v1"] = pd.DataFrame()
    if "portfolio_holding_cache_df_v1" not in st.session_state:
        st.session_state["portfolio_holding_cache_df_v1"] = pd.DataFrame()
    if "portfolio_option_cache_v1" not in st.session_state:
        st.session_state["portfolio_option_cache_v1"] = []

def 거래이력편집반영최적화(편집입력df):
    _거래이력캐시초기화()
    입력지문 = 거래이력비교지문(편집입력df)
    이전입력지문 = st.session_state.get("trade_editor_last_input_fp_v1", "")

    if 입력지문 == 이전입력지문 and "trade_history_editor_df_v1" in st.session_state:
        편집df = st.session_state.get("trade_history_editor_df_v1", pd.DataFrame()).copy()
        거래이력변경됨 = False
        자동저장성공 = True
        자동저장메시지 = "변경 없음"
    else:
        편집df = 거래이력편집용자동보정(편집입력df.reset_index(drop=True))
        편집df, 거래이력변경됨, 자동저장성공, 자동저장메시지 = 거래이력세션반영(
            편집df,
            저장강제=False,
            자동저장허용=True,
        )
        st.session_state["trade_editor_last_input_fp_v1"] = 입력지문
        st.session_state["trade_editor_last_output_fp_v1"] = 거래이력비교지문(편집df)

    계산용거래이력 = st.session_state.get("trade_history_calc_df_v1", 거래이력계산대상추출(편집df))
    계산지문 = 거래이력서명생성(계산용거래이력)

    if 계산지문 != st.session_state.get("trade_calc_cache_key_v1", ""):
        통합점검표 = 거래이력통합점검표캐시(계산지문)
        st.session_state["trade_calc_cache_key_v1"] = 계산지문
        st.session_state["trade_calc_cache_df_v1"] = 계산용거래이력.copy()
        st.session_state["trade_check_cache_df_v1"] = 통합점검표.copy()
    else:
        통합점검표 = st.session_state.get("trade_check_cache_df_v1", pd.DataFrame()).copy()

    포트폴리오캐시키 = 계산지문 + f"|{st.session_state.get('price_refresh_token_v51', 0)}"
    if 포트폴리오캐시키 != st.session_state.get("portfolio_cache_key_v1", ""):
        계산포트폴리오 = 포트폴리오계산캐시(
            계산지문,
            refresh_token=st.session_state.get("price_refresh_token_v51", 0)
        )
        보유계산포트폴리오 = 보유포트폴리오필터(계산포트폴리오)
        보유종목옵션 = 보유종목선택옵션생성(계산포트폴리오)
        st.session_state["portfolio_cache_key_v1"] = 포트폴리오캐시키
        st.session_state["portfolio_cache_df_v1"] = 계산포트폴리오.copy()
        st.session_state["portfolio_holding_cache_df_v1"] = 보유계산포트폴리오.copy()
        st.session_state["portfolio_option_cache_v1"] = list(보유종목옵션)
    else:
        계산포트폴리오 = st.session_state.get("portfolio_cache_df_v1", pd.DataFrame()).copy()
        보유계산포트폴리오 = st.session_state.get("portfolio_holding_cache_df_v1", pd.DataFrame()).copy()
        보유종목옵션 = list(st.session_state.get("portfolio_option_cache_v1", []))

    return {
        "편집df": 편집df,
        "거래이력변경됨": 거래이력변경됨,
        "자동저장성공": 자동저장성공,
        "자동저장메시지": 자동저장메시지,
        "계산용거래이력": 계산용거래이력,
        "통합점검표": 통합점검표,
        "계산포트폴리오": 계산포트폴리오,
        "보유계산포트폴리오": 보유계산포트폴리오,
        "보유종목옵션": 보유종목옵션,
    }

def 현재거래내역엑셀저장바이트(df):
    저장대상 = df.copy()

    for 컬럼 in ["종목코드", "종목명", "거래일자", "거래구분", "거래수량", "거래단가", "운용사", "비고"]:
        if 컬럼 not in 저장대상.columns:
            저장대상[컬럼] = ""

    저장대상 = 저장대상[["종목코드", "종목명", "거래일자", "거래구분", "거래수량", "거래단가", "운용사", "비고"]].copy()
    저장대상["거래일자"] = pd.to_datetime(저장대상["거래일자"], errors="coerce").dt.strftime("%Y-%m-%d")
    저장대상["종목코드"] = 저장대상["종목코드"].apply(lambda 값: "" if pd.isna(값) else str(값).zfill(6))
    저장대상 = 저장대상.fillna("")

    버퍼 = io.BytesIO()
    with pd.ExcelWriter(버퍼, engine="openpyxl") as writer:
        저장대상.to_excel(writer, index=False, sheet_name="거래이력")
    버퍼.seek(0)
    return 버퍼.getvalue()


def 엑셀시트문자열정리(df):
    결과 = df.copy()
    for 컬럼 in 결과.columns:
        if "일자" in str(컬럼) or "시각" in str(컬럼) or "조회" in str(컬럼):
            try:
                결과[컬럼] = pd.to_datetime(결과[컬럼], errors="coerce").dt.strftime("%Y-%m-%d")
                결과[컬럼] = 결과[컬럼].fillna("")
            except Exception:
                pass
    return 결과.fillna("")


def 백업엑셀파일명(prefix="stock_backup"):
    return f"{prefix}_{서울현재시각().strftime('%Y-%m-%d_%H%M')}.xlsx"


def 통합백업엑셀저장바이트(current_df, portfolio_df=None, holding_df=None):
    거래원장 = 거래이력편집용자동보정(current_df if current_df is not None else pd.DataFrame())
    계산포트폴리오 = portfolio_df.copy() if isinstance(portfolio_df, pd.DataFrame) else None
    보유포트폴리오 = holding_df.copy() if isinstance(holding_df, pd.DataFrame) else None

    if 계산포트폴리오 is None:
        try:
            계산포트폴리오 = 포트폴리오계산(
                거래이력계산대상추출(거래원장),
                refresh_token=st.session_state.get("price_refresh_token_v51", 0)
            )
        except Exception:
            계산포트폴리오 = pd.DataFrame()

    if 보유포트폴리오 is None:
        try:
            보유포트폴리오 = 보유포트폴리오필터(계산포트폴리오)
        except Exception:
            보유포트폴리오 = pd.DataFrame()

    거래시트 = 거래원장.copy()
    for 컬럼 in ["종목코드", "종목명", "거래일자", "거래구분", "거래수량", "거래단가", "운용사", "비고"]:
        if 컬럼 not in 거래시트.columns:
            거래시트[컬럼] = ""
    거래시트 = 거래시트[["종목코드", "종목명", "거래일자", "거래구분", "거래수량", "거래단가", "운용사", "비고"]].copy()
    거래시트["거래일자"] = pd.to_datetime(거래시트["거래일자"], errors="coerce").dt.strftime("%Y-%m-%d")
    거래시트["종목코드"] = 거래시트["종목코드"].apply(lambda 값: "" if pd.isna(값) else str(값).zfill(6))
    거래시트 = 거래시트.fillna("")

    보유열 = ["종목코드", "종목명", "최초매수일자", "최근거래일자", "총매수수량", "총매도수량", "보유수량", "매입평균단가", "현재가", "투자원금", "평가금액", "평가손익", "실현손익", "수익률", "현재비중", "데이터상태"]
    보유시트 = 보유포트폴리오.copy() if isinstance(보유포트폴리오, pd.DataFrame) else pd.DataFrame()
    if "종목코드" in 보유시트.columns:
        보유시트["종목명"] = 보유시트.apply(lambda 행: 종목명자동보정(행.get("종목코드", ""), 행.get("종목명", "")), axis=1)
    for 컬럼 in 보유열:
        if 컬럼 not in 보유시트.columns:
            보유시트[컬럼] = None
    보유시트 = 보유시트[보유열].copy()
    보유시트 = 보유시트.rename(columns={
        "최초매수일자": "최초 매수일자",
        "최근거래일자": "최근 거래일자",
        "총매수수량": "총 매수수량",
        "총매도수량": "총 매도수량",
        "매입평균단가": "매입 평균단가",
    })
    보유시트 = 엑셀시트문자열정리(보유시트)

    손익열 = ["종목코드", "종목명", "보유수량", "투자원금", "평가금액", "평가손익", "실현손익", "수익률", "최근거래일자", "데이터상태"]
    손익상세 = 계산포트폴리오.copy() if isinstance(계산포트폴리오, pd.DataFrame) else pd.DataFrame()
    if "종목코드" in 손익상세.columns:
        손익상세["종목명"] = 손익상세.apply(lambda 행: 종목명자동보정(행.get("종목코드", ""), 행.get("종목명", "")), axis=1)
    for 컬럼 in 손익열:
        if 컬럼 not in 손익상세.columns:
            손익상세[컬럼] = None
    손익상세 = 손익상세[손익열].copy()
    손익상세 = 손익상세.rename(columns={"최근거래일자": "최근 거래일자"})
    손익상세 = 엑셀시트문자열정리(손익상세)

    정상평가행 = 보유포트폴리오[보유포트폴리오["데이터상태"] == "정상"].copy() if isinstance(보유포트폴리오, pd.DataFrame) and not 보유포트폴리오.empty and "데이터상태" in 보유포트폴리오.columns else (보유포트폴리오.copy() if isinstance(보유포트폴리오, pd.DataFrame) else pd.DataFrame())
    총투자원금 = pd.to_numeric(정상평가행.get("투자원금"), errors="coerce").fillna(0).sum() if not 정상평가행.empty else 0.0
    총평가금액 = pd.to_numeric(정상평가행.get("평가금액"), errors="coerce").fillna(0).sum() if not 정상평가행.empty else 0.0
    총평가손익 = pd.to_numeric(정상평가행.get("평가손익"), errors="coerce").fillna(0).sum() if not 정상평가행.empty else 0.0
    총실현손익 = pd.to_numeric(계산포트폴리오.get("실현손익"), errors="coerce").fillna(0).sum() if isinstance(계산포트폴리오, pd.DataFrame) and not 계산포트폴리오.empty else 0.0
    총수익률 = (총평가손익 / 총투자원금 * 100) if 총투자원금 not in [0, None] else 0.0

    손익요약 = pd.DataFrame([
        {"항목": "저장 시각", "값": 서울현재시각().strftime("%Y-%m-%d %H:%M:%S")},
        {"항목": "앱 버전", "값": APP_VERSION},
        {"항목": "거래 건수", "값": len(거래시트)},
        {"항목": "보유 종목 수", "값": len(보유시트)},
        {"항목": "총 투자원금", "값": 총투자원금},
        {"항목": "총 평가금액", "값": 총평가금액},
        {"항목": "총 평가손익", "값": 총평가손익},
        {"항목": "총 실현손익", "값": 총실현손익},
        {"항목": "총 수익률(%)", "값": 총수익률},
    ])

    백업정보 = pd.DataFrame([
        {"항목": "저장 시각", "값": 서울현재시각().strftime("%Y-%m-%d %H:%M:%S")},
        {"항목": "앱 버전", "값": APP_VERSION},
        {"항목": "백업 형식", "값": "xlsx"},
        {"항목": "거래 원장 행 수", "값": len(거래시트)},
        {"항목": "보유 현황 행 수", "값": len(보유시트)},
        {"항목": "손익 상세 행 수", "값": len(손익상세)},
        {"항목": "데이터 지문", "값": 거래이력비교지문(거래원장)},
    ])

    버퍼 = io.BytesIO()
    with pd.ExcelWriter(버퍼, engine="openpyxl") as writer:
        거래시트.to_excel(writer, index=False, sheet_name="거래내역")
        보유시트.to_excel(writer, index=False, sheet_name="보유현황")
        손익요약.to_excel(writer, index=False, sheet_name="손익현황", startrow=0)
        손익상세.to_excel(writer, index=False, sheet_name="손익현황", startrow=len(손익요약) + 2)
        백업정보.to_excel(writer, index=False, sheet_name="백업정보")
    버퍼.seek(0)
    return 버퍼.getvalue()

def 거래이력표시용변환(df):
    표시 = df.copy()
    if "거래단가" in 표시.columns:
        표시["거래단가"] = 표시["거래단가"].apply(거래단가표시문자열)
    return 표시


def 손익색상(값):
    if pd.isna(값):
        return ""
    if 값 > 0:
        return "color: red; font-weight: 600;"
    if 값 < 0:
        return "color: blue; font-weight: 600;"
    return ""


def 수익률색상(값):
    return 손익색상(값)


def 손익문자열(값):
    if pd.isna(값):
        return "-"
    if 값 > 0:
        return f"+{값:,.0f}"
    return f"{값:,.0f}"


def 수익률문자열(값):
    if pd.isna(값):
        return "-"
    if 값 > 0:
        return f"+{값:.2f}%"
    return f"{값:.2f}%"




def 정렬대상숫자열여부(series, 컬럼명=""):
    이름 = str(컬럼명).strip()
    숫자키워드 = ["금액", "가격", "단가", "수량", "비중", "비율", "손익", "수익률", "평가", "합계", "총", "잔액", "점수", "값"]
    if any(키 in 이름 for 키 in 숫자키워드):
        return True

    try:
        if pd.api.types.is_numeric_dtype(series):
            return True
    except Exception:
        pass

    try:
        비결측 = pd.Series(series).dropna().astype(str).str.strip()
        if 비결측.empty:
            return False
        샘플 = 비결측.head(20)
        숫자형비율 = 0
        for 값 in 샘플:
            값정리 = str(값).replace(",", "").replace("원", "").replace("%", "").replace("주", "").replace("배", "").strip()
            값정리 = re.sub(r"[^0-9.\-+]", "", 값정리)
            if 값정리 not in ["", ".", "-", "+", "-.", "+."]:
                try:
                    float(값정리)
                    숫자형비율 += 1
                except Exception:
                    pass
        return (숫자형비율 / max(len(샘플), 1)) >= 0.7
    except Exception:
        return False





def 표데이터프레임(입력객체, use_container_width=True, hide_index=False, **kwargs):
    try:
        from pandas.io.formats.style import Styler
    except Exception:
        Styler = None

    스타일객체 = None
    원본df = None

    if Styler is not None and isinstance(입력객체, Styler):
        스타일객체 = 입력객체
        원본df = 스타일객체.data.copy()
    elif isinstance(입력객체, pd.DataFrame):
        원본df = 입력객체.copy()
        스타일객체 = 원본df.style
    else:
        st.dataframe(입력객체, use_container_width=use_container_width, hide_index=hide_index, **kwargs)
        return

    if hide_index:
        try:
            스타일객체 = 스타일객체.hide(axis="index")
        except Exception:
            pass

    try:
        모든열 = list(원본df.columns)
        좌측정렬열 = [열 for 열 in 모든열 if not 정렬대상숫자열여부(원본df[열], 열)]
        우측정렬열 = [열 for 열 in 모든열 if 정렬대상숫자열여부(원본df[열], 열)]

        if 좌측정렬열:
            스타일객체 = 스타일객체.set_properties(
                subset=좌측정렬열,
                **{"text-align": "left"}
            )
        if 우측정렬열:
            스타일객체 = 스타일객체.set_properties(
                subset=우측정렬열,
                **{
                    "text-align": "right",
                    "font-variant-numeric": "tabular-nums",
                    "font-feature-settings": '"tnum"',
                }
            )

        좁은열 = ["종목코드", "구분", "거래구분", "데이터상태", "권장방향", "판정", "현재비중", "수익률", "점수", "가격 위치"]
        넓은열 = ["종목명", "설명", "비고", "기준", "현재", "항목"]
        날짜열 = ["일자"]
        금액열 = ["금액", "가격", "단가", "현재가", "평가", "손익", "원금"]
        수량열 = ["수량", "총", "잔액", "배수"]

        for idx, 열이름 in enumerate(모든열):
            이름 = str(열이름).strip()
            폭 = None
            if 이름 in 좁은열 or any(키 == 이름 for 키 in 좁은열):
                폭 = "90px"
            elif any(키 in 이름 for 키 in 날짜열):
                폭 = "108px"
            elif any(키 in 이름 for 키 in 금액열):
                폭 = "112px"
            elif any(키 in 이름 for 키 in 수량열):
                폭 = "82px"
            elif 이름 in 넓은열 or any(키 == 이름 for 키 in 넓은열):
                폭 = "132px"

            if 폭:
                스타일객체 = 스타일객체.set_table_styles([
                    {"selector": f".col_heading.col{idx}", "props": [("min-width", 폭)]},
                    {"selector": f".data.col{idx}", "props": [("min-width", 폭)]},
                ], overwrite=False)

        스타일객체 = 스타일객체.set_table_styles([
            {"selector": "table", "props": [("width", "100%"), ("border-collapse", "collapse"), ("font-size", "0.98rem")]},
            {"selector": "thead th", "props": [("text-align", "center"), ("vertical-align", "middle"), ("font-weight", "700"), ("white-space", "normal"), ("line-height", "1.32"), ("padding", "10px 10px")]},
            {"selector": "tbody th", "props": [("text-align", "right"), ("font-variant-numeric", "tabular-nums"), ("vertical-align", "middle"), ("padding", "9px 10px"), ("width", "44px")]},
            {"selector": "td", "props": [("padding", "9px 10px"), ("vertical-align", "middle"), ("line-height", "1.38")]},
            {"selector": "td.col0", "props": [("text-align", "left")]},
            {"selector": "td.col1", "props": [("text-align", "left")]},
        ], overwrite=False)

    except Exception:
        pass

    html = 스타일객체.to_html()
    래퍼스타일 = "width:100%; overflow-x:auto;" if use_container_width else "overflow-x:auto;"
    st.markdown("<div class='oa-table-wrap' style='" + 래퍼스타일 + "'>" + html + "</div>", unsafe_allow_html=True)


def index_1부터(df):
    표시용 = df.copy()
    표시용.index = range(1, len(표시용) + 1)
    return 표시용


def 모바일차트높이(데스크탑높이=460, 모바일높이=360):
    return 모바일높이 if 모바일여부() else 데스크탑높이


def 거래이력표_컬럼선택(df):
    if 모바일여부():
        사용컬럼 = [c for c in ["거래일자", "종목명", "거래구분", "거래수량", "거래단가"] if c in df.columns]
        if 사용컬럼:
            return df[사용컬럼].copy()
    return df.copy()


def 포트폴리오표_컬럼선택(df):
    if 모바일여부():
        사용컬럼 = [c for c in ["종목명", "보유수량", "현재가", "평가금액", "수익률"] if c in df.columns]
        if 사용컬럼:
            return df[사용컬럼].copy()
    return df.copy()


def 리밸런싱표_컬럼선택(df):
    if 모바일여부():
        사용컬럼 = [c for c in ["종목명", "현재비중", "목표비중", "권장방향"] if c in df.columns]
        if 사용컬럼:
            return df[사용컬럼].copy()
    return df.copy()


def 안전실수변환(값):
    if 값 is None:
        return None
    if isinstance(값, (int, float)):
        return float(값)
    문자열 = re.sub(r"[^0-9.\-]", "", str(값))
    if 문자열 in ["", ".", "-", "-."]:
        return None
    try:
        return float(문자열)
    except Exception:
        return None



def 유효숫자인지(값):
    try:
        return 값 is not None and not pd.isna(값) and np.isfinite(float(값))
    except Exception:
        return False


def 마지막유효값시리즈(series):
    if series is None:
        return None
    try:
        정리 = pd.to_numeric(series, errors="coerce").dropna()
        if 정리.empty:
            return None
        return float(정리.iloc[-1])
    except Exception:
        return None


def 끝에서두번째유효값시리즈(series):
    if series is None:
        return None
    try:
        정리 = pd.to_numeric(series, errors="coerce").dropna()
        if len(정리) < 2:
            return None
        return float(정리.iloc[-2])
    except Exception:
        return None


def 시리즈길이맞추기(values, target_len):
    값목록 = list(values) if values is not None else []
    if len(값목록) < target_len:
        값목록 += [None] * (target_len - len(값목록))
    return 값목록[:target_len]


def OHLCV데이터정리(df):
    if df is None or df.empty:
        return pd.DataFrame()
    작업 = df.copy()
    if "날짜" in 작업.columns:
        작업["날짜"] = pd.to_datetime(작업["날짜"], errors="coerce")
        작업 = 작업.dropna(subset=["날짜"])
        작업["날짜"] = 작업["날짜"].dt.tz_localize(None)
        작업 = 작업.sort_values("날짜").drop_duplicates(subset=["날짜"], keep="last").set_index("날짜")
    for col in ["시가", "고가", "저가", "종가", "거래량"]:
        if col not in 작업.columns:
            작업[col] = np.nan
        작업[col] = pd.to_numeric(작업[col], errors="coerce")
    작업 = 작업.dropna(how="all", subset=["시가", "고가", "저가", "종가"])
    return 작업.sort_index()


def 최근유효OHLCV요약(df):
    if df is None or df.empty:
        return {"현재가": None, "전일가": None, "전일대비": None, "등락률": None, "기준일": None, "상태": "조회 실패"}
    작업 = OHLCV데이터정리(df)
    if 작업.empty or "종가" not in 작업.columns:
        return {"현재가": None, "전일가": None, "전일대비": None, "등락률": None, "기준일": None, "상태": "조회 실패"}
    현재가 = 마지막유효값시리즈(작업["종가"])
    전일가 = 끝에서두번째유효값시리즈(작업["종가"])
    if 현재가 is None:
        return {"현재가": None, "전일가": None, "전일대비": None, "등락률": None, "기준일": None, "상태": "조회 실패"}
    기준일 = pd.to_datetime(작업.index[-1]).date() if len(작업.index) > 0 else None
    전일대비 = None if 전일가 is None else 현재가 - 전일가
    등락률 = None if 전일가 in [None, 0] else (전일대비 / 전일가) * 100
    상태 = "직전 종가 반영" if 전일가 is not None else "최근 유효 종가 반영"
    return {"현재가": 현재가, "전일가": 전일가, "전일대비": 전일대비, "등락률": 등락률, "기준일": 기준일, "상태": 상태}


def 종목코드별야후심볼후보(코드):
    코드 = str(코드).zfill(6)
    return [f"{코드}.KS", f"{코드}.KQ"]


def _국내장중1분봉신선여부(기준시각):
    try:
        if 기준시각 is None or pd.isna(기준시각):
            return False
        ts = pd.Timestamp(기준시각)
        if getattr(ts, "tzinfo", None) is None:
            ts = ts.tz_localize("Asia/Seoul")
        else:
            ts = ts.tz_convert("Asia/Seoul")
        now = pd.Timestamp(서울현재시각())
        if getattr(now, "tzinfo", None) is None:
            now = now.tz_localize("Asia/Seoul")
        else:
            now = now.tz_convert("Asia/Seoul")
        if ts.date() != now.date():
            return False
        차이분 = abs((now - ts).total_seconds()) / 60.0
        return 차이분 <= 20
    except Exception:
        return False


@st.cache_data(ttl=8, show_spinner=False)
def 네이버국내현재가가져오기(구분, 코드, refresh_token=0):
    """국내 주식·ETF·지수 현재가를 Naver 기준으로 조회합니다.
    - 시세 새로고침 토큰을 인자로 받아 버튼 클릭 시 캐시가 반드시 분리됩니다.
    - 현재가, 전일대비, 등락률을 같은 출처 기준으로 정리합니다.
    - 실패 시 None을 반환하여 Yahoo/최근종가 보조 로직으로 넘어갑니다.
    """
    원코드 = str(코드).strip()
    코드문자 = 원코드.zfill(6) if 원코드.isdigit() and 구분 != "index" else 원코드

    if 구분 == "index":
        지수코드 = "KOSPI" if str(원코드) in ["1001", "KOSPI", "KS11", "^KS11"] else "KOSDAQ"
        url = f"https://finance.naver.com/sise/sise_index.naver?code={지수코드}"
    else:
        url = f"https://finance.naver.com/item/main.naver?code={코드문자}"

    def _html_text(s):
        return re.sub(r"\s+", " ", str(s or "")).strip()

    def _숫자목록(s):
        return [안전실수변환(x) for x in re.findall(r"[-+]?\d[\d,]*\.?\d*", str(s or ""))]

    try:
        응답 = 안전웹요청(url, timeout=3.5, attempts=2)
        if 응답 is None:
            return None
        html = 응답.text

        현재가 = None
        전일대비 = None
        등락률 = None
        방향 = 0

        if 구분 == "index":
            # Naver 지수 페이지: now_value / change_value_and_rate 영역 우선 사용
            now_m = re.search(r'id=["\']now_value["\'][^>]*>\s*([^<]+)\s*</', html, re.I | re.S)
            if now_m:
                현재가 = 안전실수변환(now_m.group(1))

            change_m = re.search(r'id=["\']change_value_and_rate["\'][^>]*>\s*([^<]+)\s*</', html, re.I | re.S)
            변화문구 = _html_text(change_m.group(1)) if change_m else ""
            nums = _숫자목록(변화문구)
            if nums:
                전일대비 = nums[0]
            if len(nums) >= 2:
                등락률 = nums[1]

            주변 = 변화문구 + " " + _html_text(html[max(0, (change_m.start() if change_m else 0)-300):(change_m.end() if change_m else 0)+300])
            if any(x in 주변 for x in ["하락", "▼", "nv_down", "no_down"]):
                방향 = -1
            elif any(x in 주변 for x in ["상승", "▲", "nv_up", "no_up"]):
                방향 = 1

        else:
            # Naver 종목 페이지: _nowVal, _diff, _rate 우선 사용
            for ptn in [
                r'id=["\']_nowVal["\'][^>]*>\s*([^<]+)\s*</',
                r'<p[^>]*class=["\']no_today["\'][^>]*>.*?<span[^>]*class=["\']blind["\'][^>]*>\s*([^<]+)\s*</span>',
            ]:
                m = re.search(ptn, html, re.I | re.S)
                if m:
                    현재가 = 안전실수변환(m.group(1))
                    if 현재가 not in [None, 0]:
                        break

            diff_m = re.search(r'id=["\']_diff["\'][^>]*>\s*([^<]+)\s*</', html, re.I | re.S)
            rate_m = re.search(r'id=["\']_rate["\'][^>]*>\s*([^<%]+)%?\s*</', html, re.I | re.S)
            if diff_m:
                전일대비 = 안전실수변환(diff_m.group(1))
            if rate_m:
                등락률 = 안전실수변환(rate_m.group(1))

            # 방향은 _diff 주변 클래스와 blind 텍스트를 함께 확인
            diff_주변 = ""
            if diff_m:
                diff_주변 = _html_text(html[max(0, diff_m.start()-600):diff_m.end()+600])
            blind_text = " ".join(re.findall(r'<span[^>]*class=["\']blind["\'][^>]*>([^<]+)</span>', html, re.I | re.S)[:40])
            주변 = diff_주변 + " " + blind_text
            if any(x in 주변 for x in ["하락", "▼", "nv_down", "no_down", "class=\"dn\"", "class='dn'"]):
                방향 = -1
            elif any(x in 주변 for x in ["상승", "▲", "nv_up", "no_up", "class=\"up\"", "class='up'"]):
                방향 = 1

        if 현재가 in [None, 0]:
            return None

        if 방향 < 0:
            if 전일대비 not in [None, 0]:
                전일대비 = -abs(float(전일대비))
            if 등락률 not in [None, 0]:
                등락률 = -abs(float(등락률))
        elif 방향 > 0:
            if 전일대비 not in [None, 0]:
                전일대비 = abs(float(전일대비))
            if 등락률 not in [None, 0]:
                등락률 = abs(float(등락률))

        전일가 = None
        if 전일대비 is not None:
            전일가 = float(현재가) - float(전일대비)
        elif 등락률 not in [None, 0]:
            전일가 = float(현재가) / (1 + float(등락률) / 100.0)
            전일대비 = float(현재가) - float(전일가)

        # 전일가가 확보되면 등락률은 현재가 기준으로 다시 계산해 기준 혼합을 막습니다.
        if 전일가 not in [None, 0]:
            전일대비 = float(현재가) - float(전일가)
            등락률 = (전일대비 / float(전일가)) * 100.0

        return {
            "현재가": float(현재가),
            "전일가": None if 전일가 in [None, 0] else float(전일가),
            "전일대비": None if 전일대비 is None else float(전일대비),
            "등락률": None if 등락률 is None else float(등락률),
            "기준일": 서울현재시각().date(),
            "기준시각": 서울현재시각(),
            "조회시각": 서울현재시각(),
            "상태": "실시간 현재가 반영(Naver)",
            "출처": "Naver",
            "비교기준": "전일 종가 대비",
        }
    except Exception:
        return None

def 시장지표표시문자열df(df):
    표시용 = df.copy()
    if 표시용.empty:
        return 표시용
    표시용["현재값"] = 표시용["현재값"].apply(lambda x: 숫자표시(x, 2))
    표시용["전일대비"] = 표시용["전일대비"].apply(lambda x: 증감문자열(x))
    표시용["등락률"] = 표시용["등락률"].apply(lambda x: 증감문자열(x, "%"))
    return 표시용


def 시장지표스타일적용(df):
    if df is None or df.empty:
        return df

    def 변화색상(v):
        실수값 = 안전실수변환(v)
        if 실수값 is None:
            return ""
        if 실수값 > 0:
            return "color: #ef4444; font-weight: 520;"
        if 실수값 < 0:
            return "color: #3b82f6; font-weight: 520;"
        return "color: #94a3b8; font-weight: 520;"

    styled = df.style.map(변화색상, subset=["전일대비", "등락률"])
    return styled


# -----------------------------------
# 목표비중 저장/불러오기
# -----------------------------------
def 목표비중불러오기():
    기본값 = {
        "069500": 50.0,
        "229200": 3.0,
        "471990": 0.0,
        "005930": 27.0,
        "000660": 20.0,
    }

    if os.path.exists(목표비중저장파일):
        try:
            with open(목표비중저장파일, "r", encoding="utf-8") as f:
                저장값 = json.load(f)
            return {
                "069500": float(저장값.get("069500", 50.0)),
                "229200": float(저장값.get("229200", 3.0)),
                "471990": float(저장값.get("471990", 0.0)),
                "005930": float(저장값.get("005930", 27.0)),
                "000660": float(저장값.get("000660", 20.0)),
            }
        except Exception:
            return 기본값

    return 기본값


def 목표비중저장(목표비중):
    return 안전JSON저장(목표비중, 목표비중저장파일)



# -----------------------------------
# 데이터 조회 함수
# -----------------------------------
@st.cache_data(ttl=180)
def 네이버페이지가져오기(url):
    try:
        응답 = 안전웹요청(url, timeout=10, attempts=2)
        if 응답 is None:
            return None
        return 응답.text
    except Exception:
        return None


@st.cache_data(ttl=30)

def 주요지표값정규화(이름, 현재값=None, 전일대비=None, 등락률=None, 전일가=None):
    def _f(v):
        return None if v is None else float(v)

    현재값 = _f(현재값)
    전일대비 = _f(전일대비)
    등락률 = _f(등락률)
    전일가 = _f(전일가)

    # 미국 10년물 금리(^TNX)는 최근 응답에서 이미 실제 금리 수준(예: 4.2x)으로
    # 들어오는 경우가 있어 추가로 10으로 나누면 0.42처럼 축소 표시되는 문제가 발생한다.
    # 따라서 별도 스케일 변환 없이 원본 값을 사용한다.
    return 현재값, 전일대비, 등락률, 전일가


def 야후현재가요약가져오기(심볼, 이름):
    기본결과 = {"지표": 이름, "현재값": None, "전일대비": None, "등락률": None, "전일가": None, "출처": "Yahoo"}
    if not 심볼:
        return 기본결과

    try:
        url = "https://query1.finance.yahoo.com/v7/finance/quote"
        응답 = 안전웹요청(url, params={"symbols": 심볼}, timeout=4, attempts=1)
        if 응답 is not None:
            payload = 응답.json()
            결과목록 = payload.get("quoteResponse", {}).get("result", [])
            if 결과목록:
                항목 = 결과목록[0]
                현재값 = 안전실수변환(항목.get("regularMarketPrice"))
                전일대비 = 안전실수변환(항목.get("regularMarketChange"))
                등락률 = 안전실수변환(항목.get("regularMarketChangePercent"))
                전일가 = 안전실수변환(항목.get("regularMarketPreviousClose"))
                현재값, 전일대비, 등락률, 전일가 = 주요지표값정규화(이름, 현재값, 전일대비, 등락률, 전일가)
                기준시각 = 항목.get("regularMarketTime")
                if 기준시각 is not None:
                    try:
                        기준시각 = datetime.fromtimestamp(int(기준시각), tz=KST) if KST is not None else datetime.fromtimestamp(int(기준시각))
                    except Exception:
                        기준시각 = 서울현재시각()
                else:
                    기준시각 = 서울현재시각()

                if 현재값 is not None:
                    if 전일가 in [None, 0] and 전일대비 is not None:
                        후보전일가 = float(현재값) - float(전일대비)
                        if 후보전일가 > 0:
                            전일가 = 후보전일가
                    if 전일대비 is None and 전일가 not in [None, 0]:
                        전일대비 = float(현재값) - float(전일가)
                    if 등락률 is None and 전일가 not in [None, 0]:
                        등락률 = ((float(현재값) - float(전일가)) / float(전일가)) * 100.0

                    return {
                        "지표": 이름,
                        "현재값": float(현재값),
                        "전일대비": None if 전일대비 is None else float(전일대비),
                        "등락률": None if 등락률 is None else float(등락률),
                        "전일가": None if 전일가 is None else float(전일가),
                        "출처": "Yahoo",
                        "기준시각": 기준시각,
                        "조회시각": 서울현재시각(),
                        "비교기준": "전일 종가 대비",
                    }
    except Exception:
        pass

    try:
        url = f"https://query1.finance.yahoo.com/v8/finance/chart/{심볼}"
        params = {"interval": "1d", "range": "7d", "includePrePost": "false", "events": "div,splits"}
        응답 = 안전웹요청(url, params=params, timeout=6, attempts=1)
        if 응답 is None:
            return 기본결과
        payload = 응답.json()
        결과목록 = payload.get("chart", {}).get("result", [])
        if not 결과목록:
            return 기본결과

        결과 = 결과목록[0]
        timestamps = 결과.get("timestamp", [])
        quotes = 결과.get("indicators", {}).get("quote", [{}])[0]
        종가목록 = quotes.get("close", [])
        if not timestamps or not 종가목록:
            return 기본결과

        df = pd.DataFrame({
            "날짜": pd.to_datetime(timestamps, unit="s"),
            "종가": 종가목록,
        }).dropna(subset=["종가"]).sort_values("날짜")
        if df.empty:
            return 기본결과

        현재값 = float(df.iloc[-1]["종가"])
        전일가 = float(df.iloc[-2]["종가"]) if len(df) >= 2 else None
        전일대비 = None if 전일가 in [None, 0] else 현재값 - 전일가
        등락률 = None if 전일가 in [None, 0] else (전일대비 / 전일가) * 100.0
        현재값, 전일대비, 등락률, 전일가 = 주요지표값정규화(이름, 현재값, 전일대비, 등락률, 전일가)
        return {
            "지표": 이름,
            "현재값": 현재값,
            "전일대비": 전일대비,
            "등락률": 등락률,
            "전일가": 전일가,
            "출처": "Yahoo(일봉보조)",
            "기준시각": 서울현재시각(),
            "조회시각": 서울현재시각(),
            "비교기준": "전일 종가 대비",
        }
    except Exception:
        return 기본결과


def 네이버시장지표현재가가져오기(이름, url, fallback_to_yahoo=True):
    def _부호반영(값, 부호):
        if 값 is None:
            return None
        값 = float(값)
        if 부호 == "up":
            return abs(값)
        if 부호 == "down":
            return -abs(값)
        return 값

    def _텍스트에서등락률후보(text):
        if not text:
            return None
        m = re.search(r'([+-]?\d[\d,]*\.?\d*)\s*%', str(text))
        if not m:
            return None
        값 = 안전실수변환(m.group(1))
        if 값 is None:
            return None
        if any(키 in str(text) for 키 in ["하락", "down", "dn", "▼", "minus"]):
            return -abs(float(값))
        if any(키 in str(text) for 키 in ["상승", "up", "▲", "plus"]):
            return abs(float(값))
        return float(값)

    if BS4_AVAILABLE:
        html = 네이버페이지가져오기(url)
        if html:
            try:
                soup = BeautifulSoup(html, "html.parser")

                현재값 = None
                전일대비 = None
                등락률 = None
                부호 = None

                현재선택자 = [
                    "div.head_info span.value",
                    "p.no_today span.blind",
                    "span.value",
                    "em.value",
                    "div.today_info span.value",
                ]
                for sel in 현재선택자:
                    for tag in soup.select(sel):
                        값 = 안전실수변환(tag.get_text(" ", strip=True))
                        if 값 is not None:
                            현재값 = float(값)
                            break
                    if 현재값 is not None:
                        break

                변화선택자 = [
                    "div.head_info span.change",
                    "div.head_info span.no_up",
                    "div.head_info span.no_down",
                    "div.head_info span.point",
                    "span.change",
                    "span.change_value",
                    "span.rate",
                ]
                for sel in 변화선택자:
                    for tag in soup.select(sel):
                        text = tag.get_text(" ", strip=True)
                        own_classes = " ".join(tag.get("class", []))
                        parent_classes = " ".join(tag.parent.get("class", [])) if getattr(tag, "parent", None) else ""
                        classes = f"{own_classes} {parent_classes}".lower()
                        local_sign = None
                        if any(x in classes for x in ["up", "rise", "plus", "red"]):
                            local_sign = "up"
                        elif any(x in classes for x in ["down", "fall", "minus", "blue", "dn"]):
                            local_sign = "down"
                        elif "상승" in text or "▲" in text:
                            local_sign = "up"
                        elif "하락" in text or "▼" in text:
                            local_sign = "down"

                        rate = _텍스트에서등락률후보(text)
                        if rate is not None and 등락률 is None:
                            등락률 = rate
                            부호 = "up" if rate > 0 else "down" if rate < 0 else 부호

                        값 = 안전실수변환(text)
                        if 값 is not None and "%" not in text and 전일대비 is None:
                            전일대비 = _부호반영(값, local_sign)
                            부호 = local_sign or 부호

                blind_texts = [x.get_text(" ", strip=True) for x in soup.select("span.blind, em, p, td, li")]
                for txt in blind_texts:
                    if 등락률 is None:
                        rate = _텍스트에서등락률후보(txt)
                        if rate is not None:
                            등락률 = rate
                            if 부호 is None:
                                부호 = "up" if rate > 0 else "down" if rate < 0 else None
                    if 전일대비 is None and any(key in txt for key in ["전일대비", "상승", "하락", "▲", "▼"]):
                        값 = 안전실수변환(txt)
                        if 값 is not None:
                            local_sign = None
                            if any(key in txt for key in ["상승", "▲"]):
                                local_sign = "up"
                            elif any(key in txt for key in ["하락", "▼"]):
                                local_sign = "down"
                            전일대비 = _부호반영(값, local_sign)
                            부호 = local_sign or 부호

                if 전일대비 is not None and 등락률 is None and 현재값 not in [None, 0]:
                    전일가 = float(현재값) - float(전일대비)
                    if 전일가 not in [None, 0]:
                        등락률 = (float(전일대비) / float(전일가)) * 100.0

                if 등락률 is not None and 전일대비 is None and 현재값 not in [None, 0]:
                    분모 = 1.0 + float(등락률) / 100.0
                    if 분모 != 0:
                        전일가 = float(현재값) / 분모
                        전일대비 = float(현재값) - float(전일가)

                if 현재값 is not None:
                    return {
                        "지표": 이름,
                        "현재값": float(현재값),
                        "전일대비": None if 전일대비 is None else float(전일대비),
                        "등락률": None if 등락률 is None else float(등락률),
                        "링크": url,
                        "출처": "네이버",
                        "기준시각": 서울현재시각(),
                        "조회시각": 서울현재시각(),
                        "비교기준": "전일 종가 대비",
                    }
            except Exception:
                pass

    if fallback_to_yahoo:
        심볼 = 야후주요지표심볼.get(이름)
        if 심볼:
            결과 = 야후현재가요약가져오기(심볼, 이름)
            결과["링크"] = url
            return 결과

    return {"지표": 이름, "현재값": None, "전일대비": None, "등락률": None, "링크": url, "출처": "네이버"}


def _정렬정제_OHLCV(데이터):
    if 데이터 is None or len(데이터) == 0:
        return pd.DataFrame()
    데이터 = 데이터.copy()
    try:
        데이터 = 데이터.sort_index()
    except Exception:
        pass
    return 데이터

def _야후차트OHLCV조회(심볼, 시작문자열, 종료문자열):
    if not 심볼:
        return pd.DataFrame()
    try:
        시작초 = int(datetime.strptime(시작문자열, "%Y%m%d").timestamp())
        종료초 = int((datetime.strptime(종료문자열, "%Y%m%d") + timedelta(days=1)).timestamp())
        url = f"https://query1.finance.yahoo.com/v8/finance/chart/{심볼}"
        params = {
            "period1": 시작초,
            "period2": 종료초,
            "interval": "1d",
            "includePrePost": "false",
            "events": "div,splits",
        }
        응답 = 안전웹요청(url, params=params, timeout=12, attempts=3)
        if 응답 is None:
            return pd.DataFrame()
        payload = 응답.json()
        결과목록 = payload.get("chart", {}).get("result", [])
        if not 결과목록:
            return pd.DataFrame()
        결과 = 결과목록[0]
        timestamps = 결과.get("timestamp", []) or []
        quotes = 결과.get("indicators", {}).get("quote", [{}])[0] or {}
        if not timestamps:
            return pd.DataFrame()

        target_len = len(timestamps)
        df = pd.DataFrame({
            "날짜": pd.to_datetime(timestamps, unit="s", errors="coerce"),
            "시가": 시리즈길이맞추기(quotes.get("open", []), target_len),
            "고가": 시리즈길이맞추기(quotes.get("high", []), target_len),
            "저가": 시리즈길이맞추기(quotes.get("low", []), target_len),
            "종가": 시리즈길이맞추기(quotes.get("close", []), target_len),
            "거래량": 시리즈길이맞추기(quotes.get("volume", []), target_len),
        })
        return OHLCV데이터정리(df)
    except Exception:
        return pd.DataFrame()


def _야후종목ETF_OHLCV조회(코드, 시작문자열, 종료문자열):
    for 심볼 in 종목코드별야후심볼후보(코드):
        df = _야후차트OHLCV조회(심볼, 시작문자열, 종료문자열)
        if not df.empty:
            return df
    return pd.DataFrame()


def _시장OHLCV조회(시작문자열, 종료문자열, 코드):
    return _야후종목ETF_OHLCV조회(코드, 시작문자열, 종료문자열)


def _ETF_OHLCV조회(시작문자열, 종료문자열, 코드):
    return _야후종목ETF_OHLCV조회(코드, 시작문자열, 종료문자열)


def _야후인덱스OHLCV조회(시작문자열, 종료문자열, 코드):
    심볼 = 야후인덱스심볼.get(str(코드))
    if not 심볼:
        return pd.DataFrame()

    try:
        시작초 = int(datetime.strptime(시작문자열, "%Y%m%d").timestamp())
        종료초 = int((datetime.strptime(종료문자열, "%Y%m%d") + timedelta(days=1)).timestamp())
        url = f"https://query1.finance.yahoo.com/v8/finance/chart/{심볼}"
        params = {
            "period1": 시작초,
            "period2": 종료초,
            "interval": "1d",
            "includePrePost": "false",
            "events": "div,splits",
        }
        응답 = 안전웹요청(url, params=params, timeout=10, attempts=2)
        if 응답 is None:
            return pd.DataFrame()
        payload = 응답.json()
        결과목록 = payload.get("chart", {}).get("result", [])
        if not 결과목록:
            return pd.DataFrame()

        결과 = 결과목록[0]
        timestamps = 결과.get("timestamp", [])
        quotes = 결과.get("indicators", {}).get("quote", [{}])[0]
        if not timestamps:
            return pd.DataFrame()

        df = pd.DataFrame({
            "날짜": pd.to_datetime(timestamps, unit="s"),
            "시가": quotes.get("open", []),
            "고가": quotes.get("high", []),
            "저가": quotes.get("low", []),
            "종가": quotes.get("close", []),
            "거래량": quotes.get("volume", []),
        }).dropna(subset=["종가"])

        if df.empty:
            return pd.DataFrame()

        df["날짜"] = pd.to_datetime(df["날짜"]).dt.tz_localize(None)
        df = df.set_index("날짜").sort_index()
        return df
    except Exception:
        return pd.DataFrame()


def _인덱스OHLCV조회(시작문자열, 종료문자열, 코드):
    return _야후인덱스OHLCV조회(시작문자열, 종료문자열, 코드)




@st.cache_data(ttl=60)
def 야후전일종가가져오기(심볼):
    if not 심볼:
        return None
    try:
        url = "https://query1.finance.yahoo.com/v7/finance/quote"
        응답 = 안전웹요청(url, params={"symbols": 심볼}, timeout=8, attempts=2)
        if 응답 is None:
            return None
        payload = 응답.json()
        결과목록 = payload.get("quoteResponse", {}).get("result", [])
        if not 결과목록:
            return None
        항목 = 결과목록[0]

        후보값 = [
            항목.get("regularMarketPreviousClose"),
            항목.get("previousClose"),
            항목.get("chartPreviousClose"),
        ]
        for 값 in 후보값:
            값 = 안전실수변환(값)
            if 값 is not None and 값 > 0:
                return float(값)
    except Exception:
        return None
    return None


@st.cache_data(ttl=5)
def 야후실시간요약가져오기(심볼):
    if not 심볼:
        return None
    try:
        url = "https://query1.finance.yahoo.com/v7/finance/quote"
        응답 = 안전웹요청(url, params={"symbols": 심볼}, timeout=2.5, attempts=1)
        if 응답 is None:
            return None
        payload = 응답.json()
        결과목록 = payload.get("quoteResponse", {}).get("result", [])
        if not 결과목록:
            return None
        항목 = 결과목록[0] or {}

        현재가 = None
        for 키 in ["regularMarketPrice", "postMarketPrice", "preMarketPrice"]:
            값 = 안전실수변환(항목.get(키))
            if 값 is not None and 값 > 0:
                현재가 = float(값)
                break

        전일가 = None
        for 키 in ["regularMarketPreviousClose", "previousClose", "chartPreviousClose"]:
            값 = 안전실수변환(항목.get(키))
            if 값 is not None and 값 > 0:
                전일가 = float(값)
                break

        전일대비 = None
        for 키 in ["regularMarketChange", "postMarketChange", "preMarketChange"]:
            값 = 안전실수변환(항목.get(키))
            if 값 is not None:
                전일대비 = float(값)
                break

        등락률 = None
        for 키 in ["regularMarketChangePercent", "postMarketChangePercent", "preMarketChangePercent"]:
            값 = 안전실수변환(항목.get(키))
            if 값 is not None:
                등락률 = float(값)
                break

        if 현재가 in [None, 0]:
            return None

        if 전일가 in [None, 0] and 전일대비 not in [None]:
            전일가 = float(현재가) - float(전일대비)
        if 전일대비 is None and 전일가 not in [None, 0]:
            전일대비 = float(현재가) - float(전일가)
        if 등락률 is None and 전일가 not in [None, 0]:
            등락률 = ((float(현재가) - float(전일가)) / float(전일가)) * 100.0

        return {
            "현재가": float(현재가),
            "전일가": None if 전일가 in [None, 0] else float(전일가),
            "전일대비": 전일대비,
            "등락률": 등락률,
            "기준일": 서울현재시각().date(),
            "기준시각": 서울현재시각(),
            "상태": "Yahoo quote 반영",
            "출처": "Yahoo",
        }
    except Exception:
        return None


def 야후실시간비교값보강(구분, 코드):
    심볼목록 = 자산야후심볼목록가져오기(구분, 코드)
    for 심볼 in 심볼목록:
        요약 = 야후실시간요약가져오기(심볼)
        if 요약 is not None and 요약.get("현재가") not in [None, 0]:
            return 요약
    return None

def 자산야후심볼목록가져오기(구분, 코드):
    if str(구분) == "index":
        심볼 = 야후인덱스심볼.get(str(코드))
        return [심볼] if 심볼 else []
    return 종목코드별야후심볼후보(코드)


@st.cache_data(ttl=10)
def 야후1분봉요약가져오기(심볼):
    if not 심볼:
        return None
    try:
        url = f"https://query1.finance.yahoo.com/v8/finance/chart/{심볼}"
        params = {
            "interval": "1m",
            "range": "2d",
            "includePrePost": "false",
            "events": "div,splits",
        }
        응답 = 안전웹요청(url, params=params, timeout=8, attempts=2)
        if 응답 is None:
            return None
        payload = 응답.json()
        결과목록 = payload.get("chart", {}).get("result", [])
        if not 결과목록:
            return None

        결과 = 결과목록[0]
        meta = 결과.get("meta", {}) or {}
        timestamps = 결과.get("timestamp", []) or []
        quotes = 결과.get("indicators", {}).get("quote", [{}])
        quote = quotes[0] if quotes else {}
        종가목록 = quote.get("close", []) or []
        if not timestamps or not 종가목록:
            return None

        df = pd.DataFrame({
            "날짜": pd.to_datetime(timestamps, unit="s", utc=True).tz_convert("Asia/Seoul").tz_localize(None),
            "종가": 종가목록,
        })
        df["종가"] = pd.to_numeric(df["종가"], errors="coerce")
        df = df.dropna(subset=["종가"]).sort_values("날짜")
        if df.empty:
            return None

        최신행 = df.iloc[-1]
        현재가 = 안전실수변환(최신행["종가"])
        if 현재가 is None or 현재가 <= 0:
            return None

        # 전일 종가는 1분봉 meta보다 quote API를 우선 사용해 상승/하락 색상이 뒤집히지 않도록 보정
        전일가 = 야후전일종가가져오기(심볼)
        if 전일가 in [None, 0]:
            전일가 = 안전실수변환(meta.get("regularMarketPreviousClose"))
        if 전일가 in [None, 0]:
            전일가 = 안전실수변환(meta.get("previousClose"))
        if 전일가 in [None, 0]:
            전일가 = 안전실수변환(meta.get("chartPreviousClose"))
        if 전일가 in [None, 0]:
            최신일자 = 최신행["날짜"].date()
            이전일데이터 = df[df["날짜"].dt.date < 최신일자]
            if not 이전일데이터.empty:
                전일가 = 안전실수변환(이전일데이터.iloc[-1]["종가"])

        전일대비 = None if 전일가 in [None, 0] else float(현재가) - float(전일가)
        등락률 = None if 전일가 in [None, 0] else (전일대비 / float(전일가)) * 100

        기준시각 = pd.to_datetime(최신행["날짜"])
        return {
            "현재가": float(현재가),
            "전일가": None if 전일가 in [None, 0] else float(전일가),
            "전일대비": 전일대비,
            "등락률": 등락률,
            "기준일": 기준시각.date(),
            "기준시각": 기준시각,
            "상태": "준실시간 1분봉 반영(전일종가 기준)",
            "신선여부": _국내장중1분봉신선여부(기준시각),
        }
    except Exception:
        return None
    return None


@st.cache_data(ttl=5)
def 야후실시간호가가져오기(심볼):
    if not 심볼:
        return None
    try:
        url = "https://query1.finance.yahoo.com/v7/finance/quote"
        응답 = 안전웹요청(url, params={"symbols": 심볼}, timeout=2.5, attempts=1)
        if 응답 is None:
            return None

        payload = 응답.json()
        결과목록 = payload.get("quoteResponse", {}).get("result", [])
        if not 결과목록:
            return None

        항목 = 결과목록[0]
        후보값 = [
            항목.get("regularMarketPrice"),
            항목.get("postMarketPrice"),
            항목.get("preMarketPrice"),
        ]
        for 값 in 후보값:
            값 = 안전실수변환(값)
            if 값 is not None and 값 > 0:
                return float(값)
    except Exception:
        return None
    return None



def 준실시간시세요약가져오기(구분, 코드):
    try:
        for 심볼 in 자산야후심볼목록가져오기(구분, 코드):
            요약 = 야후1분봉요약가져오기(심볼)
            if 요약 is not None and 요약.get("현재가") not in [None, 0]:
                return 요약
    except Exception:
        return None
    return None


def 실시간현재가가져오기(구분, 코드):
    try:
        if 구분 == "index":
            심볼 = 야후인덱스심볼.get(str(코드))
            return 야후실시간호가가져오기(심볼)

        for 심볼 in 종목코드별야후심볼후보(코드):
            값 = 야후실시간호가가져오기(심볼)
            if 값 is not None and 값 > 0:
                return 값
    except Exception:
        return None
    return None


def 전일기준가보강(구분, 코드, 기존전일가=None, 최근요약=None):
    값 = 안전실수변환(기존전일가)
    if 값 is not None and 값 > 0:
        return float(값)

    if isinstance(최근요약, dict):
        후보 = 안전실수변환(최근요약.get("전일가"))
        if 후보 is not None and 후보 > 0:
            return float(후보)

    심볼목록 = []
    if 구분 == "index":
        심볼 = 야후인덱스심볼.get(str(코드))
        if 심볼:
            심볼목록.append(심볼)
    else:
        for 심볼 in 종목코드별야후심볼후보(코드):
            if 심볼 and 심볼 not in 심볼목록:
                심볼목록.append(심볼)

    for 심볼 in 심볼목록:
        후보 = 안전실수변환(야후전일종가가져오기(심볼))
        if 후보 is not None and 후보 > 0:
            return float(후보)
    return None


def 비교값보강적용(요약, 구분, 코드):
    if 요약 is None:
        return {}
    결과 = dict(요약)
    현재가 = 안전실수변환(결과.get("현재가"))
    if 현재가 is None or 현재가 <= 0:
        return 결과

    기존전일대비 = 안전실수변환(결과.get("전일대비"))
    기존등락률 = 안전실수변환(결과.get("등락률"))
    전일가 = 전일기준가보강(구분, 코드, 결과.get("전일가"), 결과)

    if (전일가 is None or 전일가 <= 0) or (기존전일대비 is None and 기존등락률 is None):
        야후보강 = 야후실시간비교값보강(구분, 코드)
        if 야후보강:
            if 안전실수변환(야후보강.get("전일가")) not in [None, 0] and (전일가 is None or 전일가 <= 0):
                전일가 = float(야후보강.get("전일가"))
            if 기존전일대비 is None and 안전실수변환(야후보강.get("전일대비")) is not None:
                결과["전일대비"] = float(야후보강.get("전일대비"))
            if 기존등락률 is None and 안전실수변환(야후보강.get("등락률")) is not None:
                결과["등락률"] = float(야후보강.get("등락률"))
            if 결과.get("출처") in [None, "", "Naver"]:
                결과["비교출처"] = 야후보강.get("출처", "Yahoo")

    if 전일가 is not None and 전일가 > 0:
        결과["전일가"] = float(전일가)
        if 안전실수변환(결과.get("전일대비")) is None:
            결과["전일대비"] = float(현재가) - float(전일가)
        if 안전실수변환(결과.get("등락률")) is None:
            결과["등락률"] = ((float(현재가) - float(전일가)) / float(전일가)) * 100.0
    return 결과


@st.cache_data(ttl=60)
def 최근OHLCV가져오기(구분, 코드, lookback_days=15, refresh_token=0):
    조회일수후보 = []
    for 일수 in [lookback_days, max(lookback_days, 45), max(lookback_days, 120)]:
        if 일수 not in 조회일수후보:
            조회일수후보.append(일수)

    for 조회일수 in 조회일수후보:
        종료일 = datetime.today()
        시작일 = 종료일 - timedelta(days=조회일수)
        시작문자열 = 시작일.strftime("%Y%m%d")
        종료문자열 = 종료일.strftime("%Y%m%d")

        if 구분 == "index":
            데이터 = _인덱스OHLCV조회(시작문자열, 종료문자열, 코드)
        elif 구분 in ["etf", "stock"]:
            데이터 = _야후종목ETF_OHLCV조회(코드, 시작문자열, 종료문자열)
        else:
            데이터 = _시장OHLCV조회(시작문자열, 종료문자열, 코드)

        데이터 = OHLCV데이터정리(데이터)
        if not 데이터.empty:
            return 데이터

    return pd.DataFrame()


@st.cache_data(ttl=60)
def 최근시세요약가져오기(구분, 코드, lookback_days=15, refresh_token=0):
    데이터 = 최근OHLCV가져오기(구분, 코드, lookback_days=lookback_days, refresh_token=refresh_token)
    return 최근유효OHLCV요약(데이터)


def 시세요약안전병합(기준요약, 신규요약):
    """
    현재가는 신규값을 우선 반영하되, 신규요약에 전일가/전일대비/등락률이 비어 있으면
    기준요약의 기존 비교값을 유지한다.
    """
    기준 = dict(기준요약 or {})
    신규 = dict(신규요약 or {})

    결과 = 기준.copy()

    # 현재가/상태/출처/기준시각 등은 신규값 우선
    for 키, 값 in 신규.items():
        if 키 in ["전일가", "전일대비", "등락률"]:
            continue
        if 값 is not None:
            결과[키] = 값

    # 비교값은 신규가 유효할 때만 덮어쓰기
    for 키 in ["전일가", "전일대비", "등락률"]:
        신규값 = 신규.get(키)
        신규수치 = 안전실수변환(신규값)
        if 신규수치 is not None:
            결과[키] = float(신규수치)
        elif 키 not in 결과:
            결과[키] = 신규값

    # 신규 요약에 명시적 문자열 필드가 있으면 보존
    for 키 in ["비교출처"]:
        if 신규.get(키) not in [None, ""]:
            결과[키] = 신규.get(키)

    return 결과


@st.cache_data(ttl=8, show_spinner=False)
def 실시간포함시세요약가져오기(구분, 코드, lookback_days=15, refresh_token=0):
    """새로고침 시 실제 현재가 반영을 우선하는 시세요약입니다.
    국내 지수·주식·ETF는 Naver 현재가를 1순위로 사용하고, 실패할 때만 Yahoo/최근 종가로 후퇴합니다.
    """
    요약 = 최근시세요약가져오기(
        구분,
        코드,
        lookback_days=lookback_days,
        refresh_token=refresh_token,
    ).copy()

    장중 = 한국장중여부()

    # v5.14.1_realtime_price_fixed: 국내 지수·주식·ETF 모두 Naver 현재가 우선
    if 구분 in ["index", "etf", "stock"]:
        국내현재가 = 네이버국내현재가가져오기(구분, 코드, refresh_token=refresh_token)
        if 국내현재가 is not None and 국내현재가.get("현재가") not in [None, 0]:
            병합 = 시세요약안전병합(요약, 국내현재가)
            병합["비교기준"] = "전일 종가 대비"
            병합["상태"] = 국내현재가.get("상태", "실시간 현재가 반영(Naver)")
            병합["출처"] = "Naver"
            병합 = 비교값보강적용(병합, 구분, 코드)
            return 병합

    if 장중:
        실시간가 = 실시간현재가가져오기(구분, 코드)
        if 실시간가 not in [None, 0]:
            병합 = 시세요약안전병합(요약, {
                "현재가": float(실시간가),
                "기준시각": 서울현재시각(),
                "기준일": 서울현재시각().date(),
                "상태": "장중 현재가 반영(Yahoo 보조)",
                "출처": "Yahoo",
                "비교기준": "전일 종가 대비",
            })
            병합 = 비교값보강적용(병합, 구분, 코드)
            return 병합

    분봉요약 = 준실시간시세요약가져오기(구분, 코드)
    if 분봉요약 is not None and 분봉요약.get("현재가") not in [None, 0]:
        장중분봉허용 = (not 장중) or bool(분봉요약.get("신선여부", False))
        if 장중분봉허용:
            병합 = 시세요약안전병합(요약, {
                "현재가": float(분봉요약.get("현재가")),
                "전일가": 분봉요약.get("전일가"),
                "전일대비": 분봉요약.get("전일대비"),
                "등락률": 분봉요약.get("등락률"),
                "상태": 분봉요약.get("상태", "준실시간 1분봉 반영"),
                "기준일": 분봉요약.get("기준일"),
                "기준시각": 분봉요약.get("기준시각"),
                "출처": 분봉요약.get("출처", "Yahoo 1분봉"),
                "비교기준": "전일 종가 대비",
            })
            병합 = 비교값보강적용(병합, 구분, 코드)
            return 병합

    요약 = 비교값보강적용(요약, 구분, 코드)
    if 장중:
        요약["상태"] = "장중 현재가 조회 실패(최근 종가 대체)"
    else:
        요약["상태"] = 요약.get("상태", "최근 종가 반영") or "최근 종가 반영"
    요약["출처"] = 요약.get("출처", "최근 종가")
    요약["비교기준"] = 요약.get("비교기준", "전일 종가 대비")
    return 요약


@st.cache_data(ttl=60, show_spinner=False)
def 시세스냅샷캐시(거래이력json문자열, refresh_token=0):
    """
    한 번의 새로고침에서 상단 카드와 포트폴리오 계산이 같은 현재가 결과를 재사용하도록
    종목별 시세를 한 번만 모아 조회하는 스냅샷 캐시.
    """
    try:
        원본 = json.loads(거래이력json문자열)
        거래df = pd.DataFrame(원본)
    except Exception:
        거래df = pd.DataFrame()

    자산목록 = []
    # 주요 지수는 항상 포함
    자산목록.append(("index", "1001", "코스피"))
    자산목록.append(("index", "2001", "코스닥"))

    try:
        계산대상 = 거래이력계산대상추출(거래df)
        집계표 = 포트폴리오입력집계(계산대상)
        if 집계표 is not None and not 집계표.empty:
            집계표["보유수량"] = pd.to_numeric(집계표.get("보유수량"), errors="coerce").fillna(0)
            집계표 = 집계표[집계표["보유수량"] > 0].copy()
            for _, 행 in 집계표.iterrows():
                코드 = str(행.get("종목코드", "")).zfill(6)
                이름 = 종목명자동보정(코드, 행.get("종목명", ""))
                구분 = 종목구분판단(코드, 이름)
                if 코드:
                    자산목록.append((구분, 코드, 이름))
    except Exception:
        pass

    # 중복 제거
    고유자산 = []
    seen = set()
    for 구분, 코드, 이름 in 자산목록:
        키 = f"{구분}:{str(코드).zfill(6) if str(코드).isdigit() else str(코드)}"
        if 키 in seen:
            continue
        seen.add(키)
        고유자산.append((구분, 코드, 이름))

    결과 = {}

    def _단일시세조회(항목):
        구분, 코드, 이름 = 항목
        키 = f"{구분}:{str(코드)}"
        try:
            정보 = 실시간포함시세요약가져오기(
                구분,
                코드,
                lookback_days=15,
                refresh_token=refresh_token,
            ).copy()
            정보["자산명"] = 이름
            return 키, 정보
        except Exception:
            return 키, {
                "자산명": 이름,
                "현재가": None,
                "전일가": None,
                "전일대비": None,
                "등락률": None,
                "상태": "조회 실패",
            }

    # v5.11: 보유 종목이 여러 개일 때 시세 요청을 병렬로 처리해 버튼 클릭 후 대기시간을 줄입니다.
    # 네트워크 과부하를 막기 위해 동시 작업 수는 최대 6개로 제한합니다.
    if 고유자산:
        작업수 = min(6, max(1, len(고유자산)))
        try:
            with ThreadPoolExecutor(max_workers=작업수) as executor:
                futures = [executor.submit(_단일시세조회, 항목) for 항목 in 고유자산]
                for future in as_completed(futures):
                    키, 정보 = future.result()
                    결과[키] = 정보
        except Exception:
            # 일부 환경에서 병렬 요청이 제한되면 기존 방식으로 안전하게 후퇴합니다.
            for 항목 in 고유자산:
                키, 정보 = _단일시세조회(항목)
                결과[키] = 정보

    return 결과


def 시세스냅샷세션반영(거래df, refresh_token=0):
    try:
        서명 = 거래이력서명생성(거래df)
        스냅샷 = 시세스냅샷캐시(서명, refresh_token=refresh_token)
        st.session_state["price_snapshot_map_v1"] = 스냅샷
        st.session_state["price_snapshot_token_v1"] = refresh_token
        st.session_state["price_snapshot_signature_v1"] = 서명
        return 스냅샷
    except Exception:
        st.session_state["price_snapshot_map_v1"] = {}
        st.session_state["price_snapshot_token_v1"] = refresh_token
        return {}


def 스냅샷현재가조회(구분, 코드):
    맵 = st.session_state.get("price_snapshot_map_v1", {}) or {}
    return (맵.get(f"{구분}:{str(코드)}") or {}).get("현재가")


def 스냅샷자산정보조회(구분, 코드):
    맵 = st.session_state.get("price_snapshot_map_v1", {}) or {}
    return (맵.get(f"{구분}:{str(코드)}") or {}).copy()

@st.cache_data(ttl=8, show_spinner=False)
def 자산현재가정보(자산명, 자산정보, refresh_token=0):
    구분 = 자산정보["구분"]
    코드 = 자산정보["코드"]
    스냅샷토큰 = st.session_state.get("price_snapshot_token_v1")
    스냅샷정보 = 스냅샷자산정보조회(구분, 코드)
    # 새로고침 토큰이 같은 스냅샷만 재사용하여 과거 시세가 남는 것을 방지합니다.
    if 스냅샷토큰 == refresh_token and 스냅샷정보 and 스냅샷정보.get("현재가") not in [None, 0]:
        스냅샷정보["자산명"] = 자산명
        return 스냅샷정보
    정보 = 실시간포함시세요약가져오기(구분, 코드, lookback_days=15, refresh_token=refresh_token)
    정보["자산명"] = 자산명
    return 정보


def 모니터표시시세요약(자산명, 자산정보, refresh_token=0):
    """상단 모니터 전용 시세.
    - 기본: 최근 일봉/전일 종가 기준으로 빠르게 표시
    - 시세 새로고침 클릭 후: 장중·준실시간 요약 반영
    """
    구분 = 자산정보["구분"]
    코드 = 자산정보["코드"]
    실시간모드 = bool(st.session_state.get("monitor_realtime_mode_v1", False))
    if 실시간모드:
        정보 = 자산현재가정보(자산명, 자산정보, refresh_token=refresh_token)
        정보["모니터모드"] = "실시간/준실시간"
        return 정보
    정보 = 최근시세요약가져오기(구분, 코드, lookback_days=15, refresh_token=0).copy()
    정보 = 비교값보강적용(정보, 구분, 코드)
    정보["자산명"] = 자산명
    정보["상태"] = "최근 일봉 종가 반영"
    정보["출처"] = 정보.get("출처", "최근 일봉")
    정보["비교기준"] = "전일 종가 대비"
    정보["모니터모드"] = "전일종가"
    return 정보


@st.cache_data(ttl=8, show_spinner=False)
def 종목현재가가져오기(종목코드, refresh_token=0):
    스냅샷값 = 스냅샷현재가조회("stock", 종목코드)
    if 스냅샷값 not in [None, 0]:
        return 스냅샷값
    return 실시간포함시세요약가져오기("stock", 종목코드, lookback_days=15, refresh_token=refresh_token).get("현재가")


@st.cache_data(ttl=8, show_spinner=False)
def ETF현재가가져오기(종목코드, refresh_token=0):
    스냅샷값 = 스냅샷현재가조회("etf", 종목코드)
    if 스냅샷값 not in [None, 0]:
        return 스냅샷값
    return 실시간포함시세요약가져오기("etf", 종목코드, lookback_days=15, refresh_token=refresh_token).get("현재가")


@st.cache_data(ttl=8, show_spinner=False)
def 인덱스현재가가져오기(지수코드, refresh_token=0):
    스냅샷값 = 스냅샷현재가조회("index", 지수코드)
    if 스냅샷값 not in [None, 0]:
        return 스냅샷값
    return 실시간포함시세요약가져오기("index", 지수코드, lookback_days=15, refresh_token=refresh_token).get("현재가")


@st.cache_data(ttl=300)
def 자산과거가격가져오기(구분, 코드, 개월수=6):
    try:
        종료일 = datetime.today()
        시작일 = 종료일 - timedelta(days=30 * 개월수)
        시작문자열 = 시작일.strftime("%Y%m%d")
        종료문자열 = 종료일.strftime("%Y%m%d")

        if 구분 == "index":
            데이터 = _인덱스OHLCV조회(시작문자열, 종료문자열, 코드)
        elif 구분 == "etf":
            데이터 = _ETF_OHLCV조회(시작문자열, 종료문자열, 코드)
        else:
            데이터 = _시장OHLCV조회(시작문자열, 종료문자열, 코드)

        if 데이터 is None or 데이터.empty:
            return pd.DataFrame()

        데이터 = 데이터.copy()
        데이터.index = pd.to_datetime(데이터.index).tz_localize(None)
        데이터 = 데이터.sort_index()
        데이터 = 데이터[~데이터.index.duplicated(keep="last")]

        필수열 = ["시가", "고가", "저가", "종가"]
        for 열 in 필수열:
            if 열 not in 데이터.columns:
                return pd.DataFrame()
            데이터[열] = pd.to_numeric(데이터[열], errors="coerce")

        if "거래량" not in 데이터.columns:
            데이터["거래량"] = 0
        데이터["거래량"] = pd.to_numeric(데이터["거래량"], errors="coerce").fillna(0)
        데이터 = 데이터.dropna(subset=["종가"])

        if 데이터.empty:
            return pd.DataFrame()

        데이터["5일평균"] = 데이터["종가"].rolling(5, min_periods=1).mean()
        데이터["20일평균"] = 데이터["종가"].rolling(20, min_periods=1).mean()
        데이터["60일평균"] = 데이터["종가"].rolling(60, min_periods=1).mean()
        데이터["120일평균"] = 데이터["종가"].rolling(120, min_periods=1).mean()

        변화량 = 데이터["종가"].diff()
        상승분 = 변화량.clip(lower=0)
        하락분 = -변화량.clip(upper=0)
        평균상승 = 상승분.rolling(14, min_periods=14).mean()
        평균하락 = 하락분.rolling(14, min_periods=14).mean()
        rs = 평균상승 / 평균하락.replace(0, pd.NA)
        데이터["RSI(14)"] = 100 - (100 / (1 + rs))
        데이터.loc[(평균하락 == 0) & (평균상승 > 0), "RSI(14)"] = 100
        데이터.loc[(평균하락 == 0) & (평균상승 == 0), "RSI(14)"] = 50

        return 데이터
    except Exception:
        return pd.DataFrame()


@st.cache_data(ttl=15)
def 시장지표결과보강(결과, 이름, url=None):
    결과 = dict(결과 or {})
    결과["지표"] = 결과.get("지표", 이름)
    if url and 결과.get("링크") in [None, ""]:
        결과["링크"] = url
    결과["조회시각"] = 결과.get("조회시각", 서울현재시각())
    결과["기준시각"] = 결과.get("기준시각", 결과.get("조회시각", 서울현재시각()))
    결과["비교기준"] = 결과.get("비교기준", "전일 종가 대비")
    출처 = str(결과.get("출처", "-") or "-")
    if "Proxy" in 출처 or "프록시" in 출처:
        결과["지표구분"] = 결과.get("지표구분", "대체값")
        결과["상태"] = 결과.get("상태", "실시간 원본 부재 시 대체값")
    elif "Derived" in 출처:
        결과["지표구분"] = 결과.get("지표구분", "파생값")
        결과["상태"] = 결과.get("상태", "다른 원지표 조합으로 계산")
    elif 출처 == "Yahoo(일봉보조)":
        결과["지표구분"] = 결과.get("지표구분", "일봉보조")
        결과["상태"] = 결과.get("상태", "실시간 원본 부재 시 최근 일봉 기준 보조값")
    elif 출처 in ["Yahoo", "네이버", "Naver"]:
        결과["지표구분"] = 결과.get("지표구분", "준실시간")
    else:
        결과["지표구분"] = 결과.get("지표구분", "참고")
    return 결과


def 시장지표단건가져오기(이름, url):
    if 이름 == "두바이유":
        결과 = 두바이유현재가가져오기()
        if 결과 and 결과.get("현재값") is not None:
            return 시장지표결과보강(결과, 이름, url)
        return 시장지표결과보강({"지표": 이름, "현재값": None, "전일대비": None, "등락률": None, "링크": url, "출처": "-", "상태": "조회 실패"}, 이름, url)

    우선순위 = 지표대체우선순위.get(이름, ["naver", "yahoo"])

    for 소스 in 우선순위:
        if 소스 == "naver":
            결과 = 네이버시장지표현재가가져오기(이름, url, fallback_to_yahoo=False)
        elif 소스 == "yahoo":
            심볼 = 야후주요지표심볼.get(이름)
            결과 = 야후현재가요약가져오기(심볼, 이름) if 심볼 else None
            if 결과:
                결과["링크"] = url
        elif 소스 in ["derived_domestic_gold", "derived_dubai"]:
            결과 = 파생주요지표가져오기(이름)
            결과["링크"] = url
        else:
            결과 = None

        if 결과 and 결과.get("현재값") is not None:
            return 시장지표결과보강(결과, 이름, url)

    return 시장지표결과보강({"지표": 이름, "현재값": None, "전일대비": None, "등락률": None, "링크": url, "출처": "-", "상태": "조회 실패"}, 이름, url)



def 네이버시장지표목록가져오기():
    결과 = []
    표시순서 = ["USD/KRW", "국제 금", "WTI", "브렌트유", "미국 10년물 금리", "VIX"]
    for 이름 in 표시순서:
        url = 시장지표네이버URL.get(이름)
        if not url:
            continue
        결과.append(시장지표결과보강(시장지표단건가져오기(이름, url), 이름, url))

    df = pd.DataFrame(결과)
    if df.empty:
        return pd.DataFrame([
            {"지표": "USD/KRW", "현재값": None, "전일대비": None, "등락률": None, "출처": "-"},
            {"지표": "국제 금", "현재값": None, "전일대비": None, "등락률": None, "출처": "-"},
            {"지표": "WTI", "현재값": None, "전일대비": None, "등락률": None, "출처": "-"},
            {"지표": "브렌트유", "현재값": None, "전일대비": None, "등락률": None, "출처": "-"},
            {"지표": "미국 10년물 금리", "현재값": None, "전일대비": None, "등락률": None, "출처": "-"},
            {"지표": "VIX", "현재값": None, "전일대비": None, "등락률": None, "출처": "-"},
        ])
    return df


def 일간수익률가져오기(종목코드, 개월수=6):
    구분 = "stock"
    if 종목구분판단(종목코드) == "etf":
        구분 = "etf"
    데이터 = 자산과거가격가져오기(구분, 종목코드, 개월수)
    if 데이터.empty:
        return pd.Series(dtype=float)
    return 데이터["종가"].pct_change().dropna()


# -----------------------------------
# 계산 함수
# -----------------------------------
def 포트폴리오입력집계(원본포트폴리오):
    거래원본 = 원본포트폴리오.copy()

    if "거래일자" not in 거래원본.columns and "구입일자" in 거래원본.columns:
        거래원본["거래일자"] = 거래원본["구입일자"]
    if "거래구분" not in 거래원본.columns:
        거래원본["거래구분"] = "매수"
    if "거래수량" not in 거래원본.columns and "보유수량" in 거래원본.columns:
        거래원본["거래수량"] = 거래원본["보유수량"]
    if "거래단가" not in 거래원본.columns and "매입단가" in 거래원본.columns:
        거래원본["거래단가"] = 거래원본["매입단가"]
    if "종목명" not in 거래원본.columns:
        거래원본["종목명"] = ""
    if "비고" not in 거래원본.columns:
        거래원본["비고"] = ""

    if "_입력원본순서" not in 거래원본.columns:
        거래원본["_입력원본순서"] = range(len(거래원본))
    거래원본["_원본순서"] = pd.to_numeric(거래원본["_입력원본순서"], errors="coerce").fillna(pd.Series(range(len(거래원본)), index=거래원본.index)).astype(int)
    거래원본["종목코드"] = 거래원본["종목코드"].astype(str).str.extract(r'(\d+)')[0].fillna('').str.zfill(6)
    거래원본["종목명"] = 거래원본.apply(lambda 행: 종목명자동보정(행.get("종목코드", ""), 행.get("종목명", "")), axis=1)
    거래원본["거래수량"] = pd.to_numeric(거래원본["거래수량"], errors="coerce").fillna(0).clip(lower=0)
    거래원본["거래단가"] = 거래원본["거래단가"].apply(통화문자정리)
    거래원본["거래단가"] = pd.to_numeric(거래원본["거래단가"], errors="coerce").fillna(0).clip(lower=0)
    거래원본["거래일자"] = pd.to_datetime(거래원본["거래일자"], errors="coerce")
    거래원본["거래구분"] = 거래원본["거래구분"].astype(str).str.strip()

    거래원본 = 거래원본[거래원본["종목코드"].str.len() == 6].copy()
    거래원본 = 거래원본[거래원본["거래수량"] > 0].copy()
    거래원본 = 거래원본[거래원본["거래구분"].isin(["매수", "매도"])].copy()
    거래원본 = 거래원본.sort_values(["종목코드", "거래일자", "_원본순서"], ascending=[True, True, True]).reset_index(drop=True)

    집계결과 = []

    for 종목코드, 그룹 in 거래원본.groupby("종목코드", sort=False):
        총매수수량 = 총매수금액 = 총매도수량 = 총매도금액 = 실현손익 = 0.0
        보유수량 = 보유원가 = 과잉매도수량 = 0.0
        최초매수일자 = pd.NaT
        최근거래일자 = pd.NaT
        최근종목명 = ""

        for _, 행 in 그룹.iterrows():
            거래일자 = 행["거래일자"]
            거래구분 = str(행["거래구분"]).strip()
            수량 = float(행["거래수량"])
            단가 = float(행["거래단가"])
            최근거래일자 = 거래일자
            최근종목명 = 종목명자동보정(종목코드, 행.get("종목명", ""))

            if 거래구분 == "매수":
                if pd.isna(최초매수일자):
                    최초매수일자 = 거래일자
                총매수수량 += 수량
                총매수금액 += 수량 * 단가
                보유수량 += 수량
                보유원가 += 수량 * 단가
            else:
                총매도수량 += 수량
                총매도금액 += 수량 * 단가
                평균원가 = (보유원가 / 보유수량) if 보유수량 > 0 else 0.0
                반영매도수량 = min(수량, 보유수량)
                과잉매도수량 += max(0.0, 수량 - 보유수량)
                실현손익 += (단가 - 평균원가) * 반영매도수량
                보유원가 -= 평균원가 * 반영매도수량
                보유수량 -= 반영매도수량
                보유수량 = max(0.0, 보유수량)
                보유원가 = max(0.0, 보유원가)

        매입평균단가 = (보유원가 / 보유수량) if 보유수량 > 0 else 0.0
        집계결과.append({
            "종목코드": str(종목코드).zfill(6),
            "종목명": 종목명자동보정(종목코드, 최근종목명),
            "보유수량": 보유수량,
            "투자원금": 보유원가,
            "매입평균단가": 매입평균단가,
            "매입단가": 매입평균단가,
            "총매수수량": 총매수수량,
            "총매수금액": 총매수금액,
            "총매도수량": 총매도수량,
            "총매도금액": 총매도금액,
            "실현손익": 실현손익,
            "과잉매도수량": 과잉매도수량,
            "최초매수일자": 최초매수일자,
            "최근거래일자": 최근거래일자,
        })

    집계표 = pd.DataFrame(집계결과)
    if 집계표.empty:
        return 집계표

    집계표["최초매수일자"] = pd.to_datetime(집계표["최초매수일자"], errors="coerce").dt.date
    집계표["최근거래일자"] = pd.to_datetime(집계표["최근거래일자"], errors="coerce").dt.date
    return 집계표



def 포트폴리오계산(원본포트폴리오, refresh_token=0):
    계산표 = 포트폴리오입력집계(원본포트폴리오).copy()
    계산표["종목코드"] = 계산표["종목코드"].astype(str).str.zfill(6)
    계산표["보유수량"] = pd.to_numeric(계산표["보유수량"], errors="coerce").fillna(0).clip(lower=0)
    계산표["매입단가"] = pd.to_numeric(계산표["매입단가"], errors="coerce").fillna(0).clip(lower=0)

    def 현재가조회(code):
        if 종목구분판단(code) == "etf":
            return ETF현재가가져오기(code, refresh_token=refresh_token)
        return 종목현재가가져오기(code, refresh_token=refresh_token)

    계산표["현재가"] = pd.to_numeric(계산표["종목코드"].apply(현재가조회), errors="coerce")
    계산표["데이터상태"] = 계산표["현재가"].apply(
        lambda 값: "정상" if pd.notna(값) and 값 > 0 else "현재가 조회 실패"
    )

    계산표["평가금액"] = 계산표.apply(
        lambda 행: 행["현재가"] * 행["보유수량"]
        if pd.notna(행["현재가"]) and 행["현재가"] > 0 else None,
        axis=1,
    )

    계산표["평가손익"] = 계산표.apply(
        lambda 행: 행["평가금액"] - 행["투자원금"]
        if pd.notna(행["평가금액"]) else None,
        axis=1,
    )

    계산표["수익률"] = 계산표.apply(
        lambda 행: (행["평가손익"] / 행["투자원금"] * 100)
        if pd.notna(행["평가손익"]) and 행["투자원금"] not in [0, None] else None,
        axis=1,
    )

    정상평가금액합계 = 계산표.loc[계산표["데이터상태"] == "정상", "평가금액"].sum(min_count=1)
    if pd.isna(정상평가금액합계) or 정상평가금액합계 == 0:
        계산표["현재비중"] = 0.0
    else:
        계산표["현재비중"] = 계산표.apply(
            lambda 행: (행["평가금액"] / 정상평가금액합계 * 100)
            if pd.notna(행["평가금액"]) else 0.0,
            axis=1,
        )

    return 계산표


@st.cache_data(ttl=5, show_spinner=False)
def 포트폴리오계산캐시(거래이력json문자열, refresh_token=0):
    try:
        원본 = json.loads(거래이력json문자열)
        작업df = pd.DataFrame(원본)
    except Exception:
        return pd.DataFrame()
    return 포트폴리오계산(작업df, refresh_token=refresh_token)


# -----------------------------------
# v5.9 월간 수익률 리포트
# -----------------------------------
def 월간실현손익원장생성(거래df):
    """
    평균단가 방식으로 매도 시점의 실현손익을 계산합니다.
    - 월간 리포트용 보조 원장
    - 세금, 수수료, 배당금은 반영하지 않습니다.
    """
    표준열 = ["종목코드", "종목명", "거래일자", "거래구분", "거래수량", "거래단가", "운용사", "비고"]
    if 거래df is None or 거래df.empty:
        return pd.DataFrame(columns=표준열 + ["거래금액", "매도원가", "실현손익", "매도수익률", "년월"])

    작업 = 거래이력계산대상추출(거래df).copy()
    if 작업.empty:
        return pd.DataFrame(columns=표준열 + ["거래금액", "매도원가", "실현손익", "매도수익률", "년월"])

    for 열 in 표준열:
        if 열 not in 작업.columns:
            작업[열] = ""

    작업["거래일자"] = pd.to_datetime(작업["거래일자"], errors="coerce")
    작업["거래수량"] = pd.to_numeric(작업["거래수량"], errors="coerce").fillna(0.0)
    작업["거래단가"] = pd.to_numeric(작업["거래단가"], errors="coerce").fillna(0.0)
    작업 = 작업.dropna(subset=["거래일자"]).copy()
    작업 = 작업.sort_values(["종목코드", "거래일자"]).reset_index(drop=True)

    상태 = {}
    결과 = []
    for _, 행 in 작업.iterrows():
        코드 = str(행.get("종목코드", "")).zfill(6)
        종목명 = 종목명자동보정(코드, 행.get("종목명", ""))
        구분 = str(행.get("거래구분", "")).strip()
        수량 = float(행.get("거래수량", 0) or 0)
        단가 = float(행.get("거래단가", 0) or 0)
        거래금액 = 수량 * 단가

        보유수량, 평균단가 = 상태.get(코드, (0.0, 0.0))
        매도원가 = 0.0
        실현손익 = 0.0
        매도수익률 = None

        if 구분 == "매수":
            새수량 = 보유수량 + 수량
            새원가 = 보유수량 * 평균단가 + 거래금액
            평균단가 = (새원가 / 새수량) if 새수량 > 0 else 0.0
            보유수량 = 새수량
        elif 구분 == "매도":
            매도원가 = 평균단가 * 수량
            실현손익 = 거래금액 - 매도원가
            매도수익률 = (실현손익 / 매도원가 * 100) if 매도원가 > 0 else None
            보유수량 = max(0.0, 보유수량 - 수량)
            if 보유수량 <= 0:
                평균단가 = 0.0

        상태[코드] = (보유수량, 평균단가)
        결과.append({
            "종목코드": 코드,
            "종목명": 종목명,
            "거래일자": 행.get("거래일자"),
            "거래구분": 구분,
            "거래수량": 수량,
            "거래단가": 단가,
            "거래금액": 거래금액,
            "매도원가": 매도원가,
            "실현손익": 실현손익 if 구분 == "매도" else 0.0,
            "매도수익률": 매도수익률,
            "운용사": 행.get("운용사", ""),
            "비고": 행.get("비고", ""),
        })

    원장 = pd.DataFrame(결과)
    if 원장.empty:
        return 원장
    원장["년월"] = pd.to_datetime(원장["거래일자"], errors="coerce").dt.strftime("%Y-%m")
    return 원장


@st.cache_data(ttl=300, show_spinner=False)
def 월간실현손익원장캐시(거래이력json문자):
    try:
        데이터 = json.loads(거래이력json문자) if 거래이력json문자 else []
        df = pd.DataFrame(데이터)
    except Exception:
        df = pd.DataFrame()
    return 월간실현손익원장생성(df)


def 월간수익률리포트생성(거래df, 계산포트폴리오=None, 선택년월=None):
    if isinstance(거래df, pd.DataFrame) and {"년월", "거래금액", "매도원가", "실현손익"}.issubset(set(거래df.columns)):
        원장 = 거래df.copy()
    else:
        원장 = 월간실현손익원장생성(거래df)
    if 원장.empty or "년월" not in 원장.columns:
        return pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), []

    월목록 = sorted([x for x in 원장["년월"].dropna().unique().tolist() if x], reverse=True)
    선택년월 = 선택년월 or (월목록[0] if 월목록 else "")
    월원장 = 원장[원장["년월"] == 선택년월].copy()
    if 월원장.empty:
        return pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), 월목록

    매수표 = 월원장[월원장["거래구분"] == "매수"].copy()
    매도표 = 월원장[월원장["거래구분"] == "매도"].copy()
    매도건수 = int(len(매도표))
    수익매도건수 = int((pd.to_numeric(매도표.get("실현손익"), errors="coerce").fillna(0) > 0).sum()) if 매도건수 > 0 else 0
    승률 = (수익매도건수 / 매도건수 * 100) if 매도건수 > 0 else 0.0

    월매수금액 = float(pd.to_numeric(매수표.get("거래금액"), errors="coerce").fillna(0).sum()) if not 매수표.empty else 0.0
    월매도금액 = float(pd.to_numeric(매도표.get("거래금액"), errors="coerce").fillna(0).sum()) if not 매도표.empty else 0.0
    월실현손익 = float(pd.to_numeric(매도표.get("실현손익"), errors="coerce").fillna(0).sum()) if not 매도표.empty else 0.0
    월매도원가 = float(pd.to_numeric(매도표.get("매도원가"), errors="coerce").fillna(0).sum()) if not 매도표.empty else 0.0
    월실현수익률 = (월실현손익 / 월매도원가 * 100) if 월매도원가 > 0 else 0.0

    현재보유평가손익 = 0.0
    현재보유수익률 = 0.0
    if isinstance(계산포트폴리오, pd.DataFrame) and not 계산포트폴리오.empty:
        보유df = 보유포트폴리오필터(계산포트폴리오)
        정상df = 보유df[보유df["데이터상태"] == "정상"].copy() if "데이터상태" in 보유df.columns else 보유df.copy()
        현재투자원금 = float(pd.to_numeric(정상df.get("투자원금"), errors="coerce").fillna(0).sum()) if not 정상df.empty else 0.0
        현재보유평가손익 = float(pd.to_numeric(정상df.get("평가손익"), errors="coerce").fillna(0).sum()) if not 정상df.empty else 0.0
        현재보유수익률 = (현재보유평가손익 / 현재투자원금 * 100) if 현재투자원금 > 0 else 0.0

    요약 = pd.DataFrame([
        {"항목": "대상월", "값": 선택년월},
        {"항목": "거래건수", "값": int(len(월원장))},
        {"항목": "매수금액", "값": 월매수금액},
        {"항목": "매도금액", "값": 월매도금액},
        {"항목": "순매수금액", "값": 월매수금액 - 월매도금액},
        {"항목": "실현손익", "값": 월실현손익},
        {"항목": "실현수익률(%)", "값": 월실현수익률},
        {"항목": "매도승률(%)", "값": 승률},
        {"항목": "현재 보유 평가손익", "값": 현재보유평가손익},
        {"항목": "현재 보유 수익률(%)", "값": 현재보유수익률},
    ])

    종목별 = 월원장.groupby(["종목코드", "종목명"], dropna=False).agg(
        거래건수=("거래구분", "count"),
        매수금액=("거래금액", lambda s: s[월원장.loc[s.index, "거래구분"] == "매수"].sum()),
        매도금액=("거래금액", lambda s: s[월원장.loc[s.index, "거래구분"] == "매도"].sum()),
        실현손익=("실현손익", "sum"),
        매도원가=("매도원가", "sum"),
    ).reset_index()
    종목별["실현수익률"] = 종목별.apply(lambda 행: (행["실현손익"] / 행["매도원가"] * 100) if 행["매도원가"] > 0 else 0.0, axis=1)
    종목별 = 종목별.sort_values(["실현손익", "매수금액"], ascending=[False, False]).reset_index(drop=True)

    월원장 = 월원장.sort_values(["거래일자", "종목명"]).reset_index(drop=True)
    return 요약, 종목별, 월원장, 월목록


def 월간거래엑셀표시용정리(월원장):
    """월간 리포트 엑셀 다운로드 전용 표시 정리.
    - 거래일자는 YYYY-MM-DD 문자열로 고정하여 엑셀에서 00:00:00이 보이지 않게 합니다.
    - 이미 선택월로 필터된 다운로드 파일이므로 맨 오른쪽 '년월' 중복 컬럼은 제거합니다.
    - 종목코드는 앞자리 0이 사라지지 않도록 문자열 6자리로 유지합니다.
    """
    작업 = pd.DataFrame() if 월원장 is None else 월원장.copy()
    if 작업.empty:
        return 작업

    if "거래일자" in 작업.columns:
        작업["거래일자"] = 작업["거래일자"].apply(날짜값_YYYYMMDD문자열)

    if "종목코드" in 작업.columns:
        작업["종목코드"] = 작업["종목코드"].apply(
            lambda 값: "" if pd.isna(값) else re.sub(r"[^0-9]", "", str(값)).zfill(6)
        )

    # 월간거래 시트는 선택월 파일이므로 '년월' 컬럼이 중복 정보입니다.
    if "년월" in 작업.columns:
        작업 = 작업.drop(columns=["년월"])

    표시순서 = [
        "종목코드", "종목명", "거래일자", "거래구분", "거래수량", "거래단가",
        "거래금액", "매도원가", "실현손익", "매도수익률", "운용사", "비고"
    ]
    표시순서 = [열 for 열 in 표시순서 if 열 in 작업.columns]
    나머지열 = [열 for 열 in 작업.columns if 열 not in 표시순서]
    return 작업[표시순서 + 나머지열].copy()


@st.cache_data(ttl=300, show_spinner=False)
def 월간수익률리포트엑셀바이트(요약, 종목별, 월원장, 선택년월):
    버퍼 = io.BytesIO()
    월원장표시 = 월간거래엑셀표시용정리(월원장)

    with pd.ExcelWriter(버퍼, engine="openpyxl") as writer:
        요약.to_excel(writer, index=False, sheet_name="월간요약")
        종목별.to_excel(writer, index=False, sheet_name="종목별")
        월원장표시.to_excel(writer, index=False, sheet_name="월간거래")

        try:
            from openpyxl.styles import Alignment
            for sheet_name in ["월간요약", "종목별", "월간거래"]:
                ws = writer.book[sheet_name]
                ws.freeze_panes = "A2"
                for row in ws.iter_rows():
                    for cell in row:
                        cell.alignment = Alignment(vertical="center")
                for col_cells in ws.columns:
                    max_len = 0
                    col_letter = col_cells[0].column_letter
                    for cell in col_cells:
                        value = "" if cell.value is None else str(cell.value)
                        max_len = max(max_len, len(value))
                    ws.column_dimensions[col_letter].width = min(max(max_len + 2, 10), 45)

            ws = writer.book["월간거래"]
            header_map = {cell.value: idx for idx, cell in enumerate(ws[1], start=1)}

            if "종목코드" in header_map:
                col_idx = header_map["종목코드"]
                for row in range(2, ws.max_row + 1):
                    ws.cell(row=row, column=col_idx).number_format = "@"

            if "거래일자" in header_map:
                col_idx = header_map["거래일자"]
                for row in range(2, ws.max_row + 1):
                    ws.cell(row=row, column=col_idx).number_format = "yyyy-mm-dd"

            for name in ["거래수량", "거래단가", "거래금액", "매도원가", "실현손익"]:
                if name in header_map:
                    col_idx = header_map[name]
                    for row in range(2, ws.max_row + 1):
                        ws.cell(row=row, column=col_idx).number_format = "#,##0"

            if "매도수익률" in header_map:
                col_idx = header_map["매도수익률"]
                for row in range(2, ws.max_row + 1):
                    ws.cell(row=row, column=col_idx).number_format = "0.00"
        except Exception:
            pass

    버퍼.seek(0)
    return 버퍼.getvalue()


def 월간수익률리포트UI(거래df, 계산포트폴리오):
    거래이력json문자 = 거래이력서명생성(거래df if 거래df is not None else pd.DataFrame())
    원장 = 월간실현손익원장캐시(거래이력json문자)
    if 원장.empty or "년월" not in 원장.columns:
        st.info("월간 리포트를 만들 거래이력이 아직 없습니다.")
        return

    월목록 = sorted([x for x in 원장["년월"].dropna().unique().tolist() if x], reverse=True)
    if not 월목록:
        st.info("월간 리포트를 만들 거래월이 없습니다.")
        return

    with st.expander("v5.9 월간 수익률 리포트", expanded=True):
        선택년월 = st.selectbox("조회 월", 월목록, index=0, key="monthly_report_month_v59")
        요약, 종목별, 월원장, _ = 월간수익률리포트생성(원장, 계산포트폴리오, 선택년월)
        if 요약.empty:
            st.info("선택한 월의 거래이력이 없습니다.")
            return

        값사전 = {행["항목"]: 행["값"] for _, 행 in 요약.iterrows()}
        카드1, 카드2, 카드3, 카드4, 카드5 = st.columns(5)
        카드1.metric("매수금액", 안전정수포맷(값사전.get("매수금액", 0)))
        카드2.metric("매도금액", 안전정수포맷(값사전.get("매도금액", 0)))
        카드3.metric("실현손익", 손익문자열(값사전.get("실현손익", 0)))
        카드4.metric("실현수익률", 비율표시(값사전.get("실현수익률(%)", 0)))
        카드5.metric("매도승률", 비율표시(값사전.get("매도승률(%)", 0)))

        st.caption("실현손익은 평균단가 방식으로 계산합니다. 세금·수수료·배당금은 반영하지 않으며, 현재 보유 수익률은 월말 기준이 아니라 현재가 기준입니다.")

        if 종목별.empty:
            st.info("선택한 월의 종목별 집계가 없습니다.")
        else:
            표시종목별 = 종목별.rename(columns={"매도원가": "매도 원가", "실현수익률": "실현수익률"})
            표데이터프레임(
                index_1부터(표시종목별).style.format({
                    "매수금액": 안전정수포맷,
                    "매도금액": 안전정수포맷,
                    "실현손익": 손익문자열,
                    "매도 원가": 안전정수포맷,
                    "실현수익률": 수익률문자열,
                }).map(손익색상, subset=["실현손익"]).map(수익률색상, subset=["실현수익률"]),
                use_container_width=True,
            )

        with st.expander("실현손익 상세", expanded=False):
            전체거래표시 = st.checkbox("매수까지 포함한 전체 월간 거래 보기", value=False, key="monthly_report_show_all_trades_v59")
            표시원장 = 월원장.copy() if 전체거래표시 else 월원장[월원장["거래구분"] == "매도"].copy()
            if 표시원장.empty:
                st.info("선택한 월에는 매도 거래가 없어 실현손익 상세가 없습니다. 전체 거래가 필요하면 위 옵션을 켜 주세요.")
            else:
                표시원장["거래일자"] = pd.to_datetime(표시원장["거래일자"], errors="coerce").dt.strftime("%Y-%m-%d")
                표시원장 = 표시원장[["거래일자", "종목코드", "종목명", "거래구분", "거래수량", "거래단가", "거래금액", "실현손익", "매도수익률", "비고"]].copy()
                표데이터프레임(
                    index_1부터(표시원장).style.format({
                        "거래수량": 안전정수포맷,
                        "거래단가": 안전정수포맷,
                        "거래금액": 안전정수포맷,
                        "실현손익": 손익문자열,
                        "매도수익률": 수익률문자열,
                    }).map(손익색상, subset=["실현손익"]).map(수익률색상, subset=["매도수익률"]),
                    use_container_width=True,
                )

        try:
            st.download_button(
                "월간 리포트 xlsx 다운로드",
                data=월간수익률리포트엑셀바이트(요약, 종목별, 월원장, 선택년월),
                file_name=f"monthly_report_{선택년월}_{서울현재시각().strftime('%Y%m%d_%H%M')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="monthly_report_download_v592",
                use_container_width=True,
            )
        except Exception as e:
            st.warning(f"월간 리포트 다운로드 준비 실패: {e}")


# -----------------------------------
# v5.10 종목별 누적 성적표
# -----------------------------------
def 종목별누적성적표생성(거래df, 계산포트폴리오=None):
    if 계산포트폴리오 is not None and not 계산포트폴리오.empty:
        성적표 = 계산포트폴리오.copy()
    else:
        성적표 = 포트폴리오계산(거래이력계산대상추출(거래df if 거래df is not None else pd.DataFrame()))

    if 성적표 is None or 성적표.empty:
        return pd.DataFrame()

    for 열 in ["총매수금액", "총매도금액", "실현손익", "평가손익", "투자원금", "평가금액", "보유수량", "총매수수량", "총매도수량"]:
        if 열 not in 성적표.columns:
            성적표[열] = 0
        성적표[열] = pd.to_numeric(성적표[열], errors="coerce").fillna(0)

    성적표["총손익"] = 성적표["실현손익"] + 성적표["평가손익"]
    성적표["총수익률"] = 성적표.apply(
        lambda 행: (행["총손익"] / 행["총매수금액"] * 100) if 행.get("총매수금액", 0) not in [0, None] else 0,
        axis=1,
    )

    승률표 = pd.DataFrame(columns=["종목코드", "승률", "매도거래수"])
    try:
        원장 = 월간실현손익원장생성(거래이력계산대상추출(거래df if 거래df is not None else pd.DataFrame()))
        매도원장 = 원장[원장["거래구분"] == "매도"].copy() if not 원장.empty and "거래구분" in 원장.columns else pd.DataFrame()
        if not 매도원장.empty:
            매도원장["승리"] = pd.to_numeric(매도원장["실현손익"], errors="coerce").fillna(0) > 0
            승률표 = 매도원장.groupby("종목코드").agg(
                승률=("승리", lambda s: float(s.mean() * 100) if len(s) else 0),
                매도거래수=("승리", "count"),
            ).reset_index()
    except Exception:
        승률표 = pd.DataFrame(columns=["종목코드", "승률", "매도거래수"])

    성적표["종목코드"] = 성적표["종목코드"].astype(str).str.zfill(6)
    if not 승률표.empty:
        승률표["종목코드"] = 승률표["종목코드"].astype(str).str.zfill(6)
        성적표 = 성적표.merge(승률표, on="종목코드", how="left")
    else:
        성적표["승률"] = 0
        성적표["매도거래수"] = 0

    성적표["승률"] = pd.to_numeric(성적표.get("승률", 0), errors="coerce").fillna(0)
    성적표["매도거래수"] = pd.to_numeric(성적표.get("매도거래수", 0), errors="coerce").fillna(0).astype(int)

    표시열 = ["종목코드", "종목명", "보유수량", "총매수금액", "총매도금액", "실현손익", "평가손익", "총손익", "총수익률", "승률", "매도거래수", "최근거래일자"]
    for 열 in 표시열:
        if 열 not in 성적표.columns:
            성적표[열] = "" if 열 in ["종목코드", "종목명", "최근거래일자"] else 0
    return 성적표[표시열].copy()


def 종목별누적성적표UI(거래df, 계산포트폴리오):
    성적표 = 종목별누적성적표생성(거래df, 계산포트폴리오)
    if 성적표.empty:
        st.info("종목별 성적표를 만들 거래이력이 아직 없습니다.")
        return

    with st.expander("v5.10 종목별 누적 성적표", expanded=True):
        정렬옵션 = st.selectbox("정렬 기준", ["총손익 높은 순", "총수익률 높은 순", "손실 큰 순", "실현손익 높은 순", "보유수량 많은 순"], key="stock_scorecard_sort_v510")
        정렬컬럼 = {
            "총손익 높은 순": ("총손익", False),
            "총수익률 높은 순": ("총수익률", False),
            "손실 큰 순": ("총손익", True),
            "실현손익 높은 순": ("실현손익", False),
            "보유수량 많은 순": ("보유수량", False),
        }.get(정렬옵션, ("총손익", False))
        성적표 = 성적표.sort_values(정렬컬럼[0], ascending=정렬컬럼[1]).reset_index(drop=True)

        수익1위 = 성적표.iloc[성적표["총손익"].idxmax()] if not 성적표.empty else None
        수익하위 = 성적표.iloc[성적표["총손익"].idxmin()] if not 성적표.empty else None
        수익률1위 = 성적표.iloc[성적표["총수익률"].idxmax()] if not 성적표.empty else None
        보유종목수 = int((pd.to_numeric(성적표["보유수량"], errors="coerce").fillna(0) > 0).sum())

        카드1, 카드2, 카드3, 카드4 = st.columns(4)
        카드1.metric("수익 1위", 수익1위["종목명"] if 수익1위 is not None else "-", 손익문자열(수익1위["총손익"]) if 수익1위 is not None else "")
        카드2.metric("수익 하위", 수익하위["종목명"] if 수익하위 is not None else "-", 손익문자열(수익하위["총손익"]) if 수익하위 is not None else "")
        카드3.metric("수익률 1위", 수익률1위["종목명"] if 수익률1위 is not None else "-", 수익률문자열(수익률1위["총수익률"]) if 수익률1위 is not None else "")
        카드4.metric("현재 보유 종목", f"{보유종목수}개")

        st.caption("총손익은 실현손익과 현재 보유분 평가손익을 합산한 값입니다. 세금·수수료·배당금은 반영하지 않습니다. ‘수익 하위’는 손실 종목이라는 뜻이 아니라, 현재 성적표 안에서 총손익이 가장 낮은 종목을 의미합니다.")
        표시표 = 성적표.rename(columns={"총매수금액": "누적 매수금액", "총매도금액": "누적 매도금액", "총수익률": "총수익률(%)", "매도거래수": "매도 거래수", "최근거래일자": "최근 거래일자"})
        표데이터프레임(
            index_1부터(표시표).style.format({
                "보유수량": 안전정수포맷,
                "누적 매수금액": 안전정수포맷,
                "누적 매도금액": 안전정수포맷,
                "실현손익": 손익문자열,
                "평가손익": 손익문자열,
                "총손익": 손익문자열,
                "총수익률(%)": 수익률문자열,
                "승률": 비율표시,
                "매도 거래수": 안전정수포맷,
            }).map(손익색상, subset=["실현손익", "평가손익", "총손익"]).map(수익률색상, subset=["총수익률(%)"]),
            use_container_width=True,
        )
        with st.expander("성적표 읽는 법", expanded=False):
            st.markdown("""
            - **실현손익**: 매도해서 이미 확정된 손익입니다.
            - **평가손익**: 현재 보유 중인 수량의 현재가 기준 손익입니다.
            - **총손익**: 실현손익 + 평가손익입니다.
            - **총수익률**: 누적 매수금액 대비 총손익 비율입니다.
            - **승률**: 매도 거래 중 수익으로 끝난 거래 비율입니다.
            """)

        자동분석코멘트UI(성적표, 거래df=거래df)



# -----------------------------------
# -----------------------------------
# v5.10.3b 거래기록 기반 고도화 코멘트
# -----------------------------------
def 안전숫자(값, 기본값=0.0):
    try:
        if pd.isna(값):
            return 기본값
        return float(값)
    except Exception:
        return 기본값


def 종목분류라벨(종목명, 종목코드=""):
    이름 = 종목명이름정리(종목명).upper()
    코드 = str(종목코드 or "").zfill(6)
    if any(키 in 이름 for 키 in ["KODEX", "TIGER", "ACE", "KBSTAR", "ARIRANG", "SOL", "HANARO", "KOSEF", "TIMEFOLIO", "PLUS"]):
        return "ETF"
    if 코드 in ["005930", "000660"] or any(키 in 이름 for 키 in ["삼성전자", "하이닉스", "반도체"]):
        return "반도체/대형주"
    return "개별주"


def 종목별판단등급(행):
    총수익률 = 안전숫자(행.get("총수익률", 0))
    총손익 = 안전숫자(행.get("총손익", 0))
    승률 = 안전숫자(행.get("승률", 0))
    매도거래수 = int(안전숫자(행.get("매도거래수", 0)))
    보유수량 = 안전숫자(행.get("보유수량", 0))

    점수 = 50
    if 총수익률 >= 8:
        점수 += 20
    elif 총수익률 >= 3:
        점수 += 12
    elif 총수익률 > 0:
        점수 += 6
    elif 총수익률 <= -8:
        점수 -= 22
    elif 총수익률 <= -3:
        점수 -= 12
    elif 총수익률 < 0:
        점수 -= 6

    if 매도거래수 >= 3:
        if 승률 >= 65:
            점수 += 12
        elif 승률 >= 50:
            점수 += 6
        elif 승률 < 35:
            점수 -= 10
    elif 매도거래수 == 0 and 보유수량 > 0:
        점수 -= 2

    if 총손익 > 0:
        점수 += 5
    elif 총손익 < 0:
        점수 -= 5

    점수 = max(0, min(100, int(round(점수))))
    if 점수 >= 75:
        등급 = "우수"
    elif 점수 >= 60:
        등급 = "양호"
    elif 점수 >= 45:
        등급 = "중립"
    elif 점수 >= 30:
        등급 = "주의"
    else:
        등급 = "점검"
    return 점수, 등급


def 종목거래패턴분석(거래df, 종목코드, 종목명=""):
    """v5.10.3a의 장점인 거래기록 기반 해석을 v5.10.1 구조에 맞게 이식."""
    결과 = {
        "매수패턴": "거래기록이 부족해 매수 패턴 판단이 제한됩니다.",
        "보유매도": "매도 기록이 부족해 보유·매도 습관 판단이 제한됩니다.",
        "거래메모": "거래 메모가 충분하지 않습니다.",
        "매수횟수": 0,
        "매도횟수": 0,
        "평균단가개선비율": None,
        "보유일수": None,
    }
    if 거래df is None or 거래df.empty:
        return 결과

    작업 = 거래이력계산대상추출(거래df.copy())
    if 작업 is None or 작업.empty or "종목코드" not in 작업.columns:
        return 결과

    코드 = re.sub(r"[^0-9]", "", str(종목코드 or "")).zfill(6)
    작업["종목코드"] = 작업["종목코드"].apply(lambda x: re.sub(r"[^0-9]", "", str(x)).zfill(6))
    g = 작업[작업["종목코드"] == 코드].copy()
    if g.empty:
        이름 = 종목명이름정리(종목명)
        if 이름 and "종목명" in 작업.columns:
            g = 작업[작업["종목명"].apply(종목명이름정리) == 이름].copy()
    if g.empty:
        return 결과

    g["거래일자_dt"] = pd.to_datetime(g["거래일자"], errors="coerce")
    g["거래수량_num"] = pd.to_numeric(g["거래수량"], errors="coerce").fillna(0)
    g["거래단가_num"] = pd.to_numeric(g["거래단가"], errors="coerce").fillna(0)
    g = g.sort_values("거래일자_dt")
    buys = g[g["거래구분"] == "매수"].copy()
    sells = g[g["거래구분"] == "매도"].copy()
    결과["매수횟수"] = len(buys)
    결과["매도횟수"] = len(sells)

    if len(buys) == 1:
        결과["매수패턴"] = "단일 매수 중심입니다. 아직 분할매수 효과를 판단하기에는 데이터가 부족합니다."
    elif len(buys) >= 2:
        가격목록 = list(buys["거래단가_num"].astype(float))
        낮춘횟수 = sum(1 for i in range(1, len(가격목록)) if 가격목록[i] < 가격목록[i - 1])
        높인횟수 = sum(1 for i in range(1, len(가격목록)) if 가격목록[i] > 가격목록[i - 1])
        비율 = 낮춘횟수 / max(len(가격목록) - 1, 1) * 100
        결과["평균단가개선비율"] = 비율
        if 낮춘횟수 > 높인횟수:
            결과["매수패턴"] = f"{len(buys)}회 매수 중 추가매수 가격을 낮춘 흐름이 우세합니다. 평균단가 관리 측면에서는 긍정적입니다."
        elif 높인횟수 > 낮춘횟수:
            결과["매수패턴"] = f"{len(buys)}회 매수 중 이전 매수가보다 높은 가격에 추가 진입한 흐름이 우세합니다. 추격매수 여부를 점검해야 합니다."
        else:
            결과["매수패턴"] = f"{len(buys)}회 분할매수했으나 가격을 낮춘 매수와 높인 매수가 혼재되어 있습니다. 매수 기준을 더 명확히 할 필요가 있습니다."

    유효일자 = g["거래일자_dt"].dropna()
    if not 유효일자.empty:
        보유일수 = max((유효일자.max() - 유효일자.min()).days, 0)
        결과["보유일수"] = 보유일수
    else:
        보유일수 = None

    if len(sells) == 0:
        결과["보유매도"] = "아직 매도 기록이 없습니다. 현재는 수익률 자체보다 보유 비중과 추가매수 기준 관리가 핵심입니다."
    else:
        if 보유일수 is not None and 보유일수 <= 7:
            결과["보유매도"] = "짧은 기간 안에 매도가 발생했습니다. 단기 대응 기준이 분명했는지 확인이 필요합니다."
        elif 보유일수 is not None and 보유일수 >= 30:
            결과["보유매도"] = "일정 기간 보유 후 매도가 이루어졌습니다. 수익 실현 기준을 기록하면 다음 판단에 도움이 됩니다."
        else:
            결과["보유매도"] = "중간 기간 보유 후 매도가 이루어졌습니다. 매도 이유를 함께 기록하면 분석 품질이 높아집니다."

    if "비고" in g.columns:
        메모 = " ".join([str(x) for x in g["비고"].dropna().tolist() if str(x).strip()])
        if 메모:
            키워드 = []
            for 단어 in ["폭락", "반등", "하락", "유가", "환율", "금리", "이란", "트럼프", "최저", "안정", "분할매수", "손절"]:
                if 단어 in 메모 and 단어 not in 키워드:
                    키워드.append(단어)
            if 키워드:
                결과["거래메모"] = "거래 메모에 " + ", ".join(키워드[:6]) + " 관련 판단이 기록되어 있습니다. 당시 판단 근거와 실제 결과를 함께 복기할 가치가 있습니다."
            else:
                결과["거래메모"] = "거래 메모가 남아 있어 사후 복기에 활용할 수 있습니다. 다음부터는 매수 이유·목표가·손절 기준을 더 구체적으로 남기면 좋습니다."
    return 결과


def 간단문장정리(문장, 최대길이=90):
    """표 안에서 읽기 쉽도록 문장을 짧게 정리합니다."""
    문장 = "" if 문장 is None else str(문장).strip()
    문장 = re.sub(r"\s+", " ", 문장)
    if len(문장) <= 최대길이:
        return 문장
    return 문장[:최대길이].rstrip() + "…"


def 종목별자동코멘트생성(성적표, 거래df=None):
    """
    v5.13.5 개선:
    - 개별 분석은 숫자/근거 중심으로 유지
    - 코멘트는 상태·핵심·다음점검 3줄 구조로 축소
    - 모든 종목에 반복되던 공통 안내문과 장문 해석 제거
    """
    기본열 = [
        "종목코드", "종목명", "구분", "점수", "판단등급",
        "총손익", "총수익률", "승률",
        "상태", "핵심", "다음점검", "요약코멘트", "상세근거"
    ]
    if 성적표 is None or 성적표.empty:
        return pd.DataFrame(columns=기본열)

    결과 = []
    for _, 행 in 성적표.iterrows():
        종목명 = str(행.get("종목명", "") or 행.get("종목코드", "")).strip()
        종목코드 = str(행.get("종목코드", "")).zfill(6)
        구분 = 종목분류라벨(종목명, 종목코드)

        총수익률 = 안전숫자(행.get("총수익률", 0))
        총손익 = 안전숫자(행.get("총손익", 0))
        실현손익 = 안전숫자(행.get("실현손익", 0))
        평가손익 = 안전숫자(행.get("평가손익", 0))
        승률 = 안전숫자(행.get("승률", 0))
        매도거래수 = int(안전숫자(행.get("매도거래수", 0)))
        보유수량 = 안전숫자(행.get("보유수량", 0))

        점수, 등급 = 종목별판단등급(행)
        패턴 = 종목거래패턴분석(거래df, 종목코드, 종목명)

        # 1) 상태: 수익/손실 상태만 짧게
        if 총손익 > 0 and 총수익률 >= 5:
            상태 = "수익 구간"
        elif 총손익 > 0:
            상태 = "소폭 수익"
        elif 총손익 < 0 and 보유수량 > 0:
            상태 = "평가손실 관리"
        elif 총손익 < 0:
            상태 = "누적 손실"
        else:
            상태 = "중립"

        # 2) 핵심: 종목마다 가장 다른 1가지 포인트만
        매수횟수 = int(패턴.get("매수횟수", 0) or 0)
        매도횟수 = int(패턴.get("매도횟수", 0) or 0)
        평균단가개선비율 = 패턴.get("평균단가개선비율")

        if 매수횟수 >= 2 and 평균단가개선비율 is not None:
            if 평균단가개선비율 >= 60:
                핵심 = f"{매수횟수}회 분할매수 중 가격을 낮춘 흐름이 우세합니다."
            elif 평균단가개선비율 <= 40:
                핵심 = f"{매수횟수}회 분할매수 중 높은 가격 추가 진입 비중이 큽니다."
            else:
                핵심 = f"{매수횟수}회 분할매수했으나 매수 가격 방향이 혼재되어 있습니다."
        elif 매수횟수 == 1:
            핵심 = "단일 매수 중심이라 분할매수 효과 판단은 아직 제한됩니다."
        else:
            핵심 = "거래기록이 부족해 매수 패턴 판단이 제한됩니다."

        if 매도거래수 >= 3:
            핵심 += f" 매도 승률은 {승률:.0f}%입니다."
        elif 매도횟수 == 0:
            핵심 += " 아직 매도 기록은 없습니다."

        # 3) 다음 점검: 실행 가능한 1문장만
        if 총손익 < 0 and 보유수량 > 0:
            다음점검 = "추가매수 전 손실 허용 범위와 재진입 기준을 먼저 정하세요."
        elif 총손익 > 0 and 매수횟수 >= 2:
            다음점검 = "분할매수가 효과를 낸 구간인지 확인하고 같은 기준을 반복 적용할 수 있는지 점검하세요."
        elif 실현손익 > 0 and 평가손익 < 0:
            다음점검 = "확정 수익과 보유분 손실을 분리해 최근 보유 종목의 기준가를 다시 확인하세요."
        elif 실현손익 < 0 and 평가손익 > 0:
            다음점검 = "과거 손실 후 회복 중이므로 목표가와 부분매도 기준을 정해 관찰하세요."
        else:
            다음점검 = "다음 거래부터 매수 이유, 목표가, 손절 기준을 함께 기록하세요."

        요약코멘트 = f"상태: {상태} | 핵심: {핵심} | 다음: {다음점검}"

        상세근거후보 = []
        if 패턴.get("매수패턴"):
            상세근거후보.append(패턴.get("매수패턴"))
        if 패턴.get("보유매도"):
            상세근거후보.append(패턴.get("보유매도"))
        if 패턴.get("거래메모") and "충분하지" not in str(패턴.get("거래메모")):
            상세근거후보.append(패턴.get("거래메모"))
        상세근거 = " ".join(dict.fromkeys([str(x).strip() for x in 상세근거후보 if str(x).strip()]))

        결과.append({
            "종목코드": 종목코드,
            "종목명": 종목명,
            "구분": 구분,
            "점수": 점수,
            "판단등급": 등급,
            "총손익": 총손익,
            "총수익률": 총수익률,
            "승률": 승률,
            "상태": 상태,
            "핵심": 간단문장정리(핵심, 95),
            "다음점검": 간단문장정리(다음점검, 80),
            "요약코멘트": 요약코멘트,
            "상세근거": 간단문장정리(상세근거, 220),
        })

    return pd.DataFrame(결과).sort_values(["점수", "총손익"], ascending=[False, False]).reset_index(drop=True)


def 포트폴리오자동진단표(성적표):
    if 성적표 is None or 성적표.empty:
        return pd.DataFrame(columns=["구분", "현재 해석", "점검 포인트"])
    작업 = 성적표.copy()
    for 열 in ["총손익", "총수익률", "보유수량", "실현손익", "평가손익"]:
        if 열 not in 작업.columns:
            작업[열] = 0
        작업[열] = pd.to_numeric(작업[열], errors="coerce").fillna(0)

    총손익 = 작업["총손익"].sum()
    수익종목 = int((작업["총손익"] > 0).sum())
    손실종목 = int((작업["총손익"] < 0).sum())
    보유종목 = int((작업["보유수량"] > 0).sum())
    최고 = 작업.sort_values("총손익", ascending=False).iloc[0]
    최저 = 작업.sort_values("총손익", ascending=True).iloc[0]

    if 총손익 > 0:
        전체상태 = "전체 포트폴리오는 누적 기준 수익 구간입니다."
        전체점검 = "수익 확대보다 어떤 종목과 어떤 매수 방식이 수익에 기여했는지 확인"
    elif 총손익 < 0:
        전체상태 = "전체 포트폴리오는 누적 기준 손실 구간입니다."
        전체점검 = "손실 원인을 추격매수·비중과다·시장하락으로 분리"
    else:
        전체상태 = "전체 손익은 중립 구간입니다."
        전체점검 = "거래 이유와 기준을 더 축적"

    return pd.DataFrame([
        {"구분": "전체 상태", "현재 해석": 전체상태, "점검 포인트": 전체점검},
        {"구분": "종목 분포", "현재 해석": f"수익 종목 {수익종목}개 / 손실 종목 {손실종목}개 / 보유 종목 {보유종목}개", "점검 포인트": "수익 종목과 손실 종목의 매수 패턴 비교"},
        {"구분": "최고 기여", "현재 해석": f"{최고['종목명']}이 현재 가장 크게 기여하고 있습니다.", "점검 포인트": "성과가 난 이유를 매수 가격·보유기간·매도 타이밍으로 분해"},
        {"구분": "주의 종목", "현재 해석": f"{최저['종목명']}은 상대적으로 부담이 큰 종목입니다.", "점검 포인트": "손실 또는 부진 원인을 비중과 진입 시점 중심으로 점검"},
    ])


def 포트폴리오자동진단문장(성적표):
    표 = 포트폴리오자동진단표(성적표)
    if 표.empty:
        return ["분석할 거래 데이터가 아직 충분하지 않습니다."]
    return [f"{행['구분']}: {행['현재 해석']} 점검 포인트는 {행['점검 포인트']}입니다." for _, 행 in 표.iterrows()]


def 자동분석코멘트UI(성적표, 거래df=None):
    with st.expander("거래기록 기반 요약 코멘트", expanded=True):
        st.caption("개별 분석은 숫자와 성적표 중심, 코멘트는 상태·핵심·다음 점검 중심으로 분리했습니다.")

        진단표 = 포트폴리오자동진단표(성적표)
        if not 진단표.empty:
            st.markdown("#### 포트폴리오 종합 해석")
            표데이터프레임(index_1부터(진단표), use_container_width=True)

        코멘트표 = 종목별자동코멘트생성(성적표, 거래df=거래df)
        if 코멘트표.empty:
            st.info("종목별 자동 코멘트를 만들 데이터가 아직 부족합니다.")
            return

        st.markdown("#### 종목별 요약 코멘트")
        st.caption("중복 설명을 줄이고, 각 종목별로 바로 확인할 수 있는 3가지 항목만 표시합니다.")

        보기열 = ["종목명", "구분", "점수", "판단등급", "총손익", "총수익률", "승률", "상태", "핵심", "다음점검"]
        표시 = 코멘트표[보기열].copy()
        표데이터프레임(
            index_1부터(표시).style.format({
                "점수": 안전정수포맷,
                "총손익": 손익문자열,
                "총수익률": 수익률문자열,
                "승률": 비율표시,
            }).map(손익색상, subset=["총손익"]).map(수익률색상, subset=["총수익률"]),
            use_container_width=True,
        )

        with st.expander("상세 근거 보기", expanded=False):
            상세표 = 코멘트표[["종목명", "판단등급", "요약코멘트", "상세근거"]].copy()
            표데이터프레임(index_1부터(상세표), use_container_width=True)

        with st.expander("개별 분석과 코멘트의 역할 구분", expanded=False):
            차이표 = pd.DataFrame([
                {
                    "메뉴": "개별 분석",
                    "역할": "숫자와 근거 확인",
                    "주요 내용": "보유수량, 매수금액, 평가손익, 총수익률, 승률, 거래횟수",
                    "사용 목적": "현재 성적과 원인을 데이터로 확인",
                },
                {
                    "메뉴": "요약 코멘트",
                    "역할": "해석과 다음 점검",
                    "주요 내용": "상태, 핵심 패턴, 다음 점검 포인트",
                    "사용 목적": "다음 거래 전에 확인할 행동 기준 정리",
                },
            ])
            표데이터프레임(index_1부터(차이표), use_container_width=True)


def 리밸런싱계산(계산표, 목표비중사전):
    결과표 = 계산표.copy()
    총평가금액 = 결과표.loc[결과표["데이터상태"] == "정상", "평가금액"].sum()

    결과표["목표비중"] = 결과표["종목코드"].map(목표비중사전).fillna(0.0)
    결과표["비중차이"] = 결과표["현재비중"] - 결과표["목표비중"]
    결과표["목표평가금액"] = 총평가금액 * 결과표["목표비중"] / 100
    결과표["리밸런싱금액"] = 결과표["목표평가금액"] - 결과표["평가금액"]

    결과표["정확계산수량"] = 결과표.apply(
        lambda 행: (행["리밸런싱금액"] / 행["현재가"])
        if pd.notna(행["현재가"]) and 행["현재가"] not in [0, None] else 0,
        axis=1,
    )
    결과표["주문참고수량"] = 결과표["정확계산수량"].round().astype(int)

    def 권장문구(행):
        if 행.get("데이터상태") != "정상":
            return "현재가 확인 후 판단"
        수량 = int(행["주문참고수량"])
        금액 = 행["리밸런싱금액"]
        if 수량 > 0:
            return f"{abs(수량):,}주 추가 매수 검토"
        if 수량 < 0:
            return f"{abs(수량):,}주 비중 축소 검토"
        if pd.notna(행["현재가"]) and pd.notna(금액) and abs(금액) < 행["현재가"] * 0.5:
            return "거의 적정 비중"
        return "소액 조정 가능"

    결과표["권장방향"] = 결과표.apply(권장문구, axis=1)
    return 결과표, 총평가금액



def 추가투자금배분계산(계산표, 목표비중사전, 추가투자금):
    결과표 = 계산표.copy()

    if 추가투자금 <= 0:
        결과표["부족금액"] = 0.0
        결과표["추천배정금액"] = 0.0
        결과표["추천매수수량"] = 0
        결과표["실사용금액"] = 0.0
        결과표["추가매수의견"] = "추가 투자금 없음"
        return 결과표, 추가투자금, 추가투자금

    정상행 = 결과표[결과표["데이터상태"] == "정상"].copy()
    if 정상행.empty:
        결과표["부족금액"] = 0.0
        결과표["추천배정금액"] = 0.0
        결과표["추천매수수량"] = 0
        결과표["실사용금액"] = 0.0
        결과표["추가매수의견"] = 결과표["데이터상태"].apply(lambda x: "현재가 확인 필요" if x != "정상" else "추가 매수 우선순위 낮음")
        return 결과표, 0.0, 추가투자금

    현재총평가금액 = 정상행["평가금액"].sum()
    목표총자산 = 현재총평가금액 + 추가투자금

    결과표["목표비중"] = 결과표["종목코드"].map(목표비중사전).fillna(0.0)
    결과표["추가투자후목표금액"] = 결과표["목표비중"] / 100 * 목표총자산
    결과표["부족금액"] = 결과표.apply(
        lambda 행: max(행["추가투자후목표금액"] - 행["평가금액"], 0)
        if 행.get("데이터상태") == "정상" and pd.notna(행.get("평가금액")) else 0.0,
        axis=1
    )

    부족금액합계 = 결과표["부족금액"].sum()
    if 부족금액합계 == 0:
        결과표["추천배정금액"] = 0.0
        결과표["추천매수수량"] = 0
        결과표["실사용금액"] = 0.0
        결과표["추가매수의견"] = 결과표["데이터상태"].apply(
            lambda x: "현재가 확인 필요" if x != "정상" else "현재 비중이 목표 수준과 유사"
        )
        return 결과표, 0.0, 추가투자금

    결과표["추천배정금액"] = 결과표["부족금액"] / 부족금액합계 * 추가투자금

    def 매수가능수량계산(행):
        if 행.get("데이터상태") != "정상":
            return 0
        if pd.isna(행["현재가"]) or 행["현재가"] in [0, None]:
            return 0
        return math.floor(행["추천배정금액"] / 행["현재가"])

    결과표["추천매수수량"] = 결과표.apply(매수가능수량계산, axis=1)
    결과표["실사용금액"] = 결과표.apply(
        lambda 행: 행["추천매수수량"] * 행["현재가"]
        if 행.get("데이터상태") == "정상" and pd.notna(행.get("현재가")) else 0.0,
        axis=1
    )

    총실사용금액 = 결과표["실사용금액"].sum()
    남는현금 = 추가투자금 - 총실사용금액

    while 남는현금 > 0:
        매수후보 = 결과표[
            (결과표["데이터상태"] == "정상") &
            (결과표["현재가"].notna()) &
            (결과표["현재가"] > 0)
        ].copy()
        if 매수후보.empty:
            break

        매수후보["남은부족금액"] = (결과표["부족금액"] - 결과표["실사용금액"]).clip(lower=0)
        매수후보 = 매수후보.sort_values(["남은부족금액", "현재비중"], ascending=[False, True])

        추가매수실행 = False
        for idx in 매수후보.index:
            현재가 = 결과표.loc[idx, "현재가"]
            남은부족금액 = max(결과표.loc[idx, "부족금액"] - 결과표.loc[idx, "실사용금액"], 0)
            if 남는현금 >= 현재가 and 남은부족금액 >= 현재가 * 0.5:
                결과표.loc[idx, "추천매수수량"] += 1
                결과표.loc[idx, "실사용금액"] += 현재가
                남는현금 -= 현재가
                추가매수실행 = True
                break

        if not 추가매수실행:
            break

    총실사용금액 = 결과표["실사용금액"].sum()
    남는현금 = 추가투자금 - 총실사용금액

    def 추가매수의견생성(행):
        if 행.get("데이터상태") != "정상":
            return "현재가 확인 필요"
        수량 = int(행["추천매수수량"])
        if 수량 > 0:
            return f"{수량:,}주 추가 매수 추천"
        if 행["추천배정금액"] > 0:
            return "배정금액은 있으나 1주 매수 금액 부족"
        return "추가 매수 우선순위 낮음"

    결과표["추가매수의견"] = 결과표.apply(추가매수의견생성, axis=1)
    return 결과표, 총실사용금액, 남는현금


def 포트폴리오위험도분석(계산포트폴리오, 목표비중사전, 개월수=6):
    분석표 = 계산포트폴리오.copy()
    총평가금액 = 분석표["평가금액"].sum()
    if 총평가금액 == 0:
        return {
            "변동성": 0.0,
            "최대낙폭": 0.0,
            "집중도": 0.0,
            "비중이탈도": 0.0,
            "위험수준": "계산 불가",
            "위험코멘트": "포트폴리오 평가금액이 없어 위험도 분석이 어렵습니다.",
        }

    일간수익률목록 = []
    가중치목록 = []
    for _, 행 in 분석표.iterrows():
        종목코드 = 행["종목코드"]
        현재비중 = 행["현재비중"] / 100
        수익률시리즈 = 일간수익률가져오기(종목코드, 개월수)
        if not 수익률시리즈.empty:
            일간수익률목록.append(수익률시리즈.rename(종목코드))
            가중치목록.append((종목코드, 현재비중))

    if not 일간수익률목록:
        return {
            "변동성": 0.0,
            "최대낙폭": 0.0,
            "집중도": float(분석표["현재비중"].max()),
            "비중이탈도": 0.0,
            "위험수준": "계산 제한",
            "위험코멘트": "충분한 과거 수익률 데이터가 없어 위험도 계산이 제한됩니다.",
        }

    수익률데이터 = pd.concat(일간수익률목록, axis=1).fillna(0)
    포트폴리오일간수익률 = pd.Series(0.0, index=수익률데이터.index, dtype=float)
    for 종목코드, 가중치 in 가중치목록:
        if 종목코드 in 수익률데이터.columns:
            포트폴리오일간수익률 += 수익률데이터[종목코드] * 가중치

    변동성 = float(포트폴리오일간수익률.std() * (252 ** 0.5) * 100)
    누적수익 = (1 + 포트폴리오일간수익률).cumprod()
    최고점 = 누적수익.cummax()
    낙폭 = (누적수익 / 최고점) - 1
    최대낙폭 = float(낙폭.min() * 100)
    집중도 = float(분석표["현재비중"].max())

    분석표["목표비중"] = 분석표["종목코드"].map(목표비중사전).fillna(0.0)
    분석표["비중절대차"] = (분석표["현재비중"] - 분석표["목표비중"]).abs()
    비중이탈도 = float(분석표["비중절대차"].sum())

    위험점수 = 0
    if 변동성 >= 35:
        위험점수 += 3
    elif 변동성 >= 20:
        위험점수 += 2
    elif 변동성 >= 10:
        위험점수 += 1

    if abs(최대낙폭) >= 25:
        위험점수 += 3
    elif abs(최대낙폭) >= 15:
        위험점수 += 2
    elif abs(최대낙폭) >= 8:
        위험점수 += 1

    if 집중도 >= 60:
        위험점수 += 3
    elif 집중도 >= 45:
        위험점수 += 2
    elif 집중도 >= 35:
        위험점수 += 1

    if 비중이탈도 >= 30:
        위험점수 += 3
    elif 비중이탈도 >= 15:
        위험점수 += 2
    elif 비중이탈도 >= 8:
        위험점수 += 1

    if 위험점수 >= 9:
        위험수준 = "높음"
        위험코멘트 = "포트폴리오 변동성과 쏠림이 큰 편이어서 보수적 점검이 필요합니다."
    elif 위험점수 >= 5:
        위험수준 = "보통"
        위험코멘트 = "포트폴리오 위험은 중간 수준이며, 비중 조정 여부를 점검할 필요가 있습니다."
    else:
        위험수준 = "낮음"
        위험코멘트 = "현재 구조는 비교적 안정적인 편입니다."

    return {
        "변동성": 변동성,
        "최대낙폭": 최대낙폭,
        "집중도": 집중도,
        "비중이탈도": 비중이탈도,
        "위험수준": 위험수준,
        "위험코멘트": 위험코멘트,
    }


def 오늘의요약생성(계산포트폴리오, 리밸런싱표, 추가배분표, 총수익률, 위험분석결과, 추가투자금):
    요약문 = []

    if 총수익률 > 5:
        요약문.append(f"포트폴리오 전체 수익률은 {총수익률:.2f}%로 매우 양호한 상태입니다.")
    elif 총수익률 > 0:
        요약문.append(f"포트폴리오 전체 수익률은 {총수익률:.2f}%로 안정적인 수익 구간입니다.")
    elif 총수익률 < 0:
        요약문.append(f"포트폴리오 전체 수익률은 {총수익률:.2f}%로 단기 손실 구간입니다.")
    else:
        요약문.append("포트폴리오 수익률은 현재 보합 수준입니다.")

    요약문.append(f"현재 포트폴리오 위험 수준은 '{위험분석결과['위험수준']}'으로 평가되며, {위험분석결과['위험코멘트']}")

    if not 계산포트폴리오.empty:
        최대비중종목 = 계산포트폴리오.sort_values("현재비중", ascending=False).iloc[0]
        최고수익종목 = 계산포트폴리오.sort_values("수익률", ascending=False).iloc[0]
        최저수익종목 = 계산포트폴리오.sort_values("수익률", ascending=True).iloc[0]
        요약문.append(f"현재 비중이 가장 큰 종목은 {최대비중종목['종목명']}이며 비중은 {최대비중종목['현재비중']:.2f}%입니다.")
        요약문.append(f"수익률이 가장 높은 종목은 {최고수익종목['종목명']}({최고수익종목['수익률']:.2f}%)입니다.")
        요약문.append(f"수익률이 가장 낮은 종목은 {최저수익종목['종목명']}({최저수익종목['수익률']:.2f}%)입니다.")

    과대 = 리밸런싱표[리밸런싱표["비중차이"] > 0]
    if not 과대.empty:
        종목 = 과대.sort_values("비중차이", ascending=False).iloc[0]
        요약문.append(f"{종목['종목명']} 비중이 목표보다 {종목['비중차이']:.2f}%p 높아 상대적으로 비중이 큰 상태입니다.")

    부족 = 리밸런싱표[리밸런싱표["비중차이"] < 0]
    if not 부족.empty:
        종목 = 부족.sort_values("비중차이").iloc[0]
        요약문.append(f"{종목['종목명']} 비중은 목표보다 {abs(종목['비중차이']):.2f}%p 낮아 보완 우선순위가 높습니다.")

    if 추가투자금 > 0:
        추천 = 추가배분표[추가배분표["추천배정금액"] > 0].sort_values("추천배정금액", ascending=False)
        if not 추천.empty:
            종목 = 추천.iloc[0]
            요약문.append(f"추가 투자금은 {종목['종목명']} 중심으로 배분하는 것이 목표 비중에 더 가까워지는 전략입니다.")

    return 요약문


# -----------------------------------
# 그래프/분석 함수
# -----------------------------------
def 가격그래프(데이터, 제목):
    x값 = pd.to_datetime(데이터.index)
    그림 = go.Figure()

    그림.add_trace(
        go.Scatter(
            x=x값, y=데이터["종가"], mode="lines", name="종가",
            line=dict(color="#7cc4ff", width=2.2),
            hovertemplate="종가: %{y:,.0f}<extra></extra>"
        )
    )

    이동평균설정 = [
        ("5일평균", "5일 평균", "#f59e0b"),
        ("20일평균", "20일 평균", "#3b82f6"),
        ("60일평균", "60일 평균", "#6b8f5a"),
        ("120일평균", "120일 평균", "#34d399"),
    ]
    for 컬럼, 이름, 색상 in 이동평균설정:
        if 컬럼 in 데이터.columns:
            그림.add_trace(
                go.Scatter(
                    x=x값, y=데이터[컬럼], mode="lines", name=이름,
                    line=dict(color=색상, width=2),
                    hovertemplate=f"{이름}: %{{y:,.0f}}<extra></extra>"
                )
            )

    그림.update_layout(
        title=제목,
        height=모바일차트높이(460, 340),
        margin=dict(l=20, r=20, t=55, b=20),
        legend=dict(orientation="v", yanchor="top", y=0.98, xanchor="left", x=1.01),
        hovermode="x unified",
        xaxis_title="날짜",
        yaxis_title="가격",
        xaxis=dict(showgrid=True, gridcolor="rgba(148,163,184,0.15)"),
        yaxis=dict(showgrid=True, gridcolor="rgba(148,163,184,0.15)", tickformat=","),
    )
    return 그림


def 캔들차트그래프(데이터, 제목):
    x값 = pd.to_datetime(데이터.index)
    표시데이터 = 데이터.copy()
    그림 = make_subplots(
        rows=2, cols=1, shared_xaxes=True, vertical_spacing=0.03,
        row_heights=[0.74, 0.26], specs=[[{"secondary_y": False}], [{"secondary_y": False}]]
    )

    그림.add_trace(go.Candlestick(
        x=x값,
        open=표시데이터["시가"],
        high=표시데이터["고가"],
        low=표시데이터["저가"],
        close=표시데이터["종가"],
        name="캔들",
        showlegend=False,
        increasing_line_color="#e05a63",
        decreasing_line_color="#4f86d9",
        increasing_fillcolor="#e05a63",
        decreasing_fillcolor="#4f86d9",
    ), row=1, col=1)

    이동평균설정 = []
    if len(표시데이터) >= 3:
        이동평균설정.append(("5일평균", "5일", "#f59e0b"))
    if len(표시데이터) >= 8:
        이동평균설정.append(("20일평균", "20일", "#3b82f6"))
    if len(표시데이터) >= 20:
        이동평균설정.append(("60일평균", "60일", "#6b8f5a"))
    if len(표시데이터) >= 40:
        이동평균설정.append(("120일평균", "120일", "#34d399"))
    for 컬럼, 이름, 색상 in 이동평균설정:
        if 컬럼 in 표시데이터.columns:
            그림.add_trace(
                go.Scatter(
                    x=x값,
                    y=표시데이터[컬럼],
                    mode="lines",
                    name=이름,
                    showlegend=False,
                    line=dict(color=색상, width=2),
                    hovertemplate=f"{이름}선: %{{y:,.0f}}<extra></extra>"
                ),
                row=1, col=1
            )

    if "거래량" in 표시데이터.columns:
        거래량색 = ["#e05a63" if 종가 >= 시가 else "#4f86d9" for 종가, 시가 in zip(표시데이터["종가"], 표시데이터["시가"])]
        그림.add_trace(
            go.Bar(
                x=x값,
                y=표시데이터["거래량"],
                name="거래량",
                showlegend=False,
                marker_color=거래량색,
                opacity=0.9,
                hovertemplate="거래량: %{y:,.0f}<extra></extra>"
            ),
            row=2, col=1
        )

    최고행 = 표시데이터["고가"].idxmax()
    최저행 = 표시데이터["저가"].idxmin()
    최고값 = float(표시데이터.loc[최고행, "고가"])
    최저값 = float(표시데이터.loc[최저행, "저가"])
    최저대비상승률 = ((최고값 - 최저값) / 최저값 * 100) if 최저값 else 0
    최고대비하락률 = ((표시데이터.iloc[-1]["종가"] - 최고값) / 최고값 * 100) if 최고값 else 0
    날짜포맷 = '%y.%m.%d' if len(표시데이터) <= 8 else ('%y.%m' if len(표시데이터) <= 18 else '%Y')

    그림.add_annotation(
        x=pd.to_datetime(최고행), y=최고값,
        text=f"↗ {최고값:,.0f}({pd.to_datetime(최고행).strftime(날짜포맷)}), {최고대비하락률:+.2f}%",
        showarrow=False, font=dict(color="#e05a63", size=12), xanchor="left", yanchor="bottom", row=1, col=1,
        bgcolor="rgba(255,255,255,0.85)"
    )
    그림.add_annotation(
        x=pd.to_datetime(최저행), y=최저값,
        text=f"↘ {최저값:,.0f}({pd.to_datetime(최저행).strftime(날짜포맷)}), {최저대비상승률:+.2f}%",
        showarrow=False, font=dict(color="#4f86d9", size=12), xanchor="left", yanchor="top", row=1, col=1,
        bgcolor="rgba(255,255,255,0.85)"
    )

    범례항목 = ["캔들"] + [이름 for _, 이름, _ in 이동평균설정] + (["거래량"] if "거래량" in 표시데이터.columns else [])
    범례문구 = "범례: " + " · ".join(범례항목)

    그림.update_layout(
        title=dict(text=제목, x=0.01, xanchor="left", y=0.98, yanchor="top", font=dict(size=16)),
        height=660,
        margin=dict(l=30, r=30, t=110, b=56),
        xaxis_rangeslider_visible=False,
        showlegend=False,
        bargap=0.14,
        hovermode="x unified",
        plot_bgcolor="#ffffff",
        paper_bgcolor="#ffffff",
        font=dict(color="#111827"),
        autosize=True,
        annotations=list(그림.layout.annotations) + [
            dict(
                text=범례문구,
                x=0.01,
                y=1.08,
                xref="paper",
                yref="paper",
                showarrow=False,
                align="left",
                font=dict(size=12, color="#374151"),
                bgcolor="rgba(255,255,255,0.92)",
                bordercolor="#d1d5db",
                borderwidth=1,
                borderpad=4,
            )
        ]
    )
    그림.update_yaxes(
        side="right", tickformat=",", row=1, col=1, showgrid=True, gridcolor="#d1d5db", zeroline=False,
        tickfont=dict(color="#374151", size=12), title_font=dict(color="#111827"), automargin=True
    )
    그림.update_yaxes(
        side="right", tickformat="~s", row=2, col=1, showgrid=True, gridcolor="#e5e7eb", zeroline=False,
        tickfont=dict(color="#374151", size=12), title_font=dict(color="#111827"), automargin=True
    )
    그림.update_xaxes(
        showgrid=True, gridcolor="#e5e7eb", tickfont=dict(color="#374151", size=12),
        automargin=True, tickangle=0, showline=False, zeroline=False
    )
    return 그림


def 클릭캔들행가져오기(데이터, clicked_points):
    if 데이터.empty:
        return None, None

    기본날짜 = 데이터.index[-1]
    기본행 = 데이터.iloc[-1]

    if not clicked_points:
        return 기본날짜, 기본행

    try:
        클릭x = pd.to_datetime(clicked_points[0].get("x"))
        인덱스 = pd.to_datetime(pd.Index(데이터.index))
        위치 = (인덱스 - 클릭x).asi8
        절대차이 = pd.Series(위치, index=데이터.index).abs()
        선택날짜 = 절대차이.idxmin()
        return 선택날짜, 데이터.loc[선택날짜]
    except Exception:
        return 기본날짜, 기본행


def 캔들분석결과가져오기(데이터, 선택날짜, 선택행):
    if 선택행 is None or 데이터.empty:
        return {
            "캔들유형": "분석 불가",
            "방향": "중립",
            "설명": "캔들 데이터를 선택하지 못했습니다.",
            "상세": [],
            "체크표": pd.DataFrame(),
        }

    시가 = float(선택행["시가"])
    고가 = float(선택행["고가"])
    저가 = float(선택행["저가"])
    종가 = float(선택행["종가"])
    거래량 = float(선택행["거래량"]) if pd.notna(선택행.get("거래량")) else None

    전체범위 = max(고가 - 저가, 1e-9)
    몸통 = abs(종가 - 시가)
    윗꼬리 = max(고가 - max(시가, 종가), 0)
    아랫꼬리 = max(min(시가, 종가) - 저가, 0)

    몸통비율 = 몸통 / 전체범위 * 100
    윗꼬리비율 = 윗꼬리 / 전체범위 * 100
    아랫꼬리비율 = 아랫꼬리 / 전체범위 * 100

    if 몸통비율 <= 15:
        캔들유형 = "도지형"
    elif 종가 > 시가 and 아랫꼬리비율 >= 35 and 몸통비율 <= 45:
        캔들유형 = "망치형 가능성"
    elif 종가 < 시가 and 윗꼬리비율 >= 35 and 몸통비율 <= 45:
        캔들유형 = "슈팅스타 가능성"
    elif 종가 > 시가:
        캔들유형 = "양봉"
    else:
        캔들유형 = "음봉"

    방향 = "상승 우세" if 종가 > 시가 else "하락 우세" if 종가 < 시가 else "중립"

    보정선택날짜 = 인덱스기준가까운날짜찾기(데이터, 선택날짜)
    if 보정선택날짜 is None:
        보정선택날짜 = 데이터.index[-1]
    위치 = 데이터.index.get_loc(보정선택날짜)
    최근20 = 데이터.iloc[max(0, 위치 - 19): 위치 + 1]
    최근거래량20 = 최근20["거래량"].mean() if "거래량" in 최근20.columns and len(최근20) > 0 else None

    최근20고가 = 최근20["고가"].max() if len(최근20) > 0 else 고가
    최근20저가 = 최근20["저가"].min() if len(최근20) > 0 else 저가

    돌파판정 = "중립"
    if 종가 >= 최근20고가:
        돌파판정 = "20일 고점 돌파 시도"
    elif 종가 <= 최근20저가:
        돌파판정 = "20일 저점 이탈 경계"

    거래량판정 = "판정 제한"
    if 거래량 is not None and 최근거래량20 not in [None, 0] and pd.notna(최근거래량20):
        배수 = 거래량 / 최근거래량20
        if 배수 >= 1.5:
            거래량판정 = f"평균 대비 {배수:.2f}배 급증"
        elif 배수 >= 1.0:
            거래량판정 = f"평균 대비 {배수:.2f}배로 보통 이상"
        else:
            거래량판정 = f"평균 대비 {배수:.2f}배로 차분"

    if 캔들유형 == "망치형 가능성":
        설명 = "하단에서 매수 유입이 들어오며 종가를 끌어올린 형태로 해석할 수 있습니다. 다음 1~2일 안에 고점 돌파가 이어지는지 확인이 중요합니다."
    elif 캔들유형 == "슈팅스타 가능성":
        설명 = "상단에서 매도 압력이 강하게 나온 흔적으로 볼 수 있습니다. 다음 봉에서 저점 이탈이 나오면 단기 조정 신호가 강화됩니다."
    elif 캔들유형 == "도지형":
        설명 = "매수와 매도가 팽팽하게 맞선 날입니다. 추세 전환의 단서가 될 수 있으므로 다음 봉 방향 확인이 중요합니다."
    elif 캔들유형 == "양봉":
        설명 = "당일 종가가 시가보다 높아 매수 우위가 확인된 날입니다. 다만 윗꼬리 길이에 따라 상단 저항도 함께 점검해야 합니다."
    else:
        설명 = "당일 종가가 시가보다 낮아 매도 우위가 나타난 날입니다. 아랫꼬리가 길면 저가 매수 유입도 일부 있었다고 볼 수 있습니다."

    체크표 = pd.DataFrame([
        {"항목": "선택 날짜", "현재": str(pd.to_datetime(보정선택날짜 if "보정선택날짜" in locals() else 선택날짜).date()), "기준": "클릭한 캔들", "판정": "선택됨"},
        {"항목": "시가/종가", "현재": f"{시가:,.0f} / {종가:,.0f}", "기준": "종가 > 시가면 양봉", "판정": 방향},
        {"항목": "고가/저가", "현재": f"{고가:,.0f} / {저가:,.0f}", "기준": "당일 변동폭", "판정": f"{전체범위:,.0f}"},
        {"항목": "몸통 비율", "현재": f"{몸통비율:.1f}%", "기준": "15% 이하면 도지형", "판정": 캔들유형},
        {"항목": "윗꼬리/아랫꼬리", "현재": f"{윗꼬리비율:.1f}% / {아랫꼬리비율:.1f}%", "기준": "꼬리 길이 비교", "판정": 돌파판정},
        {"항목": "거래량", "현재": 숫자표시(거래량), "기준": "20일 평균 대비", "판정": 거래량판정},
    ])

    상세 = [
        f"선택한 캔들은 **{캔들유형}**으로 볼 수 있고, 당일 방향은 **{방향}**입니다.",
        f"몸통 비율은 {몸통비율:.1f}%이며 윗꼬리 {윗꼬리비율:.1f}%, 아랫꼬리 {아랫꼬리비율:.1f}%입니다.",
        f"가격 위치는 **{돌파판정}**으로 해석됩니다.",
        f"거래량은 **{거래량판정}**입니다.",
    ]

    return {
        "캔들유형": 캔들유형,
        "방향": 방향,
        "설명": 설명,
        "상세": 상세,
        "체크표": 체크표,
    }


def 비중그래프(계산표):
    작업 = 계산표.copy()
    if 작업 is None or 작업.empty:
        그림 = go.Figure()
        그림.update_layout(height=360, margin=dict(l=10, r=10, t=40, b=10), title="현재 포트폴리오 비중")
        return 그림

    작업 = 작업.copy()
    작업["평가금액"] = pd.to_numeric(작업["평가금액"], errors="coerce").fillna(0)
    작업 = 작업[작업["평가금액"] > 0].copy()
    작업 = 작업.sort_values("평가금액", ascending=False)

    그림 = go.Figure(
        go.Pie(
            labels=작업["종목명"],
            values=작업["평가금액"],
            hole=0.52,
            sort=False,
            direction="clockwise",
            textinfo="percent",
            textposition="inside",
            insidetextorientation="auto",
            hovertemplate="%{label}<br>평가금액: %{value:,.0f}원<br>비중: %{percent}<extra></extra>",
        )
    )
    그림.update_traces(marker=dict(line=dict(color="#0b1220", width=1.2)))
    그림.update_layout(
        title=dict(text="현재 포트폴리오 비중", x=0.02, xanchor="left", y=0.97),
        height=430,
        margin=dict(l=10, r=10, t=52, b=10),
        legend=dict(orientation="v", yanchor="top", y=0.98, xanchor="left", x=1.02, font=dict(size=13)),
    )
    return 그림


def 목표비중비교그래프(리밸런싱표):
    그림 = go.Figure()
    그림.add_trace(go.Bar(x=리밸런싱표["종목명"], y=리밸런싱표["현재비중"], name="현재 비중"))
    그림.add_trace(go.Bar(x=리밸런싱표["종목명"], y=리밸런싱표["목표비중"], name="목표 비중"))
    그림.update_layout(title="현재 비중 vs 목표 비중", xaxis_title="종목", yaxis_title="비중(%)", barmode="group", height=420)
    return 그림


def 지표라인그래프(df, 값열, 제목):
    그림 = go.Figure()
    그림.add_trace(go.Scatter(x=pd.to_datetime(df.index), y=df[값열], mode="lines", name=제목))
    그림.update_layout(title=제목, xaxis_title="날짜", yaxis_title=값열, height=320, margin=dict(l=10, r=10, t=50, b=10))
    return 그림


def 차트분석문구(자산명, 데이터):
    if 데이터.empty or len(데이터) < 20:
        부족문구 = "현재 데이터 길이만으로는 신뢰도 있는 해석을 제시하기 어렵습니다. 최소 20거래일 이상 확보한 뒤 추세와 RSI를 함께 보시는 편이 좋습니다."
        return {
            "ChatGPT": {
                "한줄요약": "데이터가 충분하지 않아 차트 해석이 제한됩니다.",
                "현재신호": 부족문구,
                "근거": ["시계열 길이가 짧음", "이동평균선 비교 제한", "RSI 해석 신뢰도 낮음"],
                "리스크": ["성급한 추세 판단 가능성", "단기 변동성 과대해석 가능성"],
                "보유자관점": "기존 보유자라면 추가 대응보다 데이터 축적을 먼저 확인하는 편이 좋습니다.",
                "신규진입관점": "신규 진입은 조금 더 긴 시계열 확보 후 검토하는 것이 바람직합니다.",
            },
            "Gemini": {
                "한줄요약": "데이터 부족으로 빠른 판단의 신뢰도가 낮습니다.",
                "현재신호": 부족문구,
                "근거": ["최근 데이터 부족", "모멘텀 확인 제한", "추세 지속성 판단 어려움"],
                "리스크": ["짧은 데이터에 따른 오판 가능성", "반등·하락 신호 왜곡 가능성"],
                "보유자관점": "보유 중이라면 섣부른 비중 조정보다 추세 확인이 우선입니다.",
                "신규진입관점": "신규 진입은 확인 가능한 데이터가 더 쌓인 뒤가 좋습니다.",
            },
            "Claude": {
                "한줄요약": "현재는 해석보다 관찰이 우선인 구간입니다.",
                "현재신호": 부족문구,
                "근거": ["데이터 길이 부족", "추세·RSI 동시 검증 어려움", "신호 지속성 확인 제한"],
                "리스크": ["짧은 구간을 추세로 오해할 수 있음", "가격 신호만 보고 대응할 위험"],
                "보유자관점": "기존 보유자라면 성급한 판단보다 추가 데이터 확인이 더 신중합니다.",
                "신규진입관점": "신규 진입은 최소 20거래일 이상 확보 후 검토하는 편이 적절합니다.",
            },
        }

    최신 = 데이터.iloc[-1]
    종가 = float(최신["종가"])
    ma20 = 최신.get("20일평균")
    ma60 = 최신.get("60일평균")
    rsi = 최신.get("RSI(14)")

    추세 = "중립"
    if pd.notna(ma20) and pd.notna(ma60):
        if 종가 > ma20 > ma60:
            추세 = "상승 추세"
        elif 종가 < ma20 < ma60:
            추세 = "하락 추세"

    rsi해석 = "중립권"
    if pd.notna(rsi):
        if rsi >= 70:
            rsi해석 = "과열권"
        elif rsi <= 30:
            rsi해석 = "침체권"

    변동률20일 = None
    if len(데이터) >= 21:
        기준가 = 데이터.iloc[-21]["종가"]
        if 기준가 not in [0, None]:
            변동률20일 = (종가 / 기준가 - 1) * 100

    변동문구 = f"최근 20거래일 수익률은 {변동률20일:.2f}% 수준입니다." if 변동률20일 is not None else "최근 20거래일 변화율 계산은 제한됩니다."

    지지저항문구 = "20일선 부근 공방 여부를 확인할 필요가 있습니다."
    if pd.notna(ma20):
        if 종가 > ma20:
            지지저항문구 = "단기적으로는 20일선 위에 있어 지지력은 아직 유지되는 편입니다."
        else:
            지지저항문구 = "20일선 아래에 있어 단기 반등이 나와도 저항 확인이 필요합니다."

    거래량문구 = "거래량 정보가 충분치 않다면 가격 신호만으로 성급히 판단하지 않는 편이 좋습니다."

    방향성문구 = "양호" if 추세 == "상승 추세" else "둔화 또는 중립"
    chatgpt_signal = f"{자산명}은 현재 {추세}로 해석됩니다. 종가와 20일·60일 이동평균선의 배열을 보면 방향성은 {방향성문구}합니다. RSI는 {rsi해석}이며, {지지저항문구}"
    gemini_signal = f"기술적으로 보면 {자산명}의 핵심 포인트는 이동평균선 정렬과 RSI입니다. 현재 판독은 {추세}, RSI는 {rsi해석}입니다. {변동문구}"
    claude_signal = f"종가와 이동평균선의 위치를 기준으로 보면 {추세}에 가깝고, RSI는 {rsi해석} 구간으로 읽힙니다. {지지저항문구}"

    return {
        "ChatGPT": {
            "한줄요약": f"{자산명}은 현재 {추세}로 보되, 단기 추격 대응보다 지지 확인이 우선입니다.",
            "현재신호": chatgpt_signal,
            "근거": [
                "종가와 20일·60일 이동평균선 배열",
                f"RSI는 {rsi해석} 수준",
                변동문구,
            ],
            "리스크": [
                "20일선 이탈이 이어지면 단기 약세가 재확대될 수 있음",
                "거래량 확인 없이 가격만 보고 대응하면 오판 가능성",
            ],
            "보유자관점": "기존 보유자라면 단기 추격 매수보다 지지선 유지 여부를 먼저 확인하는 접근이 더 안정적입니다.",
            "신규진입관점": "신규 진입은 한 번에 들어가기보다 분할 접근이 더 적절합니다.",
        },
        "Gemini": {
            "한줄요약": f"{자산명}은 현재 빠른 판단 기준으로 {추세} 쪽 신호가 우세합니다.",
            "현재신호": gemini_signal,
            "근거": [
                "이동평균선 정렬 상태",
                f"RSI 판독 결과는 {rsi해석}",
                "최근 20거래일 변동률 반영",
            ],
            "리스크": [
                "20일선 회복 실패 시 반등 지속성이 약해질 수 있음",
                "최근 변동성이 큰 구간이면 신호 왜곡 가능성 존재",
            ],
            "보유자관점": "보유 중이라면 단기 추세 유지 여부와 20일선 회복·이탈을 함께 보는 것이 좋습니다.",
            "신규진입관점": "신규 진입은 추세 확인 이후 소규모로 시작하는 편이 유리합니다.",
        },
        "Claude": {
            "한줄요약": f"{자산명}은 현재 {추세} 성격이 우세하지만, 이를 곧바로 강한 추세 전환으로 단정하기에는 아직 확인할 요소가 남아 있습니다.",
            "현재신호": claude_signal,
            "근거": [
                "종가의 이동평균선 대비 위치",
                f"RSI의 과열·침체 여부는 {rsi해석}",
                변동문구,
            ],
            "리스크": [
                "단기 반등이 나오더라도 20일선 안착이 실패하면 다시 변동성이 커질 수 있음",
                거래량문구,
            ],
            "보유자관점": "기존 보유자라면 성급한 비중 확대보다 지지선 유지 여부를 먼저 확인하는 접근이 더 신중합니다.",
            "신규진입관점": "신규 진입은 추세 확인 이후 분할 접근이 더 적절합니다.",
        },
    }


def 분석카드표시(분석데이터):
    st.markdown(f"### 한줄 요약")
    st.info(분석데이터["한줄요약"])

    col1, col2 = st.columns([1.6, 1])
    with col1:
        st.markdown("### 현재 신호")
        st.write(분석데이터["현재신호"])

        st.markdown("### 근거")
        for 항목 in 분석데이터["근거"]:
            st.markdown(f"- {항목}")

    with col2:
        st.markdown("### 주의할 리스크")
        for 항목 in 분석데이터["리스크"]:
            st.markdown(f"- {항목}")

    st.markdown("### 보유자 관점")
    st.write(분석데이터["보유자관점"])

    st.markdown("### 신규 진입 관점")
    st.write(분석데이터["신규진입관점"])


def 종목거래이력표생성(거래df, 종목코드=None):
    작업 = 거래이력정규화(거래df)
    if 작업.empty:
        return pd.DataFrame(columns=["거래일자", "종목코드", "종목명", "거래구분", "거래수량", "거래단가", "거래금액", "누적보유수량", "운용사", "비고"])

    if "_입력원본순서" not in 작업.columns:
        작업["_입력원본순서"] = range(len(작업))
    작업["_원본순서"] = pd.to_numeric(작업["_입력원본순서"], errors="coerce").fillna(pd.Series(range(len(작업)), index=작업.index)).astype(int)

    if 종목코드:
        코드 = str(종목코드).zfill(6)
        작업 = 작업[작업["종목코드"].astype(str).str.zfill(6) == 코드].copy()
        if 작업.empty:
            return pd.DataFrame(columns=["거래일자", "종목코드", "종목명", "거래구분", "거래수량", "거래단가", "거래금액", "누적보유수량", "운용사", "비고"])

    작업["거래일자"] = pd.to_datetime(작업["거래일자"], errors="coerce")
    작업["거래수량"] = pd.to_numeric(작업["거래수량"], errors="coerce").fillna(0)
    작업["거래단가"] = pd.to_numeric(작업["거래단가"], errors="coerce").fillna(0)

    # 누적보유수량 계산도 입력 원본 순서를 기준으로 수행
    작업 = 작업.sort_values(["_원본순서"], ascending=[True], na_position="last", kind="stable").reset_index(drop=True)
    작업["거래금액"] = 작업["거래수량"] * 작업["거래단가"]
    작업["signed_qty"] = 작업["거래수량"].where(작업["거래구분"] == "매수", -작업["거래수량"])
    작업["누적보유수량"] = 작업.groupby("종목코드", sort=False)["signed_qty"].cumsum()
    작업["거래일자"] = 작업["거래일자"].dt.date
    return 작업[["거래일자", "종목코드", "종목명", "거래구분", "거래수량", "거래단가", "거래금액", "누적보유수량", "운용사", "비고"]]


def 거래기록표시용서식(df):
    if df is None or df.empty:
        return df
    return df.style.format({
        "거래수량": 안전정수포맷,
        "거래단가": 안전정수포맷,
        "거래금액": 안전정수포맷,
        "누적보유수량": 안전정수포맷,
    }).map(lambda v: "color: #dc2626; font-weight: 520;" if v == "매수" else "color: #2563eb; font-weight: 520;", subset=["거래구분"])


def 보유종목상세해설생성(종목명, 가격데이터, 포트폴리오행=None):
    if 가격데이터 is None or 가격데이터.empty:
        return []

    최신 = 가격데이터.iloc[-1]
    이전 = 가격데이터.iloc[-2] if len(가격데이터) >= 2 else 최신
    종가 = float(최신["종가"])
    시가 = float(최신.get("시가", 종가))
    고가 = float(최신.get("고가", 종가))
    저가 = float(최신.get("저가", 종가))
    거래량 = float(최신.get("거래량", 0)) if pd.notna(최신.get("거래량")) else 0
    ma5 = 최신.get("5일평균")
    ma20 = 최신.get("20일평균")
    ma60 = 최신.get("60일평균")
    rsi = 최신.get("RSI(14)")

    최근20 = 가격데이터.tail(20).copy()
    최근20고가 = float(최근20["고가"].max()) if not 최근20.empty else 고가
    최근20저가 = float(최근20["저가"].min()) if not 최근20.empty else 저가
    평균거래량20 = float(최근20["거래량"].mean()) if "거래량" in 최근20.columns and len(최근20) > 0 else 0
    전일대비 = 종가 - float(이전["종가"])
    전일등락률 = (전일대비 / float(이전["종가"]) * 100) if float(이전["종가"]) != 0 else 0
    고점대비 = ((종가 / 최근20고가) - 1) * 100 if 최근20고가 else 0
    저점대비 = ((종가 / 최근20저가) - 1) * 100 if 최근20저가 else 0
    거래량배수 = (거래량 / 평균거래량20) if 평균거래량20 not in [0, None] else None

    if pd.notna(ma5) and pd.notna(ma20) and pd.notna(ma60):
        if 종가 > ma5 > ma20 > ma60:
            배열판정 = "정배열에 가까운 상승 구조"
        elif 종가 < ma5 < ma20 < ma60:
            배열판정 = "역배열에 가까운 하락 구조"
        elif 종가 >= ma20 and ma20 >= ma60:
            배열판정 = "중기 기준으로는 버티는 구조"
        else:
            배열판정 = "추세 전환 확인이 더 필요한 혼조 구조"
    else:
        배열판정 = "이동평균선 데이터가 충분하지 않아 배열 판정이 제한됩니다"

    if pd.isna(rsi):
        rsi판정 = "판정 제한"
    elif rsi >= 70:
        rsi판정 = "단기 과열권"
    elif rsi <= 30:
        rsi판정 = "단기 침체권"
    elif rsi < 45:
        rsi판정 = "약세 우위"
    else:
        rsi판정 = "중립 또는 완만한 회복 구간"

    최신캔들 = 캔들분석결과가져오기(가격데이터, 가격데이터.index[-1], 최신)
    문장목록 = [
        f"최근 종가는 **{종가:,.0f}원**이며 전일 대비 **{전일대비:,.0f}원 ({전일등락률:+.2f}%)** 움직였습니다.",
        f"최신 캔들은 **{최신캔들['캔들유형']} / {최신캔들['방향']}**으로 해석할 수 있고, 당일 가격 범위는 **{저가:,.0f}원 ~ {고가:,.0f}원**입니다.",
        f"이동평균선 기준으로는 **{배열판정}**입니다. 5일선 {숫자표시(ma5, 2)}, 20일선 {숫자표시(ma20, 2)}, 60일선 {숫자표시(ma60, 2)} 수준입니다.",
        f"최근 20거래일 기준 고점은 **{최근20고가:,.0f}원**, 저점은 **{최근20저가:,.0f}원**이며 현재가는 고점 대비 **{고점대비:.2f}%**, 저점 대비 **{저점대비:.2f}%** 위치입니다.",
        f"RSI(14)는 **{숫자표시(rsi, 2)}**로 **{rsi판정}**으로 읽을 수 있습니다.",
    ]

    if 거래량배수 is not None:
        문장목록.append(f"최근 거래량은 **{숫자표시(거래량)}주**로 20일 평균 대비 **{거래량배수:.2f}배** 수준입니다.")

    if 포트폴리오행 is not None and len(포트폴리오행) > 0:
        행 = 포트폴리오행.iloc[0]
        문장목록.append(
            f"Jone님 계좌 기준 현재 **보유수량 {숫자표시(행.get('보유수량'))}주**, 매입 평균단가는 **{금액표시(행.get('매입평균단가'))}**, 평가손익은 **{손익문자열(행.get('평가손익'))}원**입니다."
        )

    return 문장목록


def 미니차트그래프(데이터, 제목):
    그림 = go.Figure()
    그림.add_trace(go.Scatter(x=pd.to_datetime(데이터.index), y=데이터["종가"], mode="lines", name=제목))
    그림.update_layout(height=180, margin=dict(l=10, r=10, t=25, b=10), showlegend=False)
    그림.update_xaxes(visible=False)
    그림.update_yaxes(visible=False)
    return 그림


def 신호판정계산(데이터):
    if 데이터.empty or len(데이터) < 20:
        return {
            "종합신호": "데이터 부족",
            "색상": "#6b7280",
            "추세점수": 0,
            "추세설명": "데이터 부족",
            "모멘텀": "판정 제한",
            "RSI판정": "판정 제한",
            "실행의견": "데이터를 더 확인하세요.",
            "체크표": pd.DataFrame([
                {"항목": "데이터 길이", "현재": len(데이터), "기준": "20거래일 이상", "판정": "부족"}
            ]),
        }

    최신 = 데이터.iloc[-1]
    종가 = float(최신["종가"])
    ma20 = 최신.get("20일평균")
    ma60 = 최신.get("60일평균")
    rsi = 최신.get("RSI(14)")
    최근20수익률 = None
    if len(데이터) >= 21 and 데이터.iloc[-21]["종가"] not in [0, None]:
        최근20수익률 = (종가 / float(데이터.iloc[-21]["종가"]) - 1) * 100

    score = 0
    if pd.notna(ma20) and 종가 > ma20:
        score += 1
    if pd.notna(ma60) and 종가 > ma60:
        score += 1
    if pd.notna(ma20) and pd.notna(ma60) and ma20 > ma60:
        score += 1

    if score >= 3:
        추세설명 = "상승 배열"
    elif score == 2:
        추세설명 = "완만한 상승"
    elif score == 1:
        추세설명 = "중립"
    else:
        추세설명 = "약세"

    if 최근20수익률 is None:
        모멘텀 = "판정 제한"
    elif 최근20수익률 >= 8:
        모멘텀 = "강함"
    elif 최근20수익률 >= 0:
        모멘텀 = "보통"
    else:
        모멘텀 = "약함"

    if pd.isna(rsi):
        rsi판정 = "판정 제한"
    elif rsi >= 70:
        rsi판정 = "과열"
    elif rsi <= 35:
        rsi판정 = "저점권 관심"
    else:
        rsi판정 = "중립"

    if score >= 3 and (pd.isna(rsi) or 40 <= rsi <= 68):
        종합신호, 색상, 실행의견 = "매수 관심", "#16a34a", "추세가 양호해 분할매수 관심 구간으로 볼 수 있습니다."
    elif score >= 2 and pd.notna(rsi) and rsi > 68:
        종합신호, 색상, 실행의견 = "보유", "#2563eb", "상승 흐름은 유지되지만 단기 과열 가능성이 있어 추격 매수는 신중한 편이 좋습니다."
    elif score <= 1 and pd.notna(rsi) and rsi <= 35:
        종합신호, 색상, 실행의견 = "관찰", "#f59e0b", "낙폭 이후 반등 후보 구간일 수 있어 지지 확인 후 접근이 좋습니다."
    else:
        종합신호, 색상, 실행의견 = "관망", "#6b7280", "추세와 모멘텀이 애매해 방향 확인이 우선입니다."

    체크표 = pd.DataFrame([
        {"항목": "종가 vs 20일선", "현재": f"{종가:,.0f} / {ma20:,.0f}" if pd.notna(ma20) else f"{종가:,.0f} / -", "기준": "종가 > 20일선", "판정": "양호" if pd.notna(ma20) and 종가 > ma20 else "보통"},
        {"항목": "종가 vs 60일선", "현재": f"{종가:,.0f} / {ma60:,.0f}" if pd.notna(ma60) else f"{종가:,.0f} / -", "기준": "종가 > 60일선", "판정": "양호" if pd.notna(ma60) and 종가 > ma60 else "보통"},
        {"항목": "20일선 vs 60일선", "현재": f"{ma20:,.0f} / {ma60:,.0f}" if pd.notna(ma20) and pd.notna(ma60) else "-", "기준": "20일선 > 60일선", "판정": "양호" if pd.notna(ma20) and pd.notna(ma60) and ma20 > ma60 else "보통"},
        {"항목": "RSI(14)", "현재": f"{rsi:.2f}" if pd.notna(rsi) else "-", "기준": "40~70 중립권", "판정": rsi판정},
        {"항목": "최근 20거래일 수익률", "현재": f"{최근20수익률:.2f}%" if 최근20수익률 is not None else "-", "기준": "> 0%", "판정": 모멘텀},
    ])

    return {
        "종합신호": 종합신호,
        "색상": 색상,
        "추세점수": score,
        "추세설명": 추세설명,
        "모멘텀": 모멘텀,
        "RSI판정": rsi판정,
        "실행의견": 실행의견,
        "체크표": 체크표,
    }


def 기술분석진단계산(데이터):
    if 데이터 is None or 데이터.empty:
        return {
            "요약문장": ["가격 데이터가 없어 기술적 분석을 계산할 수 없습니다."],
            "핵심표": pd.DataFrame(columns=["항목", "값", "판정"]),
            "레벨표": pd.DataFrame(columns=["항목", "가격", "설명"]),
            "추세배열": "판정 불가",
            "지지": None,
            "저항": None,
        }

    최신 = 데이터.iloc[-1]
    종가 = float(최신.get("종가", 0) or 0)
    ma5 = float(최신.get("5일평균", 0) or 0)
    ma20 = float(최신.get("20일평균", 0) or 0)
    ma60 = float(최신.get("60일평균", 0) or 0)
    ma120 = float(최신.get("120일평균", 0) or 0)
    rsi = 최신.get("RSI(14)")
    거래량 = float(최신.get("거래량", 0) or 0)

    최근20 = 데이터.tail(min(len(데이터), 20)).copy()
    최근60 = 데이터.tail(min(len(데이터), 60)).copy()
    평균거래량20 = float(최근20["거래량"].mean()) if "거래량" in 최근20.columns and not 최근20.empty else 0
    거래량배수 = (거래량 / 평균거래량20) if 평균거래량20 not in [0, None] else None
    변동성20 = float(최근20["종가"].pct_change().std() * 100) if len(최근20) >= 2 else None

    지지 = float(최근20["저가"].min()) if not 최근20.empty else None
    저항 = float(최근20["고가"].max()) if not 최근20.empty else None
    장기지지 = float(최근60["저가"].min()) if not 최근60.empty else None
    장기저항 = float(최근60["고가"].max()) if not 최근60.empty else None

    if 종가 > ma20 > ma60 > 0:
        추세배열 = "상승 배열"
    elif 종가 > ma20 and ma20 > 0:
        추세배열 = "단기 우위"
    elif 종가 < ma20 < ma60 and ma20 > 0 and ma60 > 0:
        추세배열 = "하락 배열"
    else:
        추세배열 = "혼조"

    if pd.isna(rsi):
        rsi판정 = "판정 제한"
    elif rsi >= 70:
        rsi판정 = "과열"
    elif rsi <= 30:
        rsi판정 = "강한 침체"
    elif rsi <= 40:
        rsi판정 = "저점권 관심"
    else:
        rsi판정 = "중립"

    if 거래량배수 is None:
        거래량판정 = "판정 제한"
    elif 거래량배수 >= 1.8:
        거래량판정 = "강한 유입"
    elif 거래량배수 >= 1.2:
        거래량판정 = "증가"
    elif 거래량배수 >= 0.8:
        거래량판정 = "보통"
    else:
        거래량판정 = "감소"

    지지괴리 = ((종가 / 지지) - 1) * 100 if 지지 not in [0, None] else None
    저항괴리 = ((저항 / 종가) - 1) * 100 if 저항 not in [0, None] and 종가 != 0 else None
    ma20괴리 = ((종가 / ma20) - 1) * 100 if ma20 not in [0, None] else None
    ma60괴리 = ((종가 / ma60) - 1) * 100 if ma60 not in [0, None] else None

    요약문장 = []
    요약문장.append(f"현재 배열은 {추세배열}이며, 종가는 20일선 대비 {증감문자열(ma20괴리, '%') if ma20괴리 is not None else '-'} 수준입니다.")
    if 지지괴리 is not None and 저항괴리 is not None:
        요약문장.append(f"최근 20일 기준 지지선까지는 {지지괴리:.2f}%, 저항선까지는 {저항괴리:.2f}% 거리입니다.")
    요약문장.append(f"RSI는 {숫자표시(rsi, 2) if pd.notna(rsi) else '-'}로 {rsi판정}, 거래량은 20일 평균 대비 {숫자표시(거래량배수, 2) if 거래량배수 is not None else '-'}배로 {거래량판정}입니다.")
    if 변동성20 is not None:
        요약문장.append(f"최근 20거래일 일간 변동성은 {변동성20:.2f}%입니다.")

    핵심표 = pd.DataFrame([
        {"항목": "추세 배열", "값": 추세배열, "판정": "핵심"},
        {"항목": "종가 vs 20일선", "값": 증감문자열(ma20괴리, "%") if ma20괴리 is not None else "-", "판정": "상회" if ma20괴리 is not None and ma20괴리 >= 0 else "하회"},
        {"항목": "종가 vs 60일선", "값": 증감문자열(ma60괴리, "%") if ma60괴리 is not None else "-", "판정": "상회" if ma60괴리 is not None and ma60괴리 >= 0 else "하회"},
        {"항목": "RSI(14)", "값": 숫자표시(rsi, 2), "판정": rsi판정},
        {"항목": "거래량 배수", "값": f"{거래량배수:.2f}배" if 거래량배수 is not None else "-", "판정": 거래량판정},
        {"항목": "20일 변동성", "값": f"{변동성20:.2f}%" if 변동성20 is not None else "-", "판정": "참고"},
    ])

    레벨표 = pd.DataFrame([
        {"항목": "단기 지지", "가격": 지지, "설명": "최근 20일 저가 기준"},
        {"항목": "단기 저항", "가격": 저항, "설명": "최근 20일 고가 기준"},
        {"항목": "중기 지지", "가격": 장기지지, "설명": "최근 60일 저가 기준"},
        {"항목": "중기 저항", "가격": 장기저항, "설명": "최근 60일 고가 기준"},
        {"항목": "5일선", "가격": ma5 if ma5 > 0 else None, "설명": "단기 추세"},
        {"항목": "20일선", "가격": ma20 if ma20 > 0 else None, "설명": "기준선"},
        {"항목": "60일선", "가격": ma60 if ma60 > 0 else None, "설명": "중기 추세"},
        {"항목": "120일선", "가격": ma120 if ma120 > 0 else None, "설명": "장기 추세"},
    ])

    return {
        "요약문장": 요약문장,
        "핵심표": 핵심표,
        "레벨표": 레벨표,
        "추세배열": 추세배열,
        "지지": 지지,
        "저항": 저항,
    }




def 자동판정기준표():
    return pd.DataFrame([
        {"구분": "추세", "기준": "5개 조건", "설명": "종가≥5일선, 종가≥20일선, 종가≥60일선, 5일선≥20일선, 20일선≥60일선의 충족 개수(0~5점)"},
        {"구분": "가격 위치", "기준": "최근 20거래일", "설명": "최근 20거래일 고가·저가 범위에서 현재 종가가 어디에 있는지 백분율로 계산"},
        {"구분": "RSI(14)", "기준": "과매도/과열", "설명": "28 이하 강한 과매도, 38 이하 과매도 관심, 68 이상 과열 경계, 78 이상 강한 과열로 해석"},
        {"구분": "거래량", "기준": "20일 평균 대비", "설명": "1.8배 이상 강한 거래, 1.2배 이상 유효한 확인 신호, 0.7배 이하는 힘이 약한 구간으로 해석"},
        {"구분": "점수 합계", "기준": "복합 점수", "설명": "추세·위치·RSI·거래량·당일 흐름 점수를 합산해 최종 판정을 산출"},
        {"구분": "최종 판정", "기준": "7단계", "설명": "강매수 → 분할매수 → 반등매수 → 보유 → 관망 → 비중축소 → 차익실현 순으로 변환"},
    ])



def 고급매매코멘트생성(종목명, 가격데이터, 포트폴리오행=None):
    기준표 = 자동판정기준표()

    if 가격데이터 is None or 가격데이터.empty or len(가격데이터) < 20:
        return {
            "판정": "판정 보류",
            "실행": "데이터 보강 필요",
            "강도": 0,
            "핵심문구": "최근 20거래일 이상 데이터가 부족해 자동 매매 코멘트를 보류합니다.",
            "세부코멘트": ["가격 데이터가 충분하지 않아 추세·위치·모멘텀을 함께 판정하기 어렵습니다."],
            "근거": ["최소 20거래일 이상 데이터가 있어야 20일 범위, 이동평균선, RSI 해석이 가능합니다."],
            "근거표": pd.DataFrame([
                {"항목": "데이터 길이", "현재": len(가격데이터) if 가격데이터 is not None else 0, "기준": "20거래일 이상", "판정": "부족"}
            ]),
            "기준표": 기준표,
            "위험문구": "데이터 부족 상태에서는 자동 코멘트보다 직접 차트 확인이 우선입니다.",
            "추세판정": "판정 제한",
            "위치판정": "판정 제한",
            "RSI판정": "판정 제한",
            "거래량판정": "판정 제한",
            "총점": None,
        }

    최신 = 가격데이터.iloc[-1]
    이전 = 가격데이터.iloc[-2] if len(가격데이터) >= 2 else 최신
    종가 = float(최신["종가"])
    전일종가 = float(이전["종가"]) if pd.notna(이전.get("종가")) else 종가
    ma5 = float(최신["5일평균"]) if pd.notna(최신.get("5일평균")) else None
    ma20 = float(최신["20일평균"]) if pd.notna(최신.get("20일평균")) else None
    ma60 = float(최신["60일평균"]) if pd.notna(최신.get("60일평균")) else None
    rsi = float(최신["RSI(14)"]) if pd.notna(최신.get("RSI(14)")) else None
    거래량 = float(최신.get("거래량", 0)) if pd.notna(최신.get("거래량")) else 0

    최근20 = 가격데이터.tail(20).copy()
    최근20고가 = float(최근20["고가"].max()) if not 최근20.empty else 종가
    최근20저가 = float(최근20["저가"].min()) if not 최근20.empty else 종가
    평균거래량20 = float(최근20["거래량"].mean()) if "거래량" in 최근20.columns and len(최근20) > 0 else 0
    거래량배수 = (거래량 / 평균거래량20) if 평균거래량20 not in [0, None] else None
    전일등락률 = ((종가 - 전일종가) / 전일종가 * 100) if 전일종가 not in [0, None] else 0

    if 최근20고가 > 최근20저가:
        위치백분율 = ((종가 - 최근20저가) / (최근20고가 - 최근20저가)) * 100
    else:
        위치백분율 = 50.0

    추세점수 = 0
    if ma5 is not None and 종가 >= ma5:
        추세점수 += 1
    if ma20 is not None and 종가 >= ma20:
        추세점수 += 1
    if ma60 is not None and 종가 >= ma60:
        추세점수 += 1
    if ma5 is not None and ma20 is not None and ma5 >= ma20:
        추세점수 += 1
    if ma20 is not None and ma60 is not None and ma20 >= ma60:
        추세점수 += 1

    if 추세점수 >= 5:
        추세판정 = "강한 상승 추세"
    elif 추세점수 == 4:
        추세판정 = "상승 우위"
    elif 추세점수 == 3:
        추세판정 = "중립 이상"
    elif 추세점수 == 2:
        추세판정 = "약세 압력"
    else:
        추세판정 = "하락 추세"

    if rsi is None:
        rsi판정 = "판정 제한"
        rsi점수 = 0.0
    elif rsi <= 28:
        rsi판정 = "강한 과매도"
        rsi점수 = 2.0
    elif rsi <= 38:
        rsi판정 = "과매도 관심"
        rsi점수 = 1.2
    elif rsi < 48:
        rsi판정 = "약세권"
        rsi점수 = 0.2
    elif rsi < 68:
        rsi판정 = "중립권"
        rsi점수 = 0.5
    elif rsi < 78:
        rsi판정 = "과열 경계"
        rsi점수 = -1.0
    else:
        rsi판정 = "강한 과열"
        rsi점수 = -2.0

    if 위치백분율 <= 15:
        위치판정 = "저점권"
        위치점수 = 2.0
    elif 위치백분율 <= 30:
        위치판정 = "저점 근처"
        위치점수 = 1.2
    elif 위치백분율 < 70:
        위치판정 = "중립 구간"
        위치점수 = 0.0
    elif 위치백분율 < 85:
        위치판정 = "고점 근처"
        위치점수 = -1.2
    else:
        위치판정 = "고점권"
        위치점수 = -2.0

    if 거래량배수 is None:
        거래량판정 = "판정 제한"
        거래량점수 = 0.0
    elif 거래량배수 >= 1.8:
        거래량판정 = "강한 거래량"
        거래량점수 = 1.0
    elif 거래량배수 >= 1.2:
        거래량판정 = "평균 이상"
        거래량점수 = 0.5
    elif 거래량배수 <= 0.7:
        거래량판정 = "평균 이하"
        거래량점수 = -0.5
    else:
        거래량판정 = "보통"
        거래량점수 = 0.0

    추세가중점수 = (추세점수 - 2.5) * 0.8
    당일흐름점수 = 0.0
    if 전일등락률 >= 2.0 and ma5 is not None and 종가 >= ma5:
        당일흐름점수 = 0.5
    elif 전일등락률 <= -2.5 and 추세점수 <= 2:
        당일흐름점수 = -0.5

    총점 = 추세가중점수 + 위치점수 + rsi점수 + 거래량점수 + 당일흐름점수

    보유수량 = 0.0
    평가손익 = None
    수익률 = None
    매입평균단가 = None
    if 포트폴리오행 is not None and len(포트폴리오행) > 0:
        행 = 포트폴리오행.iloc[0]
        보유수량 = float(행.get("보유수량", 0) or 0)
        평가손익 = 행.get("평가손익")
        수익률 = 행.get("수익률")
        매입평균단가 = 행.get("매입평균단가")

    if 수익률 is not None and not pd.isna(수익률):
        if 수익률 >= 10 and 위치백분율 >= 70:
            총점 -= 0.6
        elif 수익률 <= -8 and 위치백분율 <= 30 and rsi is not None and rsi <= 38:
            총점 += 0.4

    판정 = "관망"
    실행 = "방향 확인"
    핵심문구 = f"{종목명}은 현재 추세·가격 위치·모멘텀이 혼재되어 있어 성급한 진입보다 다음 확인 신호가 더 중요합니다."

    if 총점 >= 3.2 and 위치백분율 <= 35 and (rsi is None or rsi <= 55):
        판정 = "강매수"
        실행 = "2~3회 공격적 분할매수"
        핵심문구 = f"{종목명}은 복합 점수가 높고 가격도 고점권이 아니어서 **강매수 후보**로 해석됩니다."
    elif 총점 >= 1.8 and 위치백분율 <= 50 and (rsi is None or rsi <= 60):
        판정 = "분할매수"
        실행 = "1~3회 나눠 매수"
        핵심문구 = f"{종목명}은 추세 훼손이 크지 않은 가운데 가격 부담이 낮아 **분할매수**가 가능한 구간입니다."
    elif (추세점수 <= 2 and 위치백분율 <= 30 and rsi is not None and rsi <= 38) or (총점 >= 0.8 and 위치백분율 <= 25 and rsi is not None and rsi <= 35):
        판정 = "반등매수"
        실행 = "소액 시범매수"
        핵심문구 = f"{종목명}은 아직 추세형 매수보다 **기술적 반등 대응** 성격이 더 강한 구간입니다."
    elif 총점 <= -3.0 and 위치백분율 >= 75 and rsi is not None and rsi >= 68:
        판정 = "차익실현"
        실행 = "분할매도"
        핵심문구 = f"{종목명}은 최근 가격 위치와 RSI가 과열권에 가까워 **차익실현 우선 구간**으로 해석됩니다."
    elif 총점 <= -1.8 and 위치백분율 >= 60:
        판정 = "비중축소"
        실행 = "반등 시 비중 축소"
        핵심문구 = f"{종목명}은 상승 탄력이 약한데 가격 부담이 남아 있어 **비중축소** 쪽이 더 유리합니다."
    elif 총점 >= 0.5 and 추세점수 >= 3 and 위치백분율 < 80:
        판정 = "보유"
        실행 = "기존 보유 유지"
        핵심문구 = f"{종목명}은 현재 복합 점수가 중립 이상이어서 신규 추격보다 **기존 보유 유지**가 자연스럽습니다."

    강도 = int(max(5, min(95, round((총점 + 4) / 8 * 100))))

    근거 = [
        f"추세 점수: {추세점수}/5 → {추세판정}",
        f"가격 위치: 최근 20거래일 범위의 {위치백분율:.1f}% → {위치판정}",
        f"RSI(14): {rsi:.2f} → {rsi판정}" if rsi is not None else "RSI(14): 데이터 부족",
        f"거래량: 20일 평균 대비 {거래량배수:.2f}배 → {거래량판정}" if 거래량배수 is not None else "거래량: 데이터 부족",
        f"복합 총점: {총점:.2f}점",
    ]

    세부코멘트 = [
        "이번 자동 판정은 한 가지 지표가 아니라 추세·위치·RSI·거래량을 합산해 계산합니다.",
        f"현재는 추세가 **{추세판정}**, 가격 위치는 **{위치판정}**, RSI는 **{rsi판정}**으로 해석됩니다.",
    ]

    if 판정 in ["강매수", "분할매수"]:
        세부코멘트.append("추세가 완전히 무너지지 않은 상태에서 가격 부담이 낮아 한 번에 몰아 사기보다 2~3회 분할 접근이 적절합니다.")
        if ma20 is not None:
            세부코멘트.append(f"단기 확인선은 20일선 부근({ma20:,.0f}원 전후)이며, 이 구간을 지키는지 함께 보시는 것이 좋습니다.")
    elif 판정 == "반등매수":
        세부코멘트.append("추세 자체는 아직 약하므로 중장기 추세 매수보다 기술적 반등 대응 관점으로 보는 편이 안전합니다.")
        세부코멘트.append("직전 저점 이탈 시에는 빠르게 재점검하는 방식이 필요합니다.")
    elif 판정 in ["차익실현", "비중축소"]:
        세부코멘트.append("고점권에서는 조금만 흔들려도 변동성이 커질 수 있어 한 번에 전량보다 분할 대응이 더 안정적입니다.")
        if ma5 is not None:
            세부코멘트.append(f"단기 이탈 확인선은 5일선 부근({ma5:,.0f}원 전후)으로 보시면 됩니다.")
    elif 판정 == "보유":
        세부코멘트.append("신규 추격 매수보다 현재 보유분을 유지하면서 다음 돌파 또는 눌림 신호를 기다리는 편이 더 자연스럽습니다.")
    else:
        세부코멘트.append("추세와 가격 위치가 동시에 유리하지 않아 다음 방향이 정리될 때까지 관찰 비중을 높이는 편이 좋습니다.")
        if ma20 is not None and ma60 is not None:
            세부코멘트.append(f"20일선({ma20:,.0f}원)과 60일선({ma60:,.0f}원) 사이 관계가 더 분명해지는지를 확인해 보시면 좋습니다.")

    if 보유수량 > 0:
        세부코멘트.append(f"현재 보유수량은 {숫자표시(보유수량)}주입니다.")
        if 매입평균단가 is not None and not pd.isna(매입평균단가):
            세부코멘트.append(f"평균단가는 {금액표시(매입평균단가)}입니다.")
        if 평가손익 is not None and not pd.isna(평가손익):
            세부코멘트.append(f"현재 평가손익은 {손익문자열(평가손익)}원입니다.")
        if 수익률 is not None and not pd.isna(수익률):
            세부코멘트.append(f"현재 수익률은 {수익률:.2f}%입니다.")

    위험문구 = "자동 판정은 규칙 기반 참고 의견입니다. 공시, 실적, 업황, 거시 변수 같은 비가격 정보는 반드시 별도로 확인하셔야 합니다."
    if abs(전일등락률) >= 4:
        위험문구 = f"당일 변동률이 {전일등락률:+.2f}%로 커서 하루 움직임이 총점에 과하게 반영될 수 있습니다. 단일 일봉만 보고 추격 대응하는 것은 피하는 편이 좋습니다."

    근거표 = pd.DataFrame([
        {"항목": "추세", "현재": 추세판정, "기준": "5개 조건 충족 수", "판정": f"{추세점수}/5"},
        {"항목": "가격 위치", "현재": f"{위치백분율:.1f}%", "기준": "20일 범위 내 위치", "판정": 위치판정},
        {"항목": "RSI(14)", "현재": f"{rsi:.2f}" if rsi is not None else "-", "기준": "28↓ 강한 과매도 / 68↑ 과열 경계", "판정": rsi판정},
        {"항목": "거래량 배수", "현재": f"{거래량배수:.2f}배" if 거래량배수 is not None else "-", "기준": "1.2배↑ 유효 / 0.7배↓ 약함", "판정": 거래량판정},
        {"항목": "당일 등락률", "현재": f"{전일등락률:+.2f}%", "기준": "+2% 이상 가점 / -2.5% 이하 감점", "판정": f"{당일흐름점수:+.1f}점"},
        {"항목": "복합 총점", "현재": f"{총점:.2f}점", "기준": "추세+위치+RSI+거래량+당일 흐름", "판정": 판정},
        {"항목": "실행 방향", "현재": 실행, "기준": "고급형 7단계", "판정": f"강도 {강도}%"},
    ])

    return {
        "판정": 판정,
        "실행": 실행,
        "강도": 강도,
        "핵심문구": 핵심문구,
        "세부코멘트": 세부코멘트,
        "근거": 근거,
        "근거표": 근거표,
        "기준표": 기준표,
        "위험문구": 위험문구,
        "추세판정": 추세판정,
        "위치판정": 위치판정,
        "RSI판정": rsi판정,
        "거래량판정": 거래량판정,
        "총점": 총점,
    }


def 자동판정배지HTML(판정, 실행, 강도):
    색상맵 = {
        "강매수": "#15803d",
        "분할매수": "#16a34a",
        "반등매수": "#65a30d",
        "보유": "#2563eb",
        "관망": "#6b7280",
        "비중축소": "#d97706",
        "차익실현": "#dc2626",
        "판정 보류": "#6b7280",
    }
    배경 = 색상맵.get(판정, "#334155")
    return f"""
    <div style="background:{배경}; border-radius:20px; padding:16px 18px; color:white; margin:8px 0 14px 0; box-shadow:0 10px 24px rgba(15,23,42,0.18);">
        <div style="font-size:0.95rem; opacity:0.9;">자동 매수·매도 판단</div>
        <div style="font-size:1.6rem; font-weight:560; margin-top:4px;">{판정}</div>
        <div style="font-size:1rem; margin-top:6px;">실행 방향: {실행}</div>
        <div style="font-size:0.92rem; margin-top:6px; opacity:0.95;">신호 강도: {강도}/100</div>
    </div>
    """

def 현재테마기본값():
    try:
        return st.get_option("theme.base") or "dark"
    except Exception:
        return "dark"


def 대시보드스타일적용():
    테마 = 현재테마기본값()
    if 테마 == "light":
        카드배경 = "#ffffff"
        카드테두리 = "#e5e7eb"
        카드그림자 = "0 3px 10px rgba(15, 23, 42, 0.05)"
        라벨색 = "#475569"
        제목색 = "#111827"
        메타색 = "#64748b"
        보유행배경 = "#f8fafc"
        보유행테두리 = "#e2e8f0"
    else:
        카드배경 = "#111827"
        카드테두리 = "#334155"
        카드그림자 = "0 8px 18px rgba(2, 6, 23, 0.28)"
        라벨색 = "#cbd5e1"
        제목색 = "#f8fafc"
        메타색 = "#94a3b8"
        보유행배경 = "#0f172a"
        보유행테두리 = "#1e293b"

    st.markdown(f"""
    <style>
    .main .block-container {{
        padding-top: {0.7 if 모바일여부() else 1.1}rem;
        padding-bottom: 2.2rem;
        max-width: 1360px;
    }}
    .simple-market-card {{
        border: 1px solid {카드테두리};
        border-left-width: 5px;
        border-radius: 15px;
        padding: 8px 9px 7px 9px;
        background: linear-gradient(180deg, {카드배경} 0%, rgba(15,23,42,0.98) 100%);
        box-shadow: {카드그림자};
        margin-bottom: 4px;
        min-height: 112px;
        display: flex;
        flex-direction: column;
        justify-content: flex-start;
        gap: 2px;
    }}
    .simple-market-card.up {{border-left-color: #dc2626;}}
    .simple-market-card.down {{border-left-color: #2563eb;}}
    .simple-market-card.flat {{border-left-color: #94a3b8;}}
    .simple-market-label {{
        display: inline-flex;
        align-items: center;
        width: fit-content;
        font-size: 0.66rem;
        font-weight: 520;
        color: {라벨색};
        margin-bottom: 4px;
        padding: 2px 7px;
        border-radius: 999px;
        background: rgba(148,163,184,0.12);
        line-height: 1;
    }}
    .simple-market-title {{
        font-size: 0.76rem;
        font-weight: 560;
        color: {제목색};
        margin-bottom: 4px;
        line-height: 1.24;
        min-height: 1.25em;
        word-break: keep-all;
        display: -webkit-box;
        -webkit-line-clamp: 2;
        -webkit-box-orient: vertical;
        overflow: hidden;
    }}
    .simple-market-price {{
        font-size: 1.02rem;
        font-weight: 560;
        color: {제목색};
        line-height: 1.14;
        letter-spacing: -0.02em;
        margin-bottom: 4px;
        min-height: 1.25em;
        font-variant-numeric: tabular-nums;
        font-feature-settings: "tnum";
        text-align: left;
    }}
    .simple-market-delta {{
        font-size: 0.76rem;
        font-weight: 560;
        line-height: 1.24;
        min-height: 1.25em;
        display: flex;
        align-items: flex-start;
        margin-bottom: 4px;
    }}
    .simple-market-delta.up {{color: #dc2626;}}
    .simple-market-delta.down {{color: #2563eb;}}
    .simple-market-delta.flat {{color: {메타색};}}
    .simple-market-meta {{
        font-size: 0.73rem;
        color: {메타색};
        margin-top: 5px;
        line-height: 1.22;
        min-height: 0;
        display: -webkit-box;
        -webkit-line-clamp: 2;
        -webkit-box-orient: vertical;
        overflow: hidden;
        word-break: keep-all;
    }}
    .simple-market-holdings {{
        margin-top: auto;
        font-size: 0.72rem;
        color: {메타색};
        background: {보유행배경};
        border: 1px solid {보유행테두리};
        border-radius: 9px;
        padding: 4px 6px;
        white-space: nowrap;
        overflow: hidden;
        text-overflow: ellipsis;
        line-height: 1.2;
        font-variant-numeric: tabular-nums;
        font-feature-settings: "tnum";
        text-align: left;
    }}
    .top-monitor-title {{
        font-size: 1.50rem;
        font-weight: 580;
        line-height: 1.2;
        letter-spacing: -0.02em;
        margin-bottom: 0.1rem;
    }}
    .top-monitor-sub {{
        color: #94a3b8;
        font-size: 0.96rem;
        line-height: 1.35;
        margin-bottom: 0.4rem;
    }}
    .top-monitor-time {{
        padding: 8px 4px 0 8px;
        color: #93c5fd;
        font-size: 0.80rem;
        font-weight: 520;
        line-height: 1.2;
        white-space: nowrap;
    }}

    .flow-panel {{
        border: 1px solid #334155;
        border-radius: 18px;
        padding: 14px 16px;
        background: linear-gradient(180deg, #0f172a 0%, #111827 100%);
        box-shadow: 0 8px 18px rgba(2, 6, 23, 0.22);
        min-height: 112px;
        margin-bottom: 8px;
    }}
    .flow-panel-title {{font-size: 1.02rem; font-weight: 580; color: #f8fafc; margin-bottom: 2px;}}
    .flow-panel-date {{font-size: 0.80rem; color: #94a3b8; margin-bottom: 10px;}}
    .flow-row {{display: grid; grid-template-columns: 58px 1fr 96px; align-items: center; gap: 8px; margin: 8px 0;}}
    .flow-name {{font-size: 0.86rem; font-weight: 520; color: #cbd5e1;}}
    .flow-track {{position: relative; height: 12px; background: rgba(148,163,184,0.14); border-radius: 999px; overflow: hidden;}}
    .flow-zero {{position:absolute; left:50%; top:0; width:1px; height:100%; background: rgba(226,232,240,0.35);}}
    .flow-bar {{position:absolute; top:0; height:100%; border-radius:999px;}}
    .flow-value {{font-size: 0.84rem; font-weight: 580; text-align: right; font-variant-numeric: tabular-nums;}}
    .flow-value.up {{color:#ef4444;}}
    .flow-value.down {{color:#3b82f6;}}
    .flow-value.flat {{color:#94a3b8;}}
    .flow-note {{font-size: 0.77rem; color:#94a3b8; margin-top:10px; line-height:1.35;}}

    .monitor-add-card [data-testid="stButton"] > button {{
        min-height: 112px;
        height: 176px;
        border-radius: 16px;
        border: 1.5px dashed #60a5fa;
        background: linear-gradient(180deg, rgba(7,18,44,0.94) 0%, rgba(15,23,42,0.98) 100%);
        color: #dbeafe;
        font-size: 1.05rem;
        font-weight: 560;
        line-height: 1.35;
        box-shadow: none;
    }}
    .monitor-add-card [data-testid="stButton"] > button:hover {{
        border-color: #93c5fd;
        background: linear-gradient(180deg, rgba(9,26,57,0.98) 0%, rgba(15,23,42,1) 100%);
        color: #eff6ff;
    }}
.signal-box {{
        border-radius: 18px;
        padding: 14px 16px;
        color: white;
        box-shadow: 0 6px 16px rgba(15, 23, 42, 0.12);
        margin-bottom: 4px;
    }}
    .signal-title {{font-size: 0.9rem; opacity: 0.9;}}
    .signal-main {{font-size: 1.35rem; font-weight: 560; margin-top: 4px;}}
    .trade-action-row {{margin-top: 0.25rem; margin-bottom: 0.45rem;}}
    .trade-action-row [data-testid="stButton"] > button,
    .trade-action-row [data-testid="stDownloadButton"] > button {{
        min-height: 52px;
        border-radius: 14px;
        font-weight: 520;
        width: 100%;
        white-space: normal;
        line-height: 1.25;
    }}
    .trade-upload-note {{
        margin-top: 0.15rem;
        margin-bottom: 0.7rem;
        color: #94a3b8;
        font-size: 0.93rem;
    }}
    .ratio-summary-card {{
        border: 1px solid #1f2937;
        border-radius: 16px;
        padding: 14px 16px;
        background: #020817;
        margin-bottom: 4px;
    }}
    .ratio-summary-title {{
        font-size: 0.93rem;
        color: #cbd5e1;
        font-weight: 520;
        margin-bottom: 4px;
    }}
    .ratio-summary-main {{
        font-size: 1.55rem;
        color: #f8fafc;
        font-weight: 560;
        line-height: 1.15;
    }}
    .ratio-summary-sub {{
        margin-top: 6px;
        font-size: 0.92rem;
        color: #94a3b8;
    }}

    .oa-table-wrap table {{
        width: 100% !important;
        border-collapse: collapse !important;
        font-variant-numeric: tabular-nums;
    }}
    .oa-table-wrap thead th {{
        text-align: center !important;
        vertical-align: middle !important;
        line-height: 1.32 !important;
        white-space: normal !important;
        word-break: keep-all !important;
    }}
    .oa-table-wrap tbody td,
    .oa-table-wrap tbody th {{
        padding: 8px 10px !important;
        vertical-align: middle !important;
    }}
    .oa-table-wrap tbody td {{
        text-align: left;
    }}

    div[role="radiogroup"] label {{cursor: pointer !important;}}
    div[role="radiogroup"] p {{font-weight: 600;}}
    div[data-baseweb="select"] * {{cursor: pointer !important;}}
    button[role="tab"] {{cursor: pointer !important;}}
    .stTabs [data-baseweb="tab"] {{cursor: pointer !important;}}
    @media (max-width: 1200px) {{
        .simple-market-card {{
            min-height: 112px;
        }}
    }}
    @media (max-width: 768px) {{
        .main .block-container {{
            padding-top: 0.7rem;
            padding-bottom: 1.3rem;
        }}
        .top-monitor-title {{
            font-size: 1.42rem;
        }}
        .top-monitor-sub {{
            font-size: 0.80rem;
        }}
        .top-monitor-time {{
            font-size: 0.92rem;
            padding: 6px 2px 0 6px;
        }}
        .simple-market-card {{
            min-height: 126px;
            padding: 11px 12px 9px 12px;
            border-radius: 16px;
        }}
        .simple-market-title {{
            font-size: 0.96rem;
            min-height: 2.5em;
        }}
        .simple-market-price {{
            font-size: 1.05rem;
        }}
        .simple-market-delta {{
            font-size: 0.80rem;
            min-height: 1.25em;
        }}
        .simple-market-holdings {{
            font-size: 0.82rem;
        }}
        .simple-market-meta {{
            font-size: 0.72rem;
            min-height: 0;
        }}
    }}
    </style>
    """, unsafe_allow_html=True)


def 대시보드변화방향(등락률):
    if 등락률 is None or pd.isna(등락률):
        return "flat"
    if 등락률 > 0:
        return "up"
    if 등락률 < 0:
        return "down"
    return "flat"


def 심플카드HTML(이름, 현재가, 전일대비, 등락률, 보조라벨="", 하단메모="", 보유정보문자=""):
    import html

    방향 = 대시보드변화방향(등락률)
    if 현재가 is None or pd.isna(현재가):
        현재가문자 = "데이터 확인 필요"
    else:
        현재가문자 = 숫자표시(현재가, 2)

    변화문자 = "전일 비교값 없음"
    if 등락률 is not None and not pd.isna(등락률):
        if 전일대비 is not None and not pd.isna(전일대비):
            변화문자 = f"{전일대비:+,.2f} ({등락률:+.2f}%)"
        else:
            변화문자 = f"{등락률:+.2f}%"
    elif 전일대비 is not None and not pd.isna(전일대비):
        변화문자 = f"{전일대비:+,.2f}"

    이름 = html.escape(str(이름))
    보조라벨 = html.escape(str(보조라벨)) if 보조라벨 else ""
    하단메모 = html.escape(str(하단메모)).replace("\n", "<br>") if 하단메모 else ""
    보유정보문자 = html.escape(str(보유정보문자)) if 보유정보문자 else ""

    parts = [f'<div class="simple-market-card {방향}">']
    if 보조라벨:
        parts.append(f'<div class="simple-market-label">{보조라벨}</div>')
    parts.append(f'<div class="simple-market-title">{이름}</div>')
    parts.append(f'<div class="simple-market-price">{현재가문자}</div>')
    parts.append(f'<div class="simple-market-delta {방향}">{변화문자}</div>')
    if 보유정보문자:
        parts.append(f'<div class="simple-market-holdings">{보유정보문자}</div>')
    if 하단메모:
        parts.append(f'<div class="simple-market-meta">{하단메모}</div>')
    parts.append('</div>')
    return ''.join(parts)



def 수급숫자변환(값):
    if 값 is None:
        return 0.0
    try:
        if pd.isna(값):
            return 0.0
    except Exception:
        pass
    문자 = str(값).strip().replace("\xa0", "").replace(",", "").replace("+", "")
    문자 = re.sub(r"[^0-9\-\.]+", "", 문자)
    if 문자 in ["", "-", ".", "-."]:
        return 0.0
    try:
        return float(문자)
    except Exception:
        return 0.0


@st.cache_data(ttl=600, show_spinner=False)
def pykrx투자자별순매수(시장="KOSPI", refresh_token=0):
    """
    v5.8.4g 패치:
    - 최근 40일 역순 탐색
    - get_market_trading_value_by_date 우선
    - 실패 시 get_market_trading_value_by_investor 보조
    - pykrx 버전별 반환 구조 차이 대응
    - 실제 조회 기준일 표시
    """
    market = "KOSPI" if str(시장).upper() in ["KOSPI", "코스피", "1001"] else "KOSDAQ"

    결과 = {
        "시장": "코스피" if market == "KOSPI" else "코스닥",
        "날짜": "-",
        "개인": 0.0,
        "외국인": 0.0,
        "기관계": 0.0,
        "출처": "pykrx/KRX",
        "상태": "조회 실패",
    }

    try:
        from pykrx import stock
    except Exception:
        결과["상태"] = "pykrx 미설치: pip install pykrx 필요"
        return 결과

    def _억원(값):
        try:
            숫자 = pd.to_numeric(pd.Series([값]), errors="coerce").iloc[0]
            if pd.isna(숫자):
                return 0.0
            return float(숫자) / 100000000
        except Exception:
            return 0.0

    def _정규화문자(값):
        return re.sub(r"[\s_(){}\[\]/\-]+", "", str(값)).strip()

    def _df에서컬럼값(df, 후보목록):
        if df is None or df.empty:
            return 0.0

        try:
            row = df.iloc[-1]
        except Exception:
            return 0.0

        후보정규 = [_정규화문자(x) for x in 후보목록]

        for col in df.columns:
            col_text = _정규화문자(col)
            if col_text in 후보정규 or any(h in col_text for h in 후보정규):
                try:
                    return _억원(row[col])
                except Exception:
                    pass

        return 0.0

    def _df에서투자자행값(df, 후보목록):
        if df is None or df.empty:
            return 0.0

        후보정규 = [_정규화문자(x) for x in 후보목록]

        # pykrx 기본 형태: index=투자자, columns=매도/매수/순매수
        try:
            순매수컬럼 = None
            for col in df.columns:
                if "순매수" in _정규화문자(col):
                    순매수컬럼 = col
                    break

            if 순매수컬럼 is not None:
                for idx in df.index:
                    idx_text = _정규화문자(idx)
                    if idx_text in 후보정규 or any(h in idx_text for h in 후보정규):
                        return _억원(df.loc[idx, 순매수컬럼])
        except Exception:
            pass

        # index가 아니라 컬럼에 투자자명이 들어간 경우
        return _df에서컬럼값(df, 후보목록)

    def _by_date조회(날짜문자):
        호출후보 = [
            lambda: stock.get_market_trading_value_by_date(날짜문자, 날짜문자, market),
            lambda: stock.get_market_trading_value_by_date(날짜문자, 날짜문자, market, detail=True),
            lambda: stock.get_market_trading_value_by_date(날짜문자, 날짜문자, market, etf=True, etn=True, elw=True),
        ]

        오류목록 = []
        for 호출 in 호출후보:
            try:
                df = 호출()
                if df is not None and not df.empty:
                    개인 = _df에서컬럼값(df, ["개인"])
                    외국인 = _df에서컬럼값(df, ["외국인합계", "외국인"])
                    기관 = _df에서컬럼값(df, ["기관합계", "기관"])
                    return 개인, 외국인, 기관, "정상"
            except Exception as e:
                오류목록.append(str(e)[:80])

        return 0.0, 0.0, 0.0, " / ".join(오류목록[-2:]) if 오류목록 else "빈 데이터"

    def _by_investor조회(날짜문자):
        호출후보 = [
            lambda: stock.get_market_trading_value_by_investor(날짜문자, market),
            lambda: stock.get_market_trading_value_by_investor(날짜문자, market, etf=True, etn=True, elw=True),
            lambda: stock.get_market_trading_value_by_investor(날짜문자, 날짜문자, market),
        ]

        오류목록 = []
        for 호출 in 호출후보:
            try:
                df = 호출()
                if df is not None and not df.empty:
                    개인 = _df에서투자자행값(df, ["개인"])
                    외국인 = _df에서투자자행값(df, ["외국인", "외국인합계"])
                    기관 = _df에서투자자행값(df, ["기관", "기관합계"])
                    return 개인, 외국인, 기관, "정상"
            except Exception as e:
                오류목록.append(str(e)[:80])

        return 0.0, 0.0, 0.0, " / ".join(오류목록[-2:]) if 오류목록 else "빈 데이터"

    try:
        기준일 = 서울현재시각().date()
    except Exception:
        기준일 = datetime.today().date()

    마지막상태 = ""

    for i in range(0, 40):
        조회일 = 기준일 - timedelta(days=i)

        # 주말은 건너뜀
        if 조회일.weekday() >= 5:
            continue

        날짜문자 = 조회일.strftime("%Y%m%d")

        개인, 외국인, 기관, 상태 = _by_date조회(날짜문자)

        if 상태 != "정상":
            개인, 외국인, 기관, 상태 = _by_investor조회(날짜문자)

        합계절대값 = abs(개인) + abs(외국인) + abs(기관)

        if 상태 == "정상":
            결과.update({
                "날짜": 조회일.strftime("%Y-%m-%d"),
                "개인": 개인,
                "외국인": 외국인,
                "기관계": 기관,
                "상태": "정상" if 합계절대값 != 0 else "조회 성공, 수급값 0",
            })
            return 결과

        마지막상태 = 상태

    결과["상태"] = "최근 거래일 데이터 없음" + (f": {마지막상태}" if 마지막상태 else "")
    return 결과


def 네이버투자자별순매수(시장="KOSPI", refresh_token=0):
    return pykrx투자자별순매수(시장=시장, refresh_token=refresh_token)

def 수급값문자(값):
    try:
        값 = float(값)
    except Exception:
        값 = 0.0
    부호 = "+" if 값 > 0 else ""
    return f"{부호}{값:,.0f}"


def 투자자수급HTML(제목, 데이터):
    import html
    데이터 = 데이터 or {}
    항목 = [("개인", 데이터.get("개인", 0)), ("외국인", 데이터.get("외국인", 0)), ("기관", 데이터.get("기관계", 0))]
    최대값 = max([abs(float(v or 0)) for _, v in 항목] + [1])
    parts = ["<div class='flow-panel'>"]
    parts.append(f"<div class='flow-panel-title'>{html.escape(str(제목))} 투자자별 순매수</div>")
    parts.append(f"<div class='flow-panel-date'>기준 {html.escape(str(데이터.get('날짜', '-')))} · 출처 {html.escape(str(데이터.get('출처', 'pykrx/KRX')))} · 단위 억원</div>")
    for 이름, 값 in 항목:
        try:
            값 = float(값 or 0)
        except Exception:
            값 = 0.0
        비율 = min(50, abs(값) / 최대값 * 50)
        if 값 > 0:
            left = 50
            width = 비율
            방향 = "up"
            색상 = "#ef4444"
        elif 값 < 0:
            left = 50 - 비율
            width = 비율
            방향 = "down"
            색상 = "#3b82f6"
        else:
            left = 49.5
            width = 1
            방향 = "flat"
            색상 = "#94a3b8"
        parts.append(
            f"<div class='flow-row'><div class='flow-name'>{이름}</div>"
            f"<div class='flow-track'><div class='flow-zero'></div><div class='flow-bar' style='left:{left:.2f}%; width:{width:.2f}%; background:{색상};'></div></div>"
            f"<div class='flow-value {방향}'>{수급값문자(값)}</div></div>"
        )
    상태 = str(데이터.get("상태", ""))
    if 상태 and 상태 != "정상":
        parts.append(f"<div class='flow-note'>수급 데이터 상태: {html.escape(상태)}</div>")
    else:
        parts.append("<div class='flow-note'>외국인·기관 동시 순매수는 우호적 수급으로 볼 수 있으나, 환율·금리·뉴스와 함께 참고하세요.</div>")
    parts.append("</div>")
    return "".join(parts)


def 투자자수급섹션표시(refresh_token=0):
    st.markdown("#### 오늘의 수급 흐름")
    st.caption("코스피·코스닥 투자자별 최근 거래일 순매수 흐름입니다. 빨간색은 순매수, 파란색은 순매도입니다.")
    코스피수급 = 네이버투자자별순매수("KOSPI", refresh_token=refresh_token)
    코스닥수급 = 네이버투자자별순매수("KOSDAQ", refresh_token=refresh_token)
    칸1, 칸2 = st.columns(2)
    with 칸1:
        st.markdown(투자자수급HTML("코스피", 코스피수급), unsafe_allow_html=True)
    with 칸2:
        st.markdown(투자자수급HTML("코스닥", 코스닥수급), unsafe_allow_html=True)


def 카드기준시각문자열(값):
    if 값 is None or pd.isna(값):
        return "-"
    try:
        ts = pd.Timestamp(값)
        if getattr(ts, "tzinfo", None) is None:
            ts = ts.tz_localize("Asia/Seoul")
        else:
            ts = ts.tz_convert("Asia/Seoul")
        return ts.strftime("%H:%M")
    except Exception:
        try:
            return pd.to_datetime(값).strftime("%H:%M")
        except Exception:
            return str(값)


def 모니터추가카드버튼(*args, **kwargs):
    return False
def 모니터카드하단메모생성(정보):
    return ""

def 시장지표카드하단메모생성(행):
    return ""

def 대시보드보유정보사전(거래df):
    """상단 모니터 보유 표시용 요약.
    v5.11 성능 개선:
    - 주요 모니터링 화면에서는 포트폴리오 전체 현재가 계산을 호출하지 않습니다.
    - 거래원장 집계만으로 보유수량을 산출하고, 실시간 스냅샷이 이미 있을 때만 평가금액을 가볍게 보강합니다.
    """
    try:
        계산대상 = 거래이력계산대상추출(거래df)
        집계표 = 포트폴리오입력집계(계산대상)
        if 집계표 is None or 집계표.empty:
            return {}
        작업 = 집계표.copy()
        작업["보유수량"] = pd.to_numeric(작업.get("보유수량"), errors="coerce").fillna(0)
        작업 = 작업[작업["보유수량"] > 0].copy()
        if 작업.empty:
            return {}

        결과 = {}
        for _, 행 in 작업.iterrows():
            코드 = str(행.get("종목코드", "")).zfill(6)
            이름 = 종목명자동보정(코드, 행.get("종목명", ""))
            구분 = 종목구분판단(코드, 이름)
            수량 = pd.to_numeric(pd.Series([행.get("보유수량")]), errors="coerce").fillna(0).iloc[0]
            현재가 = 스냅샷현재가조회(구분, 코드)
            if 현재가 not in [None, 0]:
                평가금액 = float(현재가) * float(수량)
                결과[코드] = f"보유 {숫자표시(수량, 0)}주 · 평가 {금액표시(평가금액)}"
            else:
                결과[코드] = f"보유 {숫자표시(수량, 0)}주"
        return 결과
    except Exception:
        return {}


def 현재보유종목코드목록(거래df):
    """상단 모니터용 보유종목 추출.
    속도와 연결 안정성을 위해 현재가 조회가 필요한 포트폴리오계산캐시를 쓰지 않고,
    거래이력 원장만 집계해 보유수량이 남아 있는 종목을 계산합니다.
    """
    try:
        계산대상 = 거래이력계산대상추출(거래df)
        집계표 = 포트폴리오입력집계(계산대상)
        if 집계표 is None or 집계표.empty:
            return []
        작업 = 집계표.copy()
        작업["보유수량"] = pd.to_numeric(작업.get("보유수량"), errors="coerce").fillna(0)
        작업 = 작업[작업["보유수량"] > 0].copy()
        if 작업.empty:
            return []
        작업["_최근거래"] = pd.to_datetime(작업.get("최근거래일자"), errors="coerce")
        작업 = 작업.sort_values(["_최근거래", "종목명"], ascending=[False, True])
        return [str(x).zfill(6) for x in 작업["종목코드"].tolist() if str(x).strip()]
    except Exception:
        return []


def 주요모니터자산구성(거래df):
    """상단 모니터 표시 순서:
    코스피 → 코스닥 → ETF(현재 보유 투자원금 큰 순) → 개별종목(현재 보유 투자원금 큰 순)
    """
    동적종목매핑갱신(거래df)
    구성 = [("코스피", 주요자산["코스피"], "주요 지수"), ("코스닥", 주요자산["코스닥"], "주요 지수")]

    try:
        계산대상 = 거래이력계산대상추출(거래df)
        집계표 = 포트폴리오입력집계(계산대상)
    except Exception:
        집계표 = pd.DataFrame()

    보유항목 = []
    if 집계표 is not None and not 집계표.empty:
        작업 = 집계표.copy()
        작업["보유수량"] = pd.to_numeric(작업.get("보유수량"), errors="coerce").fillna(0)
        작업["투자원금"] = pd.to_numeric(작업.get("투자원금"), errors="coerce").fillna(0)
        작업 = 작업[작업["보유수량"] > 0].copy()
        for _, 행 in 작업.iterrows():
            코드 = str(행.get("종목코드", "")).zfill(6)
            if not 코드 or 코드 in ["001001", "002001", "1001", "2001"]:
                continue
            이름 = 종목명자동보정(코드, 행.get("종목명", "")) or 종목코드기준종목명(코드) or 코드명매핑.get(코드) or 코드
            구분 = 종목구분판단(코드, 이름)
            투자원금 = float(행.get("투자원금", 0) or 0)
            보유항목.append({"코드": 코드, "이름": 이름, "구분": 구분, "투자원금": 투자원금})

    # ETF와 개별종목을 분리한 뒤 각각 투자원금 큰 순으로 정렬합니다.
    etf목록 = sorted([x for x in 보유항목 if x.get("구분") == "etf"], key=lambda x: (x.get("투자원금", 0), x.get("이름", "")), reverse=True)
    주식목록 = sorted([x for x in 보유항목 if x.get("구분") != "etf"], key=lambda x: (x.get("투자원금", 0), x.get("이름", "")), reverse=True)

    추가된코드 = set()
    for 항목 in etf목록 + 주식목록:
        코드 = 항목["코드"]
        이름 = 항목["이름"]
        if 코드 in 추가된코드 or not 이름:
            continue
        자산정보 = 주요자산.get(이름)
        if not 자산정보:
            자산정보 = {"구분": 항목.get("구분") or 종목구분추정(이름, 코드), "코드": 코드}
            주요자산[이름] = 자산정보
        구성.append((이름, 자산정보, "보유 종목"))
        추가된코드.add(코드)

    관심코드목록 = 세션모니터관심종목가져오기()
    for 코드 in 관심코드목록:
        코드 = str(코드).zfill(6)
        if not 코드 or 코드 in 추가된코드 or 코드 in ["1001", "2001"]:
            continue
        이름 = 종목코드기준종목명(코드) or 코드명매핑.get(코드) or 코드
        if 이름 not in 주요자산:
            주요자산[이름] = {"구분": 종목구분추정(이름, 코드), "코드": 코드}
        코드명매핑[코드] = 이름
        이름코드매핑[이름] = 코드
        구성.append((이름, 주요자산[이름], "관심 종목"))
        추가된코드.add(코드)

    return 구성


def 세션선택초기화():
    사용가능주요자산 = list(주요자산.keys())
    사용가능관심종목 = list(관심종목.values())

    if "main_asset_choice_v44" not in st.session_state or st.session_state["main_asset_choice_v44"] not in 사용가능주요자산:
        st.session_state["main_asset_choice_v44"] = 사용가능주요자산[0] if 사용가능주요자산 else ""
    if "holding_asset_choice_v44" not in st.session_state or st.session_state["holding_asset_choice_v44"] not in 사용가능관심종목:
        st.session_state["holding_asset_choice_v44"] = 사용가능관심종목[0] if 사용가능관심종목 else ""


def 현재거래이력가져오기():
    """세션 상태가 아직 없더라도 가장 최근에 저장된 거래이력을 우선 반영합니다.
    우선순위는 파일 존재 여부가 아니라 실제 수정시각입니다.
    따라서 업로드 후 편집한 최종본은 자동저장본이 최신이면 자동저장본을 불러옵니다.
    """
    if "trade_history_df_v22" not in st.session_state:
        후보목록 = []

        최근업로드df = 최근업로드거래이력불러오기()
        if 최근업로드df is not None and not 최근업로드df.empty:
            try:
                수정시각 = os.path.getmtime(최근업로드거래이력파일) if os.path.exists(최근업로드거래이력파일) else 0
            except Exception:
                수정시각 = 0
            후보목록.append((수정시각, "latest_uploaded", 최근업로드df.copy()))

        자동저장df = 자동저장불러오기(거래이력자동저장파일)
        if 자동저장df is not None and not 자동저장df.empty:
            try:
                수정시각 = os.path.getmtime(거래이력자동저장파일) if os.path.exists(거래이력자동저장파일) else 0
            except Exception:
                수정시각 = 0
            후보목록.append((수정시각, "autosave", 자동저장df.copy()))

        if 후보목록:
            후보목록 = sorted(후보목록, key=lambda x: x[0], reverse=True)
            _, 출처, 초기df = 후보목록[0]
            st.session_state["trade_history_source_v1"] = 출처
        else:
            초기df = 기본포트폴리오.copy()
            st.session_state["trade_history_source_v1"] = "default"

        편집df = 거래이력편집용자동보정(초기df)
        계산df = 거래이력계산대상추출(편집df)
        st.session_state["trade_history_editor_df_v1"] = 편집df.copy()
        st.session_state["trade_history_df_v22"] = 편집df.copy()
        st.session_state["trade_history_calc_df_v1"] = 계산df.copy()
        st.session_state["trade_history_signature_v1"] = 거래이력서명생성(편집df)
        st.session_state["trade_history_last_saved_signature_v1"] = st.session_state["trade_history_signature_v1"]
    동적종목매핑갱신(st.session_state["trade_history_df_v22"])
    return st.session_state["trade_history_df_v22"]


def 포트폴리오요약지표생성(계산포트폴리오, 표시대상포트폴리오=None):
    if 계산포트폴리오 is None or 계산포트폴리오.empty:
        return {
            "총투자원금": 0.0,
            "총평가금액": 0.0,
            "총평가손익": 0.0,
            "총실현손익": 0.0,
            "총수익률": 0.0,
            "보유종목수": 0,
            "조회실패건수": 0,
            "최대비중종목명": "-",
            "최대비중": 0.0,
        }

    표시대상 = 표시대상포트폴리오.copy() if 표시대상포트폴리오 is not None else 계산포트폴리오.copy()
    정상평가행 = 표시대상[표시대상["데이터상태"] == "정상"].copy() if "데이터상태" in 표시대상.columns else 표시대상.copy()

    총투자원금 = pd.to_numeric(정상평가행.get("투자원금"), errors="coerce").fillna(0).sum() if not 정상평가행.empty else 0.0
    총평가금액 = pd.to_numeric(정상평가행.get("평가금액"), errors="coerce").fillna(0).sum() if not 정상평가행.empty else 0.0
    총평가손익 = pd.to_numeric(정상평가행.get("평가손익"), errors="coerce").fillna(0).sum() if not 정상평가행.empty else 0.0
    총실현손익 = pd.to_numeric(계산포트폴리오.get("실현손익"), errors="coerce").fillna(0).sum()
    총수익률 = (총평가손익 / 총투자원금 * 100) if 총투자원금 not in [0, None] else 0.0
    보유종목수 = int((pd.to_numeric(표시대상.get("보유수량"), errors="coerce").fillna(0) > 0).sum()) if not 표시대상.empty else 0
    조회실패건수 = int((표시대상.get("데이터상태") != "정상").sum()) if "데이터상태" in 표시대상.columns else 0

    최대비중종목명 = "-"
    최대비중 = 0.0
    if not 표시대상.empty and "현재비중" in 표시대상.columns:
        비중작업 = 표시대상.copy()
        비중작업["현재비중"] = pd.to_numeric(비중작업.get("현재비중"), errors="coerce").fillna(0)
        비중작업 = 비중작업.sort_values(["현재비중", "종목명"], ascending=[False, True])
        if not 비중작업.empty and float(비중작업.iloc[0]["현재비중"]) > 0:
            최대비중종목명 = str(비중작업.iloc[0].get("종목명", "-"))
            최대비중 = float(비중작업.iloc[0]["현재비중"])

    return {
        "총투자원금": 총투자원금,
        "총평가금액": 총평가금액,
        "총평가손익": 총평가손익,
        "총실현손익": 총실현손익,
        "총수익률": 총수익률,
        "보유종목수": 보유종목수,
        "조회실패건수": 조회실패건수,
        "최대비중종목명": 최대비중종목명,
        "최대비중": 최대비중,
    }


def 포트폴리오요약카드표시(요약정보):
    if 모바일여부():
        카드1, 카드2 = st.columns(2)
        카드1.metric("총 투자원금", 금액표시(요약정보["총투자원금"]))
        카드2.metric("총 평가금액", 금액표시(요약정보["총평가금액"]))
        카드3, 카드4 = st.columns(2)
        카드3.metric("미실현 손익", 손익문자열(요약정보["총평가손익"]) + "원")
        카드4.metric("보유 수익률", 수익률문자열(요약정보["총수익률"]))
        카드5, 카드6 = st.columns(2)
        카드5.metric("보유 종목 수", f"{요약정보['보유종목수']}개")
        카드6.metric("최대 비중 종목", 요약정보["최대비중종목명"])
    else:
        카드1, 카드2, 카드3, 카드4 = st.columns(4)
        카드1.metric("총 투자원금", 금액표시(요약정보["총투자원금"]))
        카드2.metric("총 평가금액", 금액표시(요약정보["총평가금액"]))
        카드3.metric("미실현 손익", 손익문자열(요약정보["총평가손익"]) + "원")
        카드4.metric("보유 수익률", 수익률문자열(요약정보["총수익률"]))

        카드5, 카드6, 카드7, 카드8 = st.columns(4)
        카드5.metric("실현 손익", 손익문자열(요약정보["총실현손익"]) + "원")
        카드6.metric("보유 종목 수", f"{요약정보['보유종목수']}개")
        카드7.metric("최대 비중 종목", 요약정보["최대비중종목명"], f"{요약정보['최대비중']:.2f}%")
        카드8.metric("조회 실패 종목", f"{요약정보['조회실패건수']}개")


def 선택위젯키정리():
    # 이전 버전 위젯 상태가 남아 있으면 선택 표시와 실제 값이 어긋날 수 있어 정리합니다.
    for 이전키 in ["main_asset_selector_v42", "holding_selector_v42"]:
        if 이전키 in st.session_state:
            del st.session_state[이전키]




def 인덱스기준가까운날짜찾기(데이터, 입력날짜):
    if 데이터 is None or 데이터.empty or 입력날짜 is None:
        return None
    try:
        인덱스 = pd.to_datetime(pd.Index(데이터.index))
        목표 = pd.to_datetime(입력날짜)

        # 날짜만 선택된 경우가 많으므로 먼저 날짜 기준으로 정확히 맞는 값을 찾습니다.
        인덱스정규화 = 인덱스.normalize()
        목표정규화 = 목표.normalize()
        일치위치 = np.where(인덱스정규화 == 목표정규화)[0]
        if len(일치위치) > 0:
            return 데이터.index[일치위치[-1]]

        # 정확히 일치하는 값이 없으면 가장 가까운 시점으로 보정합니다.
        차이 = np.abs((인덱스 - 목표).asi8)
        if len(차이) == 0:
            return None
        return 데이터.index[int(np.argmin(차이))]
    except Exception:
        return 데이터.index[-1] if len(데이터.index) > 0 else None

def 날짜선택옵션(데이터, 기본개수=20):
    if 데이터 is None or 데이터.empty:
        return []
    최근 = list(pd.to_datetime(데이터.index).date.astype(str))
    최근 = 최근[-기본개수:]
    최근.reverse()
    return 최근


def 캔들표시구간제한(데이터, 구간):
    if 데이터 is None or 데이터.empty:
        return pd.DataFrame()
    개수맵 = {"일": 5, "주": 5, "월": 5, "년": 5}
    표시개수 = 개수맵.get(구간, 5)
    제한데이터 = 데이터.tail(표시개수).copy()
    return 제한데이터


def 기간별OHLCV변환(데이터, 구간):
    if 데이터 is None or 데이터.empty:
        return pd.DataFrame()
    if 구간 == "일":
        return 데이터.copy()

    빈도맵 = {"주": "W", "월": "M", "년": "Y"}
    규칙 = 빈도맵.get(구간)
    if 규칙 is None:
        return 데이터.copy()

    변환 = 데이터.copy()
    변환.index = pd.to_datetime(변환.index)
    변환 = 변환.sort_index()
    기간키 = 변환.index.to_period(규칙)

    집계 = pd.DataFrame({
        "시가": 변환.groupby(기간키)["시가"].first(),
        "고가": 변환.groupby(기간키)["고가"].max(),
        "저가": 변환.groupby(기간키)["저가"].min(),
        "종가": 변환.groupby(기간키)["종가"].last(),
        "거래량": 변환.groupby(기간키)["거래량"].sum(),
        "실제말일": 변환.groupby(기간키).apply(lambda x: pd.to_datetime(x.index).max()),
    }).dropna(subset=["시가", "고가", "저가", "종가"])

    if 집계.empty:
        return pd.DataFrame()

    집계.index = pd.to_datetime(집계["실제말일"])
    집계 = 집계.drop(columns=["실제말일"]).sort_index()

    집계["5일평균"] = 집계["종가"].rolling(5, min_periods=1).mean()
    집계["20일평균"] = 집계["종가"].rolling(20, min_periods=1).mean()
    집계["60일평균"] = 집계["종가"].rolling(60, min_periods=1).mean()
    집계["120일평균"] = 집계["종가"].rolling(120, min_periods=1).mean()

    변화량 = 집계["종가"].diff()
    상승분 = 변화량.clip(lower=0)
    하락분 = -변화량.clip(upper=0)
    평균상승 = 상승분.rolling(14, min_periods=14).mean()
    평균하락 = 하락분.rolling(14, min_periods=14).mean()
    rs = 평균상승 / 평균하락.replace(0, pd.NA)
    집계["RSI(14)"] = 100 - (100 / (1 + rs))
    return 집계


def 지표변화HTML(지표명, 현재값, 전일대비):
    현재값문자 = 숫자표시(현재값, 2)
    if 전일대비 is None or pd.isna(전일대비):
        델타문자 = "-"
        델타색 = "#94a3b8"
        화살표 = ""
    elif 전일대비 > 0:
        델타문자 = 증감문자열(전일대비)
        델타색 = "#ef4444"
        화살표 = "▲ "
    elif 전일대비 < 0:
        델타문자 = 증감문자열(전일대비)
        델타색 = "#3b82f6"
        화살표 = "▼ "
    else:
        델타문자 = 증감문자열(전일대비)
        델타색 = "#94a3b8"
        화살표 = "■ "

    return f"""
    <div style="background:#020817; border:1px solid #1f2937; border-radius:18px; padding:18px 18px 14px 18px; min-height:140px;">
        <div style="font-size:0.95rem; color:#ffffff; font-weight:520; margin-bottom:8px;">{지표명}</div>
        <div style="font-size:2.1rem; color:#f8fafc; font-weight:560; line-height:1.2; margin-bottom:12px;">{현재값문자}</div>
        <div style="display:inline-block; background:rgba(15,23,42,0.65); border:1px solid {델타색}; color:{델타색}; padding:6px 12px; border-radius:999px; font-size:1rem; font-weight:520;">{화살표}{델타문자}</div>
    </div>
    """


def 캔들유형HTML(캔들유형):
    유형 = str(캔들유형)

    if "망치형" in 유형:
        색상 = "#ef4444"
        몸통배경 = "rgba(239,68,68,0.18)"
        라벨 = "망치형"
        아이콘 = "🔨"
        top_pos = "9px"
        height = "14px"
        wick_top = "2px"
        wick_height = "40px"
    elif "슈팅스타" in 유형:
        색상 = "#3b82f6"
        몸통배경 = "rgba(59,130,246,0.18)"
        라벨 = "슈팅스타"
        아이콘 = "🌠"
        top_pos = "6px"
        height = "12px"
        wick_top = "2px"
        wick_height = "40px"
    elif "도지" in 유형:
        색상 = "#f59e0b"
        몸통배경 = "rgba(245,158,11,0.10)"
        라벨 = "도지형"
        아이콘 = "✚"
        top_pos = "20px"
        height = "4px"
        wick_top = "2px"
        wick_height = "40px"
    elif "양봉" in 유형:
        색상 = "#ef4444"
        몸통배경 = "rgba(239,68,68,0.18)"
        라벨 = "양봉"
        아이콘 = "🟥"
        top_pos = "10px"
        height = "20px"
        wick_top = "2px"
        wick_height = "40px"
    else:
        색상 = "#3b82f6"
        몸통배경 = "rgba(59,130,246,0.18)"
        라벨 = "음봉"
        아이콘 = "🟦"
        top_pos = "14px"
        height = "16px"
        wick_top = "2px"
        wick_height = "40px"

    return f"""
    <div style="display:flex; align-items:center; gap:10px; padding:10px 12px; border:1px solid #334155; border-radius:14px; background:#0f172a; width:fit-content;">
        <div style="font-size:1.15rem;">{아이콘}</div>
        <div style="position:relative; width:20px; height:44px;">
            <div style="position:absolute; left:9px; top:{wick_top}; width:2px; height:{wick_height}; background:{색상};"></div>
            <div style="position:absolute; left:4px; top:{top_pos}; width:12px; height:{height}; background:{몸통배경}; border:2px solid {색상}; border-radius:2px;"></div>
        </div>
        <div style="font-size:1.05rem; font-weight:560; color:{색상};">{라벨}</div>
    </div>
    """



# -----------------------------------
# v5.14.0 분석 인사이트 고도화 / 거래원장 표시 정리
# -----------------------------------
def 거래원장조회용빈행제거(df):
    """거래 입력창의 동적 빈 행은 유지하되, 조회/분석 화면에서는 숨깁니다."""
    if df is None:
        return pd.DataFrame()
    작업 = pd.DataFrame(df).copy()
    if 작업.empty:
        return 작업

    for 열 in ["거래일자", "종목코드", "종목명", "거래구분", "거래수량", "거래단가", "거래금액", "누적보유수량"]:
        if 열 not in 작업.columns:
            작업[열] = np.nan if 열 in ["거래수량", "거래단가", "거래금액", "누적보유수량"] else ""

    코드문자 = 작업["종목코드"].apply(lambda 값: "" if pd.isna(값) else re.sub(r"[^0-9]", "", str(값)).zfill(6) if re.sub(r"[^0-9]", "", str(값)) else "")
    이름문자 = 작업["종목명"].apply(lambda 값: "" if pd.isna(값) else str(값).strip())
    구분문자 = 작업["거래구분"].apply(lambda 값: "" if pd.isna(값) else str(값).strip())
    날짜값 = pd.to_datetime(작업["거래일자"], errors="coerce")
    수량값 = pd.to_numeric(작업["거래수량"], errors="coerce").fillna(0)
    단가값 = pd.to_numeric(작업["거래단가"], errors="coerce").fillna(0)

    빈행마스크 = (
        날짜값.isna()
        & 코드문자.isin(["", "000000"])
        & 이름문자.isin(["", "000000"])
        & 구분문자.isin(["", "None", "nan", "NaT"])
        & (수량값 <= 0)
        & (단가값 <= 0)
    )
    return 작업.loc[~빈행마스크].copy().reset_index(drop=True)


def 리스크등급판단(점수):
    try:
        점수 = float(점수)
    except Exception:
        점수 = 0
    if 점수 >= 70:
        return "주의", "집중도나 손실 위험을 줄이는 점검이 필요합니다."
    if 점수 >= 40:
        return "보통", "위험 요인이 일부 있으므로 비중과 손실 구간을 정기적으로 확인하세요."
    return "양호", "현재 보유 기준으로는 위험 부담이 비교적 분산되어 있습니다."


def 보유포트폴리오리스크표생성(보유포트폴리오, 통합자산표=None):
    """현재 보유 평가금액 기준의 1차 리스크 분석표입니다.
    MDD·변동성처럼 과거 가격 이력이 필요한 지표는 다음 단계에서 확장합니다.
    """
    결과 = {
        "요약": pd.DataFrame(),
        "종목별": pd.DataFrame(),
        "자산군": pd.DataFrame(),
        "손실종목": pd.DataFrame(),
        "리스크점수": 0,
        "등급": "양호",
        "코멘트": "표시할 보유 데이터가 없습니다.",
    }

    if 보유포트폴리오 is None or pd.DataFrame(보유포트폴리오).empty:
        return 결과

    보유 = pd.DataFrame(보유포트폴리오).copy()
    if "데이터상태" in 보유.columns:
        보유 = 보유[보유["데이터상태"].astype(str) == "정상"].copy()
    if 보유.empty:
        return 결과

    for 열 in ["평가금액", "평가손익", "수익률", "현재비중", "투자원금", "보유수량"]:
        if 열 not in 보유.columns:
            보유[열] = 0
        보유[열] = pd.to_numeric(보유[열], errors="coerce").fillna(0)

    보유 = 보유[보유["평가금액"] > 0].copy()
    if 보유.empty:
        return 결과

    총평가 = float(보유["평가금액"].sum())
    보유["보유비중"] = np.where(총평가 != 0, 보유["평가금액"] / 총평가 * 100, 0)
    보유["자산군"] = 보유.apply(lambda 행: "ETF" if 종목구분판단(행.get("종목코드", ""), 행.get("종목명", "")) == "etf" else "주식", axis=1)

    종목별 = 보유[["종목코드", "종목명", "자산군", "투자원금", "평가금액", "평가손익", "수익률", "보유비중"]].copy()
    종목별 = 종목별.sort_values(["보유비중", "평가금액"], ascending=[False, False]).reset_index(drop=True)

    상위1비중 = float(종목별["보유비중"].iloc[0]) if not 종목별.empty else 0
    상위3비중 = float(종목별["보유비중"].head(3).sum()) if not 종목별.empty else 0
    손실종목수 = int((종목별["평가손익"] < 0).sum())
    보유종목수 = int(len(종목별))
    손실비중합 = float(종목별.loc[종목별["평가손익"] < 0, "보유비중"].sum())
    최저수익률 = float(종목별["수익률"].min()) if not 종목별.empty else 0

    자산군 = 종목별.groupby("자산군", as_index=False).agg({"투자원금": "sum", "평가금액": "sum", "평가손익": "sum"})
    자산군["수익률"] = np.where(자산군["투자원금"] != 0, 자산군["평가손익"] / 자산군["투자원금"] * 100, 0)
    자산군["보유비중"] = np.where(총평가 != 0, 자산군["평가금액"] / 총평가 * 100, 0)
    자산군 = 자산군.sort_values("보유비중", ascending=False).reset_index(drop=True)

    통합현금성비중 = None
    if 통합자산표 is not None and not pd.DataFrame(통합자산표).empty:
        통합 = pd.DataFrame(통합자산표).copy()
        if "자산군" in 통합.columns and "평가금액" in 통합.columns:
            통합["평가금액"] = pd.to_numeric(통합["평가금액"], errors="coerce").fillna(0)
            통합총액 = 통합["평가금액"].sum()
            현금성 = 통합[통합["자산군"].astype(str).str.contains("현금|예수금|현금성", na=False)]["평가금액"].sum()
            통합현금성비중 = float(현금성 / 통합총액 * 100) if 통합총액 else None

    집중점수 = min(40, max(0, (상위1비중 - 25) * 0.9) + max(0, (상위3비중 - 60) * 0.4))
    손실점수 = min(35, max(0, abs(min(0, 최저수익률)) * 1.2) + max(0, 손실비중합 - 30) * 0.35)
    분산점수 = 20 if 보유종목수 <= 2 else 10 if 보유종목수 <= 4 else 0
    현금점수 = 0
    if 통합현금성비중 is not None and 통합현금성비중 < 5:
        현금점수 = 5
    리스크점수 = round(min(100, 집중점수 + 손실점수 + 분산점수 + 현금점수), 1)
    등급, 코멘트 = 리스크등급판단(리스크점수)

    요약항목 = [
        {"항목": "보유 종목 수", "값": 보유종목수, "해석": "분산 정도를 보는 기본 지표"},
        {"항목": "상위 1종목 비중", "값": 상위1비중, "해석": "30% 이상이면 특정 종목 의존도가 커질 수 있음"},
        {"항목": "상위 3종목 비중", "값": 상위3비중, "해석": "60% 이상이면 포트폴리오 집중도가 높은 편"},
        {"항목": "손실 종목 수", "값": 손실종목수, "해석": "현재 평가손익 기준 손실 종목 개수"},
        {"항목": "손실 종목 비중", "값": 손실비중합, "해석": "손실 종목이 전체 평가액에서 차지하는 비중"},
        {"항목": "최저 수익률", "값": 최저수익률, "해석": "가장 부진한 보유 종목의 수익률"},
    ]
    if 통합현금성비중 is not None:
        요약항목.append({"항목": "통합 현금성 비중", "값": 통합현금성비중, "해석": "전체 자산 중 예수금·현금성 자산 비중"})

    결과.update({
        "요약": pd.DataFrame(요약항목),
        "종목별": 종목별,
        "자산군": 자산군,
        "손실종목": 종목별[종목별["평가손익"] < 0].copy(),
        "리스크점수": 리스크점수,
        "등급": 등급,
        "코멘트": 코멘트,
    })
    return 결과


def 포트폴리오리스크분석UI(보유포트폴리오, 통합자산표=None):
    st.markdown("### 포트폴리오 리스크 점검")
    st.caption("현재 보유 평가금액 기준의 1차 리스크 분석입니다. MDD·변동성은 과거 가격 이력 기반 기능으로 다음 단계에서 확장합니다.")

    분석 = 보유포트폴리오리스크표생성(보유포트폴리오, 통합자산표)
    if 분석.get("종목별", pd.DataFrame()).empty:
        st.info("리스크 분석에 사용할 정상 보유 종목 데이터가 없습니다.")
        return 분석

    등급 = 분석.get("등급", "양호")
    점수 = 분석.get("리스크점수", 0)
    코멘트 = 분석.get("코멘트", "")

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("리스크 등급", 등급)
    c2.metric("리스크 점수", f"{점수:.1f}/100")
    요약 = 분석["요약"]
    상위1 = 요약.loc[요약["항목"] == "상위 1종목 비중", "값"]
    손실비중 = 요약.loc[요약["항목"] == "손실 종목 비중", "값"]
    c3.metric("상위 1종목 비중", f"{float(상위1.iloc[0]) if not 상위1.empty else 0:.2f}%")
    c4.metric("손실 종목 비중", f"{float(손실비중.iloc[0]) if not 손실비중.empty else 0:.2f}%")

    st.info(f"{등급}: {코멘트}")

    with st.expander("리스크 상세 보기", expanded=False):
        탭1, 탭2, 탭3 = st.tabs(["요약", "종목 집중도", "자산군"])
        with 탭1:
            표시요약 = 분석["요약"].copy()
            표시요약["값"] = 표시요약.apply(lambda 행: 안전소수포맷(행["값"], 2) + "%" if "비중" in 행["항목"] or "수익률" in 행["항목"] else 안전정수포맷(행["값"]), axis=1)
            표데이터프레임(index_1부터(표시요약), use_container_width=True)
        with 탭2:
            종목표 = 분석["종목별"].copy()
            표시열 = ["종목코드", "종목명", "자산군", "평가금액", "평가손익", "수익률", "보유비중"]
            종목표 = 종목표[[열 for 열 in 표시열 if 열 in 종목표.columns]].copy()
            표데이터프레임(
                index_1부터(종목표).style.format({
                    "평가금액": 안전정수포맷,
                    "평가손익": 손익문자열,
                    "수익률": 수익률문자열,
                    "보유비중": lambda x: 안전소수포맷(x, 2),
                }).map(손익색상, subset=["평가손익"]).map(수익률색상, subset=["수익률"]),
                use_container_width=True,
            )
        with 탭3:
            자산군표 = 분석["자산군"].copy()
            표데이터프레임(
                index_1부터(자산군표).style.format({
                    "투자원금": 안전정수포맷,
                    "평가금액": 안전정수포맷,
                    "평가손익": 손익문자열,
                    "수익률": 수익률문자열,
                    "보유비중": lambda x: 안전소수포맷(x, 2),
                }).map(손익색상, subset=["평가손익"]).map(수익률색상, subset=["수익률"]),
                use_container_width=True,
            )
    return 분석


# -----------------------------------
# v5.14.0 분석 인사이트 고도화
# -----------------------------------
def _분석값가져오기(분석, 항목명, 기본값=0):
    try:
        요약 = 분석.get("요약", pd.DataFrame()) if isinstance(분석, dict) else pd.DataFrame()
        if 요약 is None or 요약.empty or "항목" not in 요약.columns or "값" not in 요약.columns:
            return 기본값
        값 = 요약.loc[요약["항목"].astype(str) == 항목명, "값"]
        if 값.empty:
            return 기본값
        return float(값.iloc[0])
    except Exception:
        return 기본값


def 포트폴리오종합인사이트생성(보유포트폴리오, 통합자산표=None, 위험분석=None):
    """수익률·집중도·현금성 비중을 함께 읽어 종합 점검 문구를 생성합니다."""
    결과 = {
        "상태": "데이터 부족",
        "핵심": "분석할 정상 보유 데이터가 충분하지 않습니다.",
        "다음점검": "거래이력과 현재가 데이터가 정상 반영되는지 먼저 확인하세요.",
        "점검표": pd.DataFrame(),
        "우선점검": pd.DataFrame(),
    }

    보유 = pd.DataFrame() if 보유포트폴리오 is None else pd.DataFrame(보유포트폴리오).copy()
    if 보유.empty:
        return 결과
    if "데이터상태" in 보유.columns:
        보유 = 보유[보유["데이터상태"].astype(str) == "정상"].copy()
    if 보유.empty:
        return 결과

    for 열 in ["평가금액", "평가손익", "수익률", "투자원금", "보유수량"]:
        if 열 not in 보유.columns:
            보유[열] = 0
        보유[열] = pd.to_numeric(보유[열], errors="coerce").fillna(0)
    보유 = 보유[보유["평가금액"] > 0].copy()
    if 보유.empty:
        return 결과

    총원금 = float(보유["투자원금"].sum())
    총평가 = float(보유["평가금액"].sum())
    총손익 = float(보유["평가손익"].sum())
    총수익률 = (총손익 / 총원금 * 100) if 총원금 else 0
    손실종목수 = int((보유["평가손익"] < 0).sum())
    보유종목수 = int(len(보유))

    if 위험분석 is None:
        위험분석 = 보유포트폴리오리스크표생성(보유, 통합자산표)
    위험등급 = 위험분석.get("등급", "양호") if isinstance(위험분석, dict) else "양호"
    위험점수 = 위험분석.get("리스크점수", 0) if isinstance(위험분석, dict) else 0
    상위1비중 = _분석값가져오기(위험분석, "상위 1종목 비중", 0)
    상위3비중 = _분석값가져오기(위험분석, "상위 3종목 비중", 0)
    손실비중 = _분석값가져오기(위험분석, "손실 종목 비중", 0)
    현금성비중 = _분석값가져오기(위험분석, "통합 현금성 비중", None)

    if 총수익률 >= 8 and 위험등급 == "양호":
        상태 = "양호"
        핵심 = "수익성과 위험 분산이 함께 양호한 상태입니다."
    elif 총수익률 >= 0 and 위험등급 in ["양호", "보통"]:
        상태 = "관찰"
        핵심 = "전체 수익은 유지되고 있으나 일부 비중 또는 손실 구간은 정기 점검이 필요합니다."
    elif 총수익률 < 0 and 위험등급 == "주의":
        상태 = "주의"
        핵심 = "수익률과 리스크가 동시에 부담되는 구간입니다. 추가 매수보다 원인 점검이 우선입니다."
    elif 총수익률 < 0:
        상태 = "관찰"
        핵심 = "전체 손익은 약세이나 위험 구조가 과도하게 나쁘지는 않은 상태입니다."
    else:
        상태 = "점검"
        핵심 = "수익률보다 비중 구조와 손실 종목 관리가 더 중요한 구간입니다."

    점검포인트 = []
    if 상위1비중 >= 35:
        점검포인트.append(f"상위 1종목 비중이 {상위1비중:.1f}%로 높아 해당 종목 변동성이 전체 성과에 크게 작용합니다.")
    if 상위3비중 >= 70:
        점검포인트.append(f"상위 3종목 비중이 {상위3비중:.1f}%로 높아 분산 효과가 제한될 수 있습니다.")
    if 손실비중 >= 30:
        점검포인트.append(f"손실 종목 비중이 {손실비중:.1f}%로 커서 손실 구간의 원인 확인이 필요합니다.")
    if 현금성비중 is not None:
        if 현금성비중 < 5:
            점검포인트.append(f"통합 현금성 비중이 {현금성비중:.1f}%로 낮아 추가 대응 여력이 제한될 수 있습니다.")
        elif 현금성비중 >= 25:
            점검포인트.append(f"통합 현금성 비중이 {현금성비중:.1f}%로 높아 분할매수 여력은 있으나 대기자금 운용 효율도 함께 확인하세요.")
    if 손실종목수 > 0:
        점검포인트.append(f"현재 손실 종목은 {손실종목수}개입니다. 단순 손절보다 매수 사유가 유지되는지 먼저 확인하세요.")

    if not 점검포인트:
        점검포인트.append("현재는 특정 위험 신호가 과도하지 않으므로 기존 원칙을 유지하며 정기 점검하면 됩니다.")

    다음점검 = 점검포인트[0]

    점검표 = pd.DataFrame([
        {"항목": "통합 수익률", "현재값": 총수익률, "판정": "양호" if 총수익률 >= 5 else "보통" if 총수익률 >= 0 else "주의", "해석": "전체 보유 주식·ETF 기준 수익률"},
        {"항목": "리스크 등급", "현재값": 위험점수, "판정": 위험등급, "해석": "집중도·손실비중·분산도를 합산한 점검 등급"},
        {"항목": "상위 1종목 비중", "현재값": 상위1비중, "판정": "주의" if 상위1비중 >= 35 else "보통" if 상위1비중 >= 25 else "양호", "해석": "특정 종목 의존도"},
        {"항목": "손실 종목 비중", "현재값": 손실비중, "판정": "주의" if 손실비중 >= 30 else "보통" if 손실비중 > 0 else "양호", "해석": "손실 종목이 전체 평가액에서 차지하는 비중"},
    ])
    if 현금성비중 is not None:
        점검표 = pd.concat([점검표, pd.DataFrame([{"항목": "통합 현금성 비중", "현재값": 현금성비중, "판정": "주의" if 현금성비중 < 5 else "보통" if 현금성비중 >= 25 else "양호", "해석": "예수금·현금성 자산 비중"}])], ignore_index=True)

    우선점검 = 보유.copy()
    우선점검["점검점수"] = 0.0
    우선점검["점검사유"] = ""
    if 총평가:
        우선점검["보유비중"] = 우선점검["평가금액"] / 총평가 * 100
    else:
        우선점검["보유비중"] = 0

    def _종목점검(row):
        점수 = 0
        사유 = []
        if row.get("보유비중", 0) >= 30:
            점수 += 35
            사유.append("비중 높음")
        elif row.get("보유비중", 0) >= 20:
            점수 += 20
            사유.append("비중 관찰")
        if row.get("수익률", 0) <= -10:
            점수 += 35
            사유.append("손실률 큼")
        elif row.get("수익률", 0) < 0:
            점수 += 15
            사유.append("손실 구간")
        if row.get("평가손익", 0) < 0 and row.get("보유비중", 0) >= 15:
            점수 += 15
            사유.append("손실+비중")
        return pd.Series({"점검점수": 점수, "점검사유": ", ".join(사유) if 사유 else "정기 점검"})

    점검결과 = 우선점검.apply(_종목점검, axis=1)
    우선점검["점검점수"] = 점검결과["점검점수"]
    우선점검["점검사유"] = 점검결과["점검사유"]
    우선점검 = 우선점검.sort_values(["점검점수", "보유비중"], ascending=[False, False]).head(5)

    결과.update({
        "상태": 상태,
        "핵심": 핵심,
        "다음점검": 다음점검,
        "점검표": 점검표,
        "우선점검": 우선점검,
        "요약문": f"현재 포트폴리오는 '{상태}' 상태입니다. 총수익률은 {총수익률:.2f}%, 리스크 등급은 {위험등급}({float(위험점수):.1f}/100)입니다.",
    })
    return 결과


def 포트폴리오종합인사이트UI(보유포트폴리오, 통합자산표=None, 위험분석=None):
    """포트폴리오 종합 인사이트 UI
    - v5.14.1: 3개 카드형 배치로 가독성 개선
    - 줄바꿈 문자가 그대로 보이는 문제 제거
    - 요약표 글자 크기와 배치 균형 개선
    """
    st.markdown(
        """
        <style>
        .insight-wrap {
            margin-top: 0.25rem;
            margin-bottom: 0.8rem;
        }
        .insight-title {
            font-size: clamp(1.35rem, 2vw, 1.85rem);
            font-weight: 560;
            letter-spacing: -0.03em;
            margin: 0 0 0.25rem 0;
        }
        .insight-subtitle {
            color: #9ca3af;
            font-size: 0.92rem;
            line-height: 1.45;
            margin-bottom: 0.9rem;
        }
        .insight-card {
            border: 1px solid rgba(148, 163, 184, 0.28);
            background: rgba(15, 23, 42, 0.44);
            border-radius: 18px;
            padding: 1.05rem 1.1rem;
            min-height: 132px;
            box-shadow: 0 8px 24px rgba(0,0,0,0.13);
        }
        .insight-card-label {
            color: #9ca3af;
            font-size: 0.82rem;
            font-weight: 500;
            margin-bottom: 0.42rem;
            letter-spacing: -0.01em;
        }
        .insight-status {
            font-size: clamp(1.65rem, 2.6vw, 2.25rem);
            font-weight: 620;
            letter-spacing: -0.04em;
            line-height: 1.05;
            margin-top: 0.15rem;
        }
        .insight-body {
            font-size: 0.98rem;
            font-weight: 450;
            line-height: 1.58;
            word-break: keep-all;
        }
        .insight-chip {
            display: inline-block;
            margin-top: 0.62rem;
            padding: 0.24rem 0.58rem;
            border-radius: 999px;
            border: 1px solid rgba(148, 163, 184, 0.32);
            color: #cbd5e1;
            font-size: 0.78rem;
            font-weight: 500;
        }
        .insight-table-title {
            margin-top: 1.05rem;
            margin-bottom: 0.35rem;
            font-size: 1.02rem;
            font-weight: 540;
            letter-spacing: -0.02em;
        }
        .insight-help-text {
            color: #9ca3af;
            font-size: 0.82rem;
            margin-bottom: 0.35rem;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    st.markdown("<div class='insight-wrap'>", unsafe_allow_html=True)
    st.markdown("<div class='insight-title'>포트폴리오 종합 인사이트</div>", unsafe_allow_html=True)
    st.markdown(
        "<div class='insight-subtitle'>수익률, 집중도, 손실비중, 현금성 비중을 함께 읽어 현재 상태와 다음 점검 포인트를 요약합니다.</div>",
        unsafe_allow_html=True,
    )

    인사이트 = 포트폴리오종합인사이트생성(보유포트폴리오, 통합자산표, 위험분석)
    상태 = str(인사이트.get("상태", "데이터 부족") or "데이터 부족").replace("\\n", "<br>")
    핵심 = str(인사이트.get("핵심", "") or "").replace("\\n", "<br>")
    다음점검 = str(인사이트.get("다음점검", "") or "").replace("\\n", "<br>")
    요약문 = str(인사이트.get("요약문", "") or "").replace("\\n", "<br>")

    상태칩 = "정기 점검 유지"
    if "주의" in 상태:
        상태칩 = "우선 점검 필요"
    elif "관찰" in 상태 or "점검" in 상태:
        상태칩 = "관찰 구간"
    elif "양호" in 상태:
        상태칩 = "안정적 관리"

    c1, c2, c3 = st.columns([0.9, 1.75, 1.75], gap="medium")
    with c1:
        st.markdown(
            f"""
            <div class='insight-card'>
                <div class='insight-card-label'>종합 상태</div>
                <div class='insight-status'>{상태}</div>
                <div class='insight-chip'>{상태칩}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    with c2:
        st.markdown(
            f"""
            <div class='insight-card'>
                <div class='insight-card-label'>핵심 요약</div>
                <div class='insight-body'>{핵심}</div>
                <div class='insight-chip'>수익률·위험 구조</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    with c3:
        st.markdown(
            f"""
            <div class='insight-card'>
                <div class='insight-card-label'>다음 점검</div>
                <div class='insight-body'>{다음점검}</div>
                <div class='insight-chip'>우선 확인 사항</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    if 요약문:
        st.caption(요약문)

    점검표 = 인사이트.get("점검표", pd.DataFrame())
    if 점검표 is not None and not 점검표.empty:
        표시 = 점검표.copy()
        표시["현재값"] = 표시.apply(
            lambda 행: f"{float(행['현재값']):.1f}/100" if 행["항목"] == "리스크 등급" else f"{float(행['현재값']):.2f}%",
            axis=1,
        )
        st.markdown("<div class='insight-table-title'>리스크 요약표</div>", unsafe_allow_html=True)
        st.markdown("<div class='insight-help-text'>각 항목은 현재 보유 포트폴리오와 통합 자산 기준으로 계산됩니다.</div>", unsafe_allow_html=True)
        try:
            표데이터프레임(
                index_1부터(표시).style
                .set_properties(**{"font-size": "0.88rem", "line-height": "1.35"})
                .set_table_styles([
                    {"selector": "th", "props": [("font-size", "0.86rem"), ("font-weight", "500")]},
                    {"selector": "td", "props": [("padding", "0.46rem 0.55rem")]},
                ]),
                use_container_width=True,
            )
        except Exception:
            표데이터프레임(index_1부터(표시), use_container_width=True)

    with st.expander("우선 점검 종목 보기", expanded=False):
        우선 = 인사이트.get("우선점검", pd.DataFrame())
        if 우선 is None or 우선.empty:
            st.info("우선 점검 종목을 표시할 데이터가 없습니다.")
        else:
            표시열 = ["종목코드", "종목명", "평가금액", "평가손익", "수익률", "보유비중", "점검점수", "점검사유"]
            우선표 = 우선[[열 for 열 in 표시열 if 열 in 우선.columns]].copy()
            표데이터프레임(
                index_1부터(우선표).style.format({
                    "평가금액": 안전정수포맷,
                    "평가손익": 손익문자열,
                    "수익률": 수익률문자열,
                    "보유비중": lambda x: 안전소수포맷(x, 2),
                    "점검점수": 안전소수포맷,
                }).map(손익색상, subset=["평가손익"]).map(수익률색상, subset=["수익률"]),
                use_container_width=True,
            )
    st.markdown("</div>", unsafe_allow_html=True)
    return 인사이트




# -----------------------------------
# v5.14.3 월간 투자 리포트 본문 고도화
# - 숫자 중심 초안에서 문장형 투자 리포트로 개선
# - 리스크/자산군/종목 기여/다음 점검 포인트를 자연어로 연결
# - 화면 리포트와 텍스트 다운로드 원문을 함께 생성
# -----------------------------------

def _리포트숫자값(df, 컬럼명, 기본값=0.0):
    try:
        if df is None or df.empty or 컬럼명 not in df.columns:
            return 기본값
        return float(pd.to_numeric(df[컬럼명], errors="coerce").fillna(0).sum())
    except Exception:
        return 기본값


def _리포트상위행(df, 기준컬럼, ascending=False):
    try:
        if df is None or df.empty or 기준컬럼 not in df.columns:
            return None
        작업 = df.copy()
        작업[기준컬럼] = pd.to_numeric(작업[기준컬럼], errors="coerce")
        작업 = 작업.dropna(subset=[기준컬럼])
        if 작업.empty:
            return None
        return 작업.sort_values(기준컬럼, ascending=ascending).iloc[0]
    except Exception:
        return None


def _리포트문장목록_문자열(문장목록):
    try:
        return "\n".join([f"- {str(x).strip()}" for x in 문장목록 if str(x).strip()])
    except Exception:
        return ""


def _리포트HTML목록(문장목록):
    try:
        항목 = []
        for 문장 in 문장목록:
            문장 = str(문장).strip()
            if 문장:
                항목.append(f"<li>{html.escape(문장)}</li>")
        return "<ul class='monthly-report-list'>" + "".join(항목) + "</ul>"
    except Exception:
        return ""


def _리포트값(row, 컬럼명, 기본값=0.0):
    try:
        if row is None:
            return 기본값
        return float(pd.to_numeric(pd.Series([row.get(컬럼명, 기본값)]), errors="coerce").fillna(기본값).iloc[0])
    except Exception:
        return 기본값


def _리포트명(row, 기본값="-"):
    try:
        if row is None:
            return 기본값
        이름 = str(row.get("종목명", "")).strip()
        return 이름 if 이름 else 기본값
    except Exception:
        return 기본값


def _리포트자산군요약(통합):
    if 통합 is None or 통합.empty or "자산군" not in 통합.columns:
        return pd.DataFrame()
    try:
        작업 = 통합.copy()
        작업["평가금액"] = pd.to_numeric(작업.get("평가금액", 0), errors="coerce").fillna(0)
        작업["평가손익"] = pd.to_numeric(작업.get("평가손익", 0), errors="coerce").fillna(0)
        총평가 = 작업["평가금액"].sum()
        요약 = 작업.groupby("자산군", as_index=False).agg({"평가금액": "sum", "평가손익": "sum"})
        요약["비중"] = np.where(총평가 != 0, 요약["평가금액"] / 총평가 * 100, 0)
        return 요약.sort_values("평가금액", ascending=False).reset_index(drop=True)
    except Exception:
        return pd.DataFrame()


def _월간거래요약문장(거래df, 기준년월):
    문장 = []
    try:
        if 거래df is None or 거래df.empty or "거래일자" not in 거래df.columns:
            return ["이번 달 거래 내역을 해석할 수 있는 자료가 충분하지 않습니다."]
        작업 = 거래이력계산대상추출(거래df).copy()
        if 작업.empty:
            return ["이번 달 계산 가능한 거래 내역은 아직 없습니다."]
        작업["거래일자"] = pd.to_datetime(작업["거래일자"], errors="coerce")
        작업 = 작업[작업["거래일자"].dt.strftime("%Y-%m") == str(기준년월)].copy()
        if 작업.empty:
            return [f"{기준년월}에는 계산 가능한 매수·매도 거래가 없습니다."]
        작업["거래금액"] = pd.to_numeric(작업.get("거래수량", 0), errors="coerce").fillna(0) * pd.to_numeric(작업.get("거래단가", 0), errors="coerce").fillna(0)
        매수금액 = 작업.loc[작업["거래구분"].astype(str) == "매수", "거래금액"].sum()
        매도금액 = 작업.loc[작업["거래구분"].astype(str) == "매도", "거래금액"].sum()
        순매수 = 매수금액 - 매도금액
        종목수 = 작업["종목명"].astype(str).str.strip().replace("", np.nan).dropna().nunique() if "종목명" in 작업.columns else 0
        문장.append(f"{기준년월}에는 총 {len(작업)}건, {종목수}개 종목의 거래가 반영되었습니다.")
        if 매수금액 > 0 or 매도금액 > 0:
            문장.append(f"월간 매수금액은 {금액표시(매수금액)}, 매도금액은 {금액표시(매도금액)}이며 순매수는 {손익문자열(순매수)}입니다.")
        if 순매수 > 0:
            문장.append("이번 달 거래 흐름은 자산 축소보다 추가 편입과 비중 확대에 가까웠습니다.")
        elif 순매수 < 0:
            문장.append("이번 달 거래 흐름은 일부 현금화 또는 비중 축소 성격이 더 강했습니다.")
        else:
            문장.append("이번 달 거래 흐름은 매수와 매도가 비교적 균형을 이룬 상태입니다.")
        return 문장
    except Exception:
        return ["월간 거래 흐름 해석 중 일부 데이터 형식이 맞지 않아 보수적으로 표시합니다."]


def 월간투자리포트초안생성(거래df=None, 계산포트폴리오=None, 보유포트폴리오=None, 통합자산표=None, 위험분석=None, 기준년월=None):
    """v5.14.3 문장형 월간 투자 리포트 데이터를 생성합니다."""
    현재시각 = 서울현재시각()
    기준년월 = 기준년월 or 현재시각.strftime("%Y-%m")

    보유 = 보유포트폴리오.copy() if isinstance(보유포트폴리오, pd.DataFrame) else pd.DataFrame()
    계산 = 계산포트폴리오.copy() if isinstance(계산포트폴리오, pd.DataFrame) else pd.DataFrame()
    통합 = 통합자산표.copy() if isinstance(통합자산표, pd.DataFrame) else pd.DataFrame()

    if not 보유.empty and "데이터상태" in 보유.columns:
        정상보유 = 보유[보유["데이터상태"].astype(str) == "정상"].copy()
    else:
        정상보유 = 보유.copy()

    if not 통합.empty:
        총원금 = _리포트숫자값(통합, "원금")
        총평가 = _리포트숫자값(통합, "평가금액")
        총손익 = 총평가 - 총원금
        총수익률 = (총손익 / 총원금 * 100) if 총원금 else 0.0
        통합기준사용 = True
    else:
        총원금 = _리포트숫자값(정상보유, "투자원금")
        총평가 = _리포트숫자값(정상보유, "평가금액")
        총손익 = _리포트숫자값(정상보유, "평가손익")
        총수익률 = (총손익 / 총원금 * 100) if 총원금 else 0.0
        통합기준사용 = False

    보유종목수 = int(len(정상보유)) if isinstance(정상보유, pd.DataFrame) else 0
    실현손익 = _리포트숫자값(계산, "실현손익")

    최대비중행 = _리포트상위행(정상보유, "현재비중", ascending=False)
    수익상위행 = _리포트상위행(정상보유, "평가손익", ascending=False)
    수익하위행 = _리포트상위행(정상보유, "평가손익", ascending=True)
    수익률상위행 = _리포트상위행(정상보유, "수익률", ascending=False)
    수익률하위행 = _리포트상위행(정상보유, "수익률", ascending=True)

    최대비중명 = _리포트명(최대비중행)
    최대비중 = _리포트값(최대비중행, "현재비중", 0)
    수익상위명 = _리포트명(수익상위행)
    수익상위손익 = _리포트값(수익상위행, "평가손익", 0)
    수익하위명 = _리포트명(수익하위행)
    수익하위손익 = _리포트값(수익하위행, "평가손익", 0)
    수익률상위명 = _리포트명(수익률상위행)
    수익률상위 = _리포트값(수익률상위행, "수익률", 0)
    수익률하위명 = _리포트명(수익률하위행)
    수익률하위 = _리포트값(수익률하위행, "수익률", 0)

    손실종목수 = 0
    손실비중 = 0.0
    if not 정상보유.empty and "평가손익" in 정상보유.columns:
        try:
            손실표 = 정상보유[pd.to_numeric(정상보유["평가손익"], errors="coerce").fillna(0) < 0].copy()
            손실종목수 = int(len(손실표))
            전체평가 = pd.to_numeric(정상보유.get("평가금액", 0), errors="coerce").fillna(0).sum()
            손실평가 = pd.to_numeric(손실표.get("평가금액", 0), errors="coerce").fillna(0).sum() if not 손실표.empty else 0
            손실비중 = (손실평가 / 전체평가 * 100) if 전체평가 else 0.0
        except Exception:
            pass

    자산군표 = _리포트자산군요약(통합)
    현금성비중 = 0.0
    주식ETF비중 = 0.0
    최대자산군 = "-"
    최대자산군비중 = 0.0
    if not 자산군표.empty:
        try:
            현금성비중 = float(자산군표.loc[자산군표["자산군"].astype(str).str.contains("현금", na=False), "비중"].sum())
        except Exception:
            현금성비중 = 0.0
        try:
            주식ETF비중 = float(자산군표.loc[자산군표["자산군"].astype(str).isin(["주식", "ETF"]), "비중"].sum())
        except Exception:
            주식ETF비중 = 0.0
        try:
            최대자산군행 = 자산군표.iloc[0]
            최대자산군 = str(최대자산군행.get("자산군", "-"))
            최대자산군비중 = float(최대자산군행.get("비중", 0))
        except Exception:
            pass

    리스크점수 = None
    리스크등급 = "보통"
    if isinstance(위험분석, dict):
        for 후보키 in ["리스크점수", "위험점수", "risk_score", "점수"]:
            if 후보키 in 위험분석:
                try:
                    리스크점수 = float(위험분석.get(후보키))
                    break
                except Exception:
                    pass
        for 후보키 in ["리스크등급", "위험등급", "등급", "상태"]:
            if 후보키 in 위험분석 and str(위험분석.get(후보키)).strip():
                리스크등급 = str(위험분석.get(후보키)).strip()
                break

    # 1. 핵심 요약
    핵심요약 = []
    if 총수익률 >= 8:
        핵심요약.append(f"전체 포트폴리오는 {수익률문자열(총수익률)}로 양호한 수익 구간에 있습니다.")
    elif 총수익률 >= 3:
        핵심요약.append(f"전체 포트폴리오는 {수익률문자열(총수익률)}로 안정적인 수익 흐름을 유지하고 있습니다.")
    elif 총수익률 >= 0:
        핵심요약.append(f"전체 포트폴리오는 {수익률문자열(총수익률)}로 보합권 이상의 흐름을 보이고 있습니다.")
    else:
        핵심요약.append(f"전체 포트폴리오는 {수익률문자열(총수익률)}로 손실 구간에 있어 방어적 점검이 필요합니다.")

    핵심요약.append(f"총 평가액은 {금액표시(총평가)}, 총 손익은 {손익문자열(총손익)}입니다.")
    if 통합기준사용:
        핵심요약.append(f"자산군 기준으로는 {최대자산군} 비중이 {최대자산군비중:.1f}%로 가장 큽니다.")
    else:
        핵심요약.append("현재 리포트는 보유 주식·ETF 데이터 중심으로 작성되었으며, 통합 자산 자료가 있으면 해석 정확도가 높아집니다.")

    # 2. 월간 거래 흐름
    월간거래문장 = _월간거래요약문장(거래df, 기준년월)

    # 3. 종목별 기여
    주요종목문장 = []
    if 수익상위행 is not None:
        주요종목문장.append(f"평가손익 기여가 가장 큰 종목은 {수익상위명}이며 현재 손익은 {손익문자열(수익상위손익)}입니다.")
    if 수익률상위행 is not None and 수익률상위명 != 수익상위명:
        주요종목문장.append(f"수익률 기준으로는 {수익률상위명}이 {수익률문자열(수익률상위)}로 가장 높습니다.")
    if 수익하위행 is not None:
        if 수익하위손익 < 0:
            주요종목문장.append(f"손익 하위 종목은 {수익하위명}이며 현재 손익은 {손익문자열(수익하위손익)}입니다.")
        else:
            주요종목문장.append(f"현재 보유 종목 중 뚜렷한 손실 종목은 제한적이며, 상대적으로 기여도가 낮은 종목은 {수익하위명}입니다.")
    if 수익률하위행 is not None and 수익률하위 < 0 and 수익률하위명 != 수익하위명:
        주요종목문장.append(f"수익률 기준 하위 종목은 {수익률하위명}이며 수익률은 {수익률문자열(수익률하위)}입니다.")

    # 4. 리스크 해석
    리스크문장 = []
    if 최대비중 >= 35:
        리스크문장.append(f"{최대비중명} 비중이 {최대비중:.1f}%로 높아 단일 종목 변동이 전체 성과에 크게 반영될 수 있습니다.")
    elif 최대비중 >= 25:
        리스크문장.append(f"{최대비중명} 비중이 {최대비중:.1f}%로 높은 편이므로 추가 매수 전 비중 점검이 필요합니다.")
    elif 보유종목수 > 0:
        리스크문장.append("단일 종목 집중도는 과도한 수준보다는 관리 가능한 범위에 가깝습니다.")
    else:
        리스크문장.append("보유 종목 데이터가 부족해 종목 집중도 평가는 제한적입니다.")

    if 손실종목수 > 0:
        리스크문장.append(f"손실 종목은 {손실종목수}개이며, 손실 종목 평가액 비중은 약 {손실비중:.1f}%입니다.")
    else:
        리스크문장.append("현재 보유 기준으로 손실 종목 부담은 크지 않습니다.")

    if 통합기준사용:
        if 현금성비중 >= 20:
            리스크문장.append(f"현금성 자산 비중은 {현금성비중:.1f}%로 충분해 조정장 대응 여력이 양호합니다.")
        elif 현금성비중 >= 8:
            리스크문장.append(f"현금성 자산 비중은 {현금성비중:.1f}%로 기본 대응 여력은 있으나 큰 조정장에서는 추가 점검이 필요합니다.")
        else:
            리스크문장.append(f"현금성 자산 비중은 {현금성비중:.1f}%로 낮아 신규 매수보다 유동성 관리가 우선될 수 있습니다.")
        if 주식ETF비중 >= 70:
            리스크문장.append(f"주식·ETF 비중이 {주식ETF비중:.1f}%로 높아 시장 변동성에 대한 민감도가 큰 편입니다.")

    # 5. 다음 점검 포인트
    점검포인트 = []
    if 최대비중 >= 30:
        점검포인트.append(f"{최대비중명} 추가 매수 전 전체 비중이 과도해지지 않는지 먼저 확인합니다.")
    if 손실종목수 > 0:
        점검포인트.append("손실 종목은 매수 당시의 투자 이유가 현재도 유효한지 다시 점검합니다.")
    if 통합기준사용 and 현금성비중 < 8:
        점검포인트.append("현금성 자산 비중이 낮으므로 단기 대응 자금 확보 여부를 확인합니다.")
    if 실현손익 < 0:
        점검포인트.append("누적 실현손익이 손실이면 매도 기준과 손절 기준을 다시 정리합니다.")
    if 총수익률 >= 5 and 최대비중 >= 25:
        점검포인트.append("수익 구간에서는 신규 매수보다 비중 관리와 일부 리밸런싱 가능성을 함께 검토합니다.")
    if not 점검포인트:
        점검포인트 = [
            "현재 구조는 급한 경고보다 정기 점검 중심으로 관리하면 됩니다.",
            "다음 새로고침 때 주요 종목의 비중 변화와 수익률 변화를 함께 확인합니다.",
        ]

    종합코멘트문단 = []
    종합코멘트문단.append("이번 월간 리포트는 현재가가 반영된 평가금액과 거래이력을 함께 기준으로 작성되었습니다.")
    if 총수익률 >= 0:
        종합코멘트문단.append("전체적으로는 자산 가치가 원금 대비 방어 또는 성장 흐름을 유지하고 있습니다.")
    else:
        종합코멘트문단.append("전체적으로는 손실 방어와 보유 기준 재점검이 우선인 구간입니다.")
    if 최대비중 >= 25:
        종합코멘트문단.append("다만 특정 종목 또는 특정 테마의 영향력이 커질 수 있으므로, 수익률만 보지 말고 비중과 변동성을 함께 확인하는 것이 좋습니다.")
    elif 통합기준사용 and 현금성비중 >= 15:
        종합코멘트문단.append("현금성 자산이 일정 수준 확보되어 있어 급한 매매보다 계획적인 분할 대응이 가능한 구조입니다.")
    종합코멘트 = " ".join(종합코멘트문단)

    요약표 = pd.DataFrame([
        {"항목": "총 평가액", "값": 금액표시(총평가)},
        {"항목": "총 손익", "값": 손익문자열(총손익)},
        {"항목": "총 수익률", "값": 수익률문자열(총수익률)},
        {"항목": "보유 종목 수", "값": f"{보유종목수}개"},
        {"항목": "최대 비중", "값": f"{최대비중명} {최대비중:.1f}%" if 최대비중명 != "-" else "-"},
        {"항목": "손실 종목", "값": f"{손실종목수}개"},
        {"항목": "현금성 비중", "값": f"{현금성비중:.1f}%" if 통합기준사용 else "-"},
    ])

    리포트본문 = f"""# 월간 투자 리포트 ({기준년월})

## 1. 핵심 요약
{_리포트문장목록_문자열(핵심요약)}

## 2. 월간 거래 흐름
{_리포트문장목록_문자열(월간거래문장)}

## 3. 주요 종목 변화
{_리포트문장목록_문자열(주요종목문장)}

## 4. 리스크 상태
{_리포트문장목록_문자열(리스크문장)}

## 5. 다음 점검 포인트
{_리포트문장목록_문자열(점검포인트)}

## 6. 종합 코멘트
{종합코멘트}

작성 기준: {서울조회문자열(현재시각, "%Y-%m-%d %H:%M")}
""".strip()

    return {
        "기준년월": 기준년월,
        "작성시각": 현재시각,
        "총원금": 총원금,
        "총평가": 총평가,
        "총손익": 총손익,
        "총수익률": 총수익률,
        "보유종목수": 보유종목수,
        "최대비중명": 최대비중명,
        "최대비중": 최대비중,
        "현금성비중": 현금성비중,
        "손실종목수": 손실종목수,
        "손실비중": 손실비중,
        "실현손익": 실현손익,
        "리스크점수": 리스크점수,
        "리스크등급": 리스크등급,
        "핵심요약": 핵심요약,
        "월간거래문장": 월간거래문장,
        "주요종목문장": 주요종목문장,
        "리스크문장": 리스크문장,
        "점검포인트": 점검포인트,
        "종합코멘트": 종합코멘트,
        "요약표": 요약표,
        "자산군표": 자산군표,
        "본문": 리포트본문,
        "통합기준사용": 통합기준사용,
    }



def _docx_한글폰트적용(run, 글꼴="맑은 고딕", 크기=None, 굵게=None, 색상=None):
    """Word 리포트에서 한글이 안정적으로 보이도록 run 단위 글꼴을 지정합니다."""
    if run is None:
        return
    try:
        run.font.name = 글꼴
        if qn is not None:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), 글꼴)
        if 크기 is not None and Pt is not None:
            run.font.size = Pt(크기)
        if 굵게 is not None:
            run.bold = bool(굵게)
        if 색상 is not None and RGBColor is not None:
            run.font.color.rgb = RGBColor(*색상)
    except Exception:
        pass


def _docx문단추가(doc, 텍스트="", 스타일=None, 크기=10.5, 굵게=False, 색상=None, 정렬=None, 앞간격=None, 뒤간격=None):
    p = doc.add_paragraph(style=스타일) if 스타일 else doc.add_paragraph()
    run = p.add_run(str(텍스트))
    _docx_한글폰트적용(run, 크기=크기, 굵게=굵게, 색상=색상)
    try:
        if 정렬 is not None:
            p.alignment = 정렬
        if 앞간격 is not None and Pt is not None:
            p.paragraph_format.space_before = Pt(앞간격)
        if 뒤간격 is not None and Pt is not None:
            p.paragraph_format.space_after = Pt(뒤간격)
        p.paragraph_format.line_spacing = 1.18
    except Exception:
        pass
    return p


def _docx불릿목록추가(doc, 문장목록):
    if not 문장목록:
        _docx문단추가(doc, "표시할 내용이 부족합니다.", 크기=10)
        return
    for 문장 in 문장목록:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(str(문장))
        _docx_한글폰트적용(run, 크기=10.2)
        try:
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
        except Exception:
            pass


def _docx표셀텍스트(cell, 텍스트, 굵게=False, 크기=9.5, 배경색=None):
    try:
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(텍스트))
        _docx_한글폰트적용(run, 크기=크기, 굵게=굵게)
        if WD_CELL_VERTICAL_ALIGNMENT is not None:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if 배경색:
            shading = cell._tc.get_or_add_tcPr()
            from docx.oxml import parse_xml
            shade = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="%s"/>' % 배경색)
            shading.append(shade)
    except Exception:
        try:
            cell.text = str(텍스트)
        except Exception:
            pass


def 월간투자리포트워드문서생성(리포트):
    """v5.14.4: 월간 투자 리포트를 수정 가능한 Word 문서(.docx)로 생성합니다."""
    if not DOCX_AVAILABLE or Document is None:
        return None, "python-docx 패키지가 설치되어 있지 않아 Word 문서를 생성할 수 없습니다. requirements.txt에 python-docx를 추가해 주세요."
    try:
        doc = Document()
        section = doc.sections[0]
        try:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        except Exception:
            pass

        # 기본 스타일 한글 글꼴
        try:
            normal = doc.styles["Normal"]
            normal.font.name = "맑은 고딕"
            normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
            normal.font.size = Pt(10)
        except Exception:
            pass

        기준년월 = str(리포트.get("기준년월", ""))
        작성시각 = 서울조회문자열(리포트.get("작성시각", 서울현재시각()), "%Y-%m-%d %H:%M")

        _docx문단추가(doc, "월간 투자 리포트", 크기=22, 굵게=True, 색상=(31, 78, 121), 정렬=WD_ALIGN_PARAGRAPH.CENTER if WD_ALIGN_PARAGRAPH else None, 뒤간격=4)
        _docx문단추가(doc, f"기준월: {기준년월}  |  작성 기준: {작성시각}", 크기=10, 색상=(90, 90, 90), 정렬=WD_ALIGN_PARAGRAPH.CENTER if WD_ALIGN_PARAGRAPH else None, 뒤간격=10)
        _docx문단추가(doc, "현재 계산된 포트폴리오, 월간 거래 흐름, 통합 자산 현황을 바탕으로 자동 작성된 초안입니다.", 크기=9.5, 색상=(90, 90, 90), 정렬=WD_ALIGN_PARAGRAPH.CENTER if WD_ALIGN_PARAGRAPH else None, 뒤간격=12)

        # 핵심 요약 표
        _docx문단추가(doc, "1. 핵심 지표 요약", 크기=14, 굵게=True, 색상=(31, 78, 121), 앞간격=6, 뒤간격=6)
        요약항목 = [
            ("총 평가액", 금액표시(리포트.get("총평가", 0))),
            ("총 손익", 손익문자열(리포트.get("총손익", 0))),
            ("총 수익률", 수익률문자열(리포트.get("총수익률", 0))),
            ("보유 종목 수", f"{int(리포트.get('보유종목수', 0))}개"),
            ("최대 비중 종목", f"{리포트.get('최대비중명', '-')} {float(리포트.get('최대비중', 0)):.1f}%"),
            ("현금성 비중", f"{float(리포트.get('현금성비중', 0)):.1f}%"),
        ]
        table = doc.add_table(rows=1, cols=2)
        try:
            table.style = "Table Grid"
            if WD_TABLE_ALIGNMENT is not None:
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
        except Exception:
            pass
        _docx표셀텍스트(table.rows[0].cells[0], "항목", 굵게=True, 배경색="D9EAF7")
        _docx표셀텍스트(table.rows[0].cells[1], "내용", 굵게=True, 배경색="D9EAF7")
        for 항목, 값 in 요약항목:
            row = table.add_row().cells
            _docx표셀텍스트(row[0], 항목, 굵게=True)
            _docx표셀텍스트(row[1], 값)

        섹션목록 = [
            ("2. 핵심 요약", 리포트.get("핵심요약", [])),
            ("3. 월간 거래 흐름", 리포트.get("월간거래문장", [])),
            ("4. 주요 종목 변화", 리포트.get("주요종목문장", [])),
            ("5. 리스크 상태", 리포트.get("리스크문장", [])),
            ("6. 다음 점검 포인트", 리포트.get("점검포인트", [])),
        ]
        for 제목, 문장목록 in 섹션목록:
            _docx문단추가(doc, 제목, 크기=14, 굵게=True, 색상=(31, 78, 121), 앞간격=10, 뒤간격=4)
            _docx불릿목록추가(doc, 문장목록)

        _docx문단추가(doc, "7. 종합 코멘트", 크기=14, 굵게=True, 색상=(31, 78, 121), 앞간격=10, 뒤간격=4)
        _docx문단추가(doc, 리포트.get("종합코멘트", ""), 크기=10.5, 뒤간격=8)

        자산군표 = 리포트.get("자산군표", pd.DataFrame())
        if isinstance(자산군표, pd.DataFrame) and not 자산군표.empty:
            _docx문단추가(doc, "8. 자산군 요약", 크기=14, 굵게=True, 색상=(31, 78, 121), 앞간격=10, 뒤간격=4)
            표시열 = [열 for 열 in ["자산군", "평가금액", "평가손익", "비중"] if 열 in 자산군표.columns]
            asset_table = doc.add_table(rows=1, cols=len(표시열))
            try:
                asset_table.style = "Table Grid"
            except Exception:
                pass
            for i, 열 in enumerate(표시열):
                _docx표셀텍스트(asset_table.rows[0].cells[i], 열, 굵게=True, 배경색="EAF2F8")
            for _, rowdata in 자산군표[표시열].head(12).iterrows():
                cells = asset_table.add_row().cells
                for i, 열 in enumerate(표시열):
                    값 = rowdata.get(열, "")
                    if 열 in ["평가금액"]:
                        값 = 금액표시(값)
                    elif 열 in ["평가손익"]:
                        값 = 손익문자열(값)
                    elif 열 == "비중":
                        값 = f"{float(값):.2f}%"
                    _docx표셀텍스트(cells[i], 값, 크기=9)

        _docx문단추가(doc, "참고", 크기=10.5, 굵게=True, 색상=(90, 90, 90), 앞간격=12, 뒤간격=2)
        _docx문단추가(doc, "본 문서는 투자 판단을 보조하기 위한 개인용 기록 초안이며, 매수·매도 권유가 아닙니다. 실제 투자 결정 전에는 최신 시세, 계좌 잔고, 세금, 수수료를 함께 확인해야 합니다.", 크기=9, 색상=(100, 100, 100))

        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio.getvalue(), ""
    except Exception as e:
        return None, f"Word 문서 생성 중 오류가 발생했습니다: {e}"

def 월간투자리포트초안UI(거래df=None, 계산포트폴리오=None, 보유포트폴리오=None, 통합자산표=None, 위험분석=None):
    st.markdown("### 월간 투자 리포트")
    st.caption("현재 계산된 포트폴리오, 월간 거래 흐름, 통합 자산 현황을 바탕으로 읽기 쉬운 보고서 본문을 자동 생성합니다.")

    기본년월 = 서울현재시각().strftime("%Y-%m")
    좌, 우 = st.columns([1.1, 4.0])
    with 좌:
        기준년월 = st.text_input("리포트 기준월", value=기본년월, key="monthly_report_base_month_v5143")
    with 우:
        st.caption("v5.14.4는 문장형 리포트를 Word 문서로 다운로드할 수 있도록 확장한 단계입니다. PDF는 Word 안정화 후 확장하는 방식이 가장 안전합니다.")

    리포트 = 월간투자리포트초안생성(
        거래df=거래df,
        계산포트폴리오=계산포트폴리오,
        보유포트폴리오=보유포트폴리오,
        통합자산표=통합자산표,
        위험분석=위험분석,
        기준년월=기준년월,
    )

    st.markdown(
        """
        <style>
        .monthly-report-wrap {border:1px solid rgba(148,163,184,.24); border-radius:24px; padding:1.35rem 1.45rem; background:rgba(15,23,42,.30); margin-top:.65rem; box-shadow:0 10px 28px rgba(0,0,0,.14);} 
        .monthly-report-title {font-size:1.28rem; font-weight:620; letter-spacing:-.025em; margin-bottom:.25rem; color:#f8fafc;}
        .monthly-report-sub {font-size:.89rem; color:#9ca3af; margin-bottom:1.05rem; line-height:1.45;}
        .monthly-report-card {border-left:4px solid #60a5fa; padding:.82rem .98rem; background:rgba(30,41,59,.38); border-radius:16px; min-height:122px;}
        .monthly-report-label {font-size:.78rem; color:#93c5fd; margin-bottom:.35rem;}
        .monthly-report-main {font-size:1.22rem; font-weight:620; letter-spacing:-.03em; line-height:1.25;}
        .monthly-report-desc {font-size:.86rem; color:#cbd5e1; line-height:1.45; margin-top:.35rem;}
        .monthly-report-section {font-size:1.04rem; font-weight:600; margin-top:1.15rem; margin-bottom:.35rem; color:#f8fafc;}
        .monthly-report-text {font-size:.97rem; line-height:1.74; color:#e5e7eb; word-break:keep-all;}
        .monthly-report-list {margin-top:.25rem; padding-left:1.25rem;}
        .monthly-report-list li {margin-bottom:.42rem;}
        .monthly-report-comment {font-size:.98rem; line-height:1.78; color:#e5e7eb; background:rgba(2,6,23,.26); border-radius:16px; padding:.95rem 1.05rem; word-break:keep-all;}
        </style>
        """,
        unsafe_allow_html=True,
    )

    st.markdown("<div class='monthly-report-wrap'>", unsafe_allow_html=True)
    st.markdown(f"<div class='monthly-report-title'>월간 투자 리포트 · {html.escape(str(리포트['기준년월']))}</div>", unsafe_allow_html=True)
    st.markdown(f"<div class='monthly-report-sub'>작성 기준 {서울조회문자열(리포트['작성시각'], '%Y-%m-%d %H:%M')} · 현재가 새로고침 결과 기준</div>", unsafe_allow_html=True)

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.markdown(f"<div class='monthly-report-card'><div class='monthly-report-label'>총 평가액</div><div class='monthly-report-main'>{금액표시(리포트['총평가'])}</div><div class='monthly-report-desc'>통합 자산 기준 평가 규모</div></div>", unsafe_allow_html=True)
    with c2:
        st.markdown(f"<div class='monthly-report-card'><div class='monthly-report-label'>총 손익</div><div class='monthly-report-main'>{손익문자열(리포트['총손익'])}</div><div class='monthly-report-desc'>평가 손익 중심 요약</div></div>", unsafe_allow_html=True)
    with c3:
        st.markdown(f"<div class='monthly-report-card'><div class='monthly-report-label'>수익률</div><div class='monthly-report-main'>{수익률문자열(리포트['총수익률'])}</div><div class='monthly-report-desc'>원금 대비 현재 성과</div></div>", unsafe_allow_html=True)
    with c4:
        st.markdown(f"<div class='monthly-report-card'><div class='monthly-report-label'>핵심 리스크</div><div class='monthly-report-main'>{html.escape(str(리포트['최대비중명']))}</div><div class='monthly-report-desc'>{리포트['최대비중']:.1f}% · 최대 비중 종목</div></div>", unsafe_allow_html=True)

    섹션목록 = [
        ("1. 핵심 요약", 리포트["핵심요약"]),
        ("2. 월간 거래 흐름", 리포트["월간거래문장"]),
        ("3. 주요 종목 변화", 리포트["주요종목문장"]),
        ("4. 리스크 상태", 리포트["리스크문장"]),
        ("5. 다음 점검 포인트", 리포트["점검포인트"]),
    ]
    for 제목, 문장목록 in 섹션목록:
        st.markdown(f"<div class='monthly-report-section'>{html.escape(제목)}</div>", unsafe_allow_html=True)
        if 문장목록:
            st.markdown("<div class='monthly-report-text'>" + _리포트HTML목록(문장목록) + "</div>", unsafe_allow_html=True)
        else:
            st.info("표시할 데이터가 부족합니다.")

    st.markdown("<div class='monthly-report-section'>6. 종합 코멘트</div>", unsafe_allow_html=True)
    st.markdown(f"<div class='monthly-report-comment'>{html.escape(str(리포트['종합코멘트']))}</div>", unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

    with st.expander("리포트 숫자 요약표 보기", expanded=False):
        표데이터프레임(index_1부터(리포트["요약표"]), use_container_width=True)
        자산군표 = 리포트.get("자산군표", pd.DataFrame())
        if 자산군표 is not None and not 자산군표.empty:
            자산군표시 = 자산군표.copy()
            표데이터프레임(
                index_1부터(자산군표시).style.format({
                    "평가금액": 안전정수포맷,
                    "평가손익": 손익문자열,
                    "비중": lambda x: 안전소수포맷(x, 2),
                }).map(손익색상, subset=["평가손익"]),
                use_container_width=True,
            )

    다운로드좌, 다운로드우 = st.columns(2)
    with 다운로드좌:
        st.download_button(
            "리포트 본문 텍스트 다운로드",
            data=리포트["본문"].encode("utf-8-sig"),
            file_name=f"월간_투자_리포트_{리포트['기준년월'].replace('-', '')}.txt",
            mime="text/plain",
            key="download_monthly_report_text_v5144",
            use_container_width=True,
        )
    with 다운로드우:
        워드바이트, 워드오류 = 월간투자리포트워드문서생성(리포트)
        if 워드바이트:
            st.download_button(
                "Word 리포트 다운로드",
                data=워드바이트,
                file_name=f"월간_투자_리포트_{리포트['기준년월'].replace('-', '')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_monthly_report_docx_v5144",
                use_container_width=True,
            )
        else:
            st.caption(워드오류)
    st.caption("PDF 리포트는 한글 폰트와 배포 환경 안정성을 확인한 뒤 다음 단계에서 확장하는 것을 권장합니다.")
    return 리포트

선택위젯키정리()
세션선택초기화()


def 시세관련캐시초기화():
    try:
        야후1분봉요약가져오기.clear()
        최근OHLCV가져오기.clear()
        최근시세요약가져오기.clear()
        실시간포함시세요약가져오기.clear()
        네이버국내현재가가져오기.clear()
        자산현재가정보.clear()
        종목현재가가져오기.clear()
        ETF현재가가져오기.clear()
        인덱스현재가가져오기.clear()
        자산과거가격가져오기.clear()
        시세스냅샷캐시.clear()
        포트폴리오계산캐시.clear()
        야후현재가요약가져오기.clear()
        네이버일별원자재표가져오기.clear()
        두바이유현재가가져오기.clear()
        파생주요지표가져오기.clear()
        네이버시장지표현재가가져오기.clear()
        네이버시장지표목록가져오기.clear()
        try:
            pykrx투자자별순매수.clear()
        except Exception:
            pass
    except Exception:
        st.cache_data.clear()


# -----------------------------------


# -----------------------------------
# 메인 화면 3섹터 구조
# -----------------------------------
st.markdown("---")

섹터목록 = ["주요 모니터링", "포트폴리오 현황", "분석 / 인사이트"]
섹터선택키 = "main_section_selector_v5106d"

# 이전 버전에서 저장된 선택값이 남아 있으면 화면이 표시되지 않을 수 있어 새 키와 유효성 검사를 함께 사용합니다.
if 섹터선택키 not in st.session_state or st.session_state.get(섹터선택키) not in 섹터목록:
    st.session_state[섹터선택키] = "주요 모니터링"

선택섹터 = st.session_state[섹터선택키]

버튼칸 = st.columns(3, gap="small")
for idx, 섹터명 in enumerate(섹터목록):
    with 버튼칸[idx]:
        if st.button(
            섹터명,
            key=f"main_section_btn_v5106d_{idx}",
            use_container_width=True,
            type="primary" if 선택섹터 == 섹터명 else "secondary",
        ):
            st.session_state[섹터선택키] = 섹터명
            st.rerun()

선택섹터 = st.session_state[섹터선택키]

if 선택섹터 == "주요 모니터링":
    # 상단: 주요 지수/대표 종목 모니터
    # -----------------------------------
    st.markdown("---")
    대시보드스타일적용()
    st.markdown(
        """
        <style>
        /* v5.13.3: 첨부 예시 스타일 - 박스형 배경 제거, 세로 라인 중심 심플 카드 */
        .simple-market-card {
            border: 0 !important;
            border-left: 4px solid #3b82f6 !important;
            border-radius: 0 !important;
            background: transparent !important;
            box-shadow: none !important;
            min-height: 118px !important;
            padding: 0 8px 0 14px !important;
            margin: 8px 0 22px 0 !important;
            gap: 0 !important;
        }
        .simple-market-card.up {border-left-color: #ef3b2d !important;}
        .simple-market-card.down {border-left-color: #3b82f6 !important;}
        .simple-market-card.flat {border-left-color: #94a3b8 !important;}
        .simple-market-label {display: none !important;}
        .simple-market-title {
            font-size: 1.28rem !important;
            font-weight: 520 !important;
            line-height: 1.12 !important;
            margin: 0 0 6px 0 !important;
            min-height: auto !important;
            -webkit-line-clamp: 2 !important;
        }
        .simple-market-price {
            font-size: 1.92rem !important;
            font-weight: 580 !important;
            line-height: 1.05 !important;
            letter-spacing: -0.04em !important;
            margin: 0 0 4px 0 !important;
        }
        .simple-market-delta {
            font-size: 1.04rem !important;
            font-weight: 560 !important;
            line-height: 1.08 !important;
            margin: 0 !important;
            min-height: auto !important;
            align-items: center !important;
            white-space: nowrap !important;
        }
        .simple-market-delta::after {
            content: "";
            display: inline-block;
            width: 8px;
            height: 8px;
            margin-left: 8px;
            border-radius: 999px;
            background: #22c55e;
            flex: 0 0 auto;
        }
        .simple-market-holdings {
            margin-top: 8px !important;
            background: transparent !important;
            border: 0 !important;
            border-radius: 0 !important;
            padding: 0 !important;
            font-size: 0.84rem !important;
        }
        .simple-market-meta {display: none !important;}
        </style>
        """,
        unsafe_allow_html=True,
    )
    st.markdown('<div class="top-monitor-title">주요 지수 및 대표 종목 모니터</div>', unsafe_allow_html=True)

    if "monitor_realtime_mode_v1" not in st.session_state:
        st.session_state["monitor_realtime_mode_v1"] = False
    if "manual_price_refresh_ts_v1" not in st.session_state:
        st.session_state["manual_price_refresh_ts_v1"] = 서울현재시각ISO()

    모니터헤더칸1, 모니터헤더칸2 = st.columns([1.25, 8.75], gap="small")
    with 모니터헤더칸1:
        새로고침클릭 = st.button("시세 새로고침", key="refresh_monitor_btn_v851g", use_container_width=True)
    if 새로고침클릭:
        시세관련캐시초기화()
        st.session_state["monitor_realtime_mode_v1"] = True
        st.session_state["manual_price_refresh_ts_v1"] = 서울현재시각ISO()
        st.session_state["price_refresh_token_v51"] = st.session_state.get("price_refresh_token_v51", 0) + 1
        st.session_state["price_snapshot_map_v1"] = {}
        st.session_state["price_snapshot_token_v1"] = st.session_state["price_refresh_token_v51"]
        st.rerun()

    with 모니터헤더칸2:
        조회모드문구 = "Naver 실시간 우선 반영 · 비교 전일 종가 기준" if st.session_state.get("monitor_realtime_mode_v1", False) else "전일 종가 기준 표시"
        조회일시문자 = 서울조회문자열(st.session_state["manual_price_refresh_ts_v1"], 포맷=f"조회 %Y-%m-%d %H:%M · {조회모드문구}")
        st.markdown(
            f"<div class='top-monitor-time'>{조회일시문자}</div>",
            unsafe_allow_html=True,
        )

    if "show_monitor_add_form_v53" not in st.session_state:
        st.session_state["show_monitor_add_form_v53"] = False

    대시보드기준거래 = 현재거래이력가져오기().copy()
    if st.session_state.get("monitor_realtime_mode_v1", False):
        시세스냅샷세션반영(대시보드기준거래, refresh_token=st.session_state.get("price_refresh_token_v51", 0))
    else:
        st.session_state["price_snapshot_map_v1"] = {}
    모니터자산목록 = 주요모니터자산구성(대시보드기준거래)
    보유정보사전 = 대시보드보유정보사전(대시보드기준거래)


    if 모바일여부():
        렌더목록 = 모니터자산목록.copy() + [("__ADD__", {"코드": ""}, "추가")]
        카드열수 = 2 if len(렌더목록) > 2 else len(렌더목록)
        for row_start in range(0, len(렌더목록), 카드열수):
            현재행 = 렌더목록[row_start:row_start + 카드열수]
            cols = st.columns(len(현재행))
            for col, (자산명, 자산정보, 구분라벨) in zip(cols, 현재행):
                with col:
                    if 자산명 == "__ADD__":
                        if 모니터추가카드버튼(key=f"monitor_add_card_mobile_{row_start}"):
                            st.session_state["show_monitor_add_form_v53"] = not st.session_state.get("show_monitor_add_form_v53", False)
                            st.rerun()
                    else:
                        정보 = 모니터표시시세요약(자산명, 자산정보, refresh_token=st.session_state.get("price_refresh_token_v51", 0))
                        종목코드 = str(자산정보['코드']).zfill(6)
                        보유정보문자 = 보유정보사전.get(종목코드, "") if 구분라벨 == "보유 종목" else ""
                        st.markdown(
                            심플카드HTML(
                                자산명,
                                정보.get("현재가"),
                                정보.get("전일대비"),
                                정보.get("등락률"),
                                보조라벨=구분라벨,
                                하단메모="",
                                보유정보문자=보유정보문자,
                            ),
                            unsafe_allow_html=True,
                        )
    else:
        # 데스크톱: 코스피·코스닥 이후 ETF/개별종목을 투자원금 큰 순서로 6열 심플 배치
        카드열수 = 6
        for row_start in range(0, len(모니터자산목록), 카드열수):
            현재행 = 모니터자산목록[row_start:row_start + 카드열수]
            cols = st.columns(카드열수, gap="small")
            for idx in range(카드열수):
                with cols[idx]:
                    if idx < len(현재행):
                        자산명, 자산정보, 구분라벨 = 현재행[idx]
                        정보 = 모니터표시시세요약(자산명, 자산정보, refresh_token=st.session_state.get("price_refresh_token_v51", 0))
                        종목코드 = str(자산정보['코드']).zfill(6)
                        보유정보문자 = 보유정보사전.get(종목코드, "") if 구분라벨 == "보유 종목" else ""
                        st.markdown(
                            심플카드HTML(
                                자산명,
                                정보.get("현재가"),
                                정보.get("전일대비"),
                                정보.get("등락률"),
                                보조라벨=구분라벨,
                                하단메모="",
                                보유정보문자=보유정보문자,
                            ),
                            unsafe_allow_html=True,
                        )

    모니터실패건수 = sum(1 for 자산명, 자산정보, _ in 모니터자산목록 if 자산현재가정보(자산명, 자산정보, refresh_token=st.session_state.get("price_refresh_token_v51", 0)).get("현재가") is None)
    if 모니터실패건수 > 0:
        st.warning(f"시세를 불러오지 못한 자산이 {모니터실패건수}개 있습니다.")

    # 수급 기능은 데이터 소스 재설계 전까지 보류합니다.
    # 투자자수급섹션표시(refresh_token=st.session_state.get("price_refresh_token_v51", 0))

    # -----------------------------------
    # 주요 경제지표
    # -----------------------------------
    st.markdown("---")
    st.subheader("주요 지표")
    시장지표df = 네이버시장지표목록가져오기()

    if 시장지표df.empty:
        st.warning("시장지표 데이터를 불러오지 못했습니다.")
    else:
        지표행목록 = list(시장지표df.iterrows())
        지표열수 = 2 if 모바일여부() else min(7, len(지표행목록))
        for row_start in range(0, len(지표행목록), 지표열수):
            cols = st.columns(지표열수, gap="small")
            for col, (_, row) in zip(cols, 지표행목록[row_start:row_start + 지표열수]):
                with col:
                    st.markdown(
                        심플카드HTML(
                            row["지표"],
                            row.get("현재값"),
                            row.get("전일대비"),
                            row.get("등락률"),
                            보조라벨="",
                            하단메모="",
                        ),
                        unsafe_allow_html=True,
                    )


    # -----------------------------------
    # -----------------------------------
    # 포트폴리오 입력/수정
    # -----------------------------------

with st.sidebar.expander("거래이력 관리", expanded=False):
    st.markdown("#### 포트폴리오 거래이력")
    st.caption("거래이력 단독 파일 또는 거래이력+비주식자산 통합 엑셀을 업로드할 수 있습니다.")

    if "trade_history_df_v22" not in st.session_state:
        현재거래이력가져오기()
    if 선택섹터 in ["포트폴리오 현황", "분석 / 인사이트"] or st.session_state.get("show_trade_editor_v5106a", False):
        자동백업일일실행(st.session_state.get("trade_history_df_v22", pd.DataFrame()))
    else:
        st.caption("주요 모니터링 화면에서는 자동백업 점검을 건너뜁니다.")

    저장파일명 = f"거래이력_저장_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
    현재거래건수 = len(st.session_state.get("trade_history_df_v22", pd.DataFrame()))
    최근업로드메타 = 최근업로드메타불러오기()
    if 최근업로드메타:
        st.caption(f"최근 업로드: {최근업로드메타.get('파일명', '-')} · {최근업로드메타.get('건수', 0)}건")
    else:
        st.caption(f"현재 거래이력: {현재거래건수}건")

    st.markdown("##### 1. 파일 불러오기")
    업로드파일 = st.file_uploader(
        "CSV · JSON · Excel(xlsx/xls) · 통합자산관리 엑셀",
        type=["csv", "json", "xlsx", "xls"],
        key="trade_history_file_uploader_v24",
        label_visibility="visible"
    )

    if st.button(
        "업로드 파일 반영",
        disabled=업로드파일 is None,
        key="apply_upload_btn_v26",
        use_container_width=True,
    ):
        try:
            불러온df = 업로드파일에서거래이력읽기(업로드파일)
            보정df = 거래이력자동보정(불러온df.copy())
            반영df, 변경됨, 저장성공, 저장메시지 = 거래이력세션반영(보정df, 저장강제=True, 자동저장허용=True)

            비주식반영건수 = 0
            비주식저장성공, 비주식저장메시지 = True, "해당 없음"
            if 업로드파일 is not None and 통합엑셀업로드여부(업로드파일):
                비주식df = 업로드파일에서비주식자산읽기(업로드파일)
                if 비주식df is not None:
                    비주식저장성공, 비주식저장메시지 = IRP비주식자산저장(비주식df)
                    비주식반영건수 = len(비주식df)

            최근업로드저장성공, 최근업로드저장메시지 = 최근업로드거래이력저장(반영df, 업로드파일.name if 업로드파일 is not None else "")
            st.session_state["trade_history_source_v1"] = "latest_uploaded"
            st.session_state["trade_history_latest_upload_name_v1"] = 업로드파일.name if 업로드파일 is not None else ""
            st.session_state["trade_history_latest_upload_time_v1"] = 서울현재시각ISO()
            if 비주식반영건수 > 0:
                st.success(f"통합 엑셀을 반영했습니다. 거래이력 {len(반영df)}건 · 비주식자산 {비주식반영건수}건")
            elif 변경됨:
                st.success(f"거래이력을 불러왔습니다. ({len(반영df)}건)")
            else:
                st.info("업로드 내용이 현재 거래이력과 동일합니다.")
            if not 저장성공:
                st.warning(f"자동저장 실패: {저장메시지}")
            if not 비주식저장성공:
                st.warning(f"비주식자산 저장 실패: {비주식저장메시지}")
            if not 최근업로드저장성공:
                st.warning(f"최근 업로드본 저장 실패: {최근업로드저장메시지}")

            시세관련캐시초기화()
            st.session_state["manual_price_refresh_ts_v1"] = 서울현재시각ISO()
            st.session_state["price_refresh_token_v51"] = st.session_state.get("price_refresh_token_v51", 0) + 1

            with st.expander("업로드 진단 정보", expanded=False):
                st.write("업로드 파일명:", 업로드파일.name)
                st.write("인식된 거래이력 컬럼:", list(불러온df.columns))
                if 비주식반영건수 > 0:
                    st.write("비주식자산 반영 건수:", 비주식반영건수)
                if not 반영df.empty:
                    표데이터프레임(반영df.head(10), use_container_width=True)
            st.rerun()
        except Exception as e:
            st.error(f"불러오기 중 오류가 발생했습니다: {e}")

    st.markdown("##### 2. 저장·백업")
    st.download_button(
        "엑셀 저장",
        data=현재거래내역엑셀저장바이트(st.session_state["trade_history_df_v22"]),
        file_name=저장파일명,
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        key="save_trade_history_btn_v26",
        use_container_width=True,
    )
    st.download_button(
        "JSON 백업",
        data=json.dumps(거래이력JSON변환(st.session_state["trade_history_df_v22"]), ensure_ascii=False, indent=2),
        file_name=f"거래이력_백업_{datetime.now().strftime('%Y%m%d_%H%M')}.json",
        mime="application/json",
        key="save_trade_history_json_btn_v26",
        use_container_width=True,
    )

    자동백업관리열기 = st.toggle(
        "자동백업 관리 열기",
        value=False,
        key="show_backup_manager_v5106b",
        help="백업 목록 조회는 파일 확인이 필요해 필요할 때만 엽니다.",
    )
    if 자동백업관리열기:
        자동백업관리UI(
            st.session_state.get("trade_history_df_v22", pd.DataFrame()),
            portfolio_df=st.session_state.get("portfolio_cache_df_v1", pd.DataFrame()),
            holding_df=st.session_state.get("portfolio_holding_cache_df_v1", pd.DataFrame()),
        )

    st.markdown("##### 3. 직접 편집")
    st.caption("기본은 접힌 상태입니다. 직접 수정이 필요할 때만 열어 첫 로딩 속도를 줄입니다.")

    편집대상거래이력 = 거래이력입력창정렬(
        st.session_state.get("trade_history_editor_df_v1", st.session_state["trade_history_df_v22"])
    )

    직접편집열기 = st.toggle(
        "거래이력 표 직접 편집 열기",
        value=False,
        key="show_trade_editor_v5106a",
        help="많은 행을 수정할 때는 엑셀에서 편집 후 업로드하는 방식이 더 빠릅니다.",
    )

    if 직접편집열기:
        수정포트폴리오 = st.data_editor(
            편집대상거래이력,
            num_rows="dynamic",
            use_container_width=True,
            hide_index=True,
            disabled=[],
            column_order=["종목코드", "종목명", "거래일자", "거래구분", "거래수량", "거래단가", "운용사", "비고"],
            column_config={
                "종목코드": st.column_config.TextColumn("종목코드", help="6자리 종목코드를 입력하면 종목명이 자동 보정됩니다."),
                "종목명": st.column_config.TextColumn("종목명", help="종목명을 입력하면 가능한 경우 종목코드가 자동 보정됩니다."),
                "거래일자": st.column_config.DateColumn("거래일자", format="YYYY-MM-DD"),
                "거래구분": st.column_config.SelectboxColumn("거래구분", options=["매수", "매도"], required=False),
                "거래수량": st.column_config.NumberColumn("거래수량", min_value=0, step=1, format="%d"),
                "거래단가": st.column_config.NumberColumn("거래단가", min_value=0, step=1, format="%d"),
                "운용사": st.column_config.TextColumn("운용사", help="예: 신한은행 IRP, 미래에셋증권"),
                "비고": st.column_config.TextColumn("비고"),
            },
            key="trade_editor_v25",
        )
    else:
        수정포트폴리오 = 편집대상거래이력.copy()
        st.caption(f"직접 편집 표를 숨겼습니다. 현재 거래이력 {len(수정포트폴리오)}건 기준으로 계산합니다.")

    # 포트폴리오/분석 화면에서만 무거운 계산을 실행합니다.
    if 선택섹터 in ["포트폴리오 현황", "분석 / 인사이트"]:
        최적화결과 = 거래이력편집반영최적화(수정포트폴리오)
        수정포트폴리오 = 최적화결과["편집df"]
        거래이력변경됨 = 최적화결과["거래이력변경됨"]
        자동저장성공 = 최적화결과["자동저장성공"]
        자동저장메시지 = 최적화결과["자동저장메시지"]
        계산용거래이력 = 최적화결과["계산용거래이력"]
        통합점검표 = 최적화결과["통합점검표"]

        if 거래이력변경됨:
            st.caption("입력 내용은 화면에 즉시 반영됩니다. 변경 전 상태는 자동백업에 저장되며, 필요하면 상단 저장 또는 자동백업 관리에서 복원할 수 있습니다.")

        if not 통합점검표.empty and "점검항목" in 통합점검표.columns:
            불일치검증표 = 통합점검표[통합점검표["점검항목"] == "종목코드-종목명 불일치"].copy()
        else:
            불일치검증표 = pd.DataFrame()
        if not 불일치검증표.empty:
            st.error(f'종목코드와 종목명이 서로 맞지 않는 입력이 {len(불일치검증표)}건 있습니다. 자동으로 다른 종목으로 바꾸지 않고 그대로 표시했습니다.')
        if 통합점검표.empty:
            st.success("거래이력 입력 점검 결과: 현재 확인된 형식 오류가 없습니다.")
        else:
            st.warning(f"거래이력 입력 점검 결과: {len(통합점검표)}건의 확인 사항이 있습니다.")
            with st.expander("입력 검증 상세 보기", expanded=False):
                표데이터프레임(index_1부터(통합점검표), use_container_width=True)
    else:
        최적화결과 = None
        st.caption("주요 모니터링 화면에서는 포트폴리오 상세 계산을 생략해 첫 로딩을 줄입니다.")


    # -----------------------------------

if 선택섹터 == "포트폴리오 현황":
    # 포트폴리오 계산 결과
    계산포트폴리오 = 최적화결과["계산포트폴리오"]
    보유계산포트폴리오 = 최적화결과["보유계산포트폴리오"]
    보유종목옵션 = 최적화결과["보유종목옵션"]

    과잉매도종목 = 계산포트폴리오[계산포트폴리오["과잉매도수량"] > 0]
    if not 과잉매도종목.empty:
        st.warning("보유수량보다 많이 매도한 거래가 있습니다. 거래이력을 확인해 주세요.")

    st.markdown("---")
    st.subheader("포트폴리오 현황")

    if 계산포트폴리오.empty:
        st.warning("포트폴리오 데이터를 계산할 수 없습니다.")
    else:
        표시대상포트폴리오 = 보유계산포트폴리오.copy()
        정상평가행 = 표시대상포트폴리오[표시대상포트폴리오["데이터상태"] == "정상"].copy()

        총투자원금 = 정상평가행["투자원금"].sum()
        총평가금액 = 정상평가행["평가금액"].sum()
        총평가손익 = 정상평가행["평가손익"].sum()
        총실현손익 = 계산포트폴리오["실현손익"].sum()
        총수익률 = (총평가손익 / 총투자원금 * 100) if 총투자원금 != 0 else 0

        조회실패건수 = (표시대상포트폴리오["데이터상태"] != "정상").sum()
        if 표시대상포트폴리오.empty:
            st.info("현재 보유수량이 0보다 큰 종목이 없습니다. 아래 거래이력을 확인해 주세요.")

        if 조회실패건수 > 0:
            st.warning(f"현재가 조회 실패 종목 {조회실패건수}건은 평가금액/비중 계산에서 제외했습니다.")

        청산종목표 = 계산포트폴리오[계산포트폴리오["보유수량"] <= 0].copy()
        if not 청산종목표.empty:
            with st.expander(f"청산 또는 보유 0주 종목 보기 ({len(청산종목표)}건)", expanded=False):
                청산표시 = 청산종목표[["종목코드", "종목명", "총매수수량", "총매도수량", "보유수량", "실현손익", "최근거래일자"]].copy()
                청산표시 = 청산표시.rename(columns={"총매수수량": "총 매수수량", "총매도수량": "총 매도수량", "최근거래일자": "최근 거래일자"})
                표데이터프레임(index_1부터(청산표시), use_container_width=True)

        요약정보 = 포트폴리오요약지표생성(계산포트폴리오, 표시대상포트폴리오)
        포트폴리오요약카드표시(요약정보)

        st.caption("포트폴리오 요약은 현재 보유 종목 기준으로 자동 계산되며, 현재가 조회 실패 종목은 평가금액·비중 계산에서 제외됩니다.")

        st.markdown("---")
        IRP비주식자산df = IRP비주식자산편집UI()
        통합자산표 = 통합자산현황UI(보유계산포트폴리오, IRP비주식자산df)

        st.markdown("---")
        위험분석결과_v514 = 포트폴리오리스크분석UI(보유계산포트폴리오, 통합자산표)
        포트폴리오종합인사이트UI(보유계산포트폴리오, 통합자산표, 위험분석결과_v514)

        st.markdown("---")
        월간투자리포트초안UI(수정포트폴리오, 계산포트폴리오, 보유계산포트폴리오, 통합자산표, 위험분석결과_v514)

        st.markdown("---")
        월간수익률리포트UI(수정포트폴리오, 계산포트폴리오)
        종목별누적성적표UI(수정포트폴리오, 계산포트폴리오)

        포트폴리오표시 = 표시대상포트폴리오[["종목코드", "종목명", "최초매수일자", "최근거래일자", "총매수수량", "총매도수량", "보유수량", "매입평균단가", "현재가", "투자원금", "평가금액", "평가손익", "실현손익", "수익률", "현재비중", "과잉매도수량", "데이터상태"]].copy()
        포트폴리오표시 = 포트폴리오표시.rename(columns={"매입평균단가": "매입 평균단가", "총매수수량": "총 매수수량", "총매도수량": "총 매도수량", "최초매수일자": "최초 매수일자", "최근거래일자": "최근 거래일자", "과잉매도수량": "과잉 매도수량"})
        포트폴리오표시 = 포트폴리오표_컬럼선택(포트폴리오표시)
        포트폴리오표시 = index_1부터(포트폴리오표시)

        if 모바일여부():
            모바일형식사전 = {}
            if "보유수량" in 포트폴리오표시.columns:
                모바일형식사전["보유수량"] = 안전정수포맷
            if "현재가" in 포트폴리오표시.columns:
                모바일형식사전["현재가"] = 안전정수포맷
            if "평가금액" in 포트폴리오표시.columns:
                모바일형식사전["평가금액"] = 안전정수포맷
            if "수익률" in 포트폴리오표시.columns:
                모바일형식사전["수익률"] = 수익률문자열
            모바일스타일 = 포트폴리오표시.style.format(모바일형식사전)
            if "수익률" in 포트폴리오표시.columns:
                모바일스타일 = 모바일스타일.map(수익률색상, subset=["수익률"])
            표데이터프레임(모바일스타일, use_container_width=True)
        else:
            표데이터프레임(
                포트폴리오표시.style.format({
                    "총 매수수량": 안전정수포맷,
                    "총 매도수량": 안전정수포맷,
                    "보유수량": 안전정수포맷,
                    "과잉 매도수량": 안전정수포맷,
                    "매입 평균단가": 안전정수포맷,
                    "현재가": 안전정수포맷,
                    "투자원금": 안전정수포맷,
                    "평가금액": 안전정수포맷,
                    "평가손익": 손익문자열,
                    "실현손익": 손익문자열,
                    "수익률": 수익률문자열,
                    "현재비중": lambda x: 안전소수포맷(x, 2),
                }).map(손익색상, subset=["평가손익", "실현손익"]).map(수익률색상, subset=["수익률"]),
                use_container_width=True,
            )
        st.markdown("### 포트폴리오 거래 원장 조회")
        st.caption("입력 원장과 같은 데이터를 누적보유수량 기준으로 정렬·필터해서 보는 조회용 표입니다. 직접 수정은 위 거래 이력 입력 표에서 진행해 주세요.")
        전체거래표 = 거래원장조회용빈행제거(종목거래이력표생성(수정포트폴리오))
        if 전체거래표.empty:
            st.info("표시할 거래기록이 없습니다.")
        else:
            with st.expander("거래 원장 조회 열기", expanded=False):
                조회대상거래표 = 전체거래표.copy()
                필터칸1, 필터칸2, 필터칸3 = st.columns(3)
                with 필터칸1:
                    종목옵션 = ["전체"] + sorted([x for x in 조회대상거래표["종목명"].dropna().astype(str).unique().tolist() if x])
                    선택종목명 = st.selectbox("종목 필터", 종목옵션, index=0, key="ledger_filter_asset_v1")
                with 필터칸2:
                    거래구분옵션 = ["전체"] + sorted([x for x in 조회대상거래표["거래구분"].dropna().astype(str).unique().tolist() if x])
                    선택거래구분 = st.selectbox("거래구분 필터", 거래구분옵션, index=0, key="ledger_filter_type_v1")
                with 필터칸3:
                    운용사옵션 = ["전체"] + sorted([x for x in 조회대상거래표["운용사"].dropna().astype(str).unique().tolist() if x])
                    선택운용사 = st.selectbox("운용사 필터", 운용사옵션, index=0, key="ledger_filter_account_v1")

                if 선택종목명 != "전체":
                    조회대상거래표 = 조회대상거래표[조회대상거래표["종목명"] == 선택종목명].copy()
                if 선택거래구분 != "전체":
                    조회대상거래표 = 조회대상거래표[조회대상거래표["거래구분"] == 선택거래구분].copy()
                if 선택운용사 != "전체":
                    조회대상거래표 = 조회대상거래표[조회대상거래표["운용사"] == 선택운용사].copy()

                st.caption(f"조회 결과 {len(조회대상거래표)}건")
                표데이터프레임(거래기록표시용서식(index_1부터(조회대상거래표)), use_container_width=True)

        오류행 = 계산포트폴리오[(계산포트폴리오["과잉매도수량"] > 0) | (계산포트폴리오["데이터상태"] != "정상")]
        if not 오류행.empty:
            st.warning("일부 종목에 과잉 매도 입력 또는 현재가 조회 실패가 있습니다. 아래 현황표의 '과잉 매도수량', '데이터상태'를 확인해 주세요.")

        현재가실패표 = 계산포트폴리오[계산포트폴리오["데이터상태"] != "정상"][ ["종목코드", "종목명", "데이터상태"] ].copy()
        if not 현재가실패표.empty:
            st.error(f"현재가 조회 실패 종목이 {len(현재가실패표)}개 있습니다. 종목코드와 장중/휴장 여부, 네트워크 상태를 확인해 주세요.")
            with st.expander("현재가 조회 실패 종목 보기", expanded=False):
                표데이터프레임(index_1부터(현재가실패표), use_container_width=True)

        비중그래프칸, 비중요약칸 = st.columns([1.45, 0.85], gap="large")
        with 비중그래프칸:
            st.plotly_chart(
                비중그래프(계산포트폴리오),
                use_container_width=True,
                config={"displaylogo": False, "responsive": True},
            )
        with 비중요약칸:
            비중요약표 = 계산포트폴리오.copy()
            비중요약표 = 비중요약표[["종목명", "현재비중", "평가금액"]].copy()
            비중요약표["현재비중"] = pd.to_numeric(비중요약표["현재비중"], errors="coerce").fillna(0)
            비중요약표["평가금액"] = pd.to_numeric(비중요약표["평가금액"], errors="coerce").fillna(0)
            비중요약표 = 비중요약표[비중요약표["평가금액"] > 0].sort_values(["현재비중", "평가금액"], ascending=[False, False]).reset_index(drop=True)

            if not 비중요약표.empty:
                최대행 = 비중요약표.iloc[0]
                st.markdown(
                    f"""
                    <div class="ratio-summary-card">
                        <div class="ratio-summary-title">최대 비중 종목</div>
                        <div class="ratio-summary-main">{최대행['종목명']}</div>
                        <div class="ratio-summary-sub">비중 {최대행['현재비중']:.2f}% · 평가금액 {금액표시(최대행['평가금액'])}</div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
                표시용비중요약표 = 비중요약표.copy()
                표시용비중요약표["현재비중"] = 표시용비중요약표["현재비중"].map(lambda x: f"{x:.2f}%")
                표시용비중요약표["평가금액"] = 표시용비중요약표["평가금액"].map(금액표시)

                표시용비중요약표_styled = index_1부터(표시용비중요약표).style.set_properties(
                    subset=["현재비중", "평가금액"],
                    **{"text-align": "right", "font-variant-numeric": "tabular-nums", "font-feature-settings": '"tnum"'}
                ).set_properties(
                    subset=["종목명"],
                    **{"text-align": "left"}
                )

                표데이터프레임(표시용비중요약표_styled, use_container_width=True, hide_index=False)
            else:
                st.info("비중 요약을 표시할 보유 종목이 없습니다.")


    # -----------------------------------

if 선택섹터 == "분석 / 인사이트":
    계산포트폴리오 = 최적화결과["계산포트폴리오"]
    보유계산포트폴리오 = 최적화결과["보유계산포트폴리오"]
    보유종목옵션 = 최적화결과["보유종목옵션"]

    # 보유 종목 개별 분석
    # -----------------------------------
    st.markdown("---")
    st.subheader("보유 종목 개별 분석")
    st.caption("아래 분석 종목 목록은 현재 거래이력으로 계산된 보유수량 기준으로 자동 생성됩니다. 따라서 거래이력·포트폴리오 현황·개별분석이 항상 같은 기준을 사용합니다.")

    try:
        if not 보유종목옵션:
            st.info("현재 보유 중인 종목이 없어 개별 분석 항목을 표시하지 않습니다.")
        else:
            옵션코드목록 = [항목["종목코드"] for 항목 in 보유종목옵션]

            # 개별분석 선택 상태는 단일 키로만 관리해 시세 새로고침 직후
            # 드롭다운 표시값과 실제 분석 기준이 엇갈리지 않도록 한다.
            현재선택코드 = st.session_state.get("holding_asset_choice_v45", "")
            if 현재선택코드 not in 옵션코드목록:
                현재선택코드 = 옵션코드목록[0]
                st.session_state["holding_asset_choice_v45"] = 현재선택코드

            선택종목코드 = st.selectbox(
                "분석할 보유 종목 선택",
                옵션코드목록,
                format_func=lambda 코드: next((항목["표시"] for 항목 in 보유종목옵션 if 항목["종목코드"] == 코드), 코드),
                key="holding_asset_choice_v45",
            )

            # 옵션 재계산/새로고침 직후에도 실제 분석 대상은 현재 위젯 값과
            # 옵션 목록을 기준으로 다시 한 번 강제 정합성 확인
            if 선택종목코드 not in 옵션코드목록:
                선택종목코드 = 옵션코드목록[0]
                st.session_state["holding_asset_choice_v45"] = 선택종목코드

            선택행 = next((항목 for 항목 in 보유종목옵션 if 항목["종목코드"] == 선택종목코드), None)
            선택종목명 = 선택행["종목명"] if 선택행 else 종목코드기준종목명(선택종목코드)
            선택종목구분 = 종목구분판단(선택종목코드, 선택종목명)
            가격데이터 = 자산과거가격가져오기(선택종목구분, 선택종목코드, 개월수=6)
            선택포트폴리오행 = 보유계산포트폴리오[보유계산포트폴리오["종목코드"] == 선택종목코드]

            st.caption(f"현재 선택 종목: {선택종목명} · 구분: {선택종목구분} · 코드: {선택종목코드}")

            if not 선택포트폴리오행.empty:
                선택행데이터 = 선택포트폴리오행.iloc[0]
                정보칸1, 정보칸2, 정보칸3 = st.columns(3)
                정보칸1.metric("보유수량", 숫자표시(선택행데이터.get("보유수량")))
                정보칸2.metric("매입 평균단가", 금액표시(선택행데이터.get("매입평균단가")))
                정보칸3.metric("현재 비중", 비율표시(선택행데이터.get("현재비중")))

            if 가격데이터.empty:
                st.warning("가격 데이터를 불러오지 못했습니다.")
            else:
                최신값 = 가격데이터.iloc[-1]
                이전값 = 가격데이터.iloc[-2] if len(가격데이터) >= 2 else 최신값
                가격변화 = 최신값["종가"] - 이전값["종가"]
                등락률 = (가격변화 / 이전값["종가"] * 100) if 이전값["종가"] != 0 else 0

                신호결과 = 신호판정계산(가격데이터)
                자동코멘트결과 = 고급매매코멘트생성(선택종목명, 가격데이터, 선택포트폴리오행)
                모델분석 = 차트분석문구(선택종목명, 가격데이터)
                종목거래표 = 종목거래이력표생성(수정포트폴리오, 선택종목코드)
                상세해설문장 = 보유종목상세해설생성(선택종목명, 가격데이터, 선택포트폴리오행)

                보유분석탭1, 보유분석탭2, 보유분석탭3, 보유분석탭4 = st.tabs(["개요", "상세 해설", "기술적 분석", "거래기록"])

                with 보유분석탭1:
                    if 모바일여부():
                        행1_1, 행1_2 = st.columns(2)
                        행1_1.metric("현재가", 금액표시(최신값["종가"]), f"{가격변화:,.0f}원")
                        행1_2.metric("등락률", 비율표시(등락률))
                        행2_1, 행2_2 = st.columns(2)
                        행2_1.metric("거래량", 숫자표시(최신값["거래량"]))
                        행2_2.metric("RSI(14)", f"{최신값['RSI(14)']:.2f}" if pd.notna(최신값["RSI(14)"]) else "-")
                    else:
                        칸1, 칸2, 칸3, 칸4 = st.columns(4)
                        칸1.metric("현재가", 금액표시(최신값["종가"]), f"{가격변화:,.0f}원")
                        칸2.metric("등락률", 비율표시(등락률))
                        칸3.metric("거래량", 숫자표시(최신값["거래량"]))
                        칸4.metric("RSI(14)", f"{최신값['RSI(14)']:.2f}" if pd.notna(최신값["RSI(14)"]) else "-")

                    st.markdown(자동판정배지HTML(자동코멘트결과["판정"], 자동코멘트결과["실행"], 자동코멘트결과["강도"]), unsafe_allow_html=True)
                    st.info(자동코멘트결과["핵심문구"])

                    요약칸1, 요약칸2, 요약칸3, 요약칸4 = st.columns(4)
                    요약칸1.metric("자동 판정", 자동코멘트결과["판정"])
                    요약칸2.metric("추세 판정", 자동코멘트결과["추세판정"])
                    요약칸3.metric("실행 방향", 자동코멘트결과["실행"])
                    요약칸4.metric("복합 총점", f"{자동코멘트결과['총점']:.2f}점" if 자동코멘트결과.get("총점") is not None else "-")

                    st.markdown("#### 핵심 판독 요약")
                    판독칸1, 판독칸2, 판독칸3, 판독칸4 = st.columns(4)
                    판독칸1.metric("가격 위치", 자동코멘트결과.get("위치판정", "-"))
                    판독칸2.metric("RSI 판정", 자동코멘트결과.get("RSI판정", "-"))
                    판독칸3.metric("거래량 판정", 자동코멘트결과.get("거래량판정", "-"))
                    판독칸4.metric("단순 신호", 신호결과.get("종합신호", "-"))
                    st.caption(f"간단 실행 의견: {신호결과.get('실행의견', '-')}")

                    if not 선택포트폴리오행.empty:
                        포지션행 = 선택포트폴리오행.iloc[0]
                        추가칸1, 추가칸2, 추가칸3, 추가칸4 = st.columns(4)
                        추가칸1.metric("평가금액", 금액표시(포지션행.get("평가금액")))
                        추가칸2.metric("평가손익", 금액표시(포지션행.get("평가손익")))
                        추가칸3.metric("보유 수익률", 비율표시(포지션행.get("수익률")))
                        최근거래일 = "-"
                        if not 종목거래표.empty and "거래일자" in 종목거래표.columns:
                            try:
                                최근거래일 = str(pd.to_datetime(종목거래표["거래일자"], errors="coerce").max().date())
                            except Exception:
                                최근거래일 = "-"
                        추가칸4.metric("최근 거래일", 최근거래일)

                    st.plotly_chart(가격그래프(가격데이터, f"{선택종목명} 주가 추이"), use_container_width=True, config={"displaylogo": False, "responsive": True})

                    최근거래요약 = pd.DataFrame()
                    if not 종목거래표.empty:
                        최근거래요약 = 종목거래표.copy().tail(5)
                        표시컬럼 = [c for c in ["거래일자", "거래구분", "거래수량", "거래단가", "거래금액", "누적보유수량"] if c in 최근거래요약.columns]
                        최근거래요약 = 최근거래요약[표시컬럼].copy()
                        if "거래단가" in 최근거래요약.columns:
                            최근거래요약["거래단가"] = 최근거래요약["거래단가"].map(금액표시)
                        if "거래금액" in 최근거래요약.columns:
                            최근거래요약["거래금액"] = 최근거래요약["거래금액"].map(금액표시)
                        if "거래수량" in 최근거래요약.columns:
                            최근거래요약["거래수량"] = 최근거래요약["거래수량"].map(숫자표시)
                        if "누적보유수량" in 최근거래요약.columns:
                            최근거래요약["누적보유수량"] = 최근거래요약["누적보유수량"].map(숫자표시)

                    개요왼쪽, 개요오른쪽 = st.columns([1.5, 1.1])
                    with 개요왼쪽:
                        st.markdown("#### 최근 거래 흐름")
                        if 최근거래요약.empty:
                            st.info("이 종목의 최근 거래 요약을 표시할 데이터가 없습니다.")
                        else:
                            표데이터프레임(index_1부터(최근거래요약.reset_index(drop=True)), use_container_width=True)
                    with 개요오른쪽:
                        st.markdown("#### 빠른 해석")
                        st.markdown(f"- 현재 자동 판정은 **{자동코멘트결과['판정']}**입니다.")
                        st.markdown(f"- 추세는 **{자동코멘트결과.get('추세판정', '-')}**, 가격 위치는 **{자동코멘트결과.get('위치판정', '-')}**입니다.")
                        st.markdown(f"- RSI는 **{자동코멘트결과.get('RSI판정', '-')}**, 거래량은 **{자동코멘트결과.get('거래량판정', '-')}**입니다.")
                        st.markdown(f"- 현재 실행 방향은 **{자동코멘트결과['실행']}** 쪽으로 해석됩니다.")

                    개요체크표 = pd.concat([
                        자동코멘트결과["근거표"],
                        신호결과["체크표"]
                    ], axis=0, ignore_index=True)
                    with st.expander("세부 체크표 보기", expanded=False):
                        표데이터프레임(index_1부터(개요체크표), use_container_width=True)

                with 보유분석탭2:
                    st.markdown("#### 자동 매수·매도 코멘트")
                    st.success(f"**{자동코멘트결과['판정']} / {자동코멘트결과['실행']}**")
                    st.write(자동코멘트결과["핵심문구"])
                    for 문장 in 자동코멘트결과["세부코멘트"]:
                        st.markdown(f"- {문장}")
                    with st.expander("자동 판단 근거 보기", expanded=True):
                        for 항목 in 자동코멘트결과["근거"]:
                            st.markdown(f"- {항목}")
                        표데이터프레임(index_1부터(자동코멘트결과["근거표"]), use_container_width=True)
                        st.warning(자동코멘트결과["위험문구"])

                    with st.expander("자동 판정 기준 설명", expanded=True):
                        st.markdown("- 자동 판정은 **추세·가격 위치·RSI·거래량·당일 흐름**을 함께 점수화해 계산합니다.")
                        st.markdown("- 따라서 한 항목이 같아도 다른 항목이 변하면 최종 판정이 바뀔 수 있습니다.")
                        표데이터프레임(index_1부터(자동코멘트결과["기준표"]), use_container_width=True)

                    st.markdown("#### 차트 해설")
                    for 문장 in 상세해설문장:
                        st.markdown(f"- {문장}")

                    st.markdown("#### 모델별 해석 비교")
                    모델탭1, 모델탭2, 모델탭3 = st.tabs(["ChatGPT", "Gemini 스타일", "Claude"])
                    with 모델탭1:
                        분석카드표시(모델분석["ChatGPT"])
                    with 모델탭2:
                        분석카드표시(모델분석["Gemini"])
                    with 모델탭3:
                        분석카드표시(모델분석["Claude"])

                with 보유분석탭3:
                    기술진단 = 기술분석진단계산(가격데이터)
                    기술차트탭0, 기술차트탭1 = st.tabs(["실전 요약", "체크표"])

                    with 기술차트탭0:
                        실전칸1, 실전칸2, 실전칸3, 실전칸4 = st.columns(4)
                        실전칸1.metric("추세 배열", 기술진단.get("추세배열", "-"))
                        실전칸2.metric("20일 지지", 금액표시(기술진단.get("지지")))
                        실전칸3.metric("20일 저항", 금액표시(기술진단.get("저항")))
                        실전칸4.metric("단순 신호", 신호결과.get("종합신호", "-"))

                        st.markdown("#### 실전 해석 요약")
                        for 문장 in 기술진단.get("요약문장", []):
                            st.markdown(f"- {문장}")

                        요약왼쪽, 요약오른쪽 = st.columns([1.05, 1.15])
                        with 요약왼쪽:
                            st.markdown("#### 핵심 체크")
                            표데이터프레임(index_1부터(기술진단["핵심표"]), use_container_width=True, hide_index=False)
                        with 요약오른쪽:
                            st.markdown("#### 지지·저항 및 기준선")
                            레벨표표시 = 기술진단["레벨표"].copy()
                            if not 레벨표표시.empty:
                                레벨표표시["가격"] = 레벨표표시["가격"].map(금액표시)
                            표데이터프레임(index_1부터(레벨표표시), use_container_width=True, hide_index=False)

                        with st.expander("실전 체크포인트", expanded=True):
                            st.markdown("1. **상승 배열**이면 20일선 이탈 여부를 먼저 확인합니다.")
                            st.markdown("2. **지지선 근처**에서는 분할 접근, **저항선 근처**에서는 추격 매수보다 확인이 우선입니다.")
                            st.markdown("3. RSI가 과열인데 거래량까지 급증하면 단기 과열 가능성을 함께 봅니다.")
                            st.markdown("4. 가격이 20일선과 60일선을 동시에 하회하면 단기보다 중기 흐름 훼손 여부를 더 중요하게 봅니다.")
                        st.info("캔들차트는 반복 오류로 인해 이번 버전에서 제외했습니다. 개요 탭의 주가 추이와 아래 체크표를 함께 참고해 주세요.")

                    with 기술차트탭1:
                        기술체크 = pd.DataFrame([
                            {"항목": "최근 기준일", "값": str(pd.to_datetime(가격데이터.index[-1]).date())},
                            {"항목": "5일 이동평균", "값": 숫자표시(최신값.get("5일평균"), 2)},
                            {"항목": "20일 이동평균", "값": 숫자표시(최신값.get("20일평균"), 2)},
                            {"항목": "60일 이동평균", "값": 숫자표시(최신값.get("60일평균"), 2)},
                            {"항목": "120일 이동평균", "값": 숫자표시(최신값.get("120일평균"), 2)},
                            {"항목": "RSI(14)", "값": 숫자표시(최신값.get("RSI(14)"), 2)},
                            {"항목": "최근 거래량", "값": 숫자표시(최신값.get("거래량"))},
                            {"항목": "20일 고가", "값": 숫자표시(가격데이터.tail(20)["고가"].max(), 2)},
                            {"항목": "20일 저가", "값": 숫자표시(가격데이터.tail(20)["저가"].min(), 2)},
                            {"항목": "60일 고가", "값": 숫자표시(가격데이터.tail(60)["고가"].max(), 2)},
                            {"항목": "60일 저가", "값": 숫자표시(가격데이터.tail(60)["저가"].min(), 2)},
                        ])
                        표데이터프레임(index_1부터(기술체크), use_container_width=True)
                        with st.expander("신호 체크표 함께 보기", expanded=False):
                            표데이터프레임(index_1부터(신호결과["체크표"]), use_container_width=True)

                with 보유분석탭4:
                    st.markdown("#### 선택 종목 매수·매도 전체 기록")
                    if 종목거래표.empty:
                        st.info("선택 종목의 거래기록이 없습니다.")
                    else:
                        표데이터프레임(거래기록표시용서식(index_1부터(종목거래표)), use_container_width=True)
    except Exception as e:
        st.error(f"보유 종목 개별 분석 영역 오류: {e}")



st.markdown("---")
st.markdown('개발자 조현웅 <a href="mailto:hwcho@me.com">hwcho@me.com</a>', unsafe_allow_html=True)

if "trade_history_df_v22" in st.session_state and "trade_history_last_saved_fingerprint_v43" not in st.session_state:
    st.session_state["trade_history_last_saved_fingerprint_v43"] = 거래이력비교지문(st.session_state["trade_history_df_v22"])
if "trade_history_df_v22" in st.session_state and "trade_history_last_calc_fingerprint_v43" not in st.session_state:
    st.session_state["trade_history_last_calc_fingerprint_v43"] = 거래이력비교지문(st.session_state["trade_history_df_v22"])
